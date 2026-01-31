"""
================================================================================
AASHTO 1993 Flexible Pavement Design - Streamlit Application (Version 4)
================================================================================
แอปพลิเคชันสำหรับออกแบบ Flexible Pavement ตามวิธี AASHTO 1993
ปรับปรุงตามมาตรฐานกรมทางหลวง (DOH Thailand)

Features:
- Material database ตามมาตรฐาน ทล.
- Step-by-step thickness calculation (หาความหนาแต่ละชั้น)
- Drainage coefficient default = 1.0
- ปรับปรุงรูปภาพตัดขวางให้มีรูปแบบเดียวกับ Rigid Pavement
- [NEW v4] แบ่งชั้นย่อย AC ได้ 3 ชั้น: Wearing, Binder, Base Course
- [NEW v4] Dropdown ความหนามาตรฐาน ทล. สำหรับชั้นย่อย AC
- [NEW v4] แสดง W₁₈ ด้วย font ใหญ่สีฟ้า
- [NEW v4] ปุ่ม help สำหรับ Drainage Coefficient
- [NEW v4] Export Report เป็นภาษาไทย

Author: Civil Engineering Department
Version: 4.0
================================================================================
"""

import streamlit as st
import numpy as np
import json
from scipy.optimize import brentq
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.font_manager as fm
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ================================================================================
# PAGE CONFIGURATION
# ================================================================================

st.set_page_config(
    page_title="AASHTO 1993 Flexible Pavement Design (DOH)",
    page_icon="🛣️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ================================================================================
# MATERIAL DATABASE - ตามมาตรฐานกรมทางหลวง (DOH Thailand)
# ================================================================================

MATERIALS = {
    # ============ ชั้นผิวทาง (Surface) ============
    "ผิวทางลาดยาง AC": {
        "layer_coeff": 0.40,
        "drainage_coeff": 1.0,
        "mr_psi": 362500,
        "mr_mpa": 2500,
        "layer_type": "surface",
        "color": "#1C1C1C",  # สีดำ (Black)
        "short_name": "AC",
        "english_name": "Asphalt Concrete"
    },
    "ผิวทางลาดยาง PMA": {
        "layer_coeff": 0.40,
        "drainage_coeff": 1.0,
        "mr_psi": 536500,
        "mr_mpa": 3700,
        "layer_type": "surface",
        "color": "#2C2C2C",  # สีดำเข้ม (Dark Black)
        "short_name": "PMA",
        "english_name": "Polymer Modified Asphalt"
    },
    
    # ============ ชั้นพื้นทาง (Base) ============
    "พื้นทางซีเมนต์ CTB": {
        "layer_coeff": 0.18,
        "drainage_coeff": 1.0,
        "mr_psi": 174000,
        "mr_mpa": 1200,
        "layer_type": "base",
        "color": "#78909C",  # สีเทา (Gray)
        "short_name": "CTB",
        "english_name": "Cement Treated Base"
    },
    "พื้นทางหินคลุกผสมซีเมนต์ UCS 24.5 ksc.": {
        "layer_coeff": 0.15,
        "drainage_coeff": 1.0,
        "mr_psi": 123250,
        "mr_mpa": 850,
        "layer_type": "base",
        "color": "#607D8B",  # สีเทาเข้ม
        "short_name": "MOD.CRB",
        "english_name": "Mod.Crushed Rock Base"
    },
    "พื้นทางหินคลุก CBR 80%": {
        "layer_coeff": 0.13,
        "drainage_coeff": 1.0,
        "mr_psi": 50750,
        "mr_mpa": 350,
        "layer_type": "base",
        "color": "#795548",  # สีน้ำตาล
        "short_name": "CAB",
        "english_name": "Crushed Rock Base"
    },
    "พื้นทางดินซีเมนต์ UCS 17.5 ksc.": {
        "layer_coeff": 0.13,
        "drainage_coeff": 1.0,
        "mr_psi": 50750,
        "mr_mpa": 350,
        "layer_type": "base",
        "color": "#8D6E63",  # สีน้ำตาลอ่อน
        "short_name": "SCB",
        "english_name": "Soil Cement Base"
    },
    "พื้นทางวัสดุหมุนเวียน (Recycling)": {
        "layer_coeff": 0.15,
        "drainage_coeff": 1.0,
        "mr_psi": 123250,
        "mr_mpa": 850,
        "layer_type": "base",
        "color": "#5D4037",  # สีน้ำตาลเข้ม
        "short_name": "RAP",
        "english_name": "Recycled Asphalt Pavement"
    },
    
    # ============ ชั้นรองพื้นทาง (Subbase) - วัสดุมวลรวม ============
    "รองพื้นทางวัสดุมวลรวม CBR 25%": {
        "layer_coeff": 0.10,
        "drainage_coeff": 1.0,
        "mr_psi": 21750,
        "mr_mpa": 150,
        "layer_type": "subbase",
        "color": "#FFB74D",  # สีส้มอ่อน (Light Orange) - วัสดุมวลรวม
        "short_name": "GSB",
        "english_name": "Aggregate Subbase"
    },
    
    # ============ วัสดุคัดเลือก (Selected Material) - ทราย ============
    "วัสดุคัดเลือก ก": {
        "layer_coeff": 0.08,
        "drainage_coeff": 1.0,
        "mr_psi": 14504,
        "mr_mpa": 100,
        "layer_type": "selected",
        "color": "#FFF176",  # สีเหลือง (Yellow) - ทราย
        "short_name": "SM-A",
        "english_name": "Selected Material"
    },
    
    # ============ ไม่ใช้วัสดุ (Skip layer) ============
    "ไม่ใช้วัสดุคัดเลือก (ใช้ดินทางทรพ)": {
        "layer_coeff": 0.00,
        "drainage_coeff": 1.0,
        "mr_psi": 0,
        "mr_mpa": 0,
        "layer_type": "none",
        "color": "#D7CCC8",
        "short_name": "NONE",
        "english_name": "None"
    }
}

# ================================================================================
# RELIABILITY TABLE: Zr VALUES
# ================================================================================

RELIABILITY_ZR = {
    50: -0.000,
    60: -0.253,
    70: -0.524,
    75: -0.674,
    80: -0.841,
    85: -1.037,
    90: -1.282,
    91: -1.340,
    92: -1.405,
    93: -1.476,
    94: -1.555,
    95: -1.645,
    96: -1.751,
    97: -1.881,
    98: -2.054,
    99: -2.327,
    99.9: -3.090
}

# ================================================================================
# CORE CALCULATION FUNCTIONS
# ================================================================================

def aashto_1993_equation(SN: float, W18: float, Zr: float, So: float, 
                          delta_psi: float, Mr: float) -> float:
    """
    AASHTO 1993 Main Design Equation for Flexible Pavement
    
    สมการออกแบบหลักของ AASHTO 1993
    
    log₁₀(W₁₈) = Zr×So + 9.36×log₁₀(SN+1) - 0.20 
                 + log₁₀(ΔPSI/(4.2-1.5)) / (0.4 + 1094/(SN+1)^5.19)
                 + 2.32×log₁₀(Mr) - 8.07
    """
    log_W18 = np.log10(W18)
    
    term1 = Zr * So
    term2 = 9.36 * np.log10(SN + 1) - 0.20
    
    numerator = np.log10(delta_psi / (4.2 - 1.5))
    denominator = 0.4 + (1094 / ((SN + 1) ** 5.19))
    term3 = numerator / denominator
    
    term4 = 2.32 * np.log10(Mr) - 8.07
    
    right_side = term1 + term2 + term3 + term4
    
    return right_side - log_W18


def calculate_sn_for_layer(W18: float, Zr: float, So: float, 
                            delta_psi: float, Mr: float) -> float:
    """
    Calculate required SN for a given subgrade/layer modulus
    
    คำนวณค่า SN ที่ต้องการสำหรับค่า Mr ที่กำหนด
    """
    def f(SN):
        return aashto_1993_equation(SN, W18, Zr, So, delta_psi, Mr)
    
    try:
        SN_required = brentq(f, 0.01, 25.0, xtol=1e-6, maxiter=100)
        return round(SN_required, 2)
    except ValueError:
        return None


def calculate_w18_supported(SN: float, Zr: float, So: float, 
                            delta_psi: float, Mr: float) -> float:
    """
    Calculate W18 that can be supported by a given SN
    
    คำนวณค่า W₁₈ ที่โครงสร้างรองรับได้จากค่า SN ที่ออกแบบ
    """
    term1 = Zr * So
    term2 = 9.36 * np.log10(SN + 1) - 0.20
    
    numerator = np.log10(delta_psi / (4.2 - 1.5))
    denominator = 0.4 + (1094 / ((SN + 1) ** 5.19))
    term3 = numerator / denominator
    
    term4 = 2.32 * np.log10(Mr) - 8.07
    
    log_W18 = term1 + term2 + term3 + term4
    
    W18_supported = 10 ** log_W18
    
    return W18_supported


def calculate_layer_thicknesses(W18: float, Zr: float, So: float, delta_psi: float,
                                 subgrade_mr: float, layers: list, ac_sublayers: dict = None) -> dict:
    """
    Calculate minimum thickness for each layer using AASHTO 1993 method
    
    คำนวณความหนาขั้นต่ำของแต่ละชั้น ตามวิธี AASHTO 1993
    
    Parameters:
    - ac_sublayers: dict ข้อมูลชั้นย่อย AC (wearing, binder, base, total) หรือ None
    """
    results = {
        'layers': [],
        'sn_values': [],
        'subgrade_mr': subgrade_mr,
        'total_sn_required': None,
        'total_sn_provided': 0,
        'ac_sublayers': ac_sublayers  # เก็บข้อมูลชั้นย่อยไว้
    }
    
    # Get active layers (exclude "ไม่ใช้วัสดุ")
    active_layers = [l for l in layers if l['material'] != "ไม่ใช้วัสดุคัดเลือก (ใช้ดินทางทรพ)"]
    
    if not active_layers:
        return results
    
    num_layers = len(active_layers)
    sn_values = []
    
    for i in range(num_layers):
        if i == num_layers - 1:
            mr_below = subgrade_mr
        else:
            mat_below = MATERIALS[active_layers[i + 1]['material']]
            mr_below = mat_below['mr_psi']
        
        sn_i = calculate_sn_for_layer(W18, Zr, So, delta_psi, mr_below)
        sn_values.append({
            'layer_index': i + 1,
            'mr_below': mr_below,
            'sn_required': sn_i
        })
    
    results['sn_values'] = sn_values
    results['total_sn_required'] = calculate_sn_for_layer(W18, Zr, So, delta_psi, subgrade_mr)
    
    cumulative_sn = 0
    
    for i, layer in enumerate(active_layers):
        mat = MATERIALS[layer['material']]
        # ใช้ค่า a จากผู้ใช้ถ้ามี ไม่งั้นใช้จากฐานข้อมูล
        a_i = layer.get('layer_coeff', mat['layer_coeff'])
        m_i = layer.get('drainage_coeff', 1.0)
        
        sn_required_at_layer = sn_values[i]['sn_required'] if sn_values[i]['sn_required'] else 0
        
        if a_i > 0 and m_i > 0:
            remaining_sn = max(0, sn_required_at_layer - cumulative_sn)
            min_thickness_inch = remaining_sn / (a_i * m_i)
            min_thickness_cm = min_thickness_inch * 2.54
        else:
            min_thickness_inch = 0
            min_thickness_cm = 0
        
        design_thickness_cm = layer['thickness_cm']
        design_thickness_inch = design_thickness_cm / 2.54
        
        sn_contribution = a_i * design_thickness_inch * m_i
        cumulative_sn += sn_contribution
        
        is_ok = design_thickness_cm >= min_thickness_cm
        
        # เพิ่มข้อมูลชั้นย่อย AC สำหรับชั้นที่ 1 (ถ้ามี)
        layer_ac_sublayers = None
        if i == 0 and ac_sublayers is not None:
            layer_ac_sublayers = ac_sublayers
        
        results['layers'].append({
            'layer_no': i + 1,
            'material': layer['material'],
            'short_name': mat['short_name'],
            'english_name': mat.get('english_name', mat['short_name']),
            'mr_psi': mat['mr_psi'],
            'mr_mpa': mat['mr_mpa'],
            'a_i': a_i,
            'm_i': m_i,
            'sn_required_at_layer': sn_required_at_layer,
            'min_thickness_inch': round(min_thickness_inch, 2),
            'min_thickness_cm': round(min_thickness_cm, 1),
            'design_thickness_cm': design_thickness_cm,
            'design_thickness_inch': round(design_thickness_inch, 2),
            'sn_contribution': round(sn_contribution, 4),
            'cumulative_sn': round(cumulative_sn, 2),
            'is_ok': is_ok,
            'color': mat['color'],
            'ac_sublayers': layer_ac_sublayers
        })
    
    results['total_sn_provided'] = round(cumulative_sn, 2)
    
    return results


def check_design(sn_required: float, sn_provided: float) -> dict:
    """Check if design is adequate"""
    if sn_required is None:
        return {
            'status': 'ERROR',
            'passed': False,
            'message': 'Cannot calculate SN_required',
            'safety_margin': None
        }
    
    safety_margin = sn_provided - sn_required
    passed = sn_provided >= sn_required
    
    return {
        'status': 'OK' if passed else 'NG',
        'passed': passed,
        'safety_margin': round(safety_margin, 2),
        'message': f"SN_provided ({sn_provided:.2f}) {'≥' if passed else '<'} SN_required ({sn_required:.2f})"
    }



# ================================================================================
# VISUALIZATION FUNCTIONS - ปรับปรุงใหม่ตามรูปแบบ Rigid Pavement
# ================================================================================

def plot_pavement_section(layers_result: list, subgrade_mr: float = None, 
                          subgrade_cbr: float = None) -> plt.Figure:
    """
    Draw vertical pavement section diagram - รูปแบบเดียวกับ Rigid Pavement
    (ขนาดคงที่ สัดส่วนพอดี)
    
    Layout:
    - ซ้าย: ชื่อวัสดุ (Material name)
    - กลาง: ความหนา (Thickness)
    - ขวา: E = xxx MPa
    - ขวาสุด: เส้นแสดงความหนารวม (Total thickness arrow)
    - ล่าง: กรอบ Total Pavement Thickness
    """
    
    plt.rcParams['font.family'] = 'DejaVu Sans'
    
    if not layers_result:
        fig, ax = plt.subplots(figsize=(12, 8))
        ax.text(0.5, 0.5, 'No layers defined', ha='center', va='center', fontsize=14)
        ax.axis('off')
        return fig
    
    # กรองเฉพาะชั้นที่มีความหนา > 0
    valid_layers = [l for l in layers_result if l.get('design_thickness_cm', 0) > 0]
    if not valid_layers:
        fig, ax = plt.subplots(figsize=(12, 8))
        ax.text(0.5, 0.5, 'No valid layers', ha='center', va='center', fontsize=14)
        ax.axis('off')
        return fig
    
    # คำนวณความหนารวม
    total_thickness = sum([l['design_thickness_cm'] for l in valid_layers])
    
    # ========== สร้าง figure ขนาดคงที่ (เหมือน Rigid Pavement) ==========
    fig, ax = plt.subplots(figsize=(12, 8))
    
    # พารามิเตอร์การวาด (เหมือน Rigid Pavement)
    width = 3
    x_center = 6
    x_start = x_center - width / 2
    
    # ========== คำนวณความสูงแสดงผล (minimum height เพื่อให้อ่านได้) ==========
    min_display_height = 8
    display_heights = [max(l['design_thickness_cm'], min_display_height) for l in valid_layers]
    total_display = sum(display_heights)
    
    # ชั้นที่ต้องใช้ข้อความสีขาว (พื้นหลังเข้ม)
    dark_colors = ['#1C1C1C', '#2C2C2C', '#78909C', '#607D8B', '#795548', 
                   '#8D6E63', '#5D4037', '#6D4C41', '#455A64']
    
    # วาดแต่ละชั้น
    y_current = total_display
    
    for i, layer in enumerate(valid_layers):
        thickness = layer['design_thickness_cm']
        display_h = display_heights[i]
        color = layer.get('color', '#CCCCCC')
        english_name = layer.get('english_name', layer.get('short_name', f'Layer {i+1}'))
        e_mpa = layer.get('mr_mpa', 0)
        
        # วาดสี่เหลี่ยม
        y_bottom = y_current - display_h
        rect = mpatches.Rectangle(
            (x_start, y_bottom), width, display_h,
            linewidth=2,
            edgecolor='black',
            facecolor=color
        )
        ax.add_patch(rect)
        
        # ตำแหน่งกลางชั้น
        y_center_pos = y_bottom + display_h / 2
        
        # กำหนดสีข้อความ (ขาวสำหรับพื้นเข้ม)
        text_color = 'white' if color in dark_colors else 'black'
        
        # ข้อความกลาง: ความหนา
        ax.text(x_center, y_center_pos, f'{thickness:.0f} cm',
                ha='center', va='center', fontsize=16, fontweight='bold', color=text_color)
        
        # ซ้าย: ชื่อวัสดุ (English)
        ax.text(x_start - 0.5, y_center_pos, english_name,
                ha='right', va='center', fontsize=14, fontweight='bold', color='black')
        
        # ขวา: E = xxx MPa
        if e_mpa and e_mpa > 0:
            ax.text(x_start + width + 0.5, y_center_pos, f'E = {e_mpa:,.0f} MPa',
                    ha='left', va='center', fontsize=12, color='#0066CC')
        
        y_current = y_bottom
    
    # ========== เส้นแสดงความหนารวม (ลูกศรสองหัว) ==========
    ax.annotate('', xy=(x_start + width + 3.5, total_display), xytext=(x_start + width + 3.5, 0),
                arrowprops=dict(arrowstyle='<->', color='red', lw=2))
    
    # ข้อความ Total
    ax.text(x_start + width + 4, total_display / 2, f'Total\n{total_thickness:.0f} cm',
            ha='left', va='center', fontsize=14, color='red', fontweight='bold')
    
    # ========== ตั้งค่าขอบเขต (คงที่เหมือน Rigid Pavement) ==========
    margin = 10
    ax.set_xlim(0, 14)
    ax.set_ylim(-margin, total_display + margin)
    ax.axis('off')
    
    # Title
    ax.set_title('Pavement Structure', fontsize=20, fontweight='bold', pad=20)
    
    # กรอบ Total Pavement Thickness (ล่าง)
    ax.text(x_center, -margin + 4, f'Total Pavement Thickness: {total_thickness:.0f} cm',
            ha='center', va='center', fontsize=15, fontweight='bold',
            bbox=dict(boxstyle='round', facecolor='lightyellow', alpha=0.9, edgecolor='orange'))
    
    plt.tight_layout()
    return fig


def plot_pavement_section_thai(layers_result: list, subgrade_mr: float = None,
                                subgrade_cbr: float = None) -> plt.Figure:
    """
    Draw vertical pavement section diagram - รูปแบบภาษาไทย
    (ขนาดคงที่ สัดส่วนพอดี เหมือน Rigid Pavement)
    รองรับการแสดงชั้นย่อย AC (Wearing, Binder, Base Course)
    
    Layout:
    - ซ้าย: ชื่อวัสดุ (ภาษาไทย)
    - กลาง: ความหนา (Thickness)
    - ขวา: E = xxx MPa
    - ขวาสุด: เส้นแสดงความหนารวม
    - ล่าง: กรอบ ความหนารวมโครงสร้างชั้นทาง
    """
    
    # ตั้งค่า Thai font
    thai_font_path = '/usr/share/fonts/truetype/tlwg/Garuda.ttf'
    thai_font_bold_path = '/usr/share/fonts/truetype/tlwg/Garuda-Bold.ttf'
    try:
        thai_font = fm.FontProperties(fname=thai_font_path)
        thai_font_bold = fm.FontProperties(fname=thai_font_bold_path)
    except:
        try:
            # Fallback to .otf fonts
            thai_font = fm.FontProperties(fname='/usr/share/fonts/opentype/tlwg/Loma.otf')
            thai_font_bold = fm.FontProperties(fname='/usr/share/fonts/opentype/tlwg/Loma-Bold.otf')
        except:
            thai_font = fm.FontProperties()
            thai_font_bold = fm.FontProperties(weight='bold')
    
    plt.rcParams['font.family'] = 'DejaVu Sans'
    
    if not layers_result:
        fig, ax = plt.subplots(figsize=(12, 8))
        ax.text(0.5, 0.5, 'ไม่มีข้อมูลชั้นทาง', ha='center', va='center', 
                fontsize=14, fontproperties=thai_font)
        ax.axis('off')
        return fig
    
    # กรองเฉพาะชั้นที่มีความหนา > 0
    valid_layers = [l for l in layers_result if l.get('design_thickness_cm', 0) > 0]
    if not valid_layers:
        fig, ax = plt.subplots(figsize=(12, 8))
        ax.text(0.5, 0.5, 'ไม่มีชั้นทางที่ถูกต้อง', ha='center', va='center', 
                fontsize=14, fontproperties=thai_font)
        ax.axis('off')
        return fig
    
    # ========== ขยายชั้น AC ออกเป็นชั้นย่อย (ถ้ามี) ==========
    expanded_layers = []
    for layer in valid_layers:
        ac_sublayers = layer.get('ac_sublayers', None)
        if ac_sublayers is not None and layer['layer_no'] == 1:
            # แบ่งชั้น AC ออกเป็น 3 ชั้นย่อย
            # สีสำหรับชั้นย่อย AC (ไล่เฉด)
            sublayer_colors = {
                'wearing': '#1C1C1C',   # ดำเข้ม
                'binder': '#333333',    # เทาดำ
                'base': '#4A4A4A'       # เทากลาง
            }
            sublayer_names = {
                'wearing': 'Wearing Course (ผิวทาง)',
                'binder': 'Binder Course (ยึดเกาะ)',
                'base': 'Base Course (รองผิว)'
            }
            
            # เพิ่ม Wearing Course (ถ้าความหนา > 0)
            if ac_sublayers['wearing'] > 0:
                expanded_layers.append({
                    'design_thickness_cm': ac_sublayers['wearing'],
                    'material': sublayer_names['wearing'],
                    'short_name': 'WC',
                    'color': sublayer_colors['wearing'],
                    'mr_mpa': layer['mr_mpa'],
                    'is_sublayer': True
                })
            # เพิ่ม Binder Course (ถ้าความหนา > 0)
            if ac_sublayers['binder'] > 0:
                expanded_layers.append({
                    'design_thickness_cm': ac_sublayers['binder'],
                    'material': sublayer_names['binder'],
                    'short_name': 'BC',
                    'color': sublayer_colors['binder'],
                    'mr_mpa': layer['mr_mpa'],
                    'is_sublayer': True
                })
            # เพิ่ม Base Course (ถ้าความหนา > 0)
            if ac_sublayers['base'] > 0:
                expanded_layers.append({
                    'design_thickness_cm': ac_sublayers['base'],
                    'material': sublayer_names['base'],
                    'short_name': 'ABC',
                    'color': sublayer_colors['base'],
                    'mr_mpa': layer['mr_mpa'],
                    'is_sublayer': True
                })
        else:
            expanded_layers.append(layer)
    
    # คำนวณความหนารวม
    total_thickness = sum([l['design_thickness_cm'] for l in expanded_layers])
    
    # ========== สร้าง figure ขนาดคงที่ (เหมือน Rigid Pavement) ==========
    fig, ax = plt.subplots(figsize=(12, 9))
    
    # พารามิเตอร์การวาด - เลื่อนขวาเล็กน้อยเพื่อเว้นที่สำหรับชื่อไทย
    width = 3
    x_center = 7
    x_start = x_center - width / 2
    
    # ========== คำนวณความสูงแสดงผล (minimum height เพื่อให้อ่านได้) ==========
    min_display_height = 6  # ลดลงเล็กน้อยเพื่อให้มีที่ว่างมากขึ้นสำหรับชั้นย่อย
    display_heights = [max(l['design_thickness_cm'], min_display_height) for l in expanded_layers]
    total_display = sum(display_heights)
    
    # ชั้นที่ต้องใช้ข้อความสีขาว (พื้นหลังเข้ม)
    dark_colors = ['#1C1C1C', '#2C2C2C', '#333333', '#4A4A4A', '#78909C', '#607D8B', '#795548', 
                   '#8D6E63', '#5D4037', '#6D4C41', '#455A64']
    
    # วาดแต่ละชั้น
    y_current = total_display
    
    for i, layer in enumerate(expanded_layers):
        thickness = layer['design_thickness_cm']
        display_h = display_heights[i]
        color = layer.get('color', '#CCCCCC')
        thai_name = layer.get('material', layer.get('short_name', f'ชั้นที่ {i+1}'))
        e_mpa = layer.get('mr_mpa', 0)
        is_sublayer = layer.get('is_sublayer', False)
        
        # วาดสี่เหลี่ยม
        y_bottom = y_current - display_h
        
        # ถ้าเป็นชั้นย่อย ใช้เส้นประ
        line_style = '--' if is_sublayer else '-'
        line_width = 1 if is_sublayer else 2
        
        rect = mpatches.Rectangle(
            (x_start, y_bottom), width, display_h,
            linewidth=line_width,
            linestyle=line_style,
            edgecolor='black',
            facecolor=color
        )
        ax.add_patch(rect)
        
        # ตำแหน่งกลางชั้น
        y_center_pos = y_bottom + display_h / 2
        
        # กำหนดสีข้อความ
        text_color = 'white' if color in dark_colors else 'black'
        
        # ข้อความกลาง: ความหนา
        fontsize = 14 if is_sublayer else 16
        ax.text(x_center, y_center_pos, f'{thickness:.0f} cm',
                ha='center', va='center', fontsize=fontsize, fontweight='bold', color=text_color)
        
        # ซ้าย: ชื่อวัสดุ (Thai)
        name_fontsize = 12 if is_sublayer else 14
        ax.text(x_start - 0.5, y_center_pos, thai_name,
                ha='right', va='center', fontsize=name_fontsize, fontweight='bold',
                fontproperties=thai_font_bold, color='black')
        
        # ขวา: E = xxx MPa (แสดงเฉพาะชั้นหลัก ไม่แสดงซ้ำในชั้นย่อย)
        if e_mpa and e_mpa > 0 and not is_sublayer:
            ax.text(x_start + width + 0.5, y_center_pos, f'E = {e_mpa:,.0f} MPa',
                    ha='left', va='center', fontsize=12, color='#0066CC')
        
        y_current = y_bottom
    
    # ========== เส้นแสดงความหนารวม (ลูกศรสองหัว) ==========
    ax.annotate('', xy=(x_start + width + 3.5, total_display), xytext=(x_start + width + 3.5, 0),
                arrowprops=dict(arrowstyle='<->', color='red', lw=2))
    
    # ข้อความ Total (ภาษาไทย)
    ax.text(x_start + width + 4, total_display / 2, f'รวม\n{total_thickness:.0f} cm',
            ha='left', va='center', fontsize=14, color='red', fontweight='bold',
            fontproperties=thai_font_bold)
    
    # ========== ตั้งค่าขอบเขต (คงที่เหมือน Rigid Pavement) ==========
    margin = 10
    ax.set_xlim(0, 15)
    ax.set_ylim(-margin, total_display + margin)
    ax.axis('off')
    
    # Title (ภาษาไทย)
    ax.set_title('รูปตัดโครงสร้างชั้นทาง', fontsize=20, fontweight='bold', pad=20,
                 fontproperties=thai_font_bold)
    
    # กรอบ Total Pavement Thickness (ล่าง)
    ax.text(x_center, -margin + 4, f'ความหนารวมโครงสร้างชั้นทาง: {total_thickness:.0f} cm',
            ha='center', va='center', fontsize=15, fontweight='bold',
            fontproperties=thai_font_bold,
            bbox=dict(boxstyle='round', facecolor='lightyellow', alpha=0.9, edgecolor='orange'))
    
    plt.tight_layout()
    return fig

def get_figure_as_bytes(fig: plt.Figure) -> BytesIO:
    """Convert matplotlib figure to bytes"""
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    return buf


# ================================================================================
# WORD EXPORT FUNCTION
# ================================================================================

def set_thai_font(run, size_pt=15, bold=False):
    """Set TH Sarabun New font for Thai text"""
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(size_pt)
    run.bold = bold
    # Set East Asian font
    run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cs', 'TH Sarabun New')

def set_equation_font(run, size_pt=11, bold=False, italic=True):
    """Set Times New Roman font for equations"""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic

def add_thai_paragraph(doc, text, size_pt=15, bold=False, alignment=None):
    """Add paragraph with Thai font"""
    para = doc.add_paragraph()
    if alignment:
        para.alignment = alignment
    run = para.add_run(text)
    set_thai_font(run, size_pt, bold)
    return para

def add_equation_paragraph(doc, text, size_pt=11, bold=False, italic=True):
    """Add paragraph with equation font (Times New Roman)"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_equation_font(run, size_pt, bold, italic)
    return para

def create_word_report(project_title: str, inputs: dict, calc_results: dict,
                       design_check: dict, fig: plt.Figure) -> BytesIO:
    """Create Word document report with step-by-step calculations"""
    
    doc = Document()
    
    # ========================================
    # TITLE
    # ========================================
    title = doc.add_heading('รายงานการออกแบบ Flexible Pavement', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Set Thai font for title
    for run in title.runs:
        set_thai_font(run, size_pt=24, bold=True)
    
    heading1 = doc.add_heading(f'โครงการ: {project_title}', level=1)
    for run in heading1.runs:
        set_thai_font(run, size_pt=18, bold=True)
    
    add_thai_paragraph(doc, f'วันที่ออกแบบ: {datetime.now().strftime("%d/%m/%Y %H:%M")}', size_pt=15)
    
    # ========================================
    # SECTION 1: Design Method
    # ========================================
    heading2 = doc.add_heading('1. วิธีการออกแบบ', level=2)
    for run in heading2.runs:
        set_thai_font(run, size_pt=16, bold=True)
    
    add_thai_paragraph(doc, 
        'การออกแบบโครงสร้างถนนใช้วิธี AASHTO 1993 Guide for Design of Pavement Structures '
        'ตามมาตรฐานกรมทางหลวง โดยใช้สมการหลักดังนี้:', size_pt=15)
    
    # Main equation - Times New Roman 11pt
    add_equation_paragraph(doc,
        'log₁₀(W₁₈) = Zᵣ·Sₒ + 9.36·log₁₀(SN+1) - 0.20 + '
        'log₁₀(ΔPSI/2.7) / [0.4 + 1094/(SN+1)⁵·¹⁹] + 2.32·log₁₀(Mᵣ) - 8.07',
        size_pt=11, italic=True)
    
    # ========================================
    # SECTION 2: Input Parameters
    # ========================================
    heading2_2 = doc.add_heading('2. ข้อมูลนำเข้า (Design Inputs)', level=2)
    for run in heading2_2.runs:
        set_thai_font(run, size_pt=16, bold=True)
    
    input_table = doc.add_table(rows=1, cols=3)
    input_table.style = 'Table Grid'
    
    headers = ['พารามิเตอร์', 'ค่า', 'หน่วย']
    for i, header in enumerate(headers):
        cell = input_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_thai_font(run, size_pt=15, bold=True)
    
    input_data = [
        ('Design ESALs (W₁₈)', f'{inputs["W18"]:,.0f}', '18-kip ESAL'),
        ('Reliability (R)', f'{inputs["reliability"]}', '%'),
        ('Standard Normal Deviate (Zᵣ)', f'{inputs["Zr"]:.3f}', '-'),
        ('Overall Standard Deviation (Sₒ)', f'{inputs["So"]:.2f}', '-'),
        ('Initial Serviceability (P₀)', f'{inputs["P0"]:.1f}', '-'),
        ('Terminal Serviceability (Pₜ)', f'{inputs["Pt"]:.1f}', '-'),
        ('ΔPSI = P₀ - Pₜ', f'{inputs["delta_psi"]:.1f}', '-'),
        ('Subgrade CBR', f'{inputs.get("CBR", "-")}', '%'),
        ('Subgrade Mᵣ = 1500 × CBR', f'{inputs["Mr"]:,.0f}', 'psi'),
    ]
    
    for param, value, unit in input_data:
        row = input_table.add_row()
        row.cells[0].text = param
        row.cells[1].text = value
        row.cells[2].text = unit
        # Set Thai font for table cells
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_thai_font(run, size_pt=15)
    
    # ========================================
    # SECTION 3: Material Properties
    # ========================================
    heading2_3 = doc.add_heading('3. คุณสมบัติวัสดุชั้นทาง', level=2)
    for run in heading2_3.runs:
        set_thai_font(run, size_pt=16, bold=True)
    
    mat_table = doc.add_table(rows=1, cols=6)
    mat_table.style = 'Table Grid'
    
    mat_headers = ['ชั้น', 'วัสดุ', 'aᵢ', 'mᵢ', 'Mᵣ (psi)', 'E (MPa)']
    for i, header in enumerate(mat_headers):
        cell = mat_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_thai_font(run, size_pt=15, bold=True)
    
    for layer in calc_results['layers']:
        row = mat_table.add_row()
        row.cells[0].text = str(layer['layer_no'])
        row.cells[1].text = layer['material']
        row.cells[2].text = f'{layer["a_i"]:.2f}'
        row.cells[3].text = f'{layer["m_i"]:.2f}'
        row.cells[4].text = f'{layer["mr_psi"]:,}'
        row.cells[5].text = f'{layer["mr_mpa"]:,}'
        # Set Thai font for table cells
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_thai_font(run, size_pt=15)
    
    # ========================================
    # SECTION 4: Step-by-Step Calculation
    # ========================================
    heading2_4 = doc.add_heading('4. ขั้นตอนการคำนวณความหนาชั้นทาง', level=2)
    for run in heading2_4.runs:
        set_thai_font(run, size_pt=16, bold=True)
    
    add_thai_paragraph(doc,
        'การคำนวณความหนาขั้นต่ำของแต่ละชั้น ใช้หลักการว่า Structural Number (SN) '
        'ที่จุดใดๆ ต้องมากกว่าหรือเท่ากับ SN ที่ต้องการ โดยคำนวณจากค่า Mᵣ ของชั้นถัดไป',
        size_pt=15)
    
    for layer in calc_results['layers']:
        # Layer header
        layer_heading = doc.add_heading(f'ชั้นที่ {layer["layer_no"]}: {layer["material"]}', level=3)
        for run in layer_heading.runs:
            set_thai_font(run, size_pt=15, bold=True)
        
        # Material properties
        add_thai_paragraph(doc, 'ข้อมูลวัสดุ:', size_pt=15, bold=True)
        props_para = doc.add_paragraph()
        run1 = props_para.add_run(f'    • Mᵣ = {layer["mr_psi"]:,} psi = {layer["mr_mpa"]:,} MPa\n')
        set_thai_font(run1, size_pt=15)
        run2 = props_para.add_run(f'    • Layer Coefficient (a{layer["layer_no"]}) = {layer["a_i"]:.2f}\n')
        set_thai_font(run2, size_pt=15)
        run3 = props_para.add_run(f'    • Drainage Coefficient (m{layer["layer_no"]}) = {layer["m_i"]:.2f}')
        set_thai_font(run3, size_pt=15)
        
        # SN calculation
        add_thai_paragraph(doc, 'การคำนวณ SN:', size_pt=15, bold=True)
        sn_para = doc.add_paragraph()
        sn_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sn_run = sn_para.add_run(f'จากสมการ AASHTO 1993:  SN{layer["layer_no"]} = {layer["sn_required_at_layer"]:.2f}')
        set_equation_font(sn_run, size_pt=11, bold=True, italic=False)
        
        # Thickness calculation
        add_thai_paragraph(doc, 'การคำนวณความหนาขั้นต่ำ:', size_pt=15, bold=True)
        
        if layer['layer_no'] == 1:
            formula_text = f'D₁ ≥ SN₁ / (a₁ × m₁) = {layer["sn_required_at_layer"]:.2f} / ({layer["a_i"]:.2f} × {layer["m_i"]:.2f})'
            add_equation_paragraph(doc, formula_text, size_pt=11, italic=True)
        else:
            prev_sn = calc_results['layers'][layer['layer_no']-2]['cumulative_sn']
            formula_text = f'D{layer["layer_no"]} ≥ (SN{layer["layer_no"]} - SNₚᵣₑᵥ) / (a{layer["layer_no"]} × m{layer["layer_no"]}) = ({layer["sn_required_at_layer"]:.2f} - {prev_sn:.2f}) / ({layer["a_i"]:.2f} × {layer["m_i"]:.2f})'
            add_equation_paragraph(doc, formula_text, size_pt=11, italic=True)
        
        # Results
        result_para = doc.add_paragraph()
        result_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        result_run = result_para.add_run(f'D{layer["layer_no"]}(min) = {layer["min_thickness_inch"]:.2f} นิ้ว = {layer["min_thickness_cm"]:.1f} ซม.')
        set_equation_font(result_run, size_pt=11, bold=True, italic=False)
        
        # Design thickness selection
        add_thai_paragraph(doc, 'เลือกใช้ความหนา:', size_pt=15, bold=True)
        design_para = doc.add_paragraph()
        design_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        design_run = design_para.add_run(f'D{layer["layer_no"]}(design) = {layer["design_thickness_cm"]:.0f} ซม. ({layer["design_thickness_inch"]:.2f} นิ้ว)')
        set_equation_font(design_run, size_pt=11, bold=True, italic=False)
        
        # SN contribution
        add_thai_paragraph(doc, 'SN contribution:', size_pt=15, bold=True)
        contrib_text = f'ΔSN{layer["layer_no"]} = a{layer["layer_no"]} × D{layer["layer_no"]} × m{layer["layer_no"]} = {layer["a_i"]:.2f} × {layer["design_thickness_inch"]:.2f} × {layer["m_i"]:.2f} = {layer["sn_contribution"]:.3f}'
        add_equation_paragraph(doc, contrib_text, size_pt=11, italic=False)
        
        # Cumulative SN
        cum_para = doc.add_paragraph()
        cum_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cum_run = cum_para.add_run(f'ΣSN = {layer["cumulative_sn"]:.2f}')
        set_equation_font(cum_run, size_pt=11, bold=True, italic=False)
        
        # Check status
        status_text = '✓ OK' if layer['is_ok'] else '✗ NG - ต้องเพิ่มความหนา'
        status_para = doc.add_paragraph()
        status_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        status_run = status_para.add_run(f'สถานะ: {status_text}')
        set_thai_font(status_run, size_pt=15, bold=True)
        
        doc.add_paragraph()  # Spacing
    
    # ========================================
    # SECTION 5: SN Summary Table
    # ========================================
    heading2_5 = doc.add_heading('5. ตารางสรุปการคำนวณ Structural Number', level=2)
    for run in heading2_5.runs:
        set_thai_font(run, size_pt=16, bold=True)
    
    sn_table = doc.add_table(rows=1, cols=8)
    sn_table.style = 'Table Grid'
    
    sn_headers = ['ชั้น', 'วัสดุ', 'aᵢ', 'mᵢ', 'Dᵢ (นิ้ว)', 'Dᵢ (ซม.)', 'ΔSNᵢ', 'ΣSN']
    for i, header in enumerate(sn_headers):
        cell = sn_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_thai_font(run, size_pt=15, bold=True)
    
    for layer in calc_results['layers']:
        row = sn_table.add_row()
        row.cells[0].text = str(layer['layer_no'])
        row.cells[1].text = layer['material']
        row.cells[2].text = f'{layer["a_i"]:.2f}'
        row.cells[3].text = f'{layer["m_i"]:.2f}'
        row.cells[4].text = f'{layer["design_thickness_inch"]:.2f}'
        row.cells[5].text = f'{layer["design_thickness_cm"]:.0f}'
        row.cells[6].text = f'{layer["sn_contribution"]:.3f}'
        row.cells[7].text = f'{layer["cumulative_sn"]:.2f}'
        # Set Thai font for table cells
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_thai_font(run, size_pt=15)
    
    # Formula
    doc.add_paragraph()
    add_equation_paragraph(doc, 'สูตร: SN = Σ(aᵢ × Dᵢ × mᵢ)', size_pt=11, italic=True)
    
    # ========================================
    # SECTION 6: Design Verification
    # ========================================
    heading2_6 = doc.add_heading('6. ผลการตรวจสอบการออกแบบ', level=2)
    for run in heading2_6.runs:
        set_thai_font(run, size_pt=16, bold=True)
    
    result_table = doc.add_table(rows=4, cols=2)
    result_table.style = 'Table Grid'
    
    result_data = [
        ('SN Required (จากสมการ AASHTO)', f'{calc_results["total_sn_required"]:.2f}'),
        ('SN Provided (จากชั้นทาง)', f'{calc_results["total_sn_provided"]:.2f}'),
        ('Safety Margin', f'{design_check["safety_margin"]:.2f}'),
        ('ผลการตรวจสอบ', 'ผ่าน (OK)' if design_check['passed'] else 'ไม่ผ่าน (NG)'),
    ]
    
    for i, (param, value) in enumerate(result_data):
        result_table.rows[i].cells[0].text = param
        result_table.rows[i].cells[1].text = value
        # Set Thai font for table cells
        for cell in result_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_thai_font(run, size_pt=15)
    
    # Conclusion
    doc.add_paragraph()
    if design_check['passed']:
        conclusion_text = (f'สรุป: การออกแบบผ่านเกณฑ์ เนื่องจาก SN_provided ({calc_results["total_sn_provided"]:.2f}) ≥ '
            f'SN_required ({calc_results["total_sn_required"]:.2f})')
        add_thai_paragraph(doc, conclusion_text, size_pt=15, bold=True)
    else:
        conclusion_text = 'สรุป: การออกแบบไม่ผ่านเกณฑ์ กรุณาปรับเพิ่มความหนาชั้นทาง'
        add_thai_paragraph(doc, conclusion_text, size_pt=15, bold=True)
    
    # ========================================
    # SECTION 7: Figure
    # ========================================
    heading2_7 = doc.add_heading('7. ภาพตัดขวางโครงสร้างถนน', level=2)
    for run in heading2_7.runs:
        set_thai_font(run, size_pt=16, bold=True)
    
    fig_bytes = get_figure_as_bytes(fig)
    doc.add_picture(fig_bytes, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ========================================
    # Save document
    # ========================================
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes


# ================================================================================
# STREAMLIT USER INTERFACE
# ================================================================================

def main():
    """Main Streamlit application"""
    
    # Header
    st.title("🛣️ AASHTO 1993 Flexible Pavement Design")
    st.markdown("""
    **แอปพลิเคชันออกแบบโครงสร้างทางแบบยืดหยุ่น ตามมาตรฐานกรมทางหลวง (DOH Thailand)**
    
    ✅ คำนวณความหนาขั้นต่ำแต่ละชั้น | ✅ แสดงขั้นตอนการคำนวณ | ✅ Export รายงาน Word
    """)
    
    st.markdown("---")
    
    # ========================================
    # SIDEBAR: Project Info & Material Database
    # ========================================
    with st.sidebar:
        st.header("📋 ข้อมูลโครงการ")
        
        st.markdown("---")
        
        # ========================================
        # UPLOAD/DOWNLOAD JSON
        # ========================================
        st.header("💾 บันทึก/โหลดข้อมูล")
        
        uploaded_json = st.file_uploader(
            "📂 โหลดข้อมูลจากไฟล์ JSON",
            type=['json'],
            help="อัปโหลดไฟล์ JSON ที่บันทึกไว้ก่อนหน้า"
        )
        
        # ประมวลผลไฟล์ JSON ที่อัปโหลด
        if uploaded_json is not None:
            try:
                loaded_data = json.load(uploaded_json)
                
                # ตรวจสอบว่าเป็นไฟล์ใหม่หรือไม่
                file_id = f"{uploaded_json.name}_{uploaded_json.size}"
                if st.session_state.get('last_uploaded_file') != file_id:
                    st.session_state['last_uploaded_file'] = file_id
                    st.session_state['loaded_json'] = loaded_data
                    
                    # อัพเดท session_state สำหรับทุก widget
                    # Design Inputs
                    st.session_state['input_W18'] = loaded_data.get('W18', 5000000)
                    st.session_state['input_reliability'] = loaded_data.get('reliability', 90)
                    st.session_state['input_So'] = loaded_data.get('So', 0.45)
                    st.session_state['input_P0'] = loaded_data.get('P0', 4.2)
                    st.session_state['input_Pt'] = loaded_data.get('Pt', 2.5)
                    st.session_state['input_CBR'] = loaded_data.get('CBR', 5.0)
                    st.session_state['input_num_layers'] = loaded_data.get('num_layers', 4)
                    st.session_state['input_project_title'] = loaded_data.get('project_title', 'โครงการออกแบบถนน')
                    
                    # Layer data
                    layers = loaded_data.get('layers', [])
                    for i, layer in enumerate(layers):
                        st.session_state[f'layer{i+1}_mat'] = layer.get('material', '')
                        st.session_state[f'layer{i+1}_thick'] = layer.get('thickness_cm', 5.0 if i == 0 else 15.0)
                        st.session_state[f'layer{i+1}_m'] = layer.get('drainage_coeff', 1.0)
                    
                    st.success("✅ โหลดข้อมูลสำเร็จ!")
                    st.rerun()
                    
            except Exception as e:
                st.error(f"❌ ไม่สามารถอ่านไฟล์ได้: {e}")
        
        # ชื่อโครงการ
        project_title = st.text_input(
            "ชื่อโครงการ", 
            value=st.session_state.get('input_project_title', "โครงการออกแบบถนน"),
            key="project_title_input"
        )
        
        st.markdown("---")
        
        # ตัวเลือกภาษาสำหรับรูปภาพ
        st.header("🖼️ ตั้งค่ารูปภาพ")
        figure_language = st.radio(
            "ภาษาในรูปภาพ",
            options=["English", "ภาษาไทย"],
            index=0,
            help="เลือกภาษาสำหรับแสดงในรูปตัดขวาง"
        )
        
        st.markdown("---")
        st.header("📚 ฐานข้อมูลวัสดุ (ทล.)")
        
        with st.expander("ดูค่า สปส. วัสดุทั้งหมด"):
            st.markdown("**ค่า สปส. สำหรับออกแบบ**")
            for mat_name, props in MATERIALS.items():
                if props['layer_coeff'] > 0:
                    st.markdown(f"**{mat_name}**")
                    st.markdown(f"- a = {props['layer_coeff']}, m = {props['drainage_coeff']}")
                    st.markdown(f"- MR = {props['mr_psi']:,} psi ({props['mr_mpa']:,} MPa)")
                    st.markdown("---")
    
    # ========================================
    # MAIN CONTENT: Two columns
    # ========================================
    col1, col2 = st.columns([1, 1])
    
    # ========================================
    # COLUMN 1: Design Inputs
    # ========================================
    with col1:
        st.header("📝 Design Inputs")
        
        # Traffic
        st.subheader("1️⃣ Traffic & Reliability")
        
        W18 = st.number_input(
            "Design ESALs (W₁₈)",
            min_value=100000,
            max_value=250000000,
            value=st.session_state.get('input_W18', 5000000),
            step=100000,
            format="%d",
            help="จำนวน 18-kip ESAL ตลอดอายุการใช้งาน (สูงสุด 250 ล้าน)",
            key="input_W18"
        )
        
        # แสดงค่า ESAL เป็นล้าน (สีฟ้า font ใหญ่)
        esal_million = W18 / 1000000
        st.markdown(f'<p style="color: #1E90FF; font-size: 18px; font-weight: bold;">💡 W₁₈ = {esal_million:,.2f} ล้าน ESALs</p>', unsafe_allow_html=True)
        
        # หา index ของ reliability จาก session_state
        reliability_options = list(RELIABILITY_ZR.keys())
        current_reliability = st.session_state.get('input_reliability', 90)
        default_reliability_idx = reliability_options.index(current_reliability) if current_reliability in reliability_options else reliability_options.index(90)
        
        reliability = st.selectbox(
            "Reliability Level (R)",
            options=reliability_options,
            index=default_reliability_idx,
            key="input_reliability"
        )
        Zr = RELIABILITY_ZR[reliability]
        st.info(f"Zᵣ = {Zr:.3f}")
        
        So = st.number_input(
            "Overall Standard Deviation (Sₒ)",
            min_value=0.30,
            max_value=0.60,
            value=st.session_state.get('input_So', 0.45),
            step=0.01,
            format="%.2f",
            key="input_So"
        )
        
        # Serviceability
        st.subheader("2️⃣ Serviceability")
        
        col1a, col1b = st.columns(2)
        with col1a:
            P0 = st.number_input(
                "P₀ (Initial)", 
                min_value=3.0, max_value=5.0, 
                value=st.session_state.get('input_P0', 4.2), 
                step=0.1,
                key="input_P0"
            )
        with col1b:
            Pt = st.number_input(
                "Pₜ (Terminal)", 
                min_value=1.5, max_value=3.5, 
                value=st.session_state.get('input_Pt', 2.5), 
                step=0.1,
                key="input_Pt"
            )
        
        delta_psi = P0 - Pt
        st.success(f"**ΔPSI = {delta_psi:.1f}**")
        
        # Subgrade
        st.subheader("3️⃣ Subgrade (ดินเดิม/ดินถม)")
        
        CBR = st.number_input(
            "CBR (%)",
            min_value=1.0,
            max_value=30.0,
            value=st.session_state.get('input_CBR', 5.0),
            step=0.5,
            help="ค่า CBR ของดินเดิมหรือดินถมคันทาง",
            key="input_CBR"
        )
        
        # Mr = 1500 × CBR (ตามมาตรฐาน ทล.)
        Mr = int(1500 * CBR)
        st.info(f"**Mᵣ = 1,500 × CBR = 1,500 × {CBR:.1f} = {Mr:,} psi**")
    
    # ========================================
    # COLUMN 2: Layer Configuration
    # ========================================
    with col2:
        st.header("🏗️ Layer Configuration")
        
        # จำนวนชั้นทาง
        num_layers = st.slider(
            "จำนวนชั้นทาง",
            min_value=2,
            max_value=6,
            value=st.session_state.get('input_num_layers', 4),
            help="เลือกจำนวนชั้นทาง (2-6 ชั้น)",
            key="input_num_layers"
        )
        
        # สร้าง list วัสดุทั้งหมด (ยกเว้น "ไม่ใช้")
        all_materials = [m for m, p in MATERIALS.items() if p['layer_type'] != 'none']
        
        # สร้าง list วัสดุสำหรับชั้นที่ 2-6 (รวม "ไม่ใช้ชั้นนี้")
        optional_materials = all_materials + ["ไม่ใช้ชั้นนี้"]
        
        # เก็บข้อมูลชั้นทาง
        layer_data = []
        
        # เก็บ placeholders สำหรับแสดงสถานะ
        status_placeholders = {}
        
        # ========== ชั้นที่ 1: ผิวทาง (บังคับ) ==========
        st.subheader("4️⃣ ชั้นที่ 1: ผิวทาง (Surface)")
        
        surface_materials = [m for m, p in MATERIALS.items() if p['layer_type'] == 'surface']
        
        # ดึงค่า default จาก session_state สำหรับ Layer 1
        layer1_mat_default = st.session_state.get('layer1_mat', surface_materials[0])
        layer1_mat_idx = surface_materials.index(layer1_mat_default) if layer1_mat_default in surface_materials else 0
        
        layer1_mat = st.selectbox(
            "เลือกวัสดุ",
            options=surface_materials,
            index=layer1_mat_idx,
            key="layer1_mat"
        )
        
        # ==========================================
        # ตัวเลือกแบ่งชั้นย่อย AC
        # ==========================================
        use_sublayers = st.checkbox(
            "📐 แบ่งชั้นย่อยผิวทาง AC (Wearing, Binder, Base Course)",
            value=st.session_state.get('use_ac_sublayers', False),
            help="แบ่งชั้น AC ออกเป็น 3 ชั้นย่อย ตามมาตรฐานกรมทางหลวง",
            key="use_ac_sublayers"
        )
        
        # ค่าความหนามาตรฐานกรมทางหลวงสำหรับ dropdown
        DOH_THICKNESS_STANDARDS = {
            "Wearing Course": {"min": 40, "max": 70, "options": [40, 45, 50, 55, 60, 65, 70]},
            "Binder Course": {"min": 40, "max": 80, "options": [40, 45, 50, 55, 60, 65, 70, 75, 80]},
            "Base Course": {"min": 70, "max": 100, "options": [70, 75, 80, 85, 90, 95, 100]}
        }
        
        if use_sublayers:
            st.info("📋 **ความหนามาตรฐานกรมทางหลวง:**\n"
                   "- Wearing Course: 40-70 มม.\n"
                   "- Binder Course: 40-80 มม.\n"
                   "- Base Course: 70-100 มม.")
            
            # Wearing Course
            st.markdown("**🔹 Wearing Course (ชั้นผิวทาง)**")
            col_w1, col_w2 = st.columns([1, 1])
            with col_w1:
                # Dropdown สำหรับเลือกความหนามาตรฐาน
                wearing_std_options = ["กำหนดเอง"] + [f"{t} มม." for t in DOH_THICKNESS_STANDARDS["Wearing Course"]["options"]]
                wearing_std = st.selectbox(
                    "ความหนามาตรฐาน ทล.",
                    options=wearing_std_options,
                    index=0,
                    key="wearing_std_select",
                    help="Wearing Course: 40-70 มม. ตามมาตรฐานกรมทางหลวง"
                )
            with col_w2:
                # ถ้าเลือกจาก dropdown ให้ใช้ค่านั้น
                if wearing_std != "กำหนดเอง":
                    wearing_val = int(wearing_std.replace(" มม.", "")) / 10  # แปลง มม. เป็น cm
                    wearing_thick = st.number_input(
                        "ความหนา (cm)", min_value=4, max_value=15.0,
                        value=wearing_val, step=1, key="wearing_thick", disabled=True
                    )
                else:
                    wearing_thick = st.number_input(
                        "ความหนา (cm)", min_value=4, max_value=15.0,
                        value=st.session_state.get('wearing_thick', 1), step=0.5, key="wearing_thick"
                    )
            
            # Binder Course
            st.markdown("**🔹 Binder Course (ชั้นยึดเกาะ)**")
            col_b1, col_b2 = st.columns([1, 1])
            with col_b1:
                binder_std_options = ["กำหนดเอง"] + [f"{t} มม." for t in DOH_THICKNESS_STANDARDS["Binder Course"]["options"]]
                binder_std = st.selectbox(
                    "ความหนามาตรฐาน ทล.",
                    options=binder_std_options,
                    index=0,
                    key="binder_std_select",
                    help="Binder Course: 40-80 มม. ตามมาตรฐานกรมทางหลวง"
                )
            with col_b2:
                if binder_std != "กำหนดเอง":
                    binder_val = int(binder_std.replace(" มม.", "")) / 10
                    binder_thick = st.number_input(
                        "ความหนา (cm)", min_value=1.0, max_value=15.0,
                        value=binder_val, step=0.5, key="binder_thick", disabled=True
                    )
                else:
                    binder_thick = st.number_input(
                        "ความหนา (cm)", min_value=1.0, max_value=15.0,
                        value=st.session_state.get('binder_thick', 7.0), step=0.5, key="binder_thick"
                    )
            
           # Base Course
            st.markdown("**🔹 Base Course (ชั้นรองผิวทาง)**")
            col_bc1, col_bc2 = st.columns([1, 1])
            with col_bc1:
                base_std_options = ["กำหนดเอง", "ไม่ใช้ชั้นนี้"] + [f"{t} มม." for t in DOH_THICKNESS_STANDARDS["Base Course"]["options"]]
                base_std = st.selectbox(
                    "ความหนามาตรฐาน ทล.",
                    options=base_std_options,
                    index=0,
                    key="base_std_select",
                    help="Base Course: 70-100 มม. ตามมาตรฐานกรมทางหลวง (เลือก 'ไม่ใช้ชั้นนี้' หรือใส่ 0 เพื่อข้าม)"
                )
            with col_bc2:
                if base_std == "ไม่ใช้ชั้นนี้":
                    base_course_thick = 0.0
                    st.number_input(
                        "ความหนา (cm)", min_value=0.0, max_value=15.0,
                        value=0.0, step=0.5, key="base_course_thick", disabled=True
                    )
                elif base_std != "กำหนดเอง":
                    base_val = int(base_std.replace(" มม.", "")) / 10
                    base_course_thick = st.number_input(
                        "ความหนา (cm)", min_value=0.0, max_value=15.0,
                        value=base_val, step=0.5, key="base_course_thick", disabled=True
                    )
                else:
                    base_course_thick = st.number_input(
                        "ความหนา (cm)", min_value=0.0, max_value=15.0,
                        value=st.session_state.get('base_course_thick', 10.0), step=0.5, key="base_course_thick"
                    )
            
            # คำนวณความหนารวมของ AC
            layer1_thick = wearing_thick + binder_thick + base_course_thick
            st.markdown(f'<p style="color: #1E90FF; font-size: 16px; font-weight: bold;">'
                       f'📏 ความหนารวมผิวทาง AC = {wearing_thick:.1f} + {binder_thick:.1f} + {base_course_thick:.1f} = {layer1_thick:.1f} cm</p>',
                       unsafe_allow_html=True)
            
            # เก็บข้อมูลชั้นย่อยไว้ใน session_state
            st.session_state['ac_sublayers'] = {
                'wearing': wearing_thick,
                'binder': binder_thick,
                'base': base_course_thick,
                'total': layer1_thick
            }
        else:
            # ไม่แบ่งชั้นย่อย - ใช้ความหนารวม
            mat_props = MATERIALS[layer1_mat]
            default_a1 = mat_props['layer_coeff']
            default_m1 = mat_props['drainage_coeff']
            
            # ตรวจสอบว่าวัสดุเปลี่ยนหรือไม่ - ถ้าเปลี่ยนให้ reset ค่า a และ m
            if 'layer1_prev_mat' not in st.session_state:
                st.session_state['layer1_prev_mat'] = layer1_mat
            
            if st.session_state['layer1_prev_mat'] != layer1_mat:
                st.session_state['layer1_a'] = default_a1
                st.session_state['layer1_m'] = default_m1
                st.session_state['layer1_prev_mat'] = layer1_mat
            
            col_a, col_b, col_c = st.columns(3)
            with col_a:
                layer1_thick = st.number_input(
                    "ความหนา (cm)", min_value=1.0, max_value=30.0, 
                    value=st.session_state.get('layer1_thick', 5.0), step=1.0,
                    key="layer1_thick"
                )
            with col_b:
                st.markdown(f"a₁ &nbsp;&nbsp;<span style='color: #1E90FF; font-size: 12px;'>(default = {default_a1:.2f})</span>", unsafe_allow_html=True)
                layer1_a = st.number_input(
                    "a1_input", min_value=0.10, max_value=0.50, 
                    value=st.session_state.get('layer1_a', default_a1), step=0.01,
                    key="layer1_a",
                    label_visibility="collapsed"
                )
            with col_c:
                st.markdown(f"m₁ &nbsp;&nbsp;<span style='color: #1E90FF; font-size: 12px;'>(default = {default_m1:.2f})</span>", unsafe_allow_html=True)
                layer1_m = st.number_input(
                    "m1_input", min_value=0.5, max_value=1.5, 
                    value=st.session_state.get('layer1_m', default_m1), step=0.05,
                    key="layer1_m",
                    label_visibility="collapsed"
                )
            st.session_state['ac_sublayers'] = None
        
        # ค่า a และ m สำหรับชั้น AC เมื่อใช้ sublayers
        if use_sublayers:
            st.markdown("---")
            mat_props = MATERIALS[layer1_mat]
            default_a1 = mat_props['layer_coeff']
            default_m1 = mat_props['drainage_coeff']
            
            # ตรวจสอบว่าวัสดุเปลี่ยนหรือไม่ - ถ้าเปลี่ยนให้ reset ค่า a และ m
            if 'layer1_prev_mat_sub' not in st.session_state:
                st.session_state['layer1_prev_mat_sub'] = layer1_mat
            
            if st.session_state['layer1_prev_mat_sub'] != layer1_mat:
                st.session_state['layer1_a_sublayer'] = default_a1
                st.session_state['layer1_m_sublayer'] = default_m1
                st.session_state['layer1_prev_mat_sub'] = layer1_mat
            
            col_am1, col_am2 = st.columns(2)
            with col_am1:
                st.markdown(f"a₁ (Layer Coefficient) &nbsp;&nbsp;<span style='color: #1E90FF; font-size: 12px;'>(default = {default_a1:.2f})</span>", unsafe_allow_html=True)
                layer1_a = st.number_input(
                    "a1_sublayer_input",
                    min_value=0.10, max_value=0.50,
                    value=st.session_state.get('layer1_a_sublayer', default_a1), step=0.01,
                    key="layer1_a_sublayer",
                    label_visibility="collapsed"
                )
            with col_am2:
                st.markdown(f"m₁ (Drainage Coefficient) &nbsp;&nbsp;<span style='color: #1E90FF; font-size: 12px;'>(default = {default_m1:.2f})</span>", unsafe_allow_html=True)
                layer1_m = st.number_input(
                    "m1_sublayer_input",
                    min_value=0.5, max_value=1.5,
                    value=st.session_state.get('layer1_m_sublayer', default_m1), step=0.05,
                    key="layer1_m_sublayer",
                    label_visibility="collapsed"
                )
        else:
            # กรณีไม่ใช้ sublayers ให้ใช้ค่า layer1_a ที่กำหนดไว้แล้ว
            pass
        
        mat_props = MATERIALS[layer1_mat]
        st.markdown(f'<p style="color: #1E90FF; font-size: 14px;">E = {mat_props["mr_mpa"]:,} MPa</p>', unsafe_allow_html=True)
        
        # Placeholder สำหรับแสดงสถานะชั้นที่ 1
        status_placeholders[1] = st.empty()
        
        layer_data.append({
            'material': layer1_mat,
            'thickness_cm': layer1_thick,
            'layer_coeff': layer1_a,
            'drainage_coeff': layer1_m
        })
        
        # ========== ชั้นที่ 2-6: เลือกวัสดุได้ทุกชนิด ==========
        default_materials = [
            "พื้นทางซีเมนต์ CTB",
            "รองพื้นทางวัสดุมวลรวม CBR 25%",
            "วัสดุคัดเลือก ก",
            "วัสดุคัดเลือก ก",
            "วัสดุคัดเลือก ก"
        ]
        default_thickness = [15.0, 15.0, 30.0, 30.0, 30.0]
        
        for i in range(2, num_layers + 1):
            st.markdown("---")
            layer_icons = ['5️⃣', '6️⃣', '7️⃣', '8️⃣', '9️⃣']
            st.subheader(f"{layer_icons[i-2]} ชั้นที่ {i}")
            
            # ดึงค่า default จาก session_state สำหรับ Layer i
            layer_i_mat_default = st.session_state.get(f'layer{i}_mat', default_materials[i-2])
            
            # หา index ของวัสดุ
            if layer_i_mat_default in all_materials:
                default_idx = all_materials.index(layer_i_mat_default)
            else:
                default_idx = all_materials.index(default_materials[i-2]) if default_materials[i-2] in all_materials else 0
            
            layer_mat = st.selectbox(
                f"เลือกวัสดุชั้นที่ {i}",
                options=all_materials,
                index=min(default_idx, len(all_materials)-1),
                key=f"layer{i}_mat"
            )
            
            # ดึงค่า a และ m จากฐานข้อมูลของวัสดุที่เลือก
            mat_props = MATERIALS[layer_mat]
            default_a = mat_props['layer_coeff']
            default_m = mat_props['drainage_coeff']
            
            # ตรวจสอบว่าวัสดุเปลี่ยนหรือไม่ - ถ้าเปลี่ยนให้ reset ค่า a และ m
            prev_mat_key = f'layer{i}_prev_mat'
            if prev_mat_key not in st.session_state:
                st.session_state[prev_mat_key] = layer_mat
            
            # ถ้าวัสดุเปลี่ยน ให้ reset ค่า a และ m
            if st.session_state[prev_mat_key] != layer_mat:
                st.session_state[f'layer{i}_a'] = default_a
                st.session_state[f'layer{i}_m'] = default_m
                st.session_state[prev_mat_key] = layer_mat
            
            col_c, col_d, col_e = st.columns(3)
            with col_c:
                layer_thick = st.number_input(
                    "ความหนา (cm)",
                    min_value=1.0, max_value=150.0, 
                    value=st.session_state.get(f'layer{i}_thick', default_thickness[i-2]), 
                    step=5.0,
                    key=f"layer{i}_thick"
                )
            with col_d:
                # แสดงค่า a จากวัสดุ (read-only style) และให้ผู้ใช้แก้ไขได้
                st.markdown(f"a{i} &nbsp;&nbsp;<span style='color: #1E90FF; font-size: 12px;'>(default = {default_a:.2f})</span>", unsafe_allow_html=True)
                layer_a = st.number_input(
                    f"a{i}_input",
                    min_value=0.01, max_value=0.50, 
                    value=st.session_state.get(f'layer{i}_a', default_a), 
                    step=0.01,
                    key=f"layer{i}_a",
                    label_visibility="collapsed"
                )
            with col_e:
                st.markdown(f"m{i} &nbsp;&nbsp;<span style='color: #1E90FF; font-size: 12px;'>(default = {default_m:.2f})</span>", unsafe_allow_html=True)
                layer_m = st.number_input(
                    f"m{i}_input",
                    min_value=0.5, max_value=1.5, 
                    value=st.session_state.get(f'layer{i}_m', default_m), 
                    step=0.05,
                    key=f"layer{i}_m",
                    label_visibility="collapsed"
                )
            
            # แสดงค่า E
            st.markdown(f'<p style="color: #1E90FF; font-size: 14px;">E = {mat_props["mr_mpa"]:,} MPa</p>', unsafe_allow_html=True)
            
            # Placeholder สำหรับแสดงสถานะชั้นที่ i
            status_placeholders[i] = st.empty()
            
            layer_data.append({
                'material': layer_mat,
                'thickness_cm': layer_thick,
                'layer_coeff': layer_a,
                'drainage_coeff': layer_m
            })
    
    # ========================================
    # BUILD LAYERS LIST
    # ========================================
    layers = layer_data
    
    # Store inputs
    inputs = {
        'W18': W18, 'reliability': reliability, 'Zr': Zr, 'So': So,
        'P0': P0, 'Pt': Pt, 'delta_psi': delta_psi, 'CBR': CBR, 'Mr': Mr
    }
    
    # ========================================
    # CALCULATION & RESULTS
    # ========================================
    st.markdown("---")
    st.header("📊 ผลการคำนวณ (Calculation Results)")
    
    # ดึงข้อมูลชั้นย่อย AC จาก session_state
    ac_sublayers = st.session_state.get('ac_sublayers', None)
    
    # Calculate layer thicknesses
    calc_results = calculate_layer_thicknesses(W18, Zr, So, delta_psi, Mr, layers, ac_sublayers)
    
    # Design check
    design_check = check_design(
        calc_results['total_sn_required'],
        calc_results['total_sn_provided']
    )
    
    # ========================================
    # FILL STATUS PLACEHOLDERS
    # ========================================
    for layer in calc_results['layers']:
        layer_no = layer['layer_no']
        if layer_no in status_placeholders:
            with status_placeholders[layer_no]:
                if layer['is_ok']:
                    st.success(f"✅ ผ่าน (ต้องการ ≥ {layer['min_thickness_cm']:.1f} cm)")
                else:
                    shortage = layer['min_thickness_cm'] - layer['design_thickness_cm']
                    st.error(f"❌ ไม่ผ่าน (ต้องเพิ่มอีก {shortage:.1f} cm)")
    
    # ========================================
    # STEP-BY-STEP CALCULATION DISPLAY
    # ========================================
    st.subheader("🔢 ขั้นตอนการคำนวณความหนาแต่ละชั้น")
    
    for layer in calc_results['layers']:
        with st.container():
            # Header with colored background
            layer_status = "✅" if layer['is_ok'] else "❌"
            st.markdown(f"### {layer_status} ชั้นที่ {layer['layer_no']}: {layer['material']}")
            
            # แสดงข้อมูลชั้นย่อย AC (ถ้ามี)
            layer_ac_sublayers = layer.get('ac_sublayers', None)
            if layer_ac_sublayers is not None and layer['layer_no'] == 1:
                st.info(f"**📐 แบ่งชั้นย่อย AC:**\n"
                       f"- Wearing Course: {layer_ac_sublayers['wearing']:.1f} cm\n"
                       f"- Binder Course: {layer_ac_sublayers['binder']:.1f} cm\n"
                       f"- Base Course: {layer_ac_sublayers['base']:.1f} cm\n"
                       f"- **รวม: {layer_ac_sublayers['total']:.1f} cm**")
            
            col_a, col_b = st.columns([1, 1])
            
            with col_a:
                st.markdown("**ข้อมูลวัสดุ:**")
                st.markdown(f"- E (MPa) = **{layer['mr_mpa']:,}**")
                st.markdown(f"- Mᵣ (psi) = **{layer['mr_psi']:,}**")
                st.markdown(f"- Layer Coefficient (a{layer['layer_no']}) = **{layer['a_i']:.2f}**")
                st.markdown(f"- Drain Coefficient (m{layer['layer_no']}) = **{layer['m_i']:.2f}**")
            
            with col_b:
                st.markdown("**จากสมการ AASHTO:**")
                
                # Show SN calculation
                sn_at_layer = layer['sn_required_at_layer']
                
                if layer['layer_no'] == 1:
                    st.latex(f"SN_{{{layer['layer_no']}}} = {sn_at_layer:.2f}")
                else:
                    st.latex(f"SN_{{{layer['layer_no']}}} = {sn_at_layer:.2f}")
            
            # Thickness calculation formula
            st.markdown("**คำนวณความหนาผิวทาง:**")
            
            if layer['layer_no'] == 1:
                # First layer formula
                st.latex(f"D_{{1}} \\geq \\frac{{SN_{{1}}}}{{a_{{1}} \\times m_{{1}}}} = \\frac{{{sn_at_layer:.2f}}}{{{layer['a_i']:.2f} \\times {layer['m_i']:.2f}}} = {layer['min_thickness_inch']:.2f} \\text{{ นิ้ว}}")
            else:
                # Subsequent layers
                prev_sn = calc_results['layers'][layer['layer_no']-2]['cumulative_sn']
                st.latex(f"D_{{{layer['layer_no']}}} \\geq \\frac{{SN_{{{layer['layer_no']}}} - SN_{{prev}}}}{{a_{{{layer['layer_no']}}} \\times m_{{{layer['layer_no']}}}}} = \\frac{{{sn_at_layer:.2f} - {prev_sn:.2f}}}{{{layer['a_i']:.2f} \\times {layer['m_i']:.2f}}} = {layer['min_thickness_inch']:.2f} \\text{{ นิ้ว}}")
            
            # Results table
            result_cols = st.columns(4)
            
            with result_cols[0]:
                st.metric("ความหนาขั้นต่ำ", f"{layer['min_thickness_cm']:.1f} cm")
            
            with result_cols[1]:
                st.metric("ความหนาที่เลือก", f"{layer['design_thickness_cm']:.0f} cm", 
                         delta=f"{layer['design_thickness_cm'] - layer['min_thickness_cm']:.1f} cm")
            
            with result_cols[2]:
                st.metric("SN contribution", f"{layer['sn_contribution']:.3f}")
            
            with result_cols[3]:
                st.metric("Cumulative SN", f"{layer['cumulative_sn']:.2f}")
            
            # Status
            if layer['is_ok']:
                st.success(f"✅ **OK** - ความหนาเพียงพอ ({layer['design_thickness_cm']:.0f} ≥ {layer['min_thickness_cm']:.1f} cm)")
            else:
                st.error(f"❌ **NG** - ต้องเพิ่มความหนาอีก {layer['min_thickness_cm'] - layer['design_thickness_cm']:.1f} cm")
            
            st.markdown("---")
    
    # ========================================
    # SUMMARY RESULTS
    # ========================================
    st.subheader("📈 สรุปผลการออกแบบ")
    
    res_col1, res_col2, res_col3, res_col4 = st.columns(4)
    
    with res_col1:
        st.metric("SN Required", f"{calc_results['total_sn_required']:.2f}")
    
    with res_col2:
        st.metric("SN Provided", f"{calc_results['total_sn_provided']:.2f}")
    
    with res_col3:
        if design_check['passed']:
            st.metric("Safety Margin", f"{design_check['safety_margin']:.2f}", delta="OK")
        else:
            st.metric("Safety Margin", f"{design_check['safety_margin']:.2f}", delta="NG", delta_color="inverse")
    
    with res_col4:
        if design_check['passed']:
            st.success("**PASS** ✅")
        else:
            st.error("**FAIL** ❌")
    
    # W18 Supported calculation
    w18_supported = calculate_w18_supported(
        calc_results['total_sn_provided'], Zr, So, delta_psi, Mr
    )
    w18_supported_million = w18_supported / 1_000_000
    w18_diff_percent = ((w18_supported - W18) / W18) * 100
    
    st.markdown("---")
    
    w18_col1, w18_col2 = st.columns(2)
    
    with w18_col1:
        st.metric(
            "W₁₈ ออกแบบ",
            f"{W18/1_000_000:,.2f} ล้าน"
        )
    
    with w18_col2:
        delta_str = f"{w18_diff_percent:+.1f}%"
        if w18_diff_percent >= 0:
            st.metric(
                "W₁₈ รองรับได้",
                f"{w18_supported_million:,.2f} ล้าน",
                delta=delta_str
            )
        else:
            st.metric(
                "W₁₈ รองรับได้",
                f"{w18_supported_million:,.2f} ล้าน",
                delta=delta_str,
                delta_color="inverse"
            )
    
    # Status message
    if design_check['passed']:
        st.success(f"✅ การออกแบบผ่านเกณฑ์: {design_check['message']}")
    else:
        st.error(f"❌ การออกแบบไม่ผ่าน: {design_check['message']}")
    
    # ========================================
    # PAVEMENT SECTION FIGURE
    # ========================================
    st.subheader("📐 ภาพตัดขวางโครงสร้างถนน")
    
    # เลือกฟังก์ชันวาดรูปตามภาษาที่เลือก
    if figure_language == "English":
        fig = plot_pavement_section(calc_results['layers'], Mr, CBR)
    else:
        fig = plot_pavement_section_thai(calc_results['layers'], Mr, CBR)
    
    st.pyplot(fig)
    
    # ========================================
    # SN CALCULATION TABLE
    # ========================================
    with st.expander("📋 ตารางสรุปการคำนวณ SN"):
        st.markdown("### SN Contribution Table")
        
        table_data = []
        for layer in calc_results['layers']:
            table_data.append({
                'ชั้น': layer['layer_no'],
                'วัสดุ': layer['short_name'],
                'aᵢ': layer['a_i'],
                'Dᵢ (cm)': layer['design_thickness_cm'],
                'Dᵢ (in)': layer['design_thickness_inch'],
                'mᵢ': layer['m_i'],
                'E (MPa)': layer['mr_mpa'],
                'SN contrib.': layer['sn_contribution'],
                'SN cumul.': layer['cumulative_sn']
            })
        
        st.table(table_data)
        
        st.markdown(f"""
        **สูตรการคำนวณ:**
        
        $$SN = \\sum_{{i=1}}^{{n}} a_i \\times D_i \\times m_i$$
        
        **ผลลัพธ์:**
        - SN_provided = {calc_results['total_sn_provided']:.2f}
        - SN_required = {calc_results['total_sn_required']:.2f}
        """)
    
    # ========================================
    # EXPORT
    # ========================================
    st.subheader("📄 ส่งออกรายงาน")
    
    col_exp1, col_exp2, col_exp3 = st.columns(3)
    
    with col_exp1:
        if st.button("📝 สร้างรายงาน Word", type="primary"):
            with st.spinner("กำลังสร้างรายงาน..."):
                # ใช้รูปภาษาไทยในรายงาน Word
                fig_thai = plot_pavement_section_thai(calc_results['layers'], Mr, CBR)
                doc_bytes = create_word_report(
                    project_title, inputs, calc_results, design_check, fig_thai
                )
                
                st.download_button(
                    label="⬇️ ดาวน์โหลดรายงาน Word",
                    data=doc_bytes,
                    file_name=f"AASHTO_Flexible_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    
    with col_exp2:
        fig_bytes = get_figure_as_bytes(fig)
        st.download_button(
            label="📸 ดาวน์โหลดรูปตัดขวาง (PNG)",
            data=fig_bytes,
            file_name=f"Pavement_Section_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
            mime="image/png"
        )
    
    with col_exp3:
        # สร้างข้อมูล JSON สำหรับ export (รวมข้อมูลชั้นย่อย AC)
        export_data = {
            'project_title': project_title,
            'W18': W18,
            'reliability': reliability,
            'So': So,
            'P0': P0,
            'Pt': Pt,
            'CBR': CBR,
            'num_layers': num_layers,
            'layers': layer_data,
            'ac_sublayers': st.session_state.get('ac_sublayers', None)
        }
        json_str = json.dumps(export_data, ensure_ascii=False, indent=2)
        
        st.download_button(
            label="💾 ดาวน์โหลดข้อมูล (JSON)",
            data=json_str,
            file_name=f"Flexible_Input_{datetime.now().strftime('%Y%m%d_%H%M')}.json",
            mime="application/json"
        )
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; color: gray;'>
    <p>AASHTO 1993 Flexible Pavement Design Application v4.0</p>
    <p>พัฒนาตามมาตรฐานกรมทางหลวง (DOH Thailand)</p>
    </div>
    """, unsafe_allow_html=True)


# ================================================================================
# ENTRY POINT
# ================================================================================

if __name__ == "__main__":
    main()
