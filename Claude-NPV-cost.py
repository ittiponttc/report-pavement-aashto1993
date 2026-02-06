#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
================================================================================
โปรแกรมวิเคราะห์ต้นทุนตลอดอายุการใช้งานผิวทาง (LCCA) - เวอร์ชัน 3.0
Pavement Life-Cycle Cost Analysis Program (Combined Edition)
================================================================================
รวมความสามารถจาก:
- LCCA v2.0: การวิเคราะห์ต้นทุนตลอดอายุ, มูลค่าซาก, Sensitivity Analysis, EAC
- NPV-Cost v3.0: Library ราคาวัสดุ, Layer Editor, Joint Cost, โครงสร้างสำเร็จรูป

พัฒนาสำหรับการเรียนการสอนและงานวิจัยด้านวิศวกรรมทาง
ภาควิชาครุศาสตร์โยธา มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ

ประเภทผิวทาง:
1. Flexible Pavement (AC) - ผิวทางยืดหยุ่น/แอสฟัลต์
2. JPCP - Jointed Plain Concrete Pavement
3. JRCP - Jointed Reinforced Concrete Pavement
4. CRCP - Continuously Reinforced Concrete Pavement
================================================================================
"""

import streamlit as st
import numpy as np
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from typing import Dict, List, Tuple, Optional
from dataclasses import dataclass, field
import json
import io
from datetime import datetime

# สำหรับส่งออก Word
try:
    from docx import Document as WordDocument
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ตั้งค่าหน้าเว็บ
st.set_page_config(
    page_title="โปรแกรมวิเคราะห์ LCCA ผิวทาง v3.0",
    page_icon="🛣️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS
st.markdown("""
<style>
    .main-header {
        font-size: 1.8rem;
        font-weight: bold;
        color: #1E3A5F;
        text-align: center;
        padding: 1rem;
        background: linear-gradient(90deg, #E8F4FD, #D1E9FA);
        border-radius: 10px;
        margin-bottom: 1rem;
    }
    .cost-box {
        background: #f0f8ff;
        padding: 10px;
        border-radius: 8px;
        border-left: 4px solid #2E86AB;
        margin: 5px 0;
    }
    .highlight-box {
        background: #f0fff0;
        padding: 10px;
        border-radius: 8px;
        border-left: 4px solid #28A745;
        margin: 5px 0;
    }
</style>
""", unsafe_allow_html=True)


# =============================================================================
# ส่วนที่ 1: Library ราคาวัสดุ (จาก NPV-Cost v3.0)
# =============================================================================

AC_PRICE_TABLE = {
    'PMA Wearing Course': {
        2.5: 170, 3: 203, 4: 268, 5: 333, 6: 406, 7: 471, 8: 536, 9: 601, 10: 667
    },
    'AC Wearing Course': {
        2.5: 128, 3: 152, 4: 202, 5: 250, 6: 306, 7: 355, 8: 403, 9: 452, 10: 502
    },
    'AC Binder Course': {
        2.5: 129, 3: 154, 4: 202, 5: 251, 6: 308, 7: 356, 8: 405, 9: 454, 10: 503
    },
    'AC Base Course': {
        2.5: 129, 3: 154, 4: 202, 5: 251, 6: 308, 7: 356, 8: 405, 9: 454, 10: 503
    },
}

CONCRETE_PRICE_TABLE = {
    'JRCP': {25: 924, 28: 1002, 32: 1106, 35: 1184},
    'JPCP': {25: 928, 28: 1000, 32: 1095, 35: 1167},
    'CRCP': {25: 1245, 28: 1358, 32: 1509, 35: 1622},
}

BASE_MATERIAL_PRICES = {
    'Crushed Rock Base Course': 583,
    'Cement Modified Crushed Rock Base (UCS 24.5 ksc)': 864,
    'Cement Treated Base (UCS 40 ksc)': 1096,
    'Soil Aggregate Subbase': 375,
    'Soil Cement Subbase (UCS 7 ksc)': 854,
    'Selected Material A': 375,
}


# =============================================================================
# ส่วนที่ 2: โครงสร้างข้อมูล (จาก LCCA v2.0 — Dataclass)
# =============================================================================

@dataclass
class กิจกรรมบำรุงรักษา:
    """โครงสร้างข้อมูลกิจกรรมบำรุงรักษา"""
    ชื่อกิจกรรม: str
    ต้นทุนต่อหน่วย: float  # บาท/ตร.ม.
    ปีเริ่มต้น: int
    ความถี่: int = 0  # 0 = ครั้งเดียว


@dataclass
class กิจกรรมฟื้นฟูสภาพ:
    """โครงสร้างข้อมูลกิจกรรมฟื้นฟูสภาพ"""
    ชื่อกิจกรรม: str
    ต้นทุนต่อหน่วย: float  # บาท/ตร.ม.
    ปีดำเนินการ: int


@dataclass
class ทางเลือกผิวทาง:
    """โครงสร้างข้อมูลทางเลือกผิวทาง"""
    ชื่อ: str
    ประเภท: str
    ต้นทุนก่อสร้าง: float  # บาท/ตร.ม. (คำนวณจาก Layer Editor)
    แผนบำรุงรักษา: List[กิจกรรมบำรุงรักษา]
    แผนฟื้นฟูสภาพ: List[กิจกรรมฟื้นฟูสภาพ]
    ร้อยละมูลค่าซาก: float = 20.0
    พื้นที่: float = 1000.0  # ตร.ม.
    ความหนา: float = 0.0  # ซม.
    เปิดใช้งาน: bool = True
    # ข้อมูลชั้นวัสดุ (จาก NPV-Cost)
    ชั้นวัสดุ: List[dict] = field(default_factory=list)
    รอยต่อ: List[dict] = field(default_factory=list)
    รวมรอยต่อ: bool = True


# =============================================================================
# ส่วนที่ 3: ฟังก์ชันคำนวณหลัก (จาก LCCA v2.0)
# =============================================================================

def คำนวณมูลค่าปัจจุบัน(ต้นทุน: float, ปี: int, อัตราคิดลด: float) -> float:
    """คำนวณมูลค่าปัจจุบัน (Present Worth): PW = FV × (1 + i)^(-n)"""
    if ปี < 0 or อัตราคิดลด < 0:
        return 0.0
    pwf = (1 + อัตราคิดลด) ** (-ปี)
    return ต้นทุน * pwf


def คำนวณต้นทุนเฉลี่ยรายปี(pw: float, อัตราคิดลด: float, ระยะวิเคราะห์: int) -> float:
    """แปลง PW เป็นต้นทุนเฉลี่ยรายปี (EAC)"""
    if ระยะวิเคราะห์ <= 0 or อัตราคิดลด <= 0:
        return 0.0
    ตัวเศษ = อัตราคิดลด * (1 + อัตราคิดลด) ** ระยะวิเคราะห์
    ตัวส่วน = (1 + อัตราคิดลด) ** ระยะวิเคราะห์ - 1
    crf = ตัวเศษ / ตัวส่วน
    return pw * crf


def คำนวณมูลค่าซาก(
    ต้นทุนฟื้นฟูครั้งสุดท้าย: float,
    ปีฟื้นฟูครั้งสุดท้าย: int,
    อายุใช้งานที่คาดหวัง: int,
    ระยะวิเคราะห์: int,
    ร้อยละมูลค่าซาก: float = 20.0
) -> float:
    """คำนวณมูลค่าซากโดยวิธี Straight-Line Depreciation (FHWA)"""
    อายุใช้งานที่เหลือ = อายุใช้งานที่คาดหวัง - (ระยะวิเคราะห์ - ปีฟื้นฟูครั้งสุดท้าย)
    
    if อายุใช้งานที่เหลือ <= 0:
        return ต้นทุนฟื้นฟูครั้งสุดท้าย * (ร้อยละมูลค่าซาก / 100.0)
    
    ค่าเสื่อมต่อปี = ต้นทุนฟื้นฟูครั้งสุดท้าย * (1 - ร้อยละมูลค่าซาก / 100.0) / อายุใช้งานที่คาดหวัง
    มูลค่าซาก = ต้นทุนฟื้นฟูครั้งสุดท้าย - ค่าเสื่อมต่อปี * (ระยะวิเคราะห์ - ปีฟื้นฟูครั้งสุดท้าย)
    
    return max(มูลค่าซาก, ต้นทุนฟื้นฟูครั้งสุดท้าย * ร้อยละมูลค่าซาก / 100.0)


# =============================================================================
# ส่วนที่ 4: สร้างตารางกระแสเงินสด (จาก LCCA v2.0 — Logic C)
# =============================================================================

def สร้างตารางกระแสเงินสด(
    ทางเลือก: ทางเลือกผิวทาง,
    ระยะวิเคราะห์: int,
    อัตราคิดลด: float,
    รวมมูลค่าซาก: bool = True
) -> pd.DataFrame:
    """
    สร้างตารางกระแสเงินสดรายปี
    Logic C: รีเซ็ตรอบบำรุงรักษาหลังฟื้นฟูสภาพ
    """
    รายการ = []
    พื้นที่ = ทางเลือก.พื้นที่
    
    ปีฟื้นฟูทั้งหมด = sorted([ฟ.ปีดำเนินการ for ฟ in ทางเลือก.แผนฟื้นฟูสภาพ if ฟ.ปีดำเนินการ <= ระยะวิเคราะห์])
    ปีฟื้นฟู_set = set(ปีฟื้นฟูทั้งหมด)
    
    # ปีที่ 0: ต้นทุนก่อสร้างเริ่มต้น
    ต้นทุนเริ่มต้น = ทางเลือก.ต้นทุนก่อสร้าง * พื้นที่
    รายการ.append({
        'ปี': 0, 'กิจกรรม': 'ก่อสร้างเริ่มต้น', 'ประเภท': 'ก่อสร้าง',
        'ต้นทุนต่อหน่วย': ทางเลือก.ต้นทุนก่อสร้าง,
        'ต้นทุนตามปี': ต้นทุนเริ่มต้น, 'ตัวคูณ_PW': 1.0, 'มูลค่าปัจจุบัน': ต้นทุนเริ่มต้น
    })
    
    # กิจกรรมบำรุงรักษา (รีเซ็ตรอบหลังฟื้นฟู)
    for บำรุง in ทางเลือก.แผนบำรุงรักษา:
        if บำรุง.ความถี่ > 0:
            จุดเริ่มต้นช่วง = [0] + ปีฟื้นฟูทั้งหมด
            for idx, ปีเริ่มช่วง in enumerate(จุดเริ่มต้นช่วง):
                ปีสิ้นสุดช่วง = จุดเริ่มต้นช่วง[idx + 1] if idx + 1 < len(จุดเริ่มต้นช่วง) else ระยะวิเคราะห์ + 1
                ปี = ปีเริ่มช่วง + บำรุง.ความถี่
                while ปี < ปีสิ้นสุดช่วง and ปี <= ระยะวิเคราะห์:
                    if ปี not in ปีฟื้นฟู_set:
                        ต้นทุน = บำรุง.ต้นทุนต่อหน่วย * พื้นที่
                        pwf = (1 + อัตราคิดลด) ** (-ปี)
                        รายการ.append({
                            'ปี': ปี, 'กิจกรรม': บำรุง.ชื่อกิจกรรม, 'ประเภท': 'บำรุงรักษา',
                            'ต้นทุนต่อหน่วย': บำรุง.ต้นทุนต่อหน่วย,
                            'ต้นทุนตามปี': ต้นทุน, 'ตัวคูณ_PW': pwf, 'มูลค่าปัจจุบัน': ต้นทุน * pwf
                        })
                    ปี += บำรุง.ความถี่
        else:
            if บำรุง.ปีเริ่มต้น <= ระยะวิเคราะห์ and บำรุง.ปีเริ่มต้น not in ปีฟื้นฟู_set:
                ต้นทุน = บำรุง.ต้นทุนต่อหน่วย * พื้นที่
                pwf = (1 + อัตราคิดลด) ** (-บำรุง.ปีเริ่มต้น)
                รายการ.append({
                    'ปี': บำรุง.ปีเริ่มต้น, 'กิจกรรม': บำรุง.ชื่อกิจกรรม, 'ประเภท': 'บำรุงรักษา',
                    'ต้นทุนต่อหน่วย': บำรุง.ต้นทุนต่อหน่วย,
                    'ต้นทุนตามปี': ต้นทุน, 'ตัวคูณ_PW': pwf, 'มูลค่าปัจจุบัน': ต้นทุน * pwf
                })
    
    # กิจกรรมฟื้นฟูสภาพ
    ต้นทุนฟื้นฟูสุดท้าย = ทางเลือก.ต้นทุนก่อสร้าง * พื้นที่
    ปีฟื้นฟูสุดท้าย = 0
    
    for ฟื้นฟู in ทางเลือก.แผนฟื้นฟูสภาพ:
        if ฟื้นฟู.ปีดำเนินการ <= ระยะวิเคราะห์:
            ต้นทุน = ฟื้นฟู.ต้นทุนต่อหน่วย * พื้นที่
            pwf = (1 + อัตราคิดลด) ** (-ฟื้นฟู.ปีดำเนินการ)
            รายการ.append({
                'ปี': ฟื้นฟู.ปีดำเนินการ, 'กิจกรรม': ฟื้นฟู.ชื่อกิจกรรม, 'ประเภท': 'ฟื้นฟูสภาพ',
                'ต้นทุนต่อหน่วย': ฟื้นฟู.ต้นทุนต่อหน่วย,
                'ต้นทุนตามปี': ต้นทุน, 'ตัวคูณ_PW': pwf, 'มูลค่าปัจจุบัน': ต้นทุน * pwf
            })
            ต้นทุนฟื้นฟูสุดท้าย = ต้นทุน
            ปีฟื้นฟูสุดท้าย = ฟื้นฟู.ปีดำเนินการ
    
    # มูลค่าซาก (Salvage Value)
    if รวมมูลค่าซาก:
        if 'Flexible' in ทางเลือก.ประเภท or 'AC' in ทางเลือก.ประเภท:
            อายุที่คาดหวัง = 15
        elif 'CRCP' in ทางเลือก.ประเภท:
            อายุที่คาดหวัง = 25
        else:
            อายุที่คาดหวัง = 20
        
        sv = คำนวณมูลค่าซาก(
            ต้นทุนฟื้นฟูสุดท้าย, ปีฟื้นฟูสุดท้าย, อายุที่คาดหวัง,
            ระยะวิเคราะห์, ทางเลือก.ร้อยละมูลค่าซาก
        )
        pwf = (1 + อัตราคิดลด) ** (-ระยะวิเคราะห์)
        รายการ.append({
            'ปี': ระยะวิเคราะห์, 'กิจกรรม': 'มูลค่าซาก', 'ประเภท': 'มูลค่าซาก',
            'ต้นทุนต่อหน่วย': -sv / พื้นที่,
            'ต้นทุนตามปี': -sv, 'ตัวคูณ_PW': pwf, 'มูลค่าปัจจุบัน': -sv * pwf
        })
    
    df = pd.DataFrame(รายการ)
    df = df.sort_values(['ปี', 'กิจกรรม']).reset_index(drop=True)
    return df


# =============================================================================
# ส่วนที่ 5: วิเคราะห์ LCCA และ Sensitivity (จาก LCCA v2.0)
# =============================================================================

def วิเคราะห์_LCCA(
    ทางเลือกทั้งหมด: List[ทางเลือกผิวทาง],
    ระยะวิเคราะห์: int,
    อัตราคิดลด: float,
    รวมมูลค่าซาก: bool = True
) -> Tuple[pd.DataFrame, Dict[str, pd.DataFrame]]:
    """วิเคราะห์ LCCA สำหรับทุกทางเลือก"""
    สรุป_รายการ = []
    กระแสเงินสด_dict = {}
    
    ทางเลือกที่ใช้ = [ท for ท in ทางเลือกทั้งหมด if ท.เปิดใช้งาน]
    
    for ทางเลือก in ทางเลือกที่ใช้:
        cf_table = สร้างตารางกระแสเงินสด(ทางเลือก, ระยะวิเคราะห์, อัตราคิดลด, รวมมูลค่าซาก)
        กระแสเงินสด_dict[ทางเลือก.ชื่อ] = cf_table
        
        รวม_nominal = cf_table['ต้นทุนตามปี'].sum()
        รวม_pw = cf_table['มูลค่าปัจจุบัน'].sum()
        eac = คำนวณต้นทุนเฉลี่ยรายปี(รวม_pw, อัตราคิดลด, ระยะวิเคราะห์)
        
        ก่อสร้าง = cf_table[cf_table['ประเภท'] == 'ก่อสร้าง']['มูลค่าปัจจุบัน'].sum()
        บำรุงรักษา = cf_table[cf_table['ประเภท'] == 'บำรุงรักษา']['มูลค่าปัจจุบัน'].sum()
        ฟื้นฟู = cf_table[cf_table['ประเภท'] == 'ฟื้นฟูสภาพ']['มูลค่าปัจจุบัน'].sum()
        ซาก = cf_table[cf_table['ประเภท'] == 'มูลค่าซาก']['มูลค่าปัจจุบัน'].sum()
        
        สรุป_รายการ.append({
            'ทางเลือก': ทางเลือก.ชื่อ,
            'ประเภทผิวทาง': ทางเลือก.ประเภท,
            'ความหนา_ซม': getattr(ทางเลือก, 'ความหนา', 0.0),
            'พื้นที่_ตรม': ทางเลือก.พื้นที่,
            'ต้นทุนก่อสร้าง_ตรม': ทางเลือก.ต้นทุนก่อสร้าง,
            'PW_ก่อสร้าง': ก่อสร้าง,
            'PW_บำรุงรักษา': บำรุงรักษา,
            'PW_ฟื้นฟูสภาพ': ฟื้นฟู,
            'PW_มูลค่าซาก': ซาก,
            'ต้นทุนตามปีรวม': รวม_nominal,
            'มูลค่าปัจจุบันรวม': รวม_pw,
            'ต้นทุนเฉลี่ยรายปี': eac,
            'ต้นทุนต่อตรม_ต่อปี': eac / ทางเลือก.พื้นที่ if ทางเลือก.พื้นที่ > 0 else 0
        })
    
    สรุป_df = pd.DataFrame(สรุป_รายการ)
    if len(สรุป_df) > 0:
        สรุป_df['ลำดับ'] = สรุป_df['มูลค่าปัจจุบันรวม'].rank().astype(int)
        สรุป_df = สรุป_df.sort_values('มูลค่าปัจจุบันรวม').reset_index(drop=True)
    
    return สรุป_df, กระแสเงินสด_dict


def วิเคราะห์ความไว_อัตราคิดลด(
    ทางเลือกทั้งหมด: List[ทางเลือกผิวทาง],
    ระยะวิเคราะห์: int,
    อัตราฐาน: float,
    ช่วงการเปลี่ยนแปลง: float = 0.02,
    จำนวนจุด: int = 5,
    รวมมูลค่าซาก: bool = True
) -> Tuple[pd.DataFrame, pd.DataFrame]:
    """วิเคราะห์ความไวต่ออัตราคิดลด"""
    อัตราทดสอบ = np.linspace(max(0.005, อัตราฐาน - ช่วงการเปลี่ยนแปลง),
                              อัตราฐาน + ช่วงการเปลี่ยนแปลง, จำนวนจุด)
    ผลลัพธ์ = []
    ทางเลือกที่ใช้ = [ท for ท in ทางเลือกทั้งหมด if ท.เปิดใช้งาน]
    
    for อัตรา in อัตราทดสอบ:
        for ทางเลือก in ทางเลือกที่ใช้:
            cf = สร้างตารางกระแสเงินสด(ทางเลือก, ระยะวิเคราะห์, อัตรา, รวมมูลค่าซาก)
            pw = cf['มูลค่าปัจจุบัน'].sum()
            ผลลัพธ์.append({
                'อัตราคิดลด': อัตรา,
                'อัตราคิดลด_%': f"{อัตรา * 100:.1f}%",
                'ทางเลือก': ทางเลือก.ชื่อ,
                'มูลค่าปัจจุบัน': pw,
                'ต้นทุนเฉลี่ยรายปี': คำนวณต้นทุนเฉลี่ยรายปี(pw, อัตรา, ระยะวิเคราะห์)
            })
    
    df = pd.DataFrame(ผลลัพธ์)
    pivot = df.pivot(index='อัตราคิดลด_%', columns='ทางเลือก', values='มูลค่าปัจจุบัน') if len(df) > 0 else pd.DataFrame()
    return df, pivot


# =============================================================================
# ส่วนที่ 6: โครงสร้างสำเร็จรูป (จาก NPV-Cost v3.0 + LCCA v2.0 รวมกัน)
# =============================================================================

def สร้างทางเลือกเริ่มต้น(พื้นที่: float = 22000.0) -> List[ทางเลือกผิวทาง]:
    """สร้างทางเลือกผิวทาง 6 ทางเลือก พร้อมชั้นวัสดุและแผนบำรุงรักษา"""

    # 1. AC1: แอสฟัลต์บนหินคลุก
    ac1 = ทางเลือกผิวทาง(
        ชื่อ="AC1: แอสฟัลต์บนหินคลุก",
        ประเภท="Flexible",
        ต้นทุนก่อสร้าง=0,  # จะคำนวณจาก layers
        แผนบำรุงรักษา=[
            กิจกรรมบำรุงรักษา("Seal Coating", 80.0, ปีเริ่มต้น=3, ความถี่=3),
            กิจกรรมบำรุงรักษา("ซ่อมเฉพาะจุด", 50.0, ปีเริ่มต้น=5, ความถี่=5),
        ],
        แผนฟื้นฟูสภาพ=[
            กิจกรรมฟื้นฟูสภาพ("Overlay AC 50 มม.", 400.0, ปีดำเนินการ=9),
            กิจกรรมฟื้นฟูสภาพ("Overlay AC 50 มม.", 400.0, ปีดำเนินการ=18),
        ],
        ร้อยละมูลค่าซาก=20.0, พื้นที่=พื้นที่, ความหนา=24.0, เปิดใช้งาน=True,
        ชั้นวัสดุ=[
            {'name': 'AC Wearing Course', 'thickness': 7.0, 'unit': 'cm', 'qty_unit': 'sq.m', 'unit_cost': 355.0, 'layer_type': 'surface'},
            {'name': 'AC Binder Course', 'thickness': 7.0, 'unit': 'cm', 'qty_unit': 'sq.m', 'unit_cost': 356.0, 'layer_type': 'surface'},
            {'name': 'AC Base Course', 'thickness': 10.0, 'unit': 'cm', 'qty_unit': 'sq.m', 'unit_cost': 503.0, 'layer_type': 'surface'},
            {'name': 'Tack Coat', 'thickness': 2.0, 'unit': 'Layer', 'qty_unit': 'sq.m', 'unit_cost': 20.0, 'layer_type': 'surface'},
            {'name': 'Prime Coat', 'thickness': 1.0, 'unit': 'Layer', 'qty_unit': 'sq.m', 'unit_cost': 30.0, 'layer_type': 'surface'},
            {'name': 'Crushed Rock Base Course', 'thickness': 20.0, 'unit': 'cm', 'qty_unit': 'cu.m', 'unit_cost': 583.0, 'layer_type': 'base'},
            {'name': 'Soil Aggregate Subbase', 'thickness': 30.0, 'unit': 'cm', 'qty_unit': 'cu.m', 'unit_cost': 375.0, 'layer_type': 'base'},
        ],
    )
    # คำนวณต้นทุนรวม บาท/ตร.ม.
    ac1.ต้นทุนก่อสร้าง = _คำนวณต้นทุนจากชั้นวัสดุ(ac1.ชั้นวัสดุ)

    # 2. AC2: แอสฟัลต์บนหินคลุกผสมซีเมนต์
    ac2 = ทางเลือกผิวทาง(
        ชื่อ="AC2: แอสฟัลต์บน CMCR",
        ประเภท="Flexible",
        ต้นทุนก่อสร้าง=0,
        แผนบำรุงรักษา=[
            กิจกรรมบำรุงรักษา("Seal Coating", 80.0, ปีเริ่มต้น=3, ความถี่=3),
            กิจกรรมบำรุงรักษา("ซ่อมเฉพาะจุด", 50.0, ปีเริ่มต้น=5, ความถี่=5),
        ],
        แผนฟื้นฟูสภาพ=[
            กิจกรรมฟื้นฟูสภาพ("Overlay AC 50 มม.", 400.0, ปีดำเนินการ=9),
            กิจกรรมฟื้นฟูสภาพ("Overlay AC 50 มม.", 400.0, ปีดำเนินการ=18),
        ],
        ร้อยละมูลค่าซาก=20.0, พื้นที่=พื้นที่, ความหนา=10.0, เปิดใช้งาน=True,
        ชั้นวัสดุ=[
            {'name': 'AC Wearing Course', 'thickness': 5.0, 'unit': 'cm', 'qty_unit': 'sq.m', 'unit_cost': 250.0, 'layer_type': 'surface'},
            {'name': 'AC Binder Course', 'thickness': 5.0, 'unit': 'cm', 'qty_unit': 'sq.m', 'unit_cost': 251.0, 'layer_type': 'surface'},
            {'name': 'Tack Coat', 'thickness': 1.0, 'unit': 'Layer', 'qty_unit': 'sq.m', 'unit_cost': 20.0, 'layer_type': 'surface'},
            {'name': 'Prime Coat', 'thickness': 1.0, 'unit': 'Layer', 'qty_unit': 'sq.m', 'unit_cost': 30.0, 'layer_type': 'surface'},
            {'name': 'Cement Modified Crushed Rock Base (UCS 24.5 ksc)', 'thickness': 20.0, 'unit': 'cm', 'qty_unit': 'cu.m', 'unit_cost': 864.0, 'layer_type': 'base'},
            {'name': 'Soil Aggregate Subbase', 'thickness': 20.0, 'unit': 'cm', 'qty_unit': 'cu.m', 'unit_cost': 375.0, 'layer_type': 'base'},
        ],
    )
    ac2.ต้นทุนก่อสร้าง = _คำนวณต้นทุนจากชั้นวัสดุ(ac2.ชั้นวัสดุ)

    # 3. JPCP/JRCP(1): คอนกรีตบนดินซีเมนต์
    jrcp1 = ทางเลือกผิวทาง(
        ชื่อ="JPCP(1): คอนกรีตบนดินซีเมนต์",
        ประเภท="JPCP",
        ต้นทุนก่อสร้าง=0,
        แผนบำรุงรักษา=[
            กิจกรรมบำรุงรักษา("Joint Sealing", 65.0, ปีเริ่มต้น=3, ความถี่=3),
            กิจกรรมบำรุงรักษา("ซ่อมเฉพาะจุด", 40.0, ปีเริ่มต้น=10, ความถี่=10),
        ],
        แผนฟื้นฟูสภาพ=[
            กิจกรรมฟื้นฟูสภาพ("Diamond Grinding", 180.0, ปีดำเนินการ=20),
        ],
        ร้อยละมูลค่าซาก=30.0, พื้นที่=พื้นที่, ความหนา=28.0, เปิดใช้งาน=True,
        ชั้นวัสดุ=[
            {'name': '350 Ksc. Concrete (JPCP)', 'thickness': 28.0, 'unit': 'cm', 'qty_unit': 'sq.m', 'unit_cost': 1000.0, 'layer_type': 'surface'},
            {'name': 'Non Woven Geotextile', 'thickness': 1.0, 'unit': 'ชั้น', 'qty_unit': 'sq.m', 'unit_cost': 78.0, 'layer_type': 'surface'},
            {'name': 'Soil Cement Subbase (UCS 7 ksc)', 'thickness': 20.0, 'unit': 'cm', 'qty_unit': 'cu.m', 'unit_cost': 854.0, 'layer_type': 'base'},
            {'name': 'Selected Material A', 'thickness': 60.0, 'unit': 'cm', 'qty_unit': 'cu.m', 'unit_cost': 375.0, 'layer_type': 'base'},
        ],
        รอยต่อ=[
            {'name': 'Transverse Joint @4m', 'quantity_per_km': 2200, 'qty_unit': 'm', 'unit_cost': 430},
            {'name': 'Longitudinal Joint', 'quantity_per_km': 4000, 'qty_unit': 'm', 'unit_cost': 120},
        ],
    )
    jrcp1.ต้นทุนก่อสร้าง = _คำนวณต้นทุนจากชั้นวัสดุ(jrcp1.ชั้นวัสดุ, jrcp1.รอยต่อ, jrcp1.รวมรอยต่อ, jrcp1.พื้นที่)

    # 4. JPCP/JRCP(2): คอนกรีตบน CMCR
    jrcp2 = ทางเลือกผิวทาง(
        ชื่อ="JPCP(2): คอนกรีตบน CMCR",
        ประเภท="JPCP",
        ต้นทุนก่อสร้าง=0,
        แผนบำรุงรักษา=[
            กิจกรรมบำรุงรักษา("Joint Sealing", 65.0, ปีเริ่มต้น=3, ความถี่=3),
            กิจกรรมบำรุงรักษา("ซ่อมเฉพาะจุด", 40.0, ปีเริ่มต้น=10, ความถี่=10),
        ],
        แผนฟื้นฟูสภาพ=[
            กิจกรรมฟื้นฟูสภาพ("Diamond Grinding", 180.0, ปีดำเนินการ=20),
        ],
        ร้อยละมูลค่าซาก=30.0, พื้นที่=พื้นที่, ความหนา=28.0, เปิดใช้งาน=True,
        ชั้นวัสดุ=[
            {'name': '350 Ksc. Concrete (JPCP)', 'thickness': 28.0, 'unit': 'cm', 'qty_unit': 'sq.m', 'unit_cost': 1000.0, 'layer_type': 'surface'},
            {'name': 'Non Woven Geotextile', 'thickness': 1.0, 'unit': 'ชั้น', 'qty_unit': 'sq.m', 'unit_cost': 78.0, 'layer_type': 'surface'},
            {'name': 'Cement Modified Crushed Rock Base (UCS 24.5 ksc)', 'thickness': 20.0, 'unit': 'cm', 'qty_unit': 'cu.m', 'unit_cost': 864.0, 'layer_type': 'base'},
            {'name': 'Selected Material A', 'thickness': 50.0, 'unit': 'cm', 'qty_unit': 'cu.m', 'unit_cost': 375.0, 'layer_type': 'base'},
        ],
        รอยต่อ=[
            {'name': 'Transverse Joint @4m', 'quantity_per_km': 2200, 'qty_unit': 'm', 'unit_cost': 430},
            {'name': 'Longitudinal Joint', 'quantity_per_km': 4000, 'qty_unit': 'm', 'unit_cost': 120},
        ],
    )
    jrcp2.ต้นทุนก่อสร้าง = _คำนวณต้นทุนจากชั้นวัสดุ(jrcp2.ชั้นวัสดุ, jrcp2.รอยต่อ, jrcp2.รวมรอยต่อ, jrcp2.พื้นที่)

    # 5. CRCP1: เสริมเหล็กต่อเนื่องบนดินซีเมนต์
    crcp1 = ทางเลือกผิวทาง(
        ชื่อ="CRCP1: เสริมเหล็กต่อเนื่องบนดินซีเมนต์",
        ประเภท="CRCP",
        ต้นทุนก่อสร้าง=0,
        แผนบำรุงรักษา=[
            กิจกรรมบำรุงรักษา("ซ่อมเฉพาะจุด", 30.0, ปีเริ่มต้น=10, ความถี่=10),
        ],
        แผนฟื้นฟูสภาพ=[
            กิจกรรมฟื้นฟูสภาพ("Diamond Grinding", 180.0, ปีดำเนินการ=25),
        ],
        ร้อยละมูลค่าซาก=35.0, พื้นที่=พื้นที่, ความหนา=25.0, เปิดใช้งาน=True,
        ชั้นวัสดุ=[
            {'name': '350 Ksc. Concrete (CRCP)', 'thickness': 25.0, 'unit': 'cm', 'qty_unit': 'sq.m', 'unit_cost': 1245.0, 'layer_type': 'surface'},
            {'name': 'Steel Reinforcement', 'thickness': 1.0, 'unit': 'ชั้น', 'qty_unit': 'sq.m', 'unit_cost': 150.0, 'layer_type': 'surface'},
            {'name': 'Non Woven Geotextile', 'thickness': 1.0, 'unit': 'ชั้น', 'qty_unit': 'sq.m', 'unit_cost': 78.0, 'layer_type': 'surface'},
            {'name': 'Soil Cement Subbase (UCS 7 ksc)', 'thickness': 15.0, 'unit': 'cm', 'qty_unit': 'cu.m', 'unit_cost': 854.0, 'layer_type': 'base'},
            {'name': 'Selected Material A', 'thickness': 50.0, 'unit': 'cm', 'qty_unit': 'cu.m', 'unit_cost': 375.0, 'layer_type': 'base'},
        ],
    )
    crcp1.ต้นทุนก่อสร้าง = _คำนวณต้นทุนจากชั้นวัสดุ(crcp1.ชั้นวัสดุ)

    # 6. CRCP2: เสริมเหล็กต่อเนื่องบน CMCR
    crcp2 = ทางเลือกผิวทาง(
        ชื่อ="CRCP2: เสริมเหล็กต่อเนื่องบน CMCR",
        ประเภท="CRCP",
        ต้นทุนก่อสร้าง=0,
        แผนบำรุงรักษา=[
            กิจกรรมบำรุงรักษา("ซ่อมเฉพาะจุด", 30.0, ปีเริ่มต้น=10, ความถี่=10),
        ],
        แผนฟื้นฟูสภาพ=[
            กิจกรรมฟื้นฟูสภาพ("Diamond Grinding", 180.0, ปีดำเนินการ=25),
        ],
        ร้อยละมูลค่าซาก=35.0, พื้นที่=พื้นที่, ความหนา=25.0, เปิดใช้งาน=True,
        ชั้นวัสดุ=[
            {'name': '350 Ksc. Concrete (CRCP)', 'thickness': 25.0, 'unit': 'cm', 'qty_unit': 'sq.m', 'unit_cost': 1245.0, 'layer_type': 'surface'},
            {'name': 'Steel Reinforcement', 'thickness': 1.0, 'unit': 'ชั้น', 'qty_unit': 'sq.m', 'unit_cost': 150.0, 'layer_type': 'surface'},
            {'name': 'Non Woven Geotextile', 'thickness': 1.0, 'unit': 'ชั้น', 'qty_unit': 'sq.m', 'unit_cost': 78.0, 'layer_type': 'surface'},
            {'name': 'Cement Modified Crushed Rock Base (UCS 24.5 ksc)', 'thickness': 15.0, 'unit': 'cm', 'qty_unit': 'cu.m', 'unit_cost': 864.0, 'layer_type': 'base'},
            {'name': 'Selected Material A', 'thickness': 40.0, 'unit': 'cm', 'qty_unit': 'cu.m', 'unit_cost': 375.0, 'layer_type': 'base'},
        ],
    )
    crcp2.ต้นทุนก่อสร้าง = _คำนวณต้นทุนจากชั้นวัสดุ(crcp2.ชั้นวัสดุ)

    return [ac1, ac2, jrcp1, jrcp2, crcp1, crcp2]


def _คำนวณต้นทุนจากชั้นวัสดุ(ชั้นวัสดุ: List[dict], รอยต่อ: List[dict] = None, รวมรอยต่อ: bool = True, พื้นที่: float = 22000.0) -> float:
    """คำนวณต้นทุนรวม (บาท/ตร.ม.) จากรายชั้นวัสดุ"""
    total = 0.0
    for layer in ชั้นวัสดุ:
        if layer.get('qty_unit') == 'cu.m':
            # แปลง บาท/ลบ.ม. → บาท/ตร.ม.
            total += layer['unit_cost'] * layer['thickness'] / 100.0
        else:
            # บาท/ตร.ม. อยู่แล้ว
            if layer.get('unit') == 'Layer':
                total += layer['unit_cost'] * layer['thickness']  # จำนวนชั้น × ราคา
            else:
                total += layer['unit_cost']
    
    # รวม Joints (แปลงเป็น บาท/ตร.ม.)
    if รวมรอยต่อ and รอยต่อ:
        for j in รอยต่อ:
            joint_total_per_km = j.get('quantity_per_km', 0) * j.get('unit_cost', 0)
            # สมมติ 1 กม. → พื้นที่ต่อ กม.
            if พื้นที่ > 0:
                total += joint_total_per_km / พื้นที่
    
    return round(total, 2)


# =============================================================================
# ส่วนที่ 7: JSON Import/Export
# =============================================================================

def ทางเลือก_เป็น_dict(ทางเลือก: ทางเลือกผิวทาง) -> dict:
    return {
        'ชื่อ': ทางเลือก.ชื่อ, 'ประเภท': ทางเลือก.ประเภท,
        'ต้นทุนก่อสร้าง': ทางเลือก.ต้นทุนก่อสร้าง,
        'ร้อยละมูลค่าซาก': ทางเลือก.ร้อยละมูลค่าซาก,
        'พื้นที่': ทางเลือก.พื้นที่, 'ความหนา': ทางเลือก.ความหนา,
        'เปิดใช้งาน': ทางเลือก.เปิดใช้งาน,
        'ชั้นวัสดุ': ทางเลือก.ชั้นวัสดุ,
        'รอยต่อ': ทางเลือก.รอยต่อ,
        'รวมรอยต่อ': ทางเลือก.รวมรอยต่อ,
        'แผนบำรุงรักษา': [
            {'ชื่อกิจกรรม': ม.ชื่อกิจกรรม, 'ต้นทุนต่อหน่วย': ม.ต้นทุนต่อหน่วย,
             'ปีเริ่มต้น': ม.ปีเริ่มต้น, 'ความถี่': ม.ความถี่}
            for ม in ทางเลือก.แผนบำรุงรักษา
        ],
        'แผนฟื้นฟูสภาพ': [
            {'ชื่อกิจกรรม': ฟ.ชื่อกิจกรรม, 'ต้นทุนต่อหน่วย': ฟ.ต้นทุนต่อหน่วย,
             'ปีดำเนินการ': ฟ.ปีดำเนินการ}
            for ฟ in ทางเลือก.แผนฟื้นฟูสภาพ
        ]
    }


def dict_เป็น_ทางเลือก(data: dict) -> ทางเลือกผิวทาง:
    แผนบำรุง = [
        กิจกรรมบำรุงรักษา(ม['ชื่อกิจกรรม'], ม['ต้นทุนต่อหน่วย'], ม['ปีเริ่มต้น'], ม['ความถี่'])
        for ม in data.get('แผนบำรุงรักษา', [])
    ]
    แผนฟื้นฟู = [
        กิจกรรมฟื้นฟูสภาพ(ฟ['ชื่อกิจกรรม'], ฟ['ต้นทุนต่อหน่วย'], ฟ['ปีดำเนินการ'])
        for ฟ in data.get('แผนฟื้นฟูสภาพ', [])
    ]
    return ทางเลือกผิวทาง(
        ชื่อ=data['ชื่อ'], ประเภท=data['ประเภท'],
        ต้นทุนก่อสร้าง=data['ต้นทุนก่อสร้าง'],
        แผนบำรุงรักษา=แผนบำรุง, แผนฟื้นฟูสภาพ=แผนฟื้นฟู,
        ร้อยละมูลค่าซาก=data.get('ร้อยละมูลค่าซาก', 20.0),
        พื้นที่=data.get('พื้นที่', 22000.0),
        ความหนา=data.get('ความหนา', 0.0),
        เปิดใช้งาน=data.get('เปิดใช้งาน', True),
        ชั้นวัสดุ=data.get('ชั้นวัสดุ', []),
        รอยต่อ=data.get('รอยต่อ', []),
        รวมรอยต่อ=data.get('รวมรอยต่อ', True),
    )


# =============================================================================
# ส่วนที่ 8: Word Report (รวมทั้ง BOQ และ LCCA)
# =============================================================================

def สร้างรายงาน_Word(
    สรุป: pd.DataFrame,
    กระแสเงินสด: Dict[str, pd.DataFrame],
    ระยะวิเคราะห์: int,
    อัตราคิดลด: float,
    ทางเลือกทั้งหมด: List[ทางเลือกผิวทาง]
) -> io.BytesIO:
    """สร้างรายงาน LCCA + BOQ ในรูปแบบ Word"""
    doc = WordDocument()
    style = doc.styles['Normal']
    style.font.name = 'TH Sarabun New'
    style.font.size = Pt(14)

    title = doc.add_heading('รายงานการวิเคราะห์ต้นทุนตลอดอายุการใช้งานผิวทาง (LCCA)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ข้อมูลทั่วไป
    doc.add_heading('1. ข้อมูลทั่วไป', level=1)
    doc.add_paragraph(f'วันที่วิเคราะห์: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    doc.add_paragraph(f'ระยะเวลาวิเคราะห์: {ระยะวิเคราะห์} ปี')
    doc.add_paragraph(f'อัตราคิดลด: {อัตราคิดลด * 100:.1f}%')
    doc.add_paragraph(f'จำนวนทางเลือก: {len(สรุป)} ทางเลือก')

    # ตาราง BOQ แต่ละทางเลือก
    doc.add_heading('2. รายละเอียดค่าก่อสร้างแต่ละทางเลือก', level=1)
    for ทางเลือก in ทางเลือกทั้งหมด:
        if not ทางเลือก.เปิดใช้งาน:
            continue
        doc.add_heading(f'{ทางเลือก.ชื่อ}', level=2)
        if ทางเลือก.ชั้นวัสดุ:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            for j, h in enumerate(['รายการ', 'ความหนา', 'หน่วย', 'ราคา (บาท/ตร.ม.)']):
                table.rows[0].cells[j].text = h
                table.rows[0].cells[j].paragraphs[0].runs[0].bold = True
            for layer in ทางเลือก.ชั้นวัสดุ:
                row_cells = table.add_row().cells
                row_cells[0].text = layer['name']
                row_cells[1].text = f"{layer['thickness']} {layer['unit']}"
                row_cells[2].text = layer.get('qty_unit', 'sq.m')
                if layer.get('qty_unit') == 'cu.m':
                    cost_sqm = layer['unit_cost'] * layer['thickness'] / 100
                else:
                    cost_sqm = layer['unit_cost']
                row_cells[3].text = f"{cost_sqm:,.2f}"
        doc.add_paragraph(f'ต้นทุนก่อสร้างรวม: {ทางเลือก.ต้นทุนก่อสร้าง:,.2f} บาท/ตร.ม.')
        doc.add_paragraph()

    # ผลการวิเคราะห์ LCCA
    doc.add_heading('3. ผลการวิเคราะห์ LCCA', level=1)
    table2 = doc.add_table(rows=1, cols=5)
    table2.style = 'Table Grid'
    for j, h in enumerate(['ลำดับ', 'ทางเลือก', 'มูลค่าปัจจุบันรวม (บาท)', 'EAC (บาท/ปี)', 'ต้นทุน (บาท/ตร.ม./ปี)']):
        table2.rows[0].cells[j].text = h
        table2.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for _, row in สรุป.iterrows():
        row_cells = table2.add_row().cells
        row_cells[0].text = str(int(row['ลำดับ']))
        row_cells[1].text = str(row['ทางเลือก'])
        row_cells[2].text = f"{row['มูลค่าปัจจุบันรวม']:,.0f}"
        row_cells[3].text = f"{row['ต้นทุนเฉลี่ยรายปี']:,.0f}"
        row_cells[4].text = f"{row['ต้นทุนต่อตรม_ต่อปี']:,.2f}"

    # องค์ประกอบต้นทุน
    doc.add_heading('4. องค์ประกอบต้นทุน (มูลค่าปัจจุบัน)', level=1)
    table3 = doc.add_table(rows=1, cols=6)
    table3.style = 'Table Grid'
    for j, h in enumerate(['ทางเลือก', 'ก่อสร้าง', 'บำรุงรักษา', 'ฟื้นฟูสภาพ', 'มูลค่าซาก', 'รวม (บาท)']):
        table3.rows[0].cells[j].text = h
        table3.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for _, row in สรุป.iterrows():
        r = table3.add_row().cells
        r[0].text = str(row['ทางเลือก'])
        r[1].text = f"{row['PW_ก่อสร้าง']:,.0f}"
        r[2].text = f"{row['PW_บำรุงรักษา']:,.0f}"
        r[3].text = f"{row['PW_ฟื้นฟูสภาพ']:,.0f}"
        r[4].text = f"{row['PW_มูลค่าซาก']:,.0f}"
        r[5].text = f"{row['มูลค่าปัจจุบันรวม']:,.0f}"

    # สรุปผล
    doc.add_heading('5. สรุปผล', level=1)
    if len(สรุป) > 0:
        ผู้ชนะ = สรุป.iloc[0]
        doc.add_paragraph(f'ทางเลือกที่ประหยัดที่สุด: {ผู้ชนะ["ทางเลือก"]}')
        doc.add_paragraph(f'มูลค่าปัจจุบันรวม: {ผู้ชนะ["มูลค่าปัจจุบันรวม"]:,.0f} บาท')
        doc.add_paragraph(f'ต้นทุนเฉลี่ยรายปี (EAC): {ผู้ชนะ["ต้นทุนเฉลี่ยรายปี"]:,.0f} บาท/ปี')

    # กระแสเงินสด
    doc.add_page_break()
    doc.add_heading('6. รายละเอียดกระแสเงินสด', level=1)
    for ชื่อ, cf in กระแสเงินสด.items():
        doc.add_heading(ชื่อ, level=2)
        t = doc.add_table(rows=1, cols=5)
        t.style = 'Table Grid'
        for j, h in enumerate(['ปี', 'กิจกรรม', 'ต้นทุนตามปี (บาท)', 'ตัวคูณ PW', 'มูลค่าปัจจุบัน (บาท)']):
            t.rows[0].cells[j].text = h
            t.rows[0].cells[j].paragraphs[0].runs[0].bold = True
        for _, row in cf.iterrows():
            r = t.add_row().cells
            r[0].text = str(int(row['ปี']))
            r[1].text = str(row['กิจกรรม'])
            r[2].text = f"{row['ต้นทุนตามปี']:,.0f}"
            r[3].text = f"{row['ตัวคูณ_PW']:.4f}"
            r[4].text = f"{row['มูลค่าปัจจุบัน']:,.0f}"
        # รวม
        r = t.add_row().cells
        r[1].text = 'รวม'
        r[1].paragraphs[0].runs[0].bold = True
        r[2].text = f"{cf['ต้นทุนตามปี'].sum():,.0f}"
        r[4].text = f"{cf['มูลค่าปัจจุบัน'].sum():,.0f}"
        r[4].paragraphs[0].runs[0].bold = True
        doc.add_paragraph()

    doc.add_paragraph('─' * 50)
    doc.add_paragraph('รายงานสร้างโดย: โปรแกรมวิเคราะห์ LCCA ผิวทาง v3.0')
    doc.add_paragraph('ภาควิชาครุศาสตร์โยธา มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ')

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


# =============================================================================
# ส่วนที่ 9: Main Streamlit Application
# =============================================================================

def main():
    st.markdown('<div class="main-header">🛣️ โปรแกรมวิเคราะห์ต้นทุนตลอดอายุการใช้งานผิวทาง (LCCA) v3.0</div>', unsafe_allow_html=True)
    st.markdown("##### รวม Library ราคาวัสดุ + Layer Editor + LCCA Analysis + Sensitivity | AASHTO 1993 & FHWA")

    # ===== Initialize =====
    if 'ทางเลือกทั้งหมด' not in st.session_state:
        st.session_state.ทางเลือกทั้งหมด = สร้างทางเลือกเริ่มต้น()

    # ===== Sidebar =====
    with st.sidebar:
        st.header("⚙️ พารามิเตอร์การวิเคราะห์")

        ระยะวิเคราะห์ = st.slider("ระยะเวลาวิเคราะห์ (ปี)", 20, 50, 35, 5)
        อัตราคิดลด = st.slider("อัตราคิดลด (%)", 2.0, 10.0, 4.0, 0.5) / 100.0
        รวมมูลค่าซาก = st.checkbox(
            "🔄 รวมมูลค่าซาก (Salvage Value)",
            value=True,
            help="เลือกว่าจะนำมูลค่าซาก (Straight-Line Depreciation) มาพิจารณาในการวิเคราะห์ LCCA หรือไม่"
        )

        st.divider()
        st.subheader("📐 ข้อมูลถนน")
        ความกว้าง = st.number_input("ความกว้างรวม (ม.)", value=11.0, min_value=5.0, max_value=30.0, step=0.5)
        ความยาว = st.number_input("ความยาว (กม.)", value=1.0, min_value=0.1, max_value=100.0, step=0.5)
        พื้นที่คำนวณ = ความกว้าง * ความยาว * 1000  # ตร.ม. (1 ทิศทาง × ความยาว)
        st.info(f"📏 พื้นที่: **{พื้นที่คำนวณ:,.0f} ตร.ม.**")

        if st.button("📏 ใช้พื้นที่นี้กับทุกทางเลือก", use_container_width=True):
            for ท in st.session_state.ทางเลือกทั้งหมด:
                ท.พื้นที่ = พื้นที่คำนวณ
            st.success(f"✅ อัปเดตพื้นที่เป็น {พื้นที่คำนวณ:,.0f} ตร.ม.")
            st.rerun()

        st.divider()
        st.subheader("📊 Sensitivity Analysis")
        ช่วงอัตราคิดลด = st.slider("ช่วง ±%", 1.0, 4.0, 2.0, 0.5) / 100.0

        st.divider()
        st.subheader("💾 บันทึก/โหลดข้อมูล")
        uploaded_file = st.file_uploader("โหลดจาก JSON", type=['json'])
        if uploaded_file:
            try:
                data = json.load(uploaded_file)
                st.session_state.ทางเลือกทั้งหมด = [dict_เป็น_ทางเลือก(d) for d in data['ทางเลือก']]
                st.success("✅ โหลดสำเร็จ!")
                st.rerun()
            except Exception as e:
                st.error(f"❌ {e}")

        if st.button("🔄 รีเซ็ตเป็นค่าเริ่มต้น", use_container_width=True):
            st.session_state.ทางเลือกทั้งหมด = สร้างทางเลือกเริ่มต้น(พื้นที่คำนวณ)
            st.rerun()

    # ===== Tabs =====
    tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
        "📊 Library ราคา",
        "🏗️ โครงสร้าง/บำรุงรักษา",
        "📈 ผลวิเคราะห์ LCCA",
        "💰 กระแสเงินสด",
        "📉 Sensitivity Analysis",
        "📄 รายงาน",
        "ℹ️ ทฤษฎี LCCA"
    ])

    # ===== Tab 1: Library ราคาวัสดุ =====
    with tab1:
        st.header("📊 ตารางราคาเปรียบเทียบโครงสร้างชั้นทาง")
        st.info("💡 แก้ไขราคาได้ตามต้องการ ราคาจะถูกใช้เป็นค่าอ้างอิงในทุก Tab")

        if 'price_library' not in st.session_state:
            st.session_state['price_library'] = {
                'ac_prices': {k: dict(v) for k, v in AC_PRICE_TABLE.items()},
                'concrete_prices': {k: dict(v) for k, v in CONCRETE_PRICE_TABLE.items()},
                'base_prices': dict(BASE_MATERIAL_PRICES),
            }

        # AC Prices
        st.subheader("🔵 ผิวทาง Asphalt Concrete (บาท/ตร.ม.)")
        ac_types = list(AC_PRICE_TABLE.keys())
        thicknesses = [2.5, 3, 4, 5, 6, 7, 8, 9, 10]
        ac_cols = st.columns(4)
        for col_idx, ac_type in enumerate(ac_types):
            with ac_cols[col_idx]:
                st.markdown(f"**{ac_type}**")
                for thk in thicknesses:
                    default_price = st.session_state['price_library']['ac_prices'][ac_type].get(thk, 0)
                    price = st.number_input(f"{thk} cm", value=float(default_price), key=f"lib_ac_{ac_type}_{thk}", step=10.0)
                    st.session_state['price_library']['ac_prices'][ac_type][thk] = price

        st.divider()

        # Concrete Prices
        st.subheader("🟠 ผิวทางคอนกรีต (บาท/ตร.ม.)")
        conc_types = ['JRCP', 'JPCP', 'CRCP']
        conc_thicknesses = [25, 28, 32, 35]
        conc_cols = st.columns(3)
        for col_idx, conc_type in enumerate(conc_types):
            with conc_cols[col_idx]:
                st.markdown(f"**{conc_type}**")
                for thk in conc_thicknesses:
                    default_price = st.session_state['price_library']['concrete_prices'][conc_type].get(thk, 0)
                    price = st.number_input(f"{thk} cm", value=float(default_price), key=f"lib_conc_{conc_type}_{thk}", step=10.0)
                    st.session_state['price_library']['concrete_prices'][conc_type][thk] = price

        st.divider()

        # Base Material Prices
        st.subheader("🟤 วัสดุพื้นทาง/รองพื้นทาง (บาท/ลบ.ม.)")
        base_cols = st.columns(3)
        base_list = list(BASE_MATERIAL_PRICES.keys())
        for idx, mat in enumerate(base_list):
            with base_cols[idx % 3]:
                default_price = st.session_state['price_library']['base_prices'].get(mat, 0)
                price = st.number_input(mat, value=float(default_price), key=f"lib_base_{mat}", step=10.0)
                st.session_state['price_library']['base_prices'][mat] = price

    # ===== Tab 2: โครงสร้างชั้นทาง + บำรุงรักษา =====
    with tab2:
        st.header("🏗️ โครงสร้างชั้นทาง / แผนบำรุงรักษาและฟื้นฟูสภาพ")
        st.info("💡 กำหนดชั้นวัสดุ, ต้นทุนก่อสร้าง, แผนบำรุงรักษา และแผนฟื้นฟูสภาพได้ทุกทางเลือก")

        for i, ทางเลือก in enumerate(st.session_state.ทางเลือกทั้งหมด):
            with st.expander(f"{'✅' if ทางเลือก.เปิดใช้งาน else '❌'} ทางเลือก {i+1}: {ทางเลือก.ชื่อ} | {ทางเลือก.ต้นทุนก่อสร้าง:,.2f} บาท/ตร.ม.", expanded=(i == 0)):
                
                col_enable = st.columns([3, 1])
                with col_enable[1]:
                    ทางเลือก.เปิดใช้งาน = st.checkbox("เปิดใช้งาน", value=ทางเลือก.เปิดใช้งาน, key=f"enable_{i}")

                # ข้อมูลหลัก
                st.subheader("📋 ข้อมูลหลัก")
                c1, c2, c3, c4 = st.columns(4)
                with c1:
                    ทางเลือก.ชื่อ = st.text_input("ชื่อทางเลือก", value=ทางเลือก.ชื่อ, key=f"name_{i}")
                with c2:
                    types = ["Flexible", "JPCP", "JRCP", "CRCP"]
                    idx_type = types.index(ทางเลือก.ประเภท) if ทางเลือก.ประเภท in types else 0
                    ทางเลือก.ประเภท = st.selectbox("ประเภท", types, index=idx_type, key=f"type_{i}")
                with c3:
                    ทางเลือก.พื้นที่ = st.number_input("พื้นที่ (ตร.ม.)", value=float(ทางเลือก.พื้นที่), min_value=100.0, step=1000.0, key=f"area_{i}")
                with c4:
                    ทางเลือก.ร้อยละมูลค่าซาก = st.number_input("มูลค่าซาก (%)", value=float(ทางเลือก.ร้อยละมูลค่าซาก), min_value=0.0, max_value=50.0, step=5.0, key=f"salvage_{i}")

                # ชั้นวัสดุ
                st.markdown("---")
                st.subheader("🧱 ชั้นวัสดุ")
                if ทางเลือก.ชั้นวัสดุ:
                    total_cost_sqm = 0.0
                    cols_h = st.columns([3, 1.2, 1.5, 1.5])
                    cols_h[0].markdown("**วัสดุ**")
                    cols_h[1].markdown("**หนา (cm)**")
                    cols_h[2].markdown("**ราคา/หน่วย**")
                    cols_h[3].markdown("**บาท/ตร.ม.**")

                    for j, layer in enumerate(ทางเลือก.ชั้นวัสดุ):
                        cols = st.columns([3, 1.2, 1.5, 1.5])
                        with cols[0]:
                            layer['name'] = st.text_input("วัสดุ", value=layer['name'], key=f"lyr_name_{i}_{j}", label_visibility="collapsed")
                        with cols[1]:
                            layer['thickness'] = st.number_input("หนา", value=float(layer['thickness']), min_value=0.0, step=1.0, key=f"lyr_thick_{i}_{j}", label_visibility="collapsed")
                        with cols[2]:
                            layer['unit_cost'] = st.number_input("ราคา", value=float(layer['unit_cost']), min_value=0.0, step=10.0, key=f"lyr_cost_{i}_{j}", label_visibility="collapsed")
                        with cols[3]:
                            if layer.get('qty_unit') == 'cu.m':
                                cost_sqm = layer['unit_cost'] * layer['thickness'] / 100.0
                            elif layer.get('unit') == 'Layer':
                                cost_sqm = layer['unit_cost'] * layer['thickness']
                            else:
                                cost_sqm = layer['unit_cost']
                            total_cost_sqm += cost_sqm
                            st.markdown(f"**{cost_sqm:,.2f}**")

                    # Joints
                    joint_cost_sqm = 0.0
                    if ทางเลือก.รอยต่อ:
                        st.markdown("**รอยต่อ (Joints)**")
                        for jj, joint in enumerate(ทางเลือก.รอยต่อ):
                            cols_j = st.columns([3, 1.2, 1.5, 1.5])
                            with cols_j[0]:
                                st.text(joint['name'])
                            with cols_j[1]:
                                joint['quantity_per_km'] = st.number_input("ปริมาณ/กม.", value=float(joint.get('quantity_per_km', 0)), key=f"jt_qty_{i}_{jj}", label_visibility="collapsed")
                            with cols_j[2]:
                                joint['unit_cost'] = st.number_input("ราคา/ม.", value=float(joint.get('unit_cost', 0)), key=f"jt_cost_{i}_{jj}", label_visibility="collapsed")
                            with cols_j[3]:
                                j_per_sqm = joint['quantity_per_km'] * joint['unit_cost'] / ทางเลือก.พื้นที่ if ทางเลือก.พื้นที่ > 0 else 0
                                joint_cost_sqm += j_per_sqm
                                st.markdown(f"**{j_per_sqm:,.2f}**")
                        ทางเลือก.รวมรอยต่อ = st.checkbox("รวมรอยต่อในต้นทุนก่อสร้าง", value=ทางเลือก.รวมรอยต่อ, key=f"inc_joint_{i}")

                    # อัปเดตต้นทุน
                    final_cost = total_cost_sqm + (joint_cost_sqm if ทางเลือก.รวมรอยต่อ else 0)
                    ทางเลือก.ต้นทุนก่อสร้าง = round(final_cost, 2)

                    note = "(รวม Joints)" if ทางเลือก.รวมรอยต่อ and ทางเลือก.รอยต่อ else ""
                    st.markdown(f'<div class="cost-box">💰 <b>ต้นทุนก่อสร้าง:</b> {ทางเลือก.ต้นทุนก่อสร้าง:,.2f} บาท/ตร.ม. {note} | รวม: {ทางเลือก.ต้นทุนก่อสร้าง * ทางเลือก.พื้นที่:,.0f} บาท</div>', unsafe_allow_html=True)
                else:
                    ทางเลือก.ต้นทุนก่อสร้าง = st.number_input("ต้นทุนก่อสร้าง (บาท/ตร.ม.)", value=float(ทางเลือก.ต้นทุนก่อสร้าง), min_value=0.0, step=100.0, key=f"manual_cost_{i}")

                # แผนบำรุงรักษา
                st.markdown("---")
                st.subheader("🔧 แผนบำรุงรักษา")
                cols_mh = st.columns([3, 2, 1, 1, 0.5])
                cols_mh[0].markdown("**กิจกรรม**")
                cols_mh[1].markdown("**บาท/ตร.ม.**")
                cols_mh[2].markdown("**ปีเริ่มต้น**")
                cols_mh[3].markdown("**ทุกๆ (ปี)**")
                cols_mh[4].markdown("**ลบ**")

                del_maint = []
                for j, บำรุง in enumerate(ทางเลือก.แผนบำรุงรักษา):
                    cm = st.columns([3, 2, 1, 1, 0.5])
                    with cm[0]:
                        บำรุง.ชื่อกิจกรรม = st.text_input("ก.", value=บำรุง.ชื่อกิจกรรม, key=f"m_name_{i}_{j}", label_visibility="collapsed")
                    with cm[1]:
                        บำรุง.ต้นทุนต่อหน่วย = st.number_input("บ.", value=float(บำรุง.ต้นทุนต่อหน่วย), min_value=0.0, step=5.0, key=f"m_cost_{i}_{j}", label_visibility="collapsed")
                    with cm[2]:
                        บำรุง.ปีเริ่มต้น = st.number_input("ป.", value=int(บำรุง.ปีเริ่มต้น), min_value=1, max_value=50, key=f"m_year_{i}_{j}", label_visibility="collapsed")
                    with cm[3]:
                        บำรุง.ความถี่ = st.number_input("ถ.", value=int(บำรุง.ความถี่), min_value=0, max_value=20, key=f"m_freq_{i}_{j}", label_visibility="collapsed")
                    with cm[4]:
                        if st.button("🗑️", key=f"del_m_{i}_{j}"):
                            del_maint.append(j)

                for idx in sorted(del_maint, reverse=True):
                    if len(ทางเลือก.แผนบำรุงรักษา) > 1:
                        ทางเลือก.แผนบำรุงรักษา.pop(idx)
                        st.rerun()

                if st.button(f"➕ เพิ่มบำรุงรักษา", key=f"add_m_{i}"):
                    ทางเลือก.แผนบำรุงรักษา.append(กิจกรรมบำรุงรักษา("กิจกรรมใหม่", 50.0, 5, 5))
                    st.rerun()

                # แผนฟื้นฟูสภาพ
                st.markdown("---")
                st.subheader("🏗️ แผนฟื้นฟูสภาพ (Rehabilitation)")
                cols_rh = st.columns([4, 2, 1, 0.5])
                cols_rh[0].markdown("**กิจกรรม**")
                cols_rh[1].markdown("**บาท/ตร.ม.**")
                cols_rh[2].markdown("**ปีที่**")
                cols_rh[3].markdown("**ลบ**")

                del_rehab = []
                for k, ฟื้นฟู in enumerate(ทางเลือก.แผนฟื้นฟูสภาพ):
                    cr = st.columns([4, 2, 1, 0.5])
                    with cr[0]:
                        ฟื้นฟู.ชื่อกิจกรรม = st.text_input("ก.", value=ฟื้นฟู.ชื่อกิจกรรม, key=f"r_name_{i}_{k}", label_visibility="collapsed")
                    with cr[1]:
                        ฟื้นฟู.ต้นทุนต่อหน่วย = st.number_input("บ.", value=float(ฟื้นฟู.ต้นทุนต่อหน่วย), min_value=0.0, step=10.0, key=f"r_cost_{i}_{k}", label_visibility="collapsed")
                    with cr[2]:
                        ฟื้นฟู.ปีดำเนินการ = st.number_input("ป.", value=int(ฟื้นฟู.ปีดำเนินการ), min_value=1, max_value=50, key=f"r_year_{i}_{k}", label_visibility="collapsed")
                    with cr[3]:
                        if st.button("🗑️", key=f"del_r_{i}_{k}"):
                            del_rehab.append(k)

                for idx in sorted(del_rehab, reverse=True):
                    if len(ทางเลือก.แผนฟื้นฟูสภาพ) > 1:
                        ทางเลือก.แผนฟื้นฟูสภาพ.pop(idx)
                        st.rerun()

                if st.button(f"➕ เพิ่มฟื้นฟูสภาพ", key=f"add_r_{i}"):
                    ปีล่าสุด = max([ฟ.ปีดำเนินการ for ฟ in ทางเลือก.แผนฟื้นฟูสภาพ]) if ทางเลือก.แผนฟื้นฟูสภาพ else 10
                    ทางเลือก.แผนฟื้นฟูสภาพ.append(กิจกรรมฟื้นฟูสภาพ("Overlay AC 50 มม.", 400.0, ปีล่าสุด + 10))
                    st.rerun()

        # สรุปทุกทางเลือก
        st.divider()
        st.subheader("📊 สรุปต้นทุนก่อสร้างทุกทางเลือก")
        summary_data = []
        for ท in st.session_state.ทางเลือกทั้งหมด:
            summary_data.append({
                'ทางเลือก': ท.ชื่อ,
                'ประเภท': ท.ประเภท,
                'ต้นทุน (บาท/ตร.ม.)': ท.ต้นทุนก่อสร้าง,
                'พื้นที่ (ตร.ม.)': ท.พื้นที่,
                'ต้นทุนรวม (บาท)': ท.ต้นทุนก่อสร้าง * ท.พื้นที่,
                'สถานะ': '✅' if ท.เปิดใช้งาน else '❌'
            })
        st.dataframe(pd.DataFrame(summary_data), use_container_width=True, hide_index=True)

        # JSON Export
        st.divider()
        col_s1, col_s2 = st.columns(2)
        with col_s1:
            data_export = {
                'ระยะวิเคราะห์': ระยะวิเคราะห์,
                'อัตราคิดลด': อัตราคิดลด,
                'ทางเลือก': [ทางเลือก_เป็น_dict(ท) for ท in st.session_state.ทางเลือกทั้งหมด]
            }
            st.download_button("💾 บันทึก JSON", json.dumps(data_export, ensure_ascii=False, indent=2),
                             "lcca_v3_data.json", "application/json", use_container_width=True)

    # ===== Tab 3: ผลวิเคราะห์ LCCA =====
    with tab3:
        st.header("📈 ผลการวิเคราะห์ LCCA")

        col_p1, col_p2, col_p3 = st.columns(3)
        col_p1.info(f"📅 ระยะวิเคราะห์: **{ระยะวิเคราะห์} ปี**")
        col_p2.info(f"📉 อัตราคิดลด: **{อัตราคิดลด * 100:.1f}%**")
        ทางเลือกที่ใช้ = [ท for ท in st.session_state.ทางเลือกทั้งหมด if ท.เปิดใช้งาน]
        col_p3.info(f"🛣️ ทางเลือก: **{len(ทางเลือกที่ใช้)}**")

        if not รวมมูลค่าซาก:
            st.warning("⚠️ **ไม่รวมมูลค่าซาก (Salvage Value)** — การวิเคราะห์ไม่ได้หักมูลค่าซากออกจากต้นทุน สามารถเปิดได้ที่ Sidebar")

        if len(ทางเลือกที่ใช้) == 0:
            st.warning("⚠️ กรุณาเปิดใช้งานอย่างน้อย 1 ทางเลือก")
        else:
            สรุป, กระแสเงินสด = วิเคราะห์_LCCA(st.session_state.ทางเลือกทั้งหมด, ระยะวิเคราะห์, อัตราคิดลด, รวมมูลค่าซาก)

            if len(สรุป) > 0:
                ผู้ชนะ = สรุป.iloc[0]
                c1, c2, c3, c4 = st.columns(4)
                c1.metric("🏆 ทางเลือกที่ดีที่สุด", ผู้ชนะ['ทางเลือก'])
                c2.metric("💰 PW รวม", f"{ผู้ชนะ['มูลค่าปัจจุบันรวม']:,.0f} บาท")
                c3.metric("📊 EAC", f"{ผู้ชนะ['ต้นทุนเฉลี่ยรายปี']:,.0f} บาท/ปี")
                c4.metric("📐 ต้นทุน/ตร.ม./ปี", f"{ผู้ชนะ['ต้นทุนต่อตรม_ต่อปี']:,.2f} บาท")

                st.divider()

                # ตารางสรุป
                st.subheader("📋 ตารางเปรียบเทียบ")
                display_cols = ['ลำดับ', 'ทางเลือก', 'ประเภทผิวทาง', 'ต้นทุนก่อสร้าง_ตรม', 'มูลค่าปัจจุบันรวม', 'ต้นทุนเฉลี่ยรายปี', 'ต้นทุนต่อตรม_ต่อปี']
                st.dataframe(
                    สรุป[display_cols].style.format({
                        'ต้นทุนก่อสร้าง_ตรม': '{:,.2f}',
                        'มูลค่าปัจจุบันรวม': '{:,.0f}',
                        'ต้นทุนเฉลี่ยรายปี': '{:,.0f}',
                        'ต้นทุนต่อตรม_ต่อปี': '{:,.2f}'
                    }).background_gradient(subset=['มูลค่าปัจจุบันรวม'], cmap='RdYlGn_r'),
                    use_container_width=True
                )

                # กราฟเปรียบเทียบ
                st.subheader("📊 กราฟเปรียบเทียบ")
                fig = make_subplots(rows=1, cols=2,
                    subplot_titles=('มูลค่าปัจจุบันรวม (บาท)', 'องค์ประกอบต้นทุน (PW)'))

                colors = ['#2E86AB', '#A23B72', '#F18F01', '#C73E1D', '#28A745', '#6F42C1']

                fig.add_trace(go.Bar(
                    x=สรุป['ทางเลือก'], y=สรุป['มูลค่าปัจจุบันรวม'],
                    marker_color=colors[:len(สรุป)],
                    text=สรุป['มูลค่าปัจจุบันรวม'].apply(lambda x: f'{x:,.0f}'),
                    textposition='outside', name='PW รวม'
                ), row=1, col=1)

                fig.add_trace(go.Bar(x=สรุป['ทางเลือก'], y=สรุป['PW_ก่อสร้าง'], name='ก่อสร้าง', marker_color='#2E86AB'), row=1, col=2)
                fig.add_trace(go.Bar(x=สรุป['ทางเลือก'], y=สรุป['PW_บำรุงรักษา'], name='บำรุงรักษา', marker_color='#F18F01'), row=1, col=2)
                fig.add_trace(go.Bar(x=สรุป['ทางเลือก'], y=สรุป['PW_ฟื้นฟูสภาพ'], name='ฟื้นฟูสภาพ', marker_color='#A23B72'), row=1, col=2)
                fig.add_trace(go.Bar(x=สรุป['ทางเลือก'], y=สรุป['PW_มูลค่าซาก'], name='มูลค่าซาก', marker_color='#28A745'), row=1, col=2)

                fig.update_layout(height=450, barmode='relative', legend=dict(orientation="h", y=1.1))
                st.plotly_chart(fig, use_container_width=True)

                # Cumulative NPV Timeline
                st.subheader("📈 Cumulative PW Timeline")
                fig_cum = go.Figure()
                for idx, (ชื่อ, cf) in enumerate(กระแสเงินสด.items()):
                    cf_sorted = cf.sort_values('ปี')
                    cum_pw = cf_sorted['มูลค่าปัจจุบัน'].cumsum()
                    fig_cum.add_trace(go.Scatter(
                        x=cf_sorted['ปี'], y=cum_pw, mode='lines+markers',
                        name=ชื่อ, line=dict(color=colors[idx % len(colors)], width=2)
                    ))
                fig_cum.update_layout(
                    xaxis_title='ปี', yaxis_title='Cumulative PW (บาท)',
                    height=400, hovermode='x unified'
                )
                st.plotly_chart(fig_cum, use_container_width=True)

    # ===== Tab 4: กระแสเงินสด =====
    with tab4:
        st.header("💰 ตารางกระแสเงินสดรายปี")
        ทางเลือกที่ใช้ = [ท for ท in st.session_state.ทางเลือกทั้งหมด if ท.เปิดใช้งาน]
        if len(ทางเลือกที่ใช้) == 0:
            st.warning("⚠️ กรุณาเปิดใช้งานอย่างน้อย 1 ทางเลือก")
        else:
            สรุป, กระแสเงินสด = วิเคราะห์_LCCA(st.session_state.ทางเลือกทั้งหมด, ระยะวิเคราะห์, อัตราคิดลด, รวมมูลค่าซาก)
            เลือก = st.selectbox("เลือกทางเลือก:", [ท.ชื่อ for ท in ทางเลือกที่ใช้])
            if เลือก in กระแสเงินสด:
                cf = กระแสเงินสด[เลือก].copy()
                c1, c2, c3 = st.columns(3)
                c1.metric("ต้นทุนตามปีรวม", f"{cf['ต้นทุนตามปี'].sum():,.0f} บาท")
                c2.metric("มูลค่าปัจจุบันรวม", f"{cf['มูลค่าปัจจุบัน'].sum():,.0f} บาท")
                eac = คำนวณต้นทุนเฉลี่ยรายปี(cf['มูลค่าปัจจุบัน'].sum(), อัตราคิดลด, ระยะวิเคราะห์)
                c3.metric("EAC", f"{eac:,.0f} บาท/ปี")

                cf_display = cf.copy()
                cf_display['ต้นทุนต่อหน่วย'] = cf_display['ต้นทุนต่อหน่วย'].apply(lambda x: f"{x:,.2f}")
                cf_display['ต้นทุนตามปี'] = cf_display['ต้นทุนตามปี'].apply(lambda x: f"{x:,.0f}")
                cf_display['ตัวคูณ_PW'] = cf_display['ตัวคูณ_PW'].apply(lambda x: f"{x:.4f}")
                cf_display['มูลค่าปัจจุบัน'] = cf_display['มูลค่าปัจจุบัน'].apply(lambda x: f"{x:,.0f}")
                cf_display.columns = ['ปี', 'กิจกรรม', 'ประเภท', 'ต้นทุน/หน่วย', 'ต้นทุนตามปี (บาท)', 'ตัวคูณ PW', 'มูลค่าปัจจุบัน (บาท)']
                st.dataframe(cf_display, use_container_width=True, hide_index=True, height=500)

                # Timeline
                st.subheader("📅 Timeline")
                cf_plot = cf[cf['ต้นทุนตามปี'] != 0].copy()
                cf_plot['abs_cost'] = cf_plot['ต้นทุนตามปี'].abs()
                fig_tl = px.scatter(cf_plot, x='ปี', y='มูลค่าปัจจุบัน', size='abs_cost', color='ประเภท',
                                    hover_name='กิจกรรม', title=f'Timeline - {เลือก}')
                fig_tl.update_layout(height=400)
                st.plotly_chart(fig_tl, use_container_width=True)

                csv = cf.to_csv(index=False).encode('utf-8-sig')
                st.download_button("⬇️ ดาวน์โหลด CSV", csv, f"cashflow_{เลือก}.csv", "text/csv")

    # ===== Tab 5: Sensitivity Analysis =====
    with tab5:
        st.header("📉 การวิเคราะห์ความไว (Sensitivity Analysis)")
        ทางเลือกที่ใช้ = [ท for ท in st.session_state.ทางเลือกทั้งหมด if ท.เปิดใช้งาน]
        if len(ทางเลือกที่ใช้) == 0:
            st.warning("⚠️ กรุณาเปิดใช้งานอย่างน้อย 1 ทางเลือก")
        else:
            st.subheader("ความไวต่ออัตราคิดลด")
            ผลอัตรา, pivot = วิเคราะห์ความไว_อัตราคิดลด(
                st.session_state.ทางเลือกทั้งหมด, ระยะวิเคราะห์, อัตราคิดลด, ช่วงอัตราคิดลด,
                รวมมูลค่าซาก=รวมมูลค่าซาก
            )
            if len(ผลอัตรา) > 0:
                fig_sens = px.line(ผลอัตรา, x='อัตราคิดลด', y='มูลค่าปัจจุบัน', color='ทางเลือก',
                                   markers=True, title='ผลกระทบของอัตราคิดลดต่อมูลค่าปัจจุบัน')
                fig_sens.update_layout(height=500)
                fig_sens.update_xaxes(tickformat='.1%')
                st.plotly_chart(fig_sens, use_container_width=True)

                st.markdown("**ตารางสรุป:**")
                pivot_display = pivot.copy()
                for col in pivot_display.columns:
                    pivot_display[col] = pivot_display[col].apply(lambda x: f"{x:,.0f}")
                st.dataframe(pivot_display, use_container_width=True)

                อัตราต่ำ = ผลอัตรา['อัตราคิดลด'].min()
                อัตราสูง = ผลอัตรา['อัตราคิดลด'].max()
                ชนะต่ำ = ผลอัตรา[ผลอัตรา['อัตราคิดลด'] == อัตราต่ำ].nsmallest(1, 'มูลค่าปัจจุบัน')['ทางเลือก'].values[0]
                ชนะสูง = ผลอัตรา[ผลอัตรา['อัตราคิดลด'] == อัตราสูง].nsmallest(1, 'มูลค่าปัจจุบัน')['ทางเลือก'].values[0]
                if ชนะต่ำ == ชนะสูง:
                    st.success(f"✅ **{ชนะต่ำ}** เป็นทางเลือกที่ดีที่สุดในทุกอัตราคิดลด (Robust Decision)")
                else:
                    st.warning(f"⚠️ ทางเลือกเปลี่ยน: {ชนะต่ำ} (อัตราต่ำ) vs {ชนะสูง} (อัตราสูง)")

    # ===== Tab 6: รายงาน =====
    with tab6:
        st.header("📄 รายงาน")
        ทางเลือกที่ใช้ = [ท for ท in st.session_state.ทางเลือกทั้งหมด if ท.เปิดใช้งาน]
        if len(ทางเลือกที่ใช้) == 0:
            st.warning("⚠️ กรุณาเปิดใช้งานอย่างน้อย 1 ทางเลือก")
        else:
            สรุป, กระแสเงินสด = วิเคราะห์_LCCA(st.session_state.ทางเลือกทั้งหมด, ระยะวิเคราะห์, อัตราคิดลด, รวมมูลค่าซาก)
            col_e1, col_e2, col_e3 = st.columns(3)
            with col_e1:
                if DOCX_AVAILABLE:
                    word_file = สร้างรายงาน_Word(สรุป, กระแสเงินสด, ระยะวิเคราะห์, อัตราคิดลด, st.session_state.ทางเลือกทั้งหมด)
                    st.download_button("📝 ดาวน์โหลด Word", word_file,
                        f"LCCA_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True)
                else:
                    st.warning("⚠️ ติดตั้ง python-docx: `pip install python-docx`")
            with col_e2:
                csv_summary = สรุป.to_csv(index=False).encode('utf-8-sig')
                st.download_button("📊 ดาวน์โหลดสรุป CSV", csv_summary,
                    f"LCCA_Summary_{datetime.now().strftime('%Y%m%d_%H%M')}.csv", "text/csv",
                    use_container_width=True)
            with col_e3:
                data_export = {
                    'ระยะวิเคราะห์': ระยะวิเคราะห์, 'อัตราคิดลด': อัตราคิดลด,
                    'ทางเลือก': [ทางเลือก_เป็น_dict(ท) for ท in st.session_state.ทางเลือกทั้งหมด]
                }
                st.download_button("💾 ดาวน์โหลด JSON", json.dumps(data_export, ensure_ascii=False, indent=2),
                    f"LCCA_Data_{datetime.now().strftime('%Y%m%d_%H%M')}.json", "application/json",
                    use_container_width=True)

    # ===== Tab 7: ทฤษฎี =====
    with tab7:
        st.header("ℹ️ ทฤษฎี Life-Cycle Cost Analysis (LCCA)")
        st.markdown("""
        ## 1. ประเภทผิวทางคอนกรีต
        
        | ประเภท | ชื่อเต็ม | ลักษณะเด่น |
        |--------|---------|-----------|
        | **JPCP** | Jointed Plain Concrete Pavement | คอนกรีตไม่เสริมเหล็ก มีรอยต่อทุก 4-6 ม. |
        | **JRCP** | Jointed Reinforced Concrete Pavement | คอนกรีตเสริมเหล็ก รอยต่อห่าง 8-15 ม. |
        | **CRCP** | Continuously Reinforced Concrete Pavement | เสริมเหล็กต่อเนื่อง ไม่มีรอยต่อตามขวาง |
        
        ## 2. สูตรคำนวณหลัก
        
        ### 2.1 มูลค่าปัจจุบัน (Present Worth)
        """)
        st.latex(r"PW = FV \times (1 + i)^{-n}")
        st.markdown("### 2.2 ต้นทุนเฉลี่ยรายปี (Equivalent Annual Cost)")
        st.latex(r"EAC = PW \times \frac{i(1+i)^N}{(1+i)^N - 1}")
        st.markdown("""
        ### 2.3 มูลค่าซาก (Salvage Value) — Straight-Line Depreciation
        """)
        st.latex(r"SV = C_{rehab} - \frac{C_{rehab} \times (1 - SV\%)}{N_{expected}} \times (N_{analysis} - N_{rehab})")
        st.markdown("""
        ## 3. เปรียบเทียบผิวทาง
        
        | เกณฑ์ | AC | JPCP | JRCP | CRCP |
        |------|-----|------|------|------|
        | ต้นทุนก่อสร้าง | ต่ำ | ปานกลาง | ปานกลาง-สูง | สูง |
        | ค่าบำรุงรักษา | สูง | ปานกลาง | ปานกลาง | ต่ำ |
        | อายุใช้งาน | 15-20 ปี | 20-30 ปี | 25-35 ปี | 30-40 ปี |
        
        ## 4. เอกสารอ้างอิง
        
        - FHWA-SA-98-079: Life-Cycle Cost Analysis in Pavement Design
        - AASHTO Guide for Design of Pavement Structures (1993)
        - NCHRP Report 703: Guide for Pavement-Type Selection
        - มาตรฐานกรมทางหลวง
        """)

    # Footer
    st.divider()
    st.markdown("""
    ---
    **โปรแกรมวิเคราะห์ LCCA ผิวทาง v3.0** (Combined Edition) | พัฒนาสำหรับการเรียนการสอนด้านวิศวกรรมทาง  
    ภาควิชาครุศาสตร์โยธา มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ
    """)


if __name__ == "__main__":
    main()
