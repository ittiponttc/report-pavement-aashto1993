# =========================================================
# AASHTO 1993 Nomograph – FINAL v3 (STABLE)
# Streamlit Cloud Ready | PDF + Word Export
# =========================================================

import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
from PIL import Image
from io import BytesIO

# ---------------- PDF (safe import) ----------------
try:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False

# ---------------- Word ----------------
from docx import Document
from docx.shared import Inches

# =========================================================
# Page config
# =========================================================

st.set_page_config(
    page_title="AASHTO 1993 Nomograph",
    layout="centered"
)

st.title("AASHTO 1993 – Composite Modulus of Subgrade Reaction")

# =========================================================
# Mapping functions (log scale)
# =========================================================

def log_map(v, vmin, vmax, pmin, pmax):
    return pmin + (np.log10(v) - np.log10(vmin)) / \
        (np.log10(vmax) - np.log10(vmin)) * (pmax - pmin)


def log_unmap(p, vmin, vmax, pmin, pmax):
    r = (p - pmin) / (pmax - pmin)
    return 10 ** (np.log10(vmin) + r * (np.log10(vmax) - np.log10(vmin)))

# =========================================================
# User input
# =========================================================

Mr = st.slider(
    "Roadbed Resilient Modulus, Mr (psi)",
    1000, 20000, 6000, step=500
)

DSB = st.slider(
    "Subbase Thickness, DSB (inch)",
    4, 18, 14
)

# =========================================================
# Load Nomograph Image (Cloud-safe)
# =========================================================

st.subheader("Nomograph Image")

uploaded_file = st.file_uploader(
    "Upload AASHTO Nomograph Image (PNG / JPG)",
    type=["png", "jpg", "jpeg"]
)

if uploaded_file is not None:
    img = Image.open(uploaded_file)
else:
    try:
        img = Image.open("nomograph.png")
    except FileNotFoundError:
        st.error(
            "❌ ไม่พบไฟล์ nomograph.png\n\n"
            "กรุณาอัปโหลดภาพ Nomograph ก่อน"
        )
        st.stop()

# =========================================================
# Paper template coordinates (A4 Landscape)
# =========================================================
# Logical coordinate system: x = 0–100, y = 0–70

DSB_X = (20, 75)      # แกน Subbase thickness
MR_Y  = (15, 55)      # แกน Mr
K_X   = (70, 95)      # แกน k∞

# =========================================================
# Mapping input → coordinates
# =========================================================

x_dsb = log_map(DSB, 4, 18, DSB_X[0], DSB_X[1])
y_mr  = log_map(Mr, 1000, 20000, MR_Y[1], MR_Y[0])

# อ่านค่า k∞ จากตำแหน่งเส้นแนวนอนตัดโซน k
k_inf = log_unmap(
    x_dsb,
    50, 2000,          # ช่วง k∞ ใน nomograph
    K_X[0], K_X[1]
)


# =========================================================
# Plot Nomograph with A4 Template
# =========================================================

fig, ax = plt.subplots(figsize=(11.7, 8.3))  # A4 landscape
ax.set_facecolor("white")

# วางภาพลงบนกระดาษ
# =========================================================
# Place Nomograph Image on A4 Template (Correct Scaling)
# =========================================================

IMG_W, IMG_H = img.size   # 739 x 671

paper_w, paper_h = 100, 70

scale = paper_h / IMG_H
img_w_scaled = IMG_W * scale

x0 = (paper_w - img_w_scaled) / 2
x1 = x0 + img_w_scaled
y0 = 0
y1 = paper_h

ax.imshow(
    img,
    extent=[x0, x1, y0, y1],
    aspect="auto"
)

# กรอบกระดาษ
ax.plot(
    [0, paper_w, paper_w, 0, 0],
    [0, 0, paper_h, paper_h, 0],
    color="black",
    linewidth=1.5
)


# กรอบกระดาษ
ax.plot([0,100,100,0,0], [0,0,70,70,0],
        color="black", linewidth=1.5)

# เส้นอ่านค่า
ax.plot([x_dsb, x_dsb], [y_mr, 10], color="red", linewidth=2)
ax.plot([x_dsb, K_X[1]], [y_mr, y_mr], color="red", linewidth=2)
ax.scatter(x_dsb, y_mr, color="red", s=90, zorder=5)

ax.axis("off")

st.pyplot(fig)

st.success(f"Estimated composite k∞ ≈ {k_inf:,.0f} pci")

# =========================================================
# Save figure for report
# =========================================================

FIG_BUFFER = BytesIO()
fig.savefig(FIG_BUFFER, dpi=300, bbox_inches="tight")
FIG_BUFFER.seek(0)

# =========================================================
# Export PDF
# =========================================================

def export_pdf():
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()

    content = [
        Paragraph("<b>AASHTO 1993 Nomograph Report</b>", styles["Title"]),
        Spacer(1, 12),
        Paragraph(f"Roadbed Resilient Modulus (Mr): {Mr:,.0f} psi", styles["Normal"]),
        Paragraph(f"Subbase Thickness (DSB): {DSB:.1f} inch", styles["Normal"]),
        Paragraph(f"Composite Modulus of Subgrade Reaction (k∞): {k_inf:,.0f} pci", styles["Normal"]),
    ]

    doc.build(content)
    buffer.seek(0)
    return buffer

# =========================================================
# Export Word
# =========================================================

def export_word():
    doc = Document()
    doc.add_heading("AASHTO 1993 Nomograph Report", level=1)

    doc.add_paragraph(f"Roadbed Resilient Modulus (Mr): {Mr:,.0f} psi")
    doc.add_paragraph(f"Subbase Thickness (DSB): {DSB:.1f} inch")
    doc.add_paragraph(f"Composite Modulus of Subgrade Reaction (k∞): {k_inf:,.0f} pci")

    doc.add_heading("Nomograph Interpretation", level=2)
    doc.add_picture(FIG_BUFFER, width=Inches(6.5))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# =========================================================
# Download section
# =========================================================

st.subheader("Export Report")

col1, col2 = st.columns(2)

with col1:
    if REPORTLAB_OK:
        st.download_button(
            "📄 Download PDF",
            export_pdf(),
            file_name="AASHTO1993_Nomograph_Report.pdf",
            mime="application/pdf"
        )
    else:
        st.warning("PDF export disabled (reportlab not installed)")

with col2:
    st.download_button(
        "📝 Download Word",
        export_word(),
        file_name="AASHTO1993_Nomograph_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
