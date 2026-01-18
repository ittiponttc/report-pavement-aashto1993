"""
โปรแกรมออกแบบและตรวจสอบความหนาถนนคอนกรีต (Rigid Pavement)
ตามวิธี AASHTO 1993
รองรับทั้ง JPCP, JRCP และ CRCP

ปรับปรุงเพิ่มเติม:
- Odemark Method สำหรับคำนวณ Equivalent Thickness
- **** รายงาน Word แบบละเอียดพร้อมรูปโครงสร้างชั้นทาง พร้อมขั้นตอนการคำนวณ
- สามารถใส่ชื่อโครงการได้

พัฒนาสำหรับใช้ในการเรียนการสอน
ภาควิชาครุศาสตร์โยธา มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ
"""

import streamlit as st
import math
from io import BytesIO
from datetime import datetime
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import os
import tempfile

# ============================================================
# ส่วนที่ 1: ค่าคงที่และตารางอ้างอิง
# ============================================================

ZR_TABLE = {
    50: -0.000, 60: -0.253, 70: -0.524, 75: -0.674,
    80: -0.841, 85: -1.037, 90: -1.282, 91: -1.340,
    92: -1.405, 93: -1.476, 94: -1.555, 95: -1.645,
    96: -1.751, 97: -1.881, 98: -2.054, 99: -2.327
}

J_VALUES = {"JRCP": 2.8, "JPCP": 2.8, "JRCP/JPCP": 2.8, "CRCP": 2.5}

MATERIAL_MODULUS = {
    "ผิวทางลาดยาง AC": 2500,
    "ผิวทางลาดยาง PMA": 3700,
    "พื้นทางซีเมนต์ CTB": 1200,
    "หินคลุกผสมซีเมนต์ UCS 24.5 ksc": 850,
    "หินคลุก CBR 80%": 350,
    "ดินซีเมนต์ UCS 17.5 ksc": 350,
    "วัสดุหมุนเวียน (Recycling)": 850,
    "รองพื้นทางวัสดุมวลรวม CBR 25%": 150,
    "วัสดุคัดเลือก ก": 100,
    "ดินถมคันทาง / ดินเดิม": 100,
    "กำหนดเอง...": 100,
}

LAYER_COLORS = {
    "ผิวทางลาดยาง AC": "#2C3E50",
    "ผิวทางลาดยาง PMA": "#1A252F",
    "พื้นทางซีเมนต์ CTB": "#7F8C8D",
    "หินคลุกผสมซีเมนต์ UCS 24.5 ksc": "#95A5A6",
    "หินคลุก CBR 80%": "#BDC3C7",
    "ดินซีเมนต์ UCS 17.5 ksc": "#AAB7B8",
    "วัสดุหมุนเวียน (Recycling)": "#85929E",
    "รองพื้นทางวัสดุมวลรวม CBR 25%": "#D5DBDB",
    "วัสดุคัดเลือก ก": "#E8DAEF",
    "ดินถมคันทาง / ดินเดิม": "#F5CBA7",
    "กำหนดเอง...": "#FADBD8",
    "Concrete Slab": "#5DADE2",
}

THAI_TO_ENG = {
    "ผิวทางลาดยาง AC": "AC Surface",
    "ผิวทางลาดยาง PMA": "PMA Surface",
    "พื้นทางซีเมนต์ CTB": "Cement Treated Base",
    "หินคลุกผสมซีเมนต์ UCS 24.5 ksc": "Soil Cement",
    "หินคลุก CBR 80%": "Crushed Rock Base",
    "ดินซีเมนต์ UCS 17.5 ksc": "Soil Cement",
    "วัสดุหมุนเวียน (Recycling)": "Recycled Material",
    "รองพื้นทางวัสดุมวลรวม CBR 25%": "Aggregate Subbase",
    "วัสดุคัดเลือก ก": "Selected Material",
    "ดินถมคันทาง / ดินเดิม": "Subgrade",
    "กำหนดเอง...": "Custom Material",
    "Concrete Slab": "Concrete Slab",
}

# ============================================================
# ส่วนที่ 2: ฟังก์ชันการคำนวณพื้นฐาน
# ============================================================

def convert_cube_to_cylinder(fc_cube_ksc):
    return 0.8 * fc_cube_ksc

def calculate_concrete_modulus(fc_cylinder_ksc):
    fc_psi = fc_cylinder_ksc * 14.223
    return 57000 * math.sqrt(fc_psi)

def estimate_modulus_of_rupture(fc_cylinder_ksc):
    fc_psi = fc_cylinder_ksc * 14.223
    return 10.0 * math.sqrt(fc_psi)

def get_zr_value(reliability):
    return ZR_TABLE.get(int(reliability), -1.282)

# ============================================================
# ส่วนที่ 3: Odemark Method
# ============================================================

def calculate_odemark_equivalent_thickness(layers, e_subgrade=50):
    """
    คำนวณความหนาเทียบเท่าตามวิธี Odemark
    สูตร: he = h × (E_layer/E_subgrade)^(1/3)
    """
    results = {
        "layers": [],
        "total_actual_thickness": 0,
        "total_equivalent_thickness": 0,
        "e_subgrade": e_subgrade,
        "calculation_steps": []
    }
    
    total_actual = 0
    total_equivalent = 0
    
    for i, layer in enumerate(layers):
        h_cm = layer.get("thickness_cm", 0)
        e_mpa = layer.get("E_MPa", 100)
        name = layer.get("name", f"Layer {i+1}")
        
        if h_cm <= 0:
            continue
        
        if e_subgrade > 0:
            ratio = e_mpa / e_subgrade
            factor = ratio ** (1/3)
            he_cm = h_cm * factor
        else:
            ratio = 1
            factor = 1
            he_cm = h_cm
        
        total_actual += h_cm
        total_equivalent += he_cm
        
        results["layers"].append({
            "name": name,
            "h_actual": h_cm,
            "E_MPa": e_mpa,
            "E_ratio": ratio,
            "he_equivalent": he_cm,
            "factor": factor
        })
        
        step = f"Layer {i+1}: {name}\n"
        step += f"   h = {h_cm:.1f} cm, E = {e_mpa:,} MPa\n"
        step += f"   he = {h_cm:.1f} x ({e_mpa:,}/{e_subgrade:.0f})^(1/3)\n"
        step += f"   he = {h_cm:.1f} x {factor:.4f} = {he_cm:.2f} cm"
        results["calculation_steps"].append(step)
    
    results["total_actual_thickness"] = total_actual
    results["total_equivalent_thickness"] = total_equivalent
    
    return results

# ============================================================
# ส่วนที่ 4: AASHTO 1993 Calculation
# ============================================================

def calculate_aashto_rigid_w18(d_inch, delta_psi, pt, zr, so, sc_psi, cd, j, ec_psi, k_pci):
    """คำนวณ ESAL ที่รองรับได้ตามสมการ AASHTO 1993"""
    
    term1 = zr * so
    term2 = 7.35 * math.log10(d_inch + 1) - 0.06
    
    numerator3 = math.log10(delta_psi / 3.0)
    denominator3 = 1 + (1.624e7 / ((d_inch + 1) ** 8.46))
    term3 = numerator3 / denominator3
    
    d_power = d_inch ** 0.75
    numerator4 = sc_psi * cd * (d_power - 1.132)
    ec_k_ratio = ec_psi / k_pci
    denominator4 = 215.63 * j * (d_power - 18.42 / (ec_k_ratio ** 0.25))
    
    if numerator4 <= 0 or denominator4 <= 0:
        return (float('-inf'), 0)
    
    inner_term = numerator4 / denominator4
    if inner_term <= 0:
        return (float('-inf'), 0)
    
    term4 = (4.22 - 0.32 * pt) * math.log10(inner_term)
    
    log10_w18 = term1 + term2 + term3 + term4
    w18 = 10 ** log10_w18
    
    return (log10_w18, w18)

def calculate_aashto_detailed(d_inch, delta_psi, pt, zr, so, sc_psi, cd, j, ec_psi, k_pci):
    """คำนวณ AASHTO พร้อมแสดงรายละเอียดแต่ละขั้นตอน"""
    details = {}
    
    term1 = zr * so
    details["term1"] = {"value": term1, "zr": zr, "so": so}
    
    term2 = 7.35 * math.log10(d_inch + 1) - 0.06
    details["term2"] = {"value": term2, "log_d1": math.log10(d_inch + 1)}
    
    numerator3 = math.log10(delta_psi / 3.0)
    denominator3 = 1 + (1.624e7 / ((d_inch + 1) ** 8.46))
    term3 = numerator3 / denominator3
    details["term3"] = {"value": term3, "numerator": numerator3, "denominator": denominator3}
    
    d_power = d_inch ** 0.75
    numerator4 = sc_psi * cd * (d_power - 1.132)
    ec_k_ratio = ec_psi / k_pci
    denominator4 = 215.63 * j * (d_power - 18.42 / (ec_k_ratio ** 0.25))
    
    if numerator4 > 0 and denominator4 > 0:
        inner_term = numerator4 / denominator4
        term4 = (4.22 - 0.32 * pt) * math.log10(inner_term)
    else:
        inner_term = 0
        term4 = float('-inf')
    
    details["term4"] = {
        "value": term4,
        "d_power": d_power,
        "ec_k_ratio": ec_k_ratio,
        "numerator": numerator4,
        "denominator": denominator4,
        "inner_term": inner_term,
        "coefficient": 4.22 - 0.32 * pt
    }
    
    log10_w18 = term1 + term2 + term3 + term4
    w18 = 10 ** log10_w18 if log10_w18 > float('-inf') else 0
    
    details["result"] = {"log10_w18": log10_w18, "w18": w18}
    
    return details

def check_design(w18_required, w18_capacity):
    ratio = w18_capacity / w18_required if w18_required > 0 else float('inf')
    return (w18_capacity >= w18_required, ratio)

# ============================================================
# ส่วนที่ 5: สร้างรูปโครงสร้างชั้นทาง
# ============================================================

def create_pavement_structure_figure(layers_data, concrete_thickness_cm=None):
    """สร้างรูปโครงสร้างชั้นทาง"""
    
    valid_layers = [l for l in layers_data if l.get("thickness_cm", 0) > 0]
    
    all_layers = []
    if concrete_thickness_cm and concrete_thickness_cm > 0:
        all_layers.append({
            "name": "Concrete Slab",
            "thickness_cm": concrete_thickness_cm,
            "E_MPa": None
        })
    all_layers.extend(valid_layers)
    
    if not all_layers:
        return None
    
    total_thickness = sum(l.get("thickness_cm", 0) for l in all_layers)
    min_display_height = 8
    
    fig, ax = plt.subplots(figsize=(12, 8))
    
    width = 3
    x_center = 6
    x_start = x_center - width / 2
    
    display_heights = [max(l.get("thickness_cm", 0), min_display_height) for l in all_layers]
    total_display = sum(display_heights)
    y_current = total_display
    
    for i, layer in enumerate(all_layers):
        thickness = layer.get("thickness_cm", 0)
        name = layer.get("name", f"Layer {i+1}")
        e_mpa = layer.get("E_MPa", None)
        display_h = display_heights[i]
        
        if thickness <= 0:
            continue
        
        color = LAYER_COLORS.get(name, "#CCCCCC")
        y_bottom = y_current - display_h
        
        rect = patches.Rectangle(
            (x_start, y_bottom), width, display_h,
            linewidth=2, edgecolor='black', facecolor=color
        )
        ax.add_patch(rect)
        
        y_center_pos = y_bottom + display_h / 2
        display_name = THAI_TO_ENG.get(name, name)
        
        is_dark = name in ["ผิวทางลาดยาง AC", "ผิวทางลาดยาง PMA", "Concrete Slab",
                          "พื้นทางซีเมนต์ CTB", "หินคลุกผสมซีเมนต์ UCS 24.5 ksc",
                          "วัสดุหมุนเวียน (Recycling)"]
        text_color = 'white' if is_dark else 'black'
        
        ax.text(x_center, y_center_pos, f"{thickness} cm",
                ha='center', va='center', fontsize=11, fontweight='bold', color=text_color)
        ax.text(x_start - 0.5, y_center_pos, display_name,
                ha='right', va='center', fontsize=10, fontweight='bold', color='black')
        
        if e_mpa:
            ax.text(x_start + width + 0.5, y_center_pos, f"E = {e_mpa:,} MPa",
                    ha='left', va='center', fontsize=10, color='#0066CC')
        
        y_current = y_bottom
    
    ax.annotate('', xy=(x_start + width + 3.5, total_display),
                xytext=(x_start + width + 3.5, 0),
                arrowprops=dict(arrowstyle='<->', color='red', lw=2))
    ax.text(x_start + width + 4, total_display / 2, f"Total\n{total_thickness} cm",
            ha='left', va='center', fontsize=12, color='red', fontweight='bold')
    
    margin = 10
    ax.set_xlim(0, 14)
    ax.set_ylim(-margin, total_display + margin)
    ax.axis('off')
    
    ax.set_title('Pavement Structure', fontsize=18, fontweight='bold', pad=20)
    ax.text(x_center, -margin + 4,
            f"Total Pavement Thickness: {total_thickness} cm",
            ha='center', va='center', fontsize=13, fontweight='bold',
            bbox=dict(boxstyle='round', facecolor='lightyellow', alpha=0.9, edgecolor='orange'))
    
    plt.tight_layout()
    return fig

def save_figure_to_bytes(fig, dpi=150):
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    buf.seek(0)
    return buf

def save_figure_to_file(fig, filepath, dpi=150):
    fig.savefig(filepath, format='png', dpi=dpi, bbox_inches='tight',
                facecolor='white', edgecolor='none')

# ============================================================
# ส่วนที่ 6: สร้างรายงาน Word
# ============================================================

def create_word_report(
    project_name, pavement_type, inputs, calculated_values,
    odemark_results, comparison_results, selected_d,
    main_result, layers_data, detailed_calc=None, figure_path=None
):
    """สร้างรายงาน Word โดยใช้ python-docx"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        st.error("กรุณาติดตั้ง python-docx: pip install python-docx")
        return None
    
    doc = Document()
    
    # ตั้งค่าฟอนต์
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH Sarabun New'
    font.size = Pt(14)
    
    passed, ratio = main_result
    
    # ========== หน้าปก ==========
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_heading('รายการคำนวณออกแบบความหนาถนนคอนกรีต', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('RIGID PAVEMENT THICKNESS DESIGN')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    method = doc.add_paragraph('ตามวิธี AASHTO Guide for Design of Pavement Structures (1993)')
    method.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    ptype = doc.add_paragraph(f'ประเภท: {pavement_type}')
    ptype.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    pname = doc.add_heading(f'โครงการ: {project_name}', level=1)
    pname.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    date_p = doc.add_paragraph(f'วันที่คำนวณ: {datetime.now().strftime("%d %B %Y")}')
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    dept = doc.add_paragraph('ภาควิชาครุศาสตร์โยธา')
    dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    uni = doc.add_paragraph('มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ')
    uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ========== สารบัญ ==========
    doc.add_heading('สารบัญ', level=1)
    toc_items = [
        '1. ข้อมูลโครงการ',
        '2. โครงสร้างชั้นทาง',
        '3. การวิเคราะห์ความหนาเทียบเท่า (Odemark Method)',
        '4. พารามิเตอร์การออกแบบ',
        '5. สมการ AASHTO 1993',
        '6. รายละเอียดการคำนวณ',
        '7. ผลการเปรียบเทียบความหนาต่างๆ',
        '8. สรุปผลการออกแบบ',
        '9. เอกสารอ้างอิง'
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ========== 1. ข้อมูลโครงการ ==========
    doc.add_heading('1. ข้อมูลโครงการ (Project Information)', level=1)
    
    table_proj = doc.add_table(rows=4, cols=2)
    table_proj.style = 'Table Grid'
    
    proj_data = [
        ('ชื่อโครงการ (Project Name)', project_name),
        ('ประเภทถนนคอนกรีต (Pavement Type)', pavement_type),
        ('วันที่คำนวณ (Date)', datetime.now().strftime("%d/%m/%Y %H:%M")),
        ('ผู้คำนวณ (Calculated by)', 'AASHTO 1993 Design Tool')
    ]
    
    for i, (label, value) in enumerate(proj_data):
        table_proj.rows[i].cells[0].text = label
        table_proj.rows[i].cells[1].text = str(value)
    
    doc.add_paragraph()
    
    # ========== 2. โครงสร้างชั้นทาง ==========
    doc.add_heading('2. โครงสร้างชั้นทาง (Pavement Structure)', level=1)
    
    doc.add_paragraph('ตารางแสดงชั้นวัสดุและคุณสมบัติ:')
    
    table_layers = doc.add_table(rows=1, cols=4)
    table_layers.style = 'Table Grid'
    
    hdr = table_layers.rows[0].cells
    hdr[0].text = 'ลำดับ (No.)'
    hdr[1].text = 'ชนิดวัสดุ (Material)'
    hdr[2].text = 'ความหนา (cm)'
    hdr[3].text = 'Modulus E (MPa)'
    
    total_layer_thickness = 0
    for i, layer in enumerate(layers_data):
        if layer.get("thickness_cm", 0) > 0:
            row = table_layers.add_row().cells
            row[0].text = str(i + 1)
            row[1].text = layer.get('name', '')
            row[2].text = f"{layer.get('thickness_cm', 0)}"
            row[3].text = f"{layer.get('E_MPa', 0):,}"
            total_layer_thickness += layer.get('thickness_cm', 0)
    
    # แถวรวม
    total_row = table_layers.add_row().cells
    total_row[0].text = ''
    total_row[1].text = 'รวมความหนาชั้นรองรับ (Total)'
    total_row[2].text = f'{total_layer_thickness}'
    total_row[3].text = '-'
    
    doc.add_paragraph()
    
    # เพิ่มรูปถ้ามี
    if figure_path and os.path.exists(figure_path):
        doc.add_paragraph('รูปตัดโครงสร้างชั้นทาง (Pavement Structure Cross Section):')
        doc.add_picture(figure_path, width=Inches(5.5))
        cap = doc.add_paragraph('รูปที่ 1: รูปตัดโครงสร้างชั้นทาง')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # ========== 3. Odemark Method ==========
    doc.add_heading('3. การวิเคราะห์ความหนาเทียบเท่า (Odemark Method)', level=1)
    
    doc.add_paragraph('วิธี Odemark (Method of Equivalent Thickness - MET) ใช้แปลงระบบหลายชั้นให้เป็นความหนาเทียบเท่า')
    
    doc.add_heading('3.1 ทฤษฎีและสูตรการคำนวณ', level=2)
    
    doc.add_paragraph('สูตร Odemark:')
    formula = doc.add_paragraph('he = h x (E_layer / E_subgrade)^(1/3)')
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('โดยที่:')
    doc.add_paragraph('   he = ความหนาเทียบเท่า (Equivalent Thickness, cm)')
    doc.add_paragraph('   h = ความหนาจริงของชั้นวัสดุ (cm)')
    doc.add_paragraph('   E_layer = Modulus ของชั้นวัสดุ (MPa)')
    doc.add_paragraph('   E_subgrade = Modulus ของดินฐานราก (MPa)')
    
    doc.add_paragraph(f'ค่า Modulus ดินฐานราก (E_subgrade) = {odemark_results.get("e_subgrade", 50):,} MPa')
    
    doc.add_heading('3.2 ผลการคำนวณความหนาเทียบเท่า', level=2)
    
    if odemark_results and odemark_results.get("layers"):
        table_od = doc.add_table(rows=1, cols=6)
        table_od.style = 'Table Grid'
        
        hdr_od = table_od.rows[0].cells
        hdr_od[0].text = 'ลำดับ'
        hdr_od[1].text = 'วัสดุ'
        hdr_od[2].text = 'h (cm)'
        hdr_od[3].text = 'E (MPa)'
        hdr_od[4].text = 'Factor (E/Es)^1/3'
        hdr_od[5].text = 'he (cm)'
        
        for i, layer in enumerate(odemark_results["layers"]):
            row = table_od.add_row().cells
            row[0].text = str(i + 1)
            row[1].text = layer['name']
            row[2].text = f"{layer['h_actual']:.1f}"
            row[3].text = f"{layer['E_MPa']:,}"
            row[4].text = f"{layer['factor']:.4f}"
            row[5].text = f"{layer['he_equivalent']:.2f}"
        
        # แถวรวม
        sum_row = table_od.add_row().cells
        sum_row[0].text = ''
        sum_row[1].text = 'รวม (Total)'
        sum_row[2].text = f"{odemark_results.get('total_actual_thickness', 0):.1f}"
        sum_row[3].text = '-'
        sum_row[4].text = '-'
        sum_row[5].text = f"{odemark_results.get('total_equivalent_thickness', 0):.2f}"
    
    doc.add_paragraph()
    
    doc.add_heading('3.3 ขั้นตอนการคำนวณโดยละเอียด', level=2)
    
    if odemark_results.get("calculation_steps"):
        for step in odemark_results["calculation_steps"]:
            doc.add_paragraph(step)
    
    doc.add_paragraph()
    
    summary_od = doc.add_paragraph()
    summary_od.add_run('สรุป: ').bold = True
    summary_od.add_run(f'ความหนารวมจริง = {odemark_results.get("total_actual_thickness", 0):.1f} cm, ')
    summary_od.add_run(f'ความหนาเทียบเท่า (he) = {odemark_results.get("total_equivalent_thickness", 0):.2f} cm')
    
    doc.add_page_break()
    
    # ========== 4. พารามิเตอร์การออกแบบ ==========
    doc.add_heading('4. พารามิเตอร์การออกแบบ (Design Parameters)', level=1)
    
    doc.add_heading('4.1 ปริมาณจราจร (Traffic Loading)', level=2)
    doc.add_paragraph(f'ESAL ออกแบบ (W18) = {inputs["w18_design"]:,.0f} ESALs')
    
    doc.add_heading('4.2 Serviceability', level=2)
    doc.add_paragraph(f'Initial Serviceability (Po) = 4.5')
    doc.add_paragraph(f'Terminal Serviceability (Pt) = {inputs["pt"]:.1f}')
    doc.add_paragraph(f'Serviceability Loss (ΔPSI) = {calculated_values["delta_psi"]:.1f}')
    
    doc.add_heading('4.3 Reliability', level=2)
    doc.add_paragraph(f'Reliability (R) = {inputs["reliability"]:.0f}%')
    doc.add_paragraph(f'Standard Normal Deviate (ZR) = {calculated_values["zr"]:.3f}')
    doc.add_paragraph(f'Overall Standard Deviation (So) = {inputs["so"]:.2f}')
    
    doc.add_heading('4.4 Subgrade Support', level=2)
    doc.add_paragraph(f'Effective Modulus of Subgrade Reaction (k_eff) = {inputs["k_eff"]:,} pci')
    doc.add_paragraph(f'Loss of Support (LS) = {inputs.get("ls", 1.0):.1f}')
    
    doc.add_heading('4.5 Concrete Properties', level=2)
    doc.add_paragraph(f'กำลังอัดคอนกรีต (Cube) = {inputs["fc_cube"]:.0f} ksc')
    doc.add_paragraph(f'กำลังอัดคอนกรีต (Cylinder) = {calculated_values["fc_cylinder"]:.0f} ksc')
    doc.add_paragraph(f'Modulus of Elasticity (Ec) = {calculated_values["ec"]:,.0f} psi')
    doc.add_paragraph(f'Modulus of Rupture (Sc) = {inputs["sc"]:.0f} psi')
    
    doc.add_heading('4.6 Load Transfer and Drainage', level=2)
    doc.add_paragraph(f'Load Transfer Coefficient (J) = {inputs["j"]:.1f}')
    doc.add_paragraph(f'Drainage Coefficient (Cd) = {inputs["cd"]:.2f}')
    
    doc.add_page_break()
    
    # ========== 5. สมการ AASHTO 1993 ==========
    doc.add_heading('5. สมการออกแบบ AASHTO 1993', level=1)
    
    doc.add_paragraph('สมการ AASHTO 1993 สำหรับการออกแบบความหนาผิวทางคอนกรีต:')
    
    eq1 = doc.add_paragraph()
    eq1.add_run('log10(W18) = ZR x So + 7.35 x log10(D+1) - 0.06').bold = True
    
    eq2 = doc.add_paragraph()
    eq2.add_run('              + log10(ΔPSI/3.0) / [1 + 1.624x10^7/(D+1)^8.46]').bold = True
    
    eq3 = doc.add_paragraph()
    eq3.add_run('              + (4.22-0.32xPt) x log10[(ScxCdx(D^0.75-1.132))/(215.63xJx(D^0.75-18.42/(Ec/k)^0.25))]').bold = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('โดยที่:')
    vars_desc = [
        ('W18', 'จำนวน 18-kip Equivalent Single Axle Loads'),
        ('ZR', 'Standard Normal Deviate'),
        ('So', 'Overall Standard Deviation'),
        ('D', 'ความหนาแผ่นคอนกรีต (นิ้ว)'),
        ('ΔPSI', 'การสูญเสีย Serviceability (Po - Pt)'),
        ('Pt', 'Terminal Serviceability'),
        ('Sc', 'Modulus of Rupture ของคอนกรีต (psi)'),
        ('Cd', 'Drainage Coefficient'),
        ('J', 'Load Transfer Coefficient'),
        ('Ec', 'Modulus of Elasticity ของคอนกรีต (psi)'),
        ('k', 'Effective Modulus of Subgrade Reaction (pci)'),
    ]
    for sym, desc in vars_desc:
        doc.add_paragraph(f'   {sym} = {desc}')
    
    doc.add_page_break()
    
    # ========== 6. รายละเอียดการคำนวณ ==========
    doc.add_heading('6. รายละเอียดการคำนวณ (Detailed Calculation)', level=1)
    
    doc.add_paragraph(f'สำหรับความหนา D = {selected_d} นิ้ว ({selected_d * 2.54:.1f} ซม.)')
    
    if detailed_calc:
        doc.add_heading('6.1 Term 1: ZR x So', level=2)
        doc.add_paragraph(f'Term 1 = ZR x So')
        doc.add_paragraph(f'       = ({calculated_values["zr"]:.3f}) x ({inputs["so"]:.2f})')
        doc.add_paragraph(f'       = {detailed_calc["term1"]["value"]:.4f}')
        
        doc.add_heading('6.2 Term 2: 7.35 x log10(D+1) - 0.06', level=2)
        doc.add_paragraph(f'Term 2 = 7.35 x log10({selected_d}+1) - 0.06')
        doc.add_paragraph(f'       = 7.35 x {detailed_calc["term2"]["log_d1"]:.4f} - 0.06')
        doc.add_paragraph(f'       = {detailed_calc["term2"]["value"]:.4f}')
        
        doc.add_heading('6.3 Term 3: Serviceability Loss Term', level=2)
        doc.add_paragraph(f'Term 3 = log10(ΔPSI/3.0) / [1 + 1.624x10^7/(D+1)^8.46]')
        doc.add_paragraph(f'ตัวเศษ = log10({calculated_values["delta_psi"]:.1f}/3.0) = {detailed_calc["term3"]["numerator"]:.4f}')
        doc.add_paragraph(f'ตัวส่วน = 1 + 1.624x10^7/({selected_d}+1)^8.46 = {detailed_calc["term3"]["denominator"]:.4f}')
        doc.add_paragraph(f'Term 3 = {detailed_calc["term3"]["value"]:.4f}')
        
        doc.add_heading('6.4 Term 4: Strength and Support Term', level=2)
        doc.add_paragraph(f'Term 4 = (4.22 - 0.32xPt) x log10[(ScxCdx(D^0.75-1.132))/(215.63xJx(D^0.75-18.42/(Ec/k)^0.25))]')
        doc.add_paragraph(f'D^0.75 = {selected_d}^0.75 = {detailed_calc["term4"]["d_power"]:.4f}')
        doc.add_paragraph(f'Ec/k = {calculated_values["ec"]:,.0f}/{inputs["k_eff"]} = {detailed_calc["term4"]["ec_k_ratio"]:,.2f}')
        doc.add_paragraph(f'ตัวเศษ = {inputs["sc"]} x {inputs["cd"]} x ({detailed_calc["term4"]["d_power"]:.4f} - 1.132) = {detailed_calc["term4"]["numerator"]:.2f}')
        doc.add_paragraph(f'ตัวส่วน = 215.63 x {inputs["j"]} x ({detailed_calc["term4"]["d_power"]:.4f} - 18.42/({detailed_calc["term4"]["ec_k_ratio"]:,.2f})^0.25) = {detailed_calc["term4"]["denominator"]:.2f}')
        doc.add_paragraph(f'สัมประสิทธิ์ = (4.22 - 0.32 x {inputs["pt"]}) = {detailed_calc["term4"]["coefficient"]:.3f}')
        doc.add_paragraph(f'Term 4 = {detailed_calc["term4"]["value"]:.4f}')
        
        doc.add_heading('6.5 ผลรวม', level=2)
        doc.add_paragraph(f'log10(W18) = Term1 + Term2 + Term3 + Term4')
        doc.add_paragraph(f'           = {detailed_calc["term1"]["value"]:.4f} + {detailed_calc["term2"]["value"]:.4f} + {detailed_calc["term3"]["value"]:.4f} + {detailed_calc["term4"]["value"]:.4f}')
        doc.add_paragraph(f'           = {detailed_calc["result"]["log10_w18"]:.4f}')
        doc.add_paragraph(f'W18 = 10^{detailed_calc["result"]["log10_w18"]:.4f} = {detailed_calc["result"]["w18"]:,.0f} ESALs')
    
    doc.add_page_break()
    
    # ========== 7. ผลการเปรียบเทียบ ==========
    doc.add_heading('7. ผลการเปรียบเทียบความหนาต่างๆ', level=1)
    
    table_comp = doc.add_table(rows=1, cols=6)
    table_comp.style = 'Table Grid'
    
    hdr_comp = table_comp.rows[0].cells
    hdr_comp[0].text = 'D (นิ้ว)'
    hdr_comp[1].text = 'D (ซม.)'
    hdr_comp[2].text = 'log10(W18)'
    hdr_comp[3].text = 'W18 รองรับได้'
    hdr_comp[4].text = 'อัตราส่วน'
    hdr_comp[5].text = 'ผล'
    
    for result in comparison_results:
        row = table_comp.add_row().cells
        row[0].text = f"{result['d']}"
        row[1].text = f"{result['d'] * 2.54:.1f}"
        row[2].text = f"{result['log_w18']:.4f}"
        row[3].text = f"{result['w18']:,.0f}"
        row[4].text = f"{result['ratio']:.2f}"
        row[5].text = "ผ่าน" if result['passed'] else "ไม่ผ่าน"
    
    doc.add_paragraph()
    doc.add_paragraph(f'หมายเหตุ: ESAL ที่ต้องการ = {inputs["w18_design"]:,.0f} ESALs')
    doc.add_paragraph('อัตราส่วน = W18 รองรับได้ / W18 ที่ต้องการ (ต้อง >= 1.00 จึงจะผ่าน)')
    
    doc.add_page_break()
    
    # ========== 8. สรุปผล ==========
    doc.add_heading('8. สรุปผลการออกแบบ (Design Summary)', level=1)
    
    status = "ผ่านเกณฑ์การออกแบบ (PASSED)" if passed else "ไม่ผ่านเกณฑ์การออกแบบ (NOT PASSED)"
    
    summary_p = doc.add_paragraph()
    summary_p.add_run('ผลการตรวจสอบ: ').bold = True
    summary_p.add_run(status)
    
    doc.add_paragraph()
    
    table_sum = doc.add_table(rows=5, cols=2)
    table_sum.style = 'Table Grid'
    
    w18_cap = [r for r in comparison_results if r['d'] == selected_d]
    w18_capacity = w18_cap[0]["w18"] if w18_cap else 0
    
    sum_data = [
        ('ความหนาที่เลือก', f'{selected_d} นิ้ว ({selected_d * 2.54:.1f} ซม.)'),
        ('ESAL ที่ต้องการ', f'{inputs["w18_design"]:,.0f} ESALs'),
        ('ESAL ที่รองรับได้', f'{w18_capacity:,.0f} ESALs'),
        ('อัตราส่วน (Capacity/Required)', f'{ratio:.2f}'),
        ('ผลการตรวจสอบ', status),
    ]
    
    for i, (label, value) in enumerate(sum_data):
        table_sum.rows[i].cells[0].text = label
        table_sum.rows[i].cells[1].text = str(value)
    
    doc.add_paragraph()
    
    if passed:
        doc.add_paragraph('ข้อสรุป: ความหนาคอนกรีตที่เลือกสามารถรองรับปริมาณจราจรตามที่ออกแบบได้')
    else:
        doc.add_paragraph('ข้อสรุป: ความหนาคอนกรีตที่เลือกไม่เพียงพอ กรุณาเพิ่มความหนาหรือปรับปรุงคุณสมบัติวัสดุ')
    
    doc.add_page_break()
    
    # ========== 9. เอกสารอ้างอิง ==========
    doc.add_heading('9. เอกสารอ้างอิง (References)', level=1)
    
    refs = [
        'AASHTO (1993). AASHTO Guide for Design of Pavement Structures. American Association of State Highway and Transportation Officials, Washington, D.C.',
        'Huang, Y.H. (2004). Pavement Analysis and Design. 2nd Edition, Pearson Prentice Hall, New Jersey.',
        'ACI 318-19 (2019). Building Code Requirements for Structural Concrete and Commentary. American Concrete Institute.',
        'Odemark, N. (1949). Investigations as to the Elastic Properties of Soils and Design of Pavements According to the Theory of Elasticity. Statens Vaginstitut, Stockholm.',
        'กรมทางหลวง (2013). คู่มือการออกแบบผิวทางแอสฟัลต์คอนกรีตและผิวทางคอนกรีต. กระทรวงคมนาคม.'
    ]
    
    for i, ref in enumerate(refs):
        doc.add_paragraph(f'{i+1}. {ref}')
    
    # ========== Footer ==========
    doc.add_paragraph()
    doc.add_paragraph()
    
    footer1 = doc.add_paragraph('-' * 60)
    footer1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer2 = doc.add_paragraph('Generated by AASHTO 1993 Rigid Pavement Design Tool')
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer3 = doc.add_paragraph('ภาควิชาครุศาสตร์โยธา มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ')
    footer3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # บันทึก
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


# ============================================================
# ส่วนที่ 7: Streamlit UI
# ============================================================

def main():
    st.set_page_config(
        page_title="AASHTO 1993 Rigid Pavement Design",
        page_icon="🛣️",
        layout="wide"
    )
    
    st.title("🛣️ การออกแบบความหนาถนนคอนกรีต")
    st.subheader("ตามวิธี AASHTO 1993 (Rigid Pavement Design)")
    
    st.markdown("---")
    
    # ข้อมูลโครงการ
    st.header("📋 ข้อมูลโครงการ")
    project_name = st.text_input(
        "ชื่อโครงการ",
        value="โครงการก่อสร้างถนนคอนกรีต",
        help="ระบุชื่อโครงการสำหรับใส่ในรายงาน"
    )
    
    st.markdown("---")
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("📥 ข้อมูลนำเข้า (Input)")
        
        pavement_type = st.selectbox(
            "ประเภทผิวทางคอนกรีต",
            options=list(J_VALUES.keys()),
            index=1,
            help="JRCP = Jointed Reinforced, JPCP = Jointed Plain, CRCP = Continuously Reinforced"
        )
        
        st.markdown("---")
        
        # ชั้นโครงสร้างทาง
        st.subheader("🔶 ชั้นโครงสร้างทาง (Pavement Layers)")
        
        material_options = list(MATERIAL_MODULUS.keys())
        
        num_layers = st.slider(
            "จำนวนชั้นวัสดุใต้แผ่นคอนกรีต",
            min_value=1, max_value=6, value=3,
            help="เลือกจำนวนชั้นวัสดุ 1-6 ชั้น"
        )
        
        default_layers = [
            {"name": "พื้นทางซีเมนต์ CTB", "thickness_cm": 15},
            {"name": "หินคลุก CBR 80%", "thickness_cm": 15},
            {"name": "รองพื้นทางวัสดุมวลรวม CBR 25%", "thickness_cm": 15},
            {"name": "วัสดุคัดเลือก ก", "thickness_cm": 30},
            {"name": "ดินถมคันทาง / ดินเดิม", "thickness_cm": 0},
            {"name": "กำหนดเอง...", "thickness_cm": 0},
        ]
        
        layers_data = []
        
        with st.expander("📊 ตารางค่า Modulus อ้างอิง", expanded=False):
            st.markdown("""
            | วัสดุชั้นทาง | MR (MPa) |
            |-------------|----------|
            | พื้นทางซีเมนต์ CTB | 1,200 |
            | หินคลุก CBR 80% | 350 |
            | รองพื้นทางวัสดุมวลรวม CBR 25% | 150 |
            | วัสดุคัดเลือก ก | 76 |
            """)
        
        for i in range(num_layers):
            st.markdown(f"**ชั้นที่ {i+1}**")
            col_a, col_b, col_c = st.columns([2, 1, 1])
            
            default_name = default_layers[i]["name"] if i < len(default_layers) else "กำหนดเอง..."
            default_index = material_options.index(default_name) if default_name in material_options else len(material_options) - 1
            
            with col_a:
                layer_name = st.selectbox(
                    "เลือกวัสดุ", options=material_options,
                    index=default_index, key=f"layer_name_{i}"
                )
            
            with col_b:
                layer_thickness = st.number_input(
                    "ความหนา (ซม.)", min_value=0, max_value=100,
                    value=default_layers[i]["thickness_cm"] if i < len(default_layers) else 20,
                    key=f"layer_thick_{i}"
                )
            
            recommended_modulus = MATERIAL_MODULUS.get(layer_name, 100)
            
            with col_c:
                layer_modulus = st.number_input(
                    "E (MPa)", min_value=10, max_value=10000,
                    value=recommended_modulus, key=f"layer_E_{i}_{layer_name}",
                    help=f"ค่าแนะนำ: {recommended_modulus:,} MPa"
                )
            
            layers_data.append({
                "name": layer_name,
                "thickness_cm": layer_thickness,
                "E_MPa": layer_modulus
            })
        
        st.markdown("---")
        
        # Odemark Method
        st.subheader("🔷 Odemark Method")
        
        with st.expander("📖 ทฤษฎี Odemark Method", expanded=False):
            st.markdown(r"""
            **วิธี Odemark (Method of Equivalent Thickness)**
            
            แปลงระบบหลายชั้นให้เป็นความหนาเทียบเท่า:
            
            $$h_e = h \times \left(\frac{E_{layer}}{E_{subgrade}}\right)^{1/3}$$
            """)
        
        e_subgrade = st.number_input(
            "Modulus ดินฐานราก (E_subgrade) - MPa",
            min_value=10, max_value=500, value=50, step=10,
            help="ค่า Modulus ของดินเดิม/ดินถมคันทาง"
        )
        
        odemark_results = calculate_odemark_equivalent_thickness(layers_data, e_subgrade)
        
        st.info(f"""
        **ผลการคำนวณ Odemark:**
        - ความหนารวมจริง: **{odemark_results['total_actual_thickness']:.1f} cm**
        - ความหนาเทียบเท่า: **{odemark_results['total_equivalent_thickness']:.2f} cm**
        """)
        
        st.markdown("---")
        
        # 1. ESAL
        st.subheader("1️⃣ ปริมาณจราจร")
        
        w18_design = st.number_input(
            "ESAL ที่ต้องการรองรับ (W18)",
            min_value=10_000, max_value=500_000_000, value=500_000, step=100_000,
            format="%d", help="จำนวน ESAL ตลอดอายุการใช้งาน"
        )
        
        if w18_design >= 1_000_000:
            esal_text = f"{w18_design / 1_000_000:,.2f} ล้าน"
        else:
            esal_text = f"{w18_design:,.0f}"
        
        st.markdown(f"<h3 style='color: #1E88E5;'>{esal_text} ESALs</h3>", unsafe_allow_html=True)
        
        st.markdown("---")
        
        # 2. Serviceability
        st.subheader("2️⃣ Serviceability")
        pt = st.slider("Terminal Serviceability (Pt)", 1.5, 3.0, 2.0, 0.1)
        delta_psi = 4.5 - pt
        st.info(f"ΔPSI = 4.5 - {pt:.1f} = **{delta_psi:.1f}**")
        
        st.markdown("---")
        
        # 3. Reliability
        st.subheader("3️⃣ Reliability")
        reliability = st.select_slider("Reliability (R)", options=[80, 85, 90, 95], value=90)
        zr = get_zr_value(reliability)
        st.info(f"ZR = **{zr:.3f}**")
        
        so = st.number_input("Overall Standard Deviation (So)", 0.30, 0.45, 0.35, 0.01, format="%.2f")
        
        st.markdown("---")
        
        # 4. Subgrade
        st.subheader("4️⃣ Subgrade Support")
        k_eff = st.number_input("Effective k-value (pci)", 50, 1000, 200, 25, format="%d")
        ls_value = st.number_input("Loss of Support (LS)", 0.0, 3.0, 1.0, 0.5, format="%.1f")
        
        st.markdown("---")
        
        # 5. Concrete
        st.subheader("5️⃣ Concrete Properties")
        fc_cube = st.number_input("กำลังอัด (Cube) - ksc", 200, 600, 350, 10, format="%d")
        fc_cylinder = convert_cube_to_cylinder(fc_cube)
        ec = calculate_concrete_modulus(fc_cylinder)
        
        st.info(f"f'c (Cyl) = **{fc_cylinder:.0f} ksc** | Ec = **{ec:,.0f} psi**")
        
        sc_auto = estimate_modulus_of_rupture(fc_cylinder)
        sc = st.number_input("Modulus of Rupture (Sc) - psi", 400, 1000, int(sc_auto), 10, format="%d")
        
        st.markdown("---")
        
        # 6. Load Transfer
        st.subheader("6️⃣ Load Transfer & Drainage")
        j_auto = J_VALUES[pavement_type]
        j_value = st.number_input("Load Transfer (J)", 2.0, 4.5, j_auto, 0.1, format="%.1f")
        cd = st.number_input("Drainage (Cd)", 0.7, 1.3, 1.0, 0.05, format="%.2f")
        
        st.markdown("---")
        
        # 7. ความหนาคอนกรีต
        st.subheader("7️⃣ ความหนาคอนกรีต")
        d_selected = st.slider("ความหนา D (นิ้ว)", 8, 16, 12, 1)
        st.info(f"D = {d_selected} นิ้ว = **{d_selected * 2.54:.1f} ซม.**")
        
        st.markdown("---")
        
        # รูปโครงสร้าง
        st.subheader("📐 รูปตัดโครงสร้างชั้นทาง")
        concrete_cm = d_selected * 2.54
        fig_structure = create_pavement_structure_figure(layers_data, concrete_thickness_cm=concrete_cm)
        
        if fig_structure:
            st.pyplot(fig_structure)
            img_buffer = save_figure_to_bytes(fig_structure)
            st.download_button(
                "📥 ดาวน์โหลดรูป", data=img_buffer,
                file_name=f"pavement_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
                mime="image/png"
            )
            plt.close(fig_structure)
    
    # ============================================================
    # ส่วนแสดงผล
    # ============================================================
    
    with col2:
        st.header("📊 ผลการคำนวณ (Output)")
        
        comparison_results = []
        thicknesses = [8, 9, 10, 11, 12, 13, 14, 15, 16]
        
        st.subheader("📋 ตารางเปรียบเทียบ")
        
        table_data = []
        for d in thicknesses:
            log_w18, w18_capacity = calculate_aashto_rigid_w18(
                d, delta_psi, pt, zr, so, sc, cd, j_value, ec, k_eff
            )
            passed, ratio = check_design(w18_design, w18_capacity)
            
            comparison_results.append({
                'd': d, 'log_w18': log_w18, 'w18': w18_capacity,
                'passed': passed, 'ratio': ratio
            })
            
            table_data.append({
                'D (in)': d, 'D (cm)': f"{d * 2.54:.1f}",
                'log10(W18)': f"{log_w18:.4f}",
                'W18': f"{w18_capacity:,.0f}",
                'Ratio': f"{ratio:.2f}",
                'Status': "✅" if passed else "❌"
            })
        
        import pandas as pd
        st.dataframe(pd.DataFrame(table_data), use_container_width=True, hide_index=True)
        
        st.markdown("---")
        
        # Odemark Results
        st.subheader("🔷 ผล Odemark Method")
        
        if odemark_results and odemark_results.get("layers"):
            od_table = []
            for i, layer in enumerate(odemark_results["layers"]):
                od_table.append({
                    'No': i + 1, 'Material': layer['name'][:20],
                    'h (cm)': f"{layer['h_actual']:.1f}",
                    'E (MPa)': f"{layer['E_MPa']:,}",
                    'Factor': f"{layer['factor']:.3f}",
                    'he (cm)': f"{layer['he_equivalent']:.2f}"
                })
            
            st.dataframe(pd.DataFrame(od_table), use_container_width=True, hide_index=True)
            
            col_od1, col_od2 = st.columns(2)
            with col_od1:
                st.metric("ความหนาจริง", f"{odemark_results['total_actual_thickness']:.1f} cm")
            with col_od2:
                st.metric("ความหนาเทียบเท่า", f"{odemark_results['total_equivalent_thickness']:.2f} cm")
        
        st.markdown("---")
        
        # ผลการตรวจสอบ
        st.subheader(f"🎯 ผลการตรวจสอบ D = {d_selected} นิ้ว")
        
        log_w18_selected, w18_selected = calculate_aashto_rigid_w18(
            d_selected, delta_psi, pt, zr, so, sc, cd, j_value, ec, k_eff
        )
        passed_selected, ratio_selected = check_design(w18_design, w18_selected)
        
        detailed_calc = calculate_aashto_detailed(
            d_selected, delta_psi, pt, zr, so, sc, cd, j_value, ec, k_eff
        )
        
        col_a, col_b = st.columns(2)
        with col_a:
            st.metric("log10(W18)", f"{log_w18_selected:.4f}")
            st.metric("W18 รองรับได้", f"{w18_selected:,.0f}", delta=f"{w18_selected - w18_design:+,.0f}")
        with col_b:
            st.metric("W18 ที่ต้องการ", f"{w18_design:,.0f}")
            st.metric("อัตราส่วน", f"{ratio_selected:.2f}")
        
        if passed_selected:
            st.success(f"""
            ✅ **ผ่านเกณฑ์**
            
            D = {d_selected} นิ้ว ({d_selected * 2.54:.1f} ซม.) 
            รองรับ ESAL ได้ {w18_selected:,.0f}
            อัตราส่วน = {ratio_selected:.2f}
            """)
        else:
            st.error(f"""
            ❌ **ไม่ผ่านเกณฑ์**
            
            D = {d_selected} นิ้ว ({d_selected * 2.54:.1f} ซม.) 
            รองรับ ESAL ได้เพียง {w18_selected:,.0f}
            อัตราส่วน = {ratio_selected:.2f}
            
            **กรุณาเพิ่มความหนา**
            """)
        
        st.markdown("---")
        
        # รายละเอียดการคำนวณ
        with st.expander("📝 รายละเอียดการคำนวณ", expanded=False):
            st.markdown("**สมการ AASHTO 1993:**")
            st.latex(r'\log_{10}(W_{18}) = T_1 + T_2 + T_3 + T_4')
            
            st.markdown("---")
            st.markdown(f"**Term 1:** ZR x So = ({zr:.3f}) x ({so:.2f}) = **{detailed_calc['term1']['value']:.4f}**")
            st.markdown(f"**Term 2:** 7.35 x log10(D+1) - 0.06 = **{detailed_calc['term2']['value']:.4f}**")
            st.markdown(f"**Term 3:** Serviceability = **{detailed_calc['term3']['value']:.4f}**")
            st.markdown(f"**Term 4:** Strength = **{detailed_calc['term4']['value']:.4f}**")
            st.markdown("---")
            st.markdown(f"**ผลรวม:** log10(W18) = **{log_w18_selected:.4f}**")
            st.markdown(f"**W18** = 10^{log_w18_selected:.4f} = **{w18_selected:,.0f}** ESALs")
        
        st.markdown("---")
        
        # สมการ
        st.subheader("📝 สมการ AASHTO 1993")
        st.latex(r'\log_{10}(W_{18}) = Z_R \times S_o + 7.35 \times \log_{10}(D+1) - 0.06')
        st.latex(r'+ \frac{\log_{10}(\Delta PSI / 3.0)}{1 + \frac{1.624 \times 10^7}{(D+1)^{8.46}}}')
        st.latex(r'+ (4.22-0.32P_t) \times \log_{10}\left[\frac{S_c C_d (D^{0.75}-1.132)}{215.63 J (D^{0.75}-\frac{18.42}{(E_c/k)^{0.25}})}\right]')
        
        st.markdown("---")
        
        # สร้างรายงาน
        st.subheader("📄 ส่งออกรายงาน Word")
        
        inputs_dict = {
            'w18_design': w18_design, 'pt': pt, 'reliability': reliability,
            'so': so, 'k_eff': k_eff, 'ls': ls_value, 'fc_cube': fc_cube,
            'sc': sc, 'j': j_value, 'cd': cd
        }
        
        calculated_dict = {
            'fc_cylinder': fc_cylinder, 'ec': ec, 'zr': zr, 'delta_psi': delta_psi
        }
        
        if st.button("📥 สร้างรายงาน Word", type="primary"):
            with st.spinner("กำลังสร้างรายงาน..."):
                try:
                    # บันทึกรูปชั่วคราว
                    fig_for_report = create_pavement_structure_figure(layers_data, concrete_thickness_cm=concrete_cm)
                    
                    figure_path = None
                    if fig_for_report:
                        figure_path = tempfile.NamedTemporaryFile(suffix='.png', delete=False).name
                        save_figure_to_file(fig_for_report, figure_path, dpi=150)
                        plt.close(fig_for_report)
                    
                    buffer = create_word_report(
                        project_name=project_name,
                        pavement_type=pavement_type,
                        inputs=inputs_dict,
                        calculated_values=calculated_dict,
                        odemark_results=odemark_results,
                        comparison_results=comparison_results,
                        selected_d=d_selected,
                        main_result=(passed_selected, ratio_selected),
                        layers_data=layers_data,
                        detailed_calc=detailed_calc,
                        figure_path=figure_path
                    )
                    
                    if figure_path and os.path.exists(figure_path):
                        os.unlink(figure_path)
                    
                    if buffer:
                        st.download_button(
                            "⬇️ ดาวน์โหลดรายงาน (.docx)",
                            data=buffer,
                            file_name=f"AASHTO_Rigid_{project_name}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        st.success("สร้างรายงานสำเร็จ!")
                except Exception as e:
                    st.error(f"เกิดข้อผิดพลาด: {str(e)}")
                    st.info("กรุณาติดตั้ง python-docx: `pip install python-docx`")
    
    # ============================================================
    # อ้างอิง
    # ============================================================
    
    st.markdown("---")
    st.header("📚 อ้างอิง")
    
    st.markdown("""
    **เอกสารอ้างอิง:**
    1. AASHTO (1993). *AASHTO Guide for Design of Pavement Structures*
    2. Huang, Y.H. (2004). *Pavement Analysis and Design*. Pearson
    3. ACI 318-19 (2019). *Building Code Requirements for Structural Concrete*
    4. Odemark, N. (1949). *Investigations as to the Elastic Properties of Soils*
    
    **หมายเหตุ:** โปรแกรมนี้พัฒนาเพื่อใช้ในการเรียนการสอน
    """)
    
    st.markdown("---")
    st.caption("พัฒนาโดย: รศ.ดร.อิทธิพล มีผล // ภาควิชาครุศาสตร์โยธา มจพ.")


if __name__ == "__main__":
    main()
