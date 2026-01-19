"""
โปรแกรมออกแบบและตรวจสอบความหนาถนนคอนกรีต (Rigid Pavement)
ตามวิธี AASHTO 1993
รองรับทั้ง JPCP (Jointed Plain Concrete Pavement) และ CRCP (Continuously Reinforced Concrete Pavement)

พัฒนาสำหรับใช้ในการเรียนการสอน
ภาควิชาครุศาสตร์โยธา มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ
"""

import streamlit as st
import math
from io import BytesIO
from datetime import datetime
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from matplotlib import rcParams

# ============================================================
# ส่วนที่ 1: ค่าคงที่และตารางอ้างอิง AASHTO 1993
# ============================================================

# ตารางค่า ZR (Standard Normal Deviate) ตามระดับความเชื่อมั่น
ZR_TABLE = {
    50: -0.000,
    60: -0.253,
    70: -0.524,
    75: -0.674,
    80: -0.841,
    85: -1.037,
    90: -1.282,
    91: -1.340,
    92: -1.405,
    93: -1.476,
    94: -1.555,
    95: -1.645,
    96: -1.751,
    97: -1.881,
    98: -2.054,
    99: -2.327
}

# ค่า Load Transfer Coefficient (J) ตามประเภทผิวทางคอนกรีต
# อ้างอิง: AASHTO 1993 Guide
J_VALUES = {
    "JRCP": 2.8,
    "JPCP": 2.8,
    "JRCP/JPCP": 2.8,
    "CRCP": 2.5
}

# ค่า Drainage Coefficient (Cd) มาตรฐาน
CD_DEFAULT = 1.0

# ============================================================
# ส่วนที่ 2: ฟังก์ชันการคำนวณ
# ============================================================

def convert_cube_to_cylinder(fc_cube_ksc: float) -> float:
    """
    แปลงกำลังอัดคอนกรีตจาก Cube เป็น Cylinder
    fc_cylinder ≈ 0.8 × fc_cube (โดยประมาณ)
    
    Parameters:
        fc_cube_ksc: กำลังอัดคอนกรีต Cube (ksc)
    
    Returns:
        กำลังอัดคอนกรีต Cylinder (ksc)
    """
    return 0.8 * fc_cube_ksc


def calculate_concrete_modulus(fc_cylinder_ksc: float) -> float:
    """
    คำนวณ Modulus of Elasticity ของคอนกรีต (Ec)
    ตามสูตร ACI: Ec = 57,000 × √(f'c) (psi)
    
    Parameters:
        fc_cylinder_ksc: กำลังอัดคอนกรีต Cylinder (ksc)
    
    Returns:
        Ec ในหน่วย psi
    """
    # แปลง ksc เป็น psi (1 ksc = 14.223 psi)
    fc_psi = fc_cylinder_ksc * 14.223
    
    # คำนวณ Ec ตาม ACI 318
    ec_psi = 57000 * math.sqrt(fc_psi)
    
    return ec_psi


def estimate_modulus_of_rupture(fc_cylinder_ksc: float) -> float:
    """
    ประมาณค่า Modulus of Rupture (Sc) จากกำลังอัดคอนกรีต
    ตามสูตร: Sc = (7.5 ถึง 12) × √(f'c) (ACI 318, หน่วย psi)
    
    Parameters:
        fc_cylinder_ksc: กำลังอัดคอนกรีต Cylinder (ksc)
    
    Returns:
        Sc ในหน่วย psi (ใช้ค่า 10 × √f'c)
    """
    # แปลง ksc เป็น psi
    fc_psi = fc_cylinder_ksc * 14.223
    
    # ใช้สูตร: Sc = 10 × √f'c (ค่าเหมาะสมสำหรับคอนกรีตถนน)
    sc_psi = 10.0 * math.sqrt(fc_psi)
    
    return sc_psi


def get_zr_value(reliability: float) -> float:
    """
    หาค่า ZR จากตาราง AASHTO ตามระดับความเชื่อมั่น
    
    Parameters:
        reliability: ระดับความเชื่อมั่น (%)
    
    Returns:
        ค่า ZR (Standard Normal Deviate)
    """
    return ZR_TABLE.get(int(reliability), -1.282)


def calculate_aashto_rigid_w18(
    d_inch: float,
    delta_psi: float,
    pt: float,
    zr: float,
    so: float,
    sc_psi: float,
    cd: float,
    j: float,
    ec_psi: float,
    k_pci: float
) -> tuple:
    """
    คำนวณ ESAL (W18) ที่รองรับได้ตามสมการ AASHTO 1993 สำหรับ Rigid Pavement
    
    สมการ AASHTO 1993:
    log10(W18) = ZR × So + 7.35 × log10(D+1) - 0.06 
                 + log10(ΔPSI/(4.5-1.5)) / (1 + 1.624×10^7 / (D+1)^8.46)
                 + (4.22 - 0.32×Pt) × log10[(Sc×Cd×(D^0.75-1.132)) / (215.63×J×(D^0.75 - 18.42/(Ec/k)^0.25))]
    
    Parameters:
        d_inch: ความหนาแผ่นพื้นคอนกรีต (นิ้ว)
        delta_psi: การสูญเสียค่า Serviceability (ΔPSI = 4.5 - Pt)
        pt: Terminal Serviceability
        zr: Standard Normal Deviate
        so: Overall Standard Deviation
        sc_psi: Modulus of Rupture (psi)
        cd: Drainage Coefficient
        j: Load Transfer Coefficient
        ec_psi: Modulus of Elasticity ของคอนกรีต (psi)
        k_pci: Effective Modulus of Subgrade Reaction (pci)
    
    Returns:
        tuple: (log10_w18, w18)
    """
    
    # พจน์ที่ 1: ZR × So
    term1 = zr * so
    
    # พจน์ที่ 2: 7.35 × log10(D+1) - 0.06
    term2 = 7.35 * math.log10(d_inch + 1) - 0.06
    
    # พจน์ที่ 3: การสูญเสีย Serviceability
    # log10(ΔPSI/(4.5-1.5)) / (1 + 1.624×10^7 / (D+1)^8.46)
    numerator3 = math.log10(delta_psi / (4.5 - 1.5))
    denominator3 = 1 + (1.624e7 / ((d_inch + 1) ** 8.46))
    term3 = numerator3 / denominator3
    
    # พจน์ที่ 4: กำลังของคอนกรีตและฐานราก
    # (4.22 - 0.32×Pt) × log10[(Sc×Cd×(D^0.75-1.132)) / (215.63×J×(D^0.75 - 18.42/(Ec/k)^0.25))]
    
    # คำนวณ D^0.75
    d_power = d_inch ** 0.75
    
    # คำนวณตัวเศษ: Sc × Cd × (D^0.75 - 1.132)
    numerator4 = sc_psi * cd * (d_power - 1.132)
    
    # คำนวณตัวส่วน: 215.63 × J × (D^0.75 - 18.42/(Ec/k)^0.25)
    ec_k_ratio = ec_psi / k_pci
    denominator4 = 215.63 * j * (d_power - 18.42 / (ec_k_ratio ** 0.25))
    
    # ตรวจสอบว่าค่าต้องเป็นบวก
    if numerator4 <= 0 or denominator4 <= 0:
        return (float('-inf'), 0)
    
    inner_term = numerator4 / denominator4
    
    if inner_term <= 0:
        return (float('-inf'), 0)
    
    term4 = (4.22 - 0.32 * pt) * math.log10(inner_term)
    
    # รวมทุกพจน์
    log10_w18 = term1 + term2 + term3 + term4
    
    # คำนวณ W18
    w18 = 10 ** log10_w18
    
    return (log10_w18, w18)


def check_design(w18_required: float, w18_capacity: float) -> tuple:
    """
    ตรวจสอบว่าความหนาที่กำหนดรองรับ ESAL ได้หรือไม่
    
    Parameters:
        w18_required: ESAL ที่ต้องการรองรับ
        w18_capacity: ESAL ที่รองรับได้
    
    Returns:
        tuple: (ผลการตรวจสอบ (bool), อัตราส่วน)
    """
    ratio = w18_capacity / w18_required if w18_required > 0 else float('inf')
    passed = w18_capacity >= w18_required
    return (passed, ratio)


def create_pavement_structure_figure(layers_data: list, concrete_thickness_cm: float = None):
    """
    สร้างรูปโครงสร้างชั้นทาง
    
    Parameters:
        layers_data: รายการข้อมูลชั้นวัสดุ [{"name": ..., "thickness_cm": ..., "E_MPa": ...}, ...]
        concrete_thickness_cm: ความหนาแผ่นคอนกรีต (ซม.) ถ้ามี
    
    Returns:
        matplotlib figure
    """
    # แปลงชื่อวัสดุเป็นภาษาอังกฤษสำหรับแสดงในรูป
    THAI_TO_ENG = {
        "รองผิวทางคอนกรีตด้วย AC": "AC Interlayer",
        "รองผิวทางคอนกรีตด้วย PMA(AC)": "PMA Interlayer",
        "พื้นทางซีเมนต์ CTB": "Cement Treated Base",
        "หินคลุกผสมซีเมนต์ UCS 24.5 ksc": "Soil Cement",
        "หินคลุก CBR 80%": "Crushed Rock Base",
        "ดินซีเมนต์ UCS 17.5 ksc": "Soil Cement",
        "วัสดุหมุนเวียน (Recycling)": "Recycled Material",
        "รองพื้นทางวัสดุมวลรวม CBR 25%": "Aggregate Subbase",
        "วัสดุคัดเลือก ก": "Selected Material",
        "ดินถมคันทาง / ดินเดิม": "Subgrade",
        "กำหนดเอง...": "Custom Material",
        "แผ่นคอนกรีต": "Concrete Slab",
        "Concrete Slab": "Concrete Slab",
    }
    
    # สีสำหรับแต่ละประเภทวัสดุ
    LAYER_COLORS = {
        "รองผิวทางคอนกรีตด้วย AC": "#2C3E50",
        "รองผิวทางคอนกรีตด้วย PMA(AC)": "#1A252F",
        "พื้นทางซีเมนต์ CTB": "#7F8C8D",
        "หินคลุกผสมซีเมนต์ UCS 24.5 ksc": "#95A5A6",
        "หินคลุก CBR 80%": "#BDC3C7",
        "ดินซีเมนต์ UCS 17.5 ksc": "#AAB7B8",
        "วัสดุหมุนเวียน (Recycling)": "#85929E",
        "รองพื้นทางวัสดุมวลรวม CBR 25%": "#FFCC99",
        "วัสดุคัดเลือก ก": "#E8DAEF",
        "ดินถมคันทาง / ดินเดิม": "#F5CBA7",
        "กำหนดเอง...": "#FADBD8",
        "Concrete Slab": "#808080",
    }
    
    # กรองเฉพาะชั้นที่มีความหนา > 0
    valid_layers = [l for l in layers_data if l.get("thickness_cm", 0) > 0]
    
    # เพิ่มชั้นคอนกรีตถ้ามี
    all_layers = []
    if concrete_thickness_cm and concrete_thickness_cm > 0:
        all_layers.append({
            "name": "Concrete Slab",
            "thickness_cm": concrete_thickness_cm,
            "E_MPa": None
        })
    all_layers.extend(valid_layers)
    
    if not all_layers:
        return None
    
    # คำนวณความหนารวม
    total_thickness = sum(l.get("thickness_cm", 0) for l in all_layers)
    
    # ใช้ scale factor เพื่อให้ชั้นบางๆ ยังมองเห็นได้
    min_display_height = 8  # ความสูงขั้นต่ำในการแสดงผล
    
    # สร้าง figure
    fig, ax = plt.subplots(figsize=(12, 8))
    
    # กำหนดขนาดรูป
    width = 3  # ความกว้างของชั้นทาง
    x_center = 6  # ตำแหน่ง x กึ่งกลาง
    x_start = x_center - width / 2
    
    # คำนวณความสูงที่ใช้แสดงผล
    display_heights = []
    for layer in all_layers:
        thickness = layer.get("thickness_cm", 0)
        display_h = max(thickness, min_display_height)
        display_heights.append(display_h)
    
    total_display = sum(display_heights)
    y_current = total_display
    
    # วาดแต่ละชั้น
    for i, layer in enumerate(all_layers):
        thickness = layer.get("thickness_cm", 0)
        name = layer.get("name", f"Layer {i+1}")
        e_mpa = layer.get("E_MPa", None)
        display_h = display_heights[i]
        
        if thickness <= 0:
            continue
        
        # หาสี
        color = LAYER_COLORS.get(name, "#CCCCCC")
        
        # กำหนด hatch pattern สำหรับวัสดุหมุนเวียน
        hatch_pattern = '///' if name == "วัสดุหมุนเวียน (Recycling)" else None
        
        # วาดสี่เหลี่ยม
        y_bottom = y_current - display_h
        rect = patches.Rectangle(
            (x_start, y_bottom), 
            width, 
            display_h,
            linewidth=2,
            edgecolor='black',
            facecolor=color,
            hatch=hatch_pattern
        )
        ax.add_patch(rect)
        
        # เพิ่มข้อความ
        y_center_pos = y_bottom + display_h / 2
        
        # แปลงชื่อเป็นภาษาอังกฤษ
        display_name = THAI_TO_ENG.get(name, name)
        
        # กำหนดสีข้อความตามสีพื้นหลัง
        is_dark = name in ["รองผิวทางคอนกรีตด้วย AC", "รองผิวทางคอนกรีตด้วย PMA(AC)", "Concrete Slab", 
                          "พื้นทางซีเมนต์ CTB", "หินคลุกผสมซีเมนต์ UCS 24.5 ksc",
                          "วัสดุหมุนเวียน (Recycling)"]
        text_color = 'white' if is_dark else 'black'
        
        # ข้อความในกล่อง (ความหนา)
        ax.text(x_center, y_center_pos, f"{thickness} cm",
                ha='center', va='center', fontsize=16, fontweight='bold', color=text_color)
        
        # ข้อความด้านซ้าย (ชื่อวัสดุ)
        ax.text(x_start - 0.5, y_center_pos, display_name,
                ha='right', va='center', fontsize=14, fontweight='bold', color='black')
        
        # ข้อความด้านขวา (E value)
        if e_mpa:
            ax.text(x_start + width + 0.5, y_center_pos, f"E = {e_mpa:,} MPa",
                    ha='left', va='center', fontsize=12, color='#0066CC')
        
        y_current = y_bottom
    
    # วาดเส้นบอกขนาดรวมด้านขวาสุด
    ax.annotate('', xy=(x_start + width + 3.5, total_display), 
                xytext=(x_start + width + 3.5, 0),
                arrowprops=dict(arrowstyle='<->', color='red', lw=2))
    ax.text(x_start + width + 4, total_display / 2, f"Total\n{total_thickness} cm",
            ha='left', va='center', fontsize=14, color='red', fontweight='bold')
    
    # ตั้งค่า axes
    margin = 10
    ax.set_xlim(0, 14)
    ax.set_ylim(-margin, total_display + margin)
    ax.axis('off')
    
    # หัวข้อ
    ax.set_title('Pavement Structure', 
                 fontsize=20, fontweight='bold', pad=20)
    
    # เพิ่มข้อความความหนารวมด้านล่าง
    ax.text(x_center, -margin + 4, 
            f"Total Pavement Thickness: {total_thickness} cm",
            ha='center', va='center', fontsize=15, fontweight='bold',
            bbox=dict(boxstyle='round', facecolor='lightyellow', alpha=0.9, edgecolor='orange'))
    
    plt.tight_layout()
    
    return fig


def save_figure_to_bytes(fig):
    """บันทึก matplotlib figure เป็น bytes สำหรับดาวน์โหลด"""
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', 
                facecolor='white', edgecolor='none')
    buf.seek(0)
    return buf


# ============================================================
# ส่วนที่ 3: ฟังก์ชันสร้างรายงาน Word
# ============================================================

def create_word_report(
    pavement_type: str,
    inputs: dict,
    calculated_values: dict,
    comparison_results: list,
    selected_d_cm: int,
    main_result: tuple,
    layers_data: list = None,
    project_name: str = "",
    structure_figure = None,
    subgrade_info: dict = None,
    e_equivalent_psi: float = 0
) -> BytesIO:
    """
    สร้างรายงานการคำนวณในรูปแบบไฟล์ Word (.docx)
    ใช้ python-docx library
    
    Parameters:
        selected_d_cm: ความหนาที่เลือก (ซม.)
        project_name: ชื่อโครงการ
        structure_figure: รูป matplotlib ของโครงสร้างชั้นทาง
        subgrade_info: ข้อมูลดินคันทาง {'cbr': %, 'mr_psi': psi}
        e_equivalent_psi: ค่า E_equivalent (psi)
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        st.error("กรุณาติดตั้ง python-docx: pip install python-docx")
        return None
    
    # ปัดค่านิ้วเป็นจำนวนเต็ม (เหมือนกับการคำนวณ)
    selected_d_inch = round(selected_d_cm / 2.54)
    
    # สร้างเอกสารใหม่
    doc = Document()
    
    # ตั้งค่าฟอนต์ภาษาไทย
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH Sarabun New'
    font.size = Pt(14)
    
    # หัวข้อเอกสาร
    title = doc.add_heading('รายการคำนวณออกแบบความหนาถนนคอนกรีต', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('ตามวิธี AASHTO 1993')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ข้อมูลทั่วไป
    doc.add_heading('1. ข้อมูลทั่วไป', level=1)
    if project_name:
        doc.add_paragraph(f'ชื่อโครงการ: {project_name}')
    doc.add_paragraph(f'ประเภทถนน: {pavement_type}')
    doc.add_paragraph(f'วันที่คำนวณ: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    
    # ตารางชั้นโครงสร้างทาง
    doc.add_heading('2. ชั้นโครงสร้างทาง (Pavement Layers)', level=1)
    
    table_layers = doc.add_table(rows=1, cols=4)
    table_layers.style = 'Table Grid'
    hdr_layers = table_layers.rows[0].cells
    hdr_layers[0].text = 'ลำดับ'
    hdr_layers[1].text = 'ชนิดวัสดุ'
    hdr_layers[2].text = 'ความหนา (ซม.)'
    hdr_layers[3].text = 'Modulus E (MPa)'
    
    # แถวที่ 1: ผิวทางคอนกรีต
    row_cells = table_layers.add_row().cells
    row_cells[0].text = '1'
    row_cells[1].text = f'ผิวทางคอนกรีต {pavement_type}'
    row_cells[2].text = f'{selected_d_cm}'
    row_cells[3].text = '-'
    
    # แถวถัดไป: ชั้นวัสดุอื่นๆ
    layer_count = 1
    if layers_data and len(layers_data) > 0:
        for i, layer in enumerate(layers_data):
            layer_count += 1
            row_cells = table_layers.add_row().cells
            row_cells[0].text = str(layer_count)
            row_cells[1].text = layer.get('name', f'Layer {i+1}')
            row_cells[2].text = f"{layer.get('thickness_cm', 0)}"
            row_cells[3].text = f"{layer.get('E_MPa', 0):,}"
    
    # แถวสุดท้าย: ดินคันทาง
    if subgrade_info:
        layer_count += 1
        row_cells = table_layers.add_row().cells
        row_cells[0].text = str(layer_count)
        row_cells[1].text = 'ดินคันทาง'
        row_cells[2].text = f"CBR {subgrade_info.get('cbr', 0)} %"
        row_cells[3].text = f"{subgrade_info.get('mr_mpa', 0):.0f} ({subgrade_info.get('mr_psi', 0):,.0f} psi)"
    
    doc.add_paragraph('')  # เว้นบรรทัด
    
    # เพิ่มรูปโครงสร้างชั้นทาง
    if structure_figure is not None:
        doc.add_paragraph('รูปตัดโครงสร้างชั้นทาง:')
        
        # บันทึกรูปเป็น bytes
        img_buffer = BytesIO()
        structure_figure.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight',
                                 facecolor='white', edgecolor='none')
        img_buffer.seek(0)
        
        # เพิ่มรูปในเอกสาร
        from docx.shared import Inches
        doc.add_picture(img_buffer, width=Inches(5.5))
        
        # จัดกึ่งกลางรูป
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph('')  # เว้นบรรทัด
    
    # ข้อมูลนำเข้า
    doc.add_heading('3. ข้อมูลนำเข้า (Input Parameters)', level=1)
    
    table1 = doc.add_table(rows=1, cols=4)
    table1.style = 'Table Grid'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'พารามิเตอร์'
    hdr_cells[1].text = 'สัญลักษณ์'
    hdr_cells[2].text = 'ค่า'
    hdr_cells[3].text = 'หน่วย'
    
    input_data = [
        ('ESAL ออกแบบ', 'W₁₈', f"{inputs['w18_design']:,.0f}", 'ESALs'),
        ('Terminal Serviceability', 'Pt', f"{inputs['pt']:.1f}", '-'),
        ('Reliability', 'R', f"{inputs['reliability']:.0f}", '%'),
        ('Standard Deviation', 'So', f"{inputs['so']:.2f}", '-'),
        ('Modulus of Subgrade Reaction', 'k_eff', f"{inputs['k_eff']:,.0f}", 'pci'),
        ('Loss of Support', 'LS', f"{inputs.get('ls', 1.0):.1f}", '-'),
        ('กำลังคอนกรีต', "f'c", f"{inputs['fc_cube']:.0f} Cube ({int(inputs['fc_cube']*0.8)} Cyl.)", 'ksc'),
        ('Modulus of Rupture', 'Sc', f"{inputs['sc']:.0f}", 'psi'),
        ('Load Transfer Coefficient', 'J', f"{inputs['j']:.1f}", '-'),
        ('Drainage Coefficient', 'Cd', f"{inputs['cd']:.1f}", '-'),
    ]
    
    for param, symbol, value, unit in input_data:
        row_cells = table1.add_row().cells
        row_cells[0].text = param
        row_cells[1].text = symbol
        row_cells[2].text = value
        row_cells[3].text = unit
    
    # ค่าที่คำนวณได้
    doc.add_heading('4. ค่าที่คำนวณได้ (Calculated Values)', level=1)
    
    table2 = doc.add_table(rows=1, cols=4)
    table2.style = 'Table Grid'
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = 'พารามิเตอร์'
    hdr_cells2[1].text = 'สัญลักษณ์'
    hdr_cells2[2].text = 'ค่า'
    hdr_cells2[3].text = 'หน่วย'
    
    calc_data = [
        ('Modulus of Elasticity', 'Ec', f"{calculated_values['ec']:,.0f}", 'psi'),
        ('Standard Normal Deviate', 'ZR', f"{calculated_values['zr']:.3f}", '-'),
        ('การสูญเสีย Serviceability', 'ΔPSI', f"{calculated_values['delta_psi']:.1f}", '-'),
    ]
    
    for param, symbol, value, unit in calc_data:
        row_cells = table2.add_row().cells
        row_cells[0].text = param
        row_cells[1].text = symbol
        row_cells[2].text = value
        row_cells[3].text = unit
    
    # สมการ AASHTO 1993
    doc.add_heading('5. สมการออกแบบ AASHTO 1993', level=1)
    
    equation_text = """
    log₁₀(W₁₈) = ZR × So + 7.35 × log₁₀(D+1) - 0.06 
                 + log₁₀(ΔPSI/(4.5-1.5)) / (1 + 1.624×10⁷/(D+1)^8.46)
                 + (4.22 - 0.32×Pt) × log₁₀[(Sc×Cd×(D^0.75-1.132))/(215.63×J×(D^0.75 - 18.42/(Ec/k)^0.25))]
    """
    doc.add_paragraph(equation_text)
    
    # ผลการเปรียบเทียบ
    doc.add_heading('6. ผลการเปรียบเทียบความหนาต่างๆ', level=1)
    
    table3 = doc.add_table(rows=1, cols=6)
    table3.style = 'Table Grid'
    hdr_cells3 = table3.rows[0].cells
    hdr_cells3[0].text = 'D (ซม.)'
    hdr_cells3[1].text = 'D (นิ้ว)'
    hdr_cells3[2].text = 'log₁₀(W₁₈)'
    hdr_cells3[3].text = 'W₁₈ รองรับได้'
    hdr_cells3[4].text = 'อัตราส่วน'
    hdr_cells3[5].text = 'ผลการตรวจสอบ'
    
    for result in comparison_results:
        row_cells = table3.add_row().cells
        row_cells[0].text = f"{result['d_cm']:.0f}"
        row_cells[1].text = f"{result['d_inch']:.0f}"
        row_cells[2].text = f"{result['log_w18']:.4f}"
        row_cells[3].text = f"{result['w18']:,.0f}"
        row_cells[4].text = f"{result['ratio']:.2f}"
        row_cells[5].text = "ผ่าน ✓" if result['passed'] else "ไม่ผ่าน ✗"
    
    # สรุปผล
    doc.add_heading('7. สรุปผลการออกแบบ', level=1)
    
    passed, ratio = main_result
    status = "ผ่านเกณฑ์ ✓" if passed else "ไม่ผ่านเกณฑ์ ✗"
    
    # หา ESAL ที่รองรับได้จาก comparison_results
    w18_capacity_selected = None
    for r in comparison_results:
        if r['d_cm'] == selected_d_cm:
            w18_capacity_selected = r['w18']
            break
    
    # คำนวณ E_equivalent เป็น MPa
    e_equivalent_mpa = e_equivalent_psi / 145.038 if e_equivalent_psi > 0 else 0
    
    summary = f"""
    ความหนาที่เลือก: {selected_d_cm:.0f} ซม. ({selected_d_inch:.0f} นิ้ว)
    ESAL ที่ต้องการ: {inputs['w18_design']:,.0f} ESALs
    ESAL ที่รองรับได้: {w18_capacity_selected:,.0f} ESALs (โดยประมาณ)
    อัตราส่วน: {ratio:.2f}
    E_equivalent ของวัสดุ: {e_equivalent_psi:,.0f} psi ({e_equivalent_mpa:.1f} MPa)
    ผลการตรวจสอบ: {status}
    """
    doc.add_paragraph(summary)
    
    # หมายเหตุ
    doc.add_heading('8. หมายเหตุ', level=1)
    notes = """
    - การคำนวณนี้ใช้หลักการตามคู่มือ AASHTO Guide for Design of Pavement Structures (1993)
    - สมการ: log₁₀(W₁₈) รวม term (D^0.75 - 1.132) ในตัวเศษ
    - ค่า J Default: JRCP/JPCP = 2.8, CRCP = 2.5
    - การแปลงกำลังคอนกรีต: f'c (cylinder) ≈ 0.833 × f'c (cube)
    - Ec = 57,000 × √f'c (psi) ตาม ACI 318
    - Sc ≈ 10 × √f'c (psi)
    """
    doc.add_paragraph(notes)
    
    # บันทึกไฟล์ลง BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


# ============================================================
# ส่วนที่ 4: Streamlit UI
# ============================================================

def main():
    # ตั้งค่าหน้าเว็บ
    st.set_page_config(
        page_title="AASHTO 1993 Rigid Pavement Design",
        page_icon="🛣️",
        layout="wide"
    )
    
    # หัวข้อหลัก
    st.title("🛣️ การออกแบบความหนาถนนคอนกรีต")
    st.subheader("ตามวิธี AASHTO 1993 (Rigid Pavement Design)")
    
    st.markdown("---")
    
    # แบ่งคอลัมน์
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("📥 ข้อมูลนำเข้า (Input)")
        
        # ชื่อโครงการ
        project_name = st.text_input(
            "🏗️ ชื่อโครงการ",
            value="",
            placeholder="กรอกชื่อโครงการ...",
            help="ชื่อโครงการจะแสดงในรายงาน Word"
        )
        
        st.markdown("---")
        
        # เลือกประเภทผิวทางคอนกรีต
        pavement_type = st.selectbox(
            "ประเภทผิวทางคอนกรีต",
            options=list(J_VALUES.keys()),
            index=1,
            help="JRCP = Jointed Reinforced, JPCP = Jointed Plain, CRCP = Continuously Reinforced"
        )
        
        st.markdown("---")
        
        # ชั้นโครงสร้างทาง (Pavement Layers)
        st.subheader("🔶 ชั้นโครงสร้างทาง (Pavement Layers)")
        
        # ตารางค่า Modulus ตามประเภทวัสดุ (ตามข้อมูลอาจารย์)
        MATERIAL_MODULUS = {
            "รองผิวทางคอนกรีตด้วย AC": 2500,
            "รองผิวทางคอนกรีตด้วย PMA(AC)": 3700,
            "พื้นทางซีเมนต์ CTB": 1200,
            "หินคลุกผสมซีเมนต์ UCS 24.5 ksc": 850,
            "หินคลุก CBR 80%": 350,
            "ดินซีเมนต์ UCS 17.5 ksc": 350,
            "วัสดุหมุนเวียน (Recycling)": 850,
            "รองพื้นทางวัสดุมวลรวม CBR 25%": 150,
            "วัสดุคัดเลือก ก": 76,
            "ดินถมคันทาง / ดินเดิม": 100,
            "กำหนดเอง...": 100,
        }
        
        # รายการวัสดุสำหรับ dropdown
        material_options = list(MATERIAL_MODULUS.keys())
        
        # จำนวนชั้นวัสดุ
        num_layers = st.slider(
            "จำนวนชั้นวัสดุใต้แผ่นคอนกรีต",
            min_value=1,
            max_value=6,
            value=5,
            help="เลือกจำนวนชั้นวัสดุ 1-6 ชั้น"
        )
        
        # ค่า Default สำหรับแต่ละชั้น
        default_layers = [
            {"name": "รองผิวทางคอนกรีตด้วย AC", "thickness_cm": 5},
            {"name": "พื้นทางซีเมนต์ CTB", "thickness_cm": 20},
            {"name": "หินคลุก CBR 80%", "thickness_cm": 15},
            {"name": "รองพื้นทางวัสดุมวลรวม CBR 25%", "thickness_cm": 25},
            {"name": "วัสดุคัดเลือก ก", "thickness_cm": 30},
            {"name": "ดินถมคันทาง / ดินเดิม", "thickness_cm": 0},
        ]
        
        # เก็บข้อมูลชั้นวัสดุ
        layers_data = []
        
        with st.expander("📊 ตารางค่า Modulus อ้างอิง", expanded=False):
            st.markdown("""
            | วัสดุชั้นทาง | MR (MPa) |
            |-------------|----------|
            | รองผิวทางคอนกรีตด้วย AC | 2,500 |
            | รองผิวทางคอนกรีตด้วย PMA(AC) | 3,700 |
            | พื้นทางซีเมนต์ CTB | 1,200 |
            | หินคลุกผสมซีเมนต์ UCS 24.5 ksc | 850 |
            | หินคลุก CBR 80% | 350 |
            | ดินซีเมนต์ UCS 17.5 ksc | 350 |
            | วัสดุหมุนเวียน (Recycling) | 850 |
            | รองพื้นทางวัสดุมวลรวม CBR 25% | 150 |
            | วัสดุคัดเลือก ก | 76 |
            | ดินถมคันทาง / ดินเดิม | 100 |
            """)
        
        for i in range(num_layers):
            st.markdown(f"**ชั้นที่ {i+1}**")
            col_a, col_b, col_c = st.columns([2, 1, 1])
            
            # หา default index สำหรับ dropdown
            default_name = default_layers[i]["name"] if i < len(default_layers) else "กำหนดเอง..."
            default_index = material_options.index(default_name) if default_name in material_options else len(material_options) - 1
            
            with col_a:
                layer_name = st.selectbox(
                    f"เลือกวัสดุ",
                    options=material_options,
                    index=default_index,
                    key=f"layer_name_{i}"
                )
            
            with col_b:
                layer_thickness = st.number_input(
                    f"ความหนา (ซม.)",
                    min_value=0,
                    max_value=100,
                    value=default_layers[i]["thickness_cm"] if i < len(default_layers) else 20,
                    key=f"layer_thick_{i}"
                )
            
            # หาค่า Modulus จากวัสดุที่เลือก
            recommended_modulus = MATERIAL_MODULUS.get(layer_name, 100)
            
            with col_c:
                # ใช้ key ที่รวม layer_name เพื่อให้ reset เมื่อเปลี่ยนวัสดุ
                layer_modulus = st.number_input(
                    f"E (MPa)",
                    min_value=10,
                    max_value=10000,
                    value=recommended_modulus,
                    key=f"layer_E_{i}_{layer_name}",
                    help=f"ค่าแนะนำ: {recommended_modulus:,} MPa"
                )
            
            layers_data.append({
                "name": layer_name,
                "thickness_cm": layer_thickness,
                "E_MPa": layer_modulus
            })
        
        st.markdown("---")
        
        # 1. ESAL ที่ต้องการรองรับ
        st.subheader("1️⃣ ปริมาณจราจร")
        
        # แสดงตัวช่วยประมาณ ESAL
        with st.expander("📊 ตัวช่วยประมาณ ESAL ตามประเภทถนน"):
            st.markdown("""
            | ประเภทถนน | ESAL (20 ปี) |
            |-----------|--------------|
            | ถนนในหมู่บ้าน | 50,000 - 200,000 |
            | ถนนเทศบาล | 200,000 - 500,000 |
            | ถนน อบจ. / ทางหลวงชนบท | 500,000 - 2,000,000 |
            | ทางหลวงแผ่นดิน (2 ช่องจราจร) | 2,000,000 - 10,000,000 |
            | ทางหลวงแผ่นดิน (4 ช่องจราจร) | 10,000,000 - 50,000,000 |
            """)
        
        w18_design = st.number_input(
            "ESAL ที่ต้องการรองรับ (W₁₈)",
            min_value=10_000,
            max_value=500_000_000,
            value=500_000,
            step=100_000,
            format="%d",
            help="จำนวน Equivalent Single Axle Load (18 kip) ตลอดอายุการใช้งาน"
        )
        
        # แสดงค่า ESAL เป็นภาษาไทย
        def format_thai_number(num):
            if num >= 1_000_000_000:
                return f"{num/1_000_000_000:.2f} พันล้าน ESALs"
            elif num >= 1_000_000:
                return f"{num/1_000_000:.2f} ล้าน ESALs"
            elif num >= 100_000:
                return f"{num/100_000:.2f} แสน ESALs"
            elif num >= 10_000:
                return f"{num/10_000:.2f} หมื่น ESALs"
            elif num >= 1_000:
                return f"{num/1_000:.2f} พัน ESALs"
            else:
                return f"{num:.0f} ESALs"
        
        st.markdown(f"<span style='color: #1E90FF; font-size: 1.2em; font-weight: bold;'>{format_thai_number(w18_design)}</span>", unsafe_allow_html=True)
        
        st.markdown("---")
        
        # 2. Serviceability
        st.subheader("2️⃣ Serviceability")
        pt = st.slider(
            "Terminal Serviceability (Pt)",
            min_value=1.5,
            max_value=3.0,
            value=2.0,
            step=0.1,
            help="ค่า Serviceability ที่ยอมรับได้ต่ำสุด (มาตรฐาน = 2.0)"
        )
        
        # คำนวณ ΔPSI
        delta_psi = 4.5 - pt
        st.info(f"ΔPSI = 4.5 - {pt:.1f} = **{delta_psi:.1f}**")
        
        st.markdown("---")
        
        # 3. Reliability
        st.subheader("3️⃣ ความเชื่อมั่นในการออกแบบ")
        reliability = st.select_slider(
            "Reliability (R)",
            options=[80, 85, 90, 95],
            value=90,
            help="ระดับความเชื่อมั่นในการออกแบบ (%)"
        )
        
        # หาค่า ZR
        zr = get_zr_value(reliability)
        st.info(f"ZR = **{zr:.3f}** (จากตาราง AASHTO)")
        
        # Standard Deviation
        so = st.number_input(
            "Overall Standard Deviation (So)",
            min_value=0.30,
            max_value=0.45,
            value=0.35,
            step=0.01,
            format="%.2f",
            help="ค่าเบี่ยงเบนมาตรฐานรวม (มาตรฐาน = 0.35 สำหรับ Rigid Pavement)"
        )
        
        st.markdown("---")
        
        # 4. คุณสมบัติดินฐานราก
        st.subheader("4️⃣ คุณสมบัติดินฐานราก")
        
        # ตารางประเภทดินและค่า CBR
        SUBGRADE_TYPES = {
            "ดินถมคันทาง CBR 2%": 2,
            "ดินถมคันทาง CBR 3%": 3,
            "ดินถมคันทาง CBR 4%": 4,
            "ดินถมคันทาง CBR 5%": 5,
            "ดินถมคันทาง CBR 6%": 6,
            "ดินถมคันทาง CBR 8%": 8,
            "ดินถมคันทาง CBR 10%": 10,
            "ดินถมคันทาง CBR 15%": 15,
            "ดินถมคันทาง CBR 20%": 20,
            "กำหนดเอง": 0,
        }
        
        subgrade_type = st.selectbox(
            "เลือกประเภทดินฐานราก",
            options=list(SUBGRADE_TYPES.keys()),
            index=2,  # Default: CBR 4%
            help="เลือกประเภทดินฐานรากตามค่า CBR"
        )
        
        # ค่า CBR
        if subgrade_type == "กำหนดเอง":
            cbr_value = st.number_input(
                "ค่า CBR (%)",
                min_value=1.0,
                max_value=100.0,
                value=4.0,
                step=0.5,
                format="%.2f"
            )
        else:
            cbr_value = float(SUBGRADE_TYPES[subgrade_type])
            st.number_input(
                "ค่า CBR (%)",
                min_value=1.0,
                max_value=100.0,
                value=cbr_value,
                step=0.5,
                format="%.2f",
                disabled=True
            )
        
        # คำนวณ MR (Subgrade) จาก CBR
        # สูตร: MR (psi) = 1500 × CBR (สำหรับ CBR < 10)
        # หรือ MR (psi) = 1000 + 555 × CBR (สำหรับ CBR ≥ 10)
        if cbr_value < 10:
            mr_subgrade_psi = 1500 * cbr_value
        else:
            mr_subgrade_psi = 1000 + 555 * cbr_value
        
        mr_subgrade_mpa = mr_subgrade_psi / 145.038  # แปลง psi เป็น MPa
        
        st.info(f"M_R (Subgrade) = {mr_subgrade_psi:,.0f} psi ({mr_subgrade_mpa:.0f} MPa)")
        
        with st.expander("📊 ตารางประมาณค่า k จาก CBR"):
            st.markdown("""
            | CBR (%) | k (pci) | คำอธิบาย |
            |---------|---------|----------|
            | 2-3 | 75-100 | ดินเหนียวอ่อน |
            | 4-5 | 100-130 | ดินเหนียวแข็ง |
            | 6-10 | 130-170 | ดินทรายปนดินเหนียว |
            | 10-20 | 170-230 | ดินทรายอัดแน่น |
            | 20-50 | 230-350 | หินคลุก/ลูกรัง |
            | > 50 | 350-500+ | ชั้น Base คุณภาพดี |
            
            **หมายเหตุ:** ค่า k_eff รวมผลของชั้น Subbase แล้ว
            """)
        
        k_eff = st.number_input(
            "Effective Modulus of Subgrade Reaction (k_eff)",
            min_value=50,
            max_value=1000,
            value=200,
            step=25,
            format="%d",
            help="ค่า k จากการทดสอบ Plate Bearing Test หรือประมาณจาก CBR (หน่วย: pci)"
        )
        
        # Loss of Support (LS)
        st.markdown("**Loss of Support (LS)**")
        
        with st.expander("📊 ตารางค่า Loss of Support แนะนำ (AASHTO 1993)"):
            st.markdown("""
            | ประเภทวัสดุ | Loss of Support (LS) |
            |------------|---------------------|
            | Cement Treated Granular Base | 0.0 - 1.0 |
            | Cement Aggregate Mixtures | 0.0 - 1.0 |
            | Asphalt Treated Base | 0.0 - 1.0 |
            | Bituminous Stabilized Mixtures | 0.0 - 1.0 |
            | Lime Stabilized | 1.0 - 3.0 |
            | Unbound Granular Materials | 1.0 - 3.0 |
            | Fine Grained or Natural Subgrade | 2.0 - 3.0 |
            
            **หมายเหตุ:** ค่า LS ใช้ปรับลดค่า k_eff เพื่อคำนึงถึงการสูญเสียการรองรับจากการกัดเซาะ
            """)
        
        ls_value = st.number_input(
            "ค่า Loss of Support (LS)",
            min_value=0.0,
            max_value=3.0,
            value=1.0,
            step=0.5,
            format="%.1f",
            help="ค่า LS สำหรับปรับลด k_eff (0.0-3.0)"
        )
        
        st.markdown("---")
        
        # 5. คุณสมบัติคอนกรีต
        st.subheader("5️⃣ คุณสมบัติคอนกรีต")
        
        fc_cube = st.number_input(
            "กำลังอัดคอนกรีต (Cube) - f'c",
            min_value=200,
            max_value=600,
            value=350,
            step=10,
            format="%d",
            help="กำลังอัดคอนกรีตที่ 28 วัน ทดสอบด้วย Cube 15×15×15 ซม. (หน่วย: ksc)"
        )
        
        # แปลง Cube เป็น Cylinder
        fc_cylinder = convert_cube_to_cylinder(fc_cube)
        st.info(f"f'c (Cylinder) = 0.8 × {fc_cube} = **{fc_cylinder:.0f} ksc**")
        
        # คำนวณ Ec
        ec = calculate_concrete_modulus(fc_cylinder)
        st.info(f"Ec = 57,000 × √({fc_cylinder * 14.223:.0f}) = **{ec:,.0f} psi**")
        
        # Modulus of Rupture
        st.markdown("**Modulus of Rupture (Sc)**")
        
        # คำนวณค่า Sc อัตโนมัติ
        sc_auto = estimate_modulus_of_rupture(fc_cylinder)
        st.info(f"ค่าประมาณ: Sc = 10 × √({fc_cylinder * 14.223:.0f}) = **{sc_auto:.0f} psi**")
        
        # ให้ผู้ใช้ป้อนค่าที่ต้องการใช้
        sc = st.number_input(
            "ค่า Sc ที่ใช้ในการคำนวณ (psi)",
            min_value=400,
            max_value=1000,
            value=int(round(sc_auto)),
            step=10,
            format="%d",
            help="ค่าเริ่มต้นคำนวณจาก 10×√f'c สามารถแก้ไขได้ตามผลทดสอบจริง"
        )
        
        st.markdown("---")
        
        # 6. Load Transfer และ Drainage
        st.subheader("6️⃣ Load Transfer และ Drainage")
        
        # แสดงค่า J อัตโนมัติตามประเภทผิวทาง
        j_auto = J_VALUES[pavement_type]
        st.info(f"ค่าแนะนำสำหรับ {pavement_type}: **J = {j_auto}**")
        
        # ตารางอ้างอิงค่า J
        with st.expander("📊 ตารางค่า Load Transfer Coefficient (J)"):
            st.markdown("""
            | ประเภทถนน | J (AC Shoulder_Yes) | J (AC Shoulder_No) | J (Tied P.C.C_Yes) | J (Tied P.C.C_No) |
            |-----------|---------------------|--------------------|--------------------|-------------------|
            | 1. JRCP/JPCP | 3.2 | 3.8-4.4 | 2.5-3.1 (Mid 2.8) | 3.6-4.2 |
            | 2. CRCP | 2.9-3.2 | N/A | 2.3-2.9 (Mid 2.5) | N/A |
            
            **หมายเหตุ:** ค่า J ต่ำ = การถ่ายแรงดี = รองรับ ESAL ได้มากขึ้น
            
            ค่า J สามารถปรับได้ตามเงื่อนไข:
            - มี Dowel Bar: ลดลง 0.2-0.3
            - มี Tied Shoulder: ลดลง 0.2
            - ไม่มี Dowel Bar: เพิ่มขึ้น 0.5-1.0
            """)
        
        # ให้ผู้ใช้ป้อนค่าที่ต้องการใช้
        j_value = st.number_input(
            "ค่า J ที่ใช้ในการคำนวณ",
            min_value=2.0,
            max_value=4.5,
            value=j_auto,
            step=0.1,
            format="%.1f",
            help="ค่าเริ่มต้นตามประเภทถนนที่เลือก สามารถแก้ไขได้"
        )
        
        cd = st.number_input(
            "Drainage Coefficient (Cd)",
            min_value=0.7,
            max_value=1.3,
            value=1.0,
            step=0.05,
            format="%.2f",
            help="สัมประสิทธิ์การระบายน้ำ (1.0 = การระบายน้ำปานกลาง)"
        )
        
        st.markdown("---")
        
        # 7. ความหนาคอนกรีต
        st.subheader("7️⃣ ความหนาคอนกรีตที่ต้องการตรวจสอบ")
        d_cm_selected = st.slider(
            "ความหนาคอนกรีต D (ซม.)",
            min_value=20,
            max_value=40,
            value=30,
            step=1,
            help="ความหนาแผ่นพื้นคอนกรีต (20-40 ซม.)"
        )
        # ปัดค่านิ้วเป็นจำนวนเต็ม แล้วใช้ค่านั้นในการคำนวณ
        d_inch_selected = round(d_cm_selected / 2.54)
        st.info(f"D = **{d_cm_selected} ซม.** ≈ **{d_inch_selected} นิ้ว** (ใช้ในการคำนวณ)")
        
        st.markdown("---")
        
        # ============================================================
        # ตารางสรุปข้อมูลที่กรอก และคำนวณ E_equivalent
        # ============================================================
        st.subheader("📊 สรุปข้อมูลที่กรอก")
        
        # สร้างตารางสรุป
        import pandas as pd
        
        summary_data = []
        total_thickness_cm = 0
        sum_e_h = 0  # ผลรวม E × h
        
        for i, layer in enumerate(layers_data):
            h_cm = layer.get('thickness_cm', 0)
            if h_cm > 0:
                h_inch = h_cm / 2.54
                e_mpa = layer.get('E_MPa', 0)
                total_thickness_cm += h_cm
                sum_e_h += e_mpa * h_cm
                
                summary_data.append({
                    'ชั้น': f'ชั้นที่ {i+1}',
                    'ชนิดวัสดุ': layer.get('name', ''),
                    'ความหนา (ซม.)': h_cm,
                    'ความหนา (นิ้ว)': round(h_inch, 4),
                    'MR (MPa)': e_mpa
                })
        
        # กำหนดค่า default สำหรับ e_equivalent
        e_equivalent_mpa = 0
        e_equivalent_psi = 0
        
        if summary_data:
            df_summary = pd.DataFrame(summary_data)
            st.dataframe(df_summary, use_container_width=True, hide_index=True)
            
            # คำนวณ E_equivalent ด้วยสูตร Odemark's Method
            # E_eq = (Σ(h_i × E_i^(1/3)) / Σh_i)³
            if total_thickness_cm > 0:
                sum_h_e_cbrt = 0  # ผลรวม h × E^(1/3)
                for layer in layers_data:
                    h = layer.get('thickness_cm', 0)
                    e = layer.get('E_MPa', 0)
                    if h > 0 and e > 0:
                        sum_h_e_cbrt += h * (e ** (1/3))
                
                e_equivalent_mpa = (sum_h_e_cbrt / total_thickness_cm) ** 3
                e_equivalent_psi = e_equivalent_mpa * 145.038
                
                st.info(f"โมดูลัสเทียบเท่า (E_equivalent) = **{e_equivalent_psi:,.0f} psi** ({e_equivalent_mpa:.1f} MPa)")
        
        st.markdown("---")
        
        # แสดงรูปโครงสร้างชั้นทาง (รวมชั้นคอนกรีตบนสุด)
        st.subheader("📐 รูปตัดโครงสร้างชั้นทาง")
        
        # สร้างรูป (ชั้นคอนกรีตจะอยู่บนสุด)
        fig_structure = create_pavement_structure_figure(layers_data, concrete_thickness_cm=d_cm_selected)
        
        if fig_structure:
            st.pyplot(fig_structure)
            
            # ปุ่มดาวน์โหลดรูป
            img_buffer = save_figure_to_bytes(fig_structure)
            st.download_button(
                label="📥 ดาวน์โหลดรูปโครงสร้างชั้นทาง",
                data=img_buffer,
                file_name=f"pavement_structure_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
                mime="image/png"
            )
            plt.close(fig_structure)
    
    # ============================================================
    # ส่วนแสดงผลการคำนวณ
    # ============================================================
    
    with col2:
        st.header("📊 ผลการคำนวณ (Output)")
        
        # เก็บผลการคำนวณสำหรับความหนาต่างๆ
        comparison_results = []
        
        # ใช้ความหนาเป็น cm เป็นหลัก
        thicknesses_cm = [20, 22, 25, 28, 30, 32, 35, 38, 40]
        
        # คำนวณสำหรับแต่ละความหนา
        st.subheader("📋 ตารางเปรียบเทียบความหนาต่างๆ")
        
        # สร้างตาราง
        table_data = []
        for d_cm in thicknesses_cm:
            # ปัดค่านิ้วเป็นจำนวนเต็ม แล้วใช้ค่านั้นในการคำนวณ
            d_inch = round(d_cm / 2.54)
            
            log_w18, w18_capacity = calculate_aashto_rigid_w18(
                d_inch=d_inch,  # ใช้ค่านิ้วที่ปัดแล้ว
                delta_psi=delta_psi,
                pt=pt,
                zr=zr,
                so=so,
                sc_psi=sc,
                cd=cd,
                j=j_value,
                ec_psi=ec,
                k_pci=k_eff
            )
            passed, ratio = check_design(w18_design, w18_capacity)
            
            comparison_results.append({
                'd_cm': d_cm,
                'd_inch': d_inch,
                'log_w18': log_w18,
                'w18': w18_capacity,
                'passed': passed,
                'ratio': ratio
            })
            
            table_data.append({
                'D (ซม.)': d_cm,
                'D (นิ้ว)': d_inch,
                'log₁₀(W₁₈)': f"{log_w18:.4f}",
                'W₁₈ รองรับได้': f"{w18_capacity:,.0f}",
                'อัตราส่วน': f"{ratio:.2f}",
                'ผล': "✅ ผ่าน" if passed else "❌ ไม่ผ่าน"
            })
        
        # แสดงตาราง
        import pandas as pd
        df = pd.DataFrame(table_data)
        
        # จัดรูปแบบตาราง
        st.dataframe(
            df,
            use_container_width=True,
            hide_index=True
        )
        
        st.markdown("---")
        
        # ผลการคำนวณสำหรับความหนาที่เลือก
        st.subheader(f"🎯 ผลการตรวจสอบ D = {d_cm_selected} ซม. ({d_inch_selected} นิ้ว)")
        
        log_w18_selected, w18_selected = calculate_aashto_rigid_w18(
            d_inch=d_inch_selected,  # ใช้ค่านิ้วที่ปัดแล้ว
            delta_psi=delta_psi,
            pt=pt,
            zr=zr,
            so=so,
            sc_psi=sc,
            cd=cd,
            j=j_value,
            ec_psi=ec,
            k_pci=k_eff
        )
        passed_selected, ratio_selected = check_design(w18_design, w18_selected)
        
        # แสดงผลด้วยสี
        col_a, col_b = st.columns(2)
        
        with col_a:
            st.metric(
                label="log₁₀(W₁₈)",
                value=f"{log_w18_selected:.4f}"
            )
            st.metric(
                label="W₁₈ รองรับได้",
                value=f"{w18_selected:,.0f}",
                delta=f"{w18_selected - w18_design:+,.0f}"
            )
        
        with col_b:
            st.metric(
                label="W₁₈ ที่ต้องการ",
                value=f"{w18_design:,.0f}"
            )
            st.metric(
                label="อัตราส่วน (Capacity/Required)",
                value=f"{ratio_selected:.2f}"
            )
        
        # แสดงผลผ่าน/ไม่ผ่าน
        if passed_selected:
            st.success(f"""
            ✅ **ผ่านเกณฑ์การออกแบบ**
            
            ความหนา D = {d_cm_selected} ซม. ({d_inch_selected} นิ้ว) 
            สามารถรองรับ ESAL ได้ {w18_selected:,.0f} ESALs
            ซึ่งมากกว่า ESAL ที่ต้องการ {w18_design:,.0f} ESALs
            
            อัตราส่วน = {ratio_selected:.2f} (≥ 1.00)
            """)
        else:
            st.error(f"""
            ❌ **ไม่ผ่านเกณฑ์การออกแบบ**
            
            ความหนา D = {d_cm_selected} ซม. ({d_inch_selected} นิ้ว) 
            รองรับ ESAL ได้เพียง {w18_selected:,.0f} ESALs
            ซึ่งน้อยกว่า ESAL ที่ต้องการ {w18_design:,.0f} ESALs
            
            อัตราส่วน = {ratio_selected:.2f} (< 1.00)
            
            **กรุณาเพิ่มความหนาคอนกรีต หรือปรับปรุงคุณสมบัติวัสดุ**
            """)
        
        st.markdown("---")
        
        # แสดงสมการที่ใช้
        st.subheader("📝 สมการ AASHTO 1993")
        
        st.latex(r'''
        \log_{10}(W_{18}) = Z_R \times S_o + 7.35 \times \log_{10}(D+1) - 0.06
        ''')
        
        st.latex(r'''
        + \frac{\log_{10}\left(\frac{\Delta PSI}{4.5-1.5}\right)}{1 + \frac{1.624 \times 10^7}{(D+1)^{8.46}}}
        ''')
        
        st.latex(r'''
        + (4.22 - 0.32 \times P_t) \times \log_{10}\left[\frac{S_c \times C_d \times (D^{0.75} - 1.132)}{215.63 \times J \times \left(D^{0.75} - \frac{18.42}{(E_c/k)^{0.25}}\right)}\right]
        ''')
        
        st.markdown("---")
        
        # ส่งออกรายงาน Word
        st.subheader("📄 ส่งออกรายงาน")
        
        # เตรียมข้อมูลสำหรับรายงาน
        inputs_dict = {
            'w18_design': w18_design,
            'pt': pt,
            'reliability': reliability,
            'so': so,
            'k_eff': k_eff,
            'ls': ls_value,
            'fc_cube': fc_cube,
            'sc': sc,
            'j': j_value,
            'cd': cd
        }
        
        calculated_dict = {
            'fc_cylinder': fc_cylinder,
            'ec': ec,
            'zr': zr,
            'delta_psi': delta_psi
        }
        
        # ข้อมูลดินคันทาง
        subgrade_info = {
            'cbr': cbr_value,
            'mr_psi': mr_subgrade_psi,
            'mr_mpa': mr_subgrade_mpa
        }
        
        # สร้างปุ่มดาวน์โหลด
        if st.button("📥 สร้างรายงาน Word", type="primary"):
            with st.spinner("กำลังสร้างรายงาน..."):
                try:
                    # สร้างรูปโครงสร้างชั้นทางสำหรับรายงาน Word
                    fig_for_report = create_pavement_structure_figure(layers_data, concrete_thickness_cm=d_cm_selected)
                    
                    buffer = create_word_report(
                        pavement_type=pavement_type,
                        inputs=inputs_dict,
                        calculated_values=calculated_dict,
                        comparison_results=comparison_results,
                        selected_d_cm=d_cm_selected,
                        main_result=(passed_selected, ratio_selected),
                        layers_data=layers_data,
                        project_name=project_name,
                        structure_figure=fig_for_report,
                        subgrade_info=subgrade_info,
                        e_equivalent_psi=e_equivalent_psi
                    )
                    
                    # ปิด figure หลังใช้งาน
                    if fig_for_report:
                        plt.close(fig_for_report)
                    
                    if buffer:
                        st.download_button(
                            label="⬇️ ดาวน์โหลดรายงาน (.docx)",
                            data=buffer,
                            file_name=f"AASHTO_Rigid_Pavement_Design_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        st.success("สร้างรายงานสำเร็จ!")
                except Exception as e:
                    st.error(f"เกิดข้อผิดพลาด: {str(e)}")
                    st.info("กรุณาติดตั้ง python-docx: `pip install python-docx`")
    
    # ============================================================
    # ส่วนอ้างอิง
    # ============================================================
    
    st.markdown("---")
    st.header("📚 อ้างอิง")
    
    st.markdown("""
    **เอกสารอ้างอิง:**
    1. AASHTO (1993). *AASHTO Guide for Design of Pavement Structures*. American Association of State Highway and Transportation Officials.
    2. Huang, Y.H. (2004). *Pavement Analysis and Design*. Pearson Prentice Hall.
    3. ACI 318-19 (2019). *Building Code Requirements for Structural Concrete*. American Concrete Institute.
    
    **หมายเหตุ:**
    - โปรแกรมนี้พัฒนาเพื่อใช้ในการเรียนการสอน
    - การออกแบบจริงควรพิจารณาปัจจัยอื่นๆ ร่วมด้วย เช่น สภาพแวดล้อม การก่อสร้าง และการบำรุงรักษา
    """)
    
    # Footer
    st.markdown("---")
    st.caption("พัฒนาโดย: ภาควิชาครุศาสตร์โยธา มจพ. | AASHTO 1993 Rigid Pavement Design Tool")


if __name__ == "__main__":
    main()
