"""
โปรแกรมออกแบบและตรวจสอบความหนาถนนคอนกรีต (Rigid Pavement)
ตามวิธี AASHTO 1993
รองรับทั้ง JPCP (Jointed Plain Concrete Pavement) และ CRCP (Continuously Reinforced Concrete Pavement)

รวมโปรแกรม:
1. การหาค่า k-value และปรับแก้ Loss of Support (LS) จาก Nomograph
2. การคำนวณความหนาถนนคอนกรีตตาม AASHTO 1993
การปรับปรุง
1. ย้ายรูปแบบการตรวจสอบความหนาผิวทางไปcolumn 2 ฝั่งขวา
พัฒนาสำหรับใช้ในการเรียนการสอน
ภาควิชาครุศาสตร์โยธา มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ
"""

import streamlit as st
import math
from io import BytesIO
from datetime import datetime
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from PIL import Image, ImageDraw
import io
import json

# ============================================================
# ค่าคงที่และตารางอ้างอิง AASHTO 1993
# ============================================================

ZR_TABLE = {
    50: -0.000, 60: -0.253, 70: -0.524, 75: -0.674, 80: -0.841, 85: -1.037,
    90: -1.282, 91: -1.340, 92: -1.405, 93: -1.476, 94: -1.555, 95: -1.645,
    96: -1.751, 97: -1.881, 98: -2.054, 99: -2.327
}

J_VALUES = {"JRCP": 2.8, "JPCP": 2.8, "JRCP/JPCP": 2.8, "CRCP": 2.5}
CD_DEFAULT = 1.0

MATERIAL_MODULUS = {
    "รองผิวทางคอนกรีตด้วย AC": 2500, "รองผิวทางคอนกรีตด้วย PMA(AC)": 3700,
    "พื้นทางซีเมนต์ CTB": 1200, "หินคลุกผสมซีเมนต์ UCS 24.5 ksc": 850,
    "หินคลุก CBR 80%": 350, "ดินซีเมนต์ UCS 17.5 ksc": 350,
    "วัสดุหมุนเวียน (Recycling)": 850, "รองพื้นทางวัสดุมวลรวม CBR 25%": 150,
    "วัสดุคัดเลือก ก": 76, "ดินถมคันทาง / ดินเดิม": 100, "กำหนดเอง...": 100,
}

LS_PRESETS = {
    0.0: (138, 715, 753, 84), 0.5: (129, 728, 908, 0), 1.0: (150, 718, 903, 84),
    1.5: (153, 721, 928, 138), 2.0: (164, 718, 929, 220), 3.0: (212, 719, 929, 328)
}

# ============================================================
# ฟังก์ชันการคำนวณ
# ============================================================

def convert_cube_to_cylinder(fc_cube_ksc):
    return 0.8 * fc_cube_ksc

def calculate_concrete_modulus(fc_cylinder_ksc):
    fc_psi = fc_cylinder_ksc * 14.223
    return 57000 * math.sqrt(fc_psi)

def estimate_modulus_of_rupture(fc_cylinder_ksc):
    fc_psi = fc_cylinder_ksc * 14.223
    return 10.0 * math.sqrt(fc_psi)

def get_zr_value(reliability):
    return ZR_TABLE.get(int(reliability), -1.282)

def calculate_aashto_rigid_w18(d_inch, delta_psi, pt, zr, so, sc_psi, cd, j, ec_psi, k_pci):
    term1 = zr * so
    term2 = 7.35 * math.log10(d_inch + 1) - 0.06
    numerator3 = math.log10(delta_psi / (4.5 - 1.5))
    denominator3 = 1 + (1.624e7 / ((d_inch + 1) ** 8.46))
    term3 = numerator3 / denominator3
    d_power = d_inch ** 0.75
    numerator4 = sc_psi * cd * (d_power - 1.132)
    ec_k_ratio = ec_psi / k_pci
    denominator4 = 215.63 * j * (d_power - 18.42 / (ec_k_ratio ** 0.25))
    if numerator4 <= 0 or denominator4 <= 0:
        return (float('-inf'), 0)
    inner_term = numerator4 / denominator4
    if inner_term <= 0:
        return (float('-inf'), 0)
    term4 = (4.22 - 0.32 * pt) * math.log10(inner_term)
    log10_w18 = term1 + term2 + term3 + term4
    w18 = 10 ** log10_w18
    return (log10_w18, w18)

def check_design(w18_required, w18_capacity):
    ratio = w18_capacity / w18_required if w18_required > 0 else float('inf')
    return (w18_capacity >= w18_required, ratio)

def draw_arrow_fixed(draw, start, end, color, width=4, arrow_size=15):
    draw.line([start, end], fill=color, width=width)
    dx = end[0] - start[0]
    dy = end[1] - start[1]
    length = math.sqrt(dx*dx + dy*dy)
    if length > 0:
        dx /= length
        dy /= length
        px, py = -dy, dx
        x3, y3 = end[0], end[1]
        base_x = end[0] - arrow_size * dx
        base_y = end[1] - arrow_size * dy
        x4 = base_x + arrow_size * 0.5 * px
        y4 = base_y + arrow_size * 0.5 * py
        x5 = base_x - arrow_size * 0.5 * px
        y5 = base_y - arrow_size * 0.5 * py
        draw.polygon([(x3, y3), (x4, y4), (x5, y5)], fill=color)

# ============================================================
# ฟังก์ชันสร้างรูปโครงสร้างชั้นทาง
# ============================================================

def create_pavement_structure_figure(layers_data, concrete_thickness_cm=None):
    THAI_TO_ENG = {
        "รองผิวทางคอนกรีตด้วย AC": "AC Interlayer", "รองผิวทางคอนกรีตด้วย PMA(AC)": "PMA Interlayer",
        "พื้นทางซีเมนต์ CTB": "Cement Treated Base", "หินคลุกผสมซีเมนต์ UCS 24.5 ksc": "Mod.Crushed Rock ",
        "หินคลุก CBR 80%": "Crushed Rock Base", "ดินซีเมนต์ UCS 17.5 ksc": "Soil Cement",
        "วัสดุหมุนเวียน (Recycling)": "Recycled Material", "รองพื้นทางวัสดุมวลรวม CBR 25%": "Aggregate Subbase",
        "วัสดุคัดเลือก ก": "Selected Material", "ดินถมคันทาง / ดินเดิม": "Subgrade",
        "กำหนดเอง...": "Custom Material", "แผ่นคอนกรีต": "Concrete Slab", "Concrete Slab": "Concrete Slab",
    }
    LAYER_COLORS = {
        "รองผิวทางคอนกรีตด้วย AC": "#2C3E50", "รองผิวทางคอนกรีตด้วย PMA(AC)": "#1A252F",
        "พื้นทางซีเมนต์ CTB": "#7F8C8D", "หินคลุกผสมซีเมนต์ UCS 24.5 ksc": "#95A5A6",
        "หินคลุก CBR 80%": "#BDC3C7", "ดินซีเมนต์ UCS 17.5 ksc": "#AAB7B8",
        "วัสดุหมุนเวียน (Recycling)": "#85929E", "รองพื้นทางวัสดุมวลรวม CBR 25%": "#FFCC99",
        "วัสดุคัดเลือก ก": "#E8DAEF", "ดินถมคันทาง / ดินเดิม": "#F5CBA7",
        "กำหนดเอง...": "#FADBD8", "Concrete Slab": "#808080",
    }
    
    valid_layers = [l for l in layers_data if l.get("thickness_cm", 0) > 0]
    all_layers = []
    if concrete_thickness_cm and concrete_thickness_cm > 0:
        all_layers.append({"name": "Concrete Slab", "thickness_cm": concrete_thickness_cm, "E_MPa": None})
    all_layers.extend(valid_layers)
    if not all_layers:
        return None
    
    total_thickness = sum(l.get("thickness_cm", 0) for l in all_layers)
    min_display_height = 8
    fig, ax = plt.subplots(figsize=(12, 8))
    width, x_center = 3, 6
    x_start = x_center - width / 2
    display_heights = [max(l.get("thickness_cm", 0), min_display_height) for l in all_layers]
    total_display = sum(display_heights)
    y_current = total_display
    
    for i, layer in enumerate(all_layers):
        thickness = layer.get("thickness_cm", 0)
        name = layer.get("name", f"Layer {i+1}")
        e_mpa = layer.get("E_MPa", None)
        display_h = display_heights[i]
        if thickness <= 0:
            continue
        color = LAYER_COLORS.get(name, "#CCCCCC")
        hatch_pattern = '///' if name == "วัสดุหมุนเวียน (Recycling)" else None
        y_bottom = y_current - display_h
        rect = patches.Rectangle((x_start, y_bottom), width, display_h, linewidth=2,
                                  edgecolor='black', facecolor=color, hatch=hatch_pattern)
        ax.add_patch(rect)
        y_center_pos = y_bottom + display_h / 2
        display_name = THAI_TO_ENG.get(name, name)
        is_dark = name in ["รองผิวทางคอนกรีตด้วย AC", "รองผิวทางคอนกรีตด้วย PMA(AC)", "Concrete Slab",
                          "พื้นทางซีเมนต์ CTB", "หินคลุกผสมซีเมนต์ UCS 24.5 ksc", "วัสดุหมุนเวียน (Recycling)"]
        text_color = 'white' if is_dark else 'black'
        ax.text(x_center, y_center_pos, f"{thickness} cm", ha='center', va='center', fontsize=16, fontweight='bold', color=text_color)
        ax.text(x_start - 0.5, y_center_pos, display_name, ha='right', va='center', fontsize=14, fontweight='bold', color='black')
        if e_mpa:
            ax.text(x_start + width + 0.5, y_center_pos, f"E = {e_mpa:,} MPa", ha='left', va='center', fontsize=12, color='#0066CC')
        y_current = y_bottom
    
    ax.annotate('', xy=(x_start + width + 3.5, total_display), xytext=(x_start + width + 3.5, 0),
                arrowprops=dict(arrowstyle='<->', color='red', lw=2))
    ax.text(x_start + width + 4, total_display / 2, f"Total\n{total_thickness} cm", ha='left', va='center', fontsize=14, color='red', fontweight='bold')
    margin = 10
    ax.set_xlim(0, 14)
    ax.set_ylim(-margin, total_display + margin)
    ax.axis('off')
    ax.set_title('Pavement Structure', fontsize=20, fontweight='bold', pad=20)
    ax.text(x_center, -margin + 4, f"Total Pavement Thickness: {total_thickness} cm", ha='center', va='center', fontsize=15, fontweight='bold',
            bbox=dict(boxstyle='round', facecolor='lightyellow', alpha=0.9, edgecolor='orange'))
    plt.tight_layout()
    return fig

def save_figure_to_bytes(fig):
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white', edgecolor='none')
    buf.seek(0)
    return buf

# ============================================================
# ฟังก์ชัน Save/Load JSON
# ============================================================

def save_project_to_json(project_data):
    json_str = json.dumps(project_data, ensure_ascii=False, indent=2)
    return json_str.encode('utf-8')

def load_project_from_json(uploaded_file):
    try:
        content = uploaded_file.read()
        return json.loads(content.decode('utf-8'))
    except Exception as e:
        st.error(f"เกิดข้อผิดพลาดในการอ่านไฟล์: {str(e)}")
        return None

def collect_design_data(project_name, pavement_type, num_layers, layers_data, w18_design, pt, reliability, so,
                        k_eff, ls_value, fc_cube, sc, j_value, cd, d_cm_selected, cbr_value,
                        mr_val=0, esb_val=0, dsb_val=0, k_inf_val=0, ls_select=0, k_corrected=0):
    return {
        "version": "1.0",
        "save_date": datetime.now().isoformat(),
        "project_info": {"project_name": project_name, "pavement_type": pavement_type},
        "layers": {"num_layers": num_layers, "layers_data": layers_data},
        "design_parameters": {
            "w18_design": w18_design, "pt": pt, "reliability": reliability, "so": so,
            "k_eff": k_eff, "ls_value": ls_value, "fc_cube": fc_cube, "sc": sc,
            "j_value": j_value, "cd": cd, "d_cm_selected": d_cm_selected
        },
        "subgrade": {"cbr_value": cbr_value},
        "nomograph": {"mr_val": mr_val, "esb_val": esb_val, "dsb_val": dsb_val,
                      "k_inf_val": k_inf_val, "ls_select": ls_select, "k_corrected": k_corrected}
    }

# ============================================================
# ฟังก์ชันสร้างรายงาน Word
# ============================================================

def create_word_report(pavement_type, inputs, calculated_values, comparison_results, selected_d_cm,
                       main_result, layers_data=None, project_name="", structure_figure=None,
                       subgrade_info=None, e_equivalent_psi=0):
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        st.error("กรุณาติดตั้ง python-docx: pip install python-docx")
        return None
    
    selected_d_inch = round(selected_d_cm / 2.54)
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'TH Sarabun New'
    style.font.size = Pt(14)
    
    title = doc.add_heading('รายการคำนวณออกแบบความหนาถนนคอนกรีต', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('ตามวิธี AASHTO 1993').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('1. ข้อมูลทั่วไป', level=1)
    if project_name:
        doc.add_paragraph(f'ชื่อโครงการ: {project_name}')
    doc.add_paragraph(f'ประเภทถนน: {pavement_type}')
    doc.add_paragraph(f'วันที่คำนวณ: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    
    doc.add_heading('2. ชั้นโครงสร้างทาง', level=1)
    table_layers = doc.add_table(rows=1, cols=4)
    table_layers.style = 'Table Grid'
    hdr = table_layers.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'ลำดับ', 'ชนิดวัสดุ', 'ความหนา (ซม.)', 'Modulus E (MPa)'
    row = table_layers.add_row().cells
    row[0].text, row[1].text, row[2].text, row[3].text = '1', f'ผิวทางคอนกรีต {pavement_type}', f'{selected_d_cm}', '-'
    layer_count = 1
    if layers_data:
        for i, layer in enumerate(layers_data):
            layer_count += 1
            row = table_layers.add_row().cells
            row[0].text = str(layer_count)
            row[1].text = layer.get('name', f'Layer {i+1}')
            row[2].text = f"{layer.get('thickness_cm', 0)}"
            row[3].text = f"{layer.get('E_MPa', 0):,}"
    if subgrade_info:
        layer_count += 1
        row = table_layers.add_row().cells
        row[0].text = str(layer_count)
        row[1].text = 'ดินคันทาง'
        row[2].text = f"CBR {subgrade_info.get('cbr', 0)} %"
        row[3].text = f"{subgrade_info.get('mr_mpa', 0):.0f} ({subgrade_info.get('mr_psi', 0):,.0f} psi)"
    
    if structure_figure:
        doc.add_paragraph('รูปตัดโครงสร้างชั้นทาง:')
        img_buf = BytesIO()
        structure_figure.savefig(img_buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
        img_buf.seek(0)
        doc.add_picture(img_buf, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('3. ข้อมูลนำเข้า', level=1)
    table1 = doc.add_table(rows=1, cols=4)
    table1.style = 'Table Grid'
    hdr = table1.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'พารามิเตอร์', 'สัญลักษณ์', 'ค่า', 'หน่วย'
    input_data = [
        ('ESAL ออกแบบ', 'W₁₈', f"{inputs['w18_design']:,.0f}", 'ESALs'),
        ('Terminal Serviceability', 'Pt', f"{inputs['pt']:.1f}", '-'),
        ('Reliability', 'R', f"{inputs['reliability']:.0f}", '%'),
        ('Standard Deviation', 'So', f"{inputs['so']:.2f}", '-'),
        ('Modulus of Subgrade Reaction', 'k_eff', f"{inputs['k_eff']:,.0f}", 'pci'),
        ('Loss of Support', 'LS', f"{inputs.get('ls', 1.0):.1f}", '-'),
        ('กำลังคอนกรีต', "f'c", f"{inputs['fc_cube']:.0f} Cube", 'ksc'),
        ('Modulus of Rupture', 'Sc', f"{inputs['sc']:.0f}", 'psi'),
        ('Load Transfer Coefficient', 'J', f"{inputs['j']:.1f}", '-'),
        ('Drainage Coefficient', 'Cd', f"{inputs['cd']:.1f}", '-'),
    ]
    for param, symbol, value, unit in input_data:
        row = table1.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text = param, symbol, value, unit
    
    doc.add_heading('4. ค่าที่คำนวณได้', level=1)
    table2 = doc.add_table(rows=1, cols=4)
    table2.style = 'Table Grid'
    hdr = table2.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'พารามิเตอร์', 'สัญลักษณ์', 'ค่า', 'หน่วย'
    calc_data = [
        ('Modulus of Elasticity', 'Ec', f"{calculated_values['ec']:,.0f}", 'psi'),
        ('Standard Normal Deviate', 'ZR', f"{calculated_values['zr']:.3f}", '-'),
        ('การสูญเสีย Serviceability', 'ΔPSI', f"{calculated_values['delta_psi']:.1f}", '-'),
    ]
    for param, symbol, value, unit in calc_data:
        row = table2.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text = param, symbol, value, unit
    
    doc.add_heading('5. ผลการเปรียบเทียบความหนา', level=1)
    table3 = doc.add_table(rows=1, cols=6)
    table3.style = 'Table Grid'
    hdr = table3.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = 'D (ซม.)', 'D (นิ้ว)', 'log₁₀(W₁₈)'
    hdr[3].text, hdr[4].text, hdr[5].text = 'W₁₈ รองรับได้', 'อัตราส่วน', 'ผล'
    for r in comparison_results:
        row = table3.add_row().cells
        row[0].text = f"{r['d_cm']:.0f}"
        row[1].text = f"{r['d_inch']:.0f}"
        row[2].text = f"{r['log_w18']:.4f}"
        row[3].text = f"{r['w18']:,.0f}"
        row[4].text = f"{r['ratio']:.2f}"
        row[5].text = "ผ่าน ✓" if r['passed'] else "ไม่ผ่าน ✗"
    
    doc.add_heading('6. สรุปผล', level=1)
    passed, ratio = main_result
    w18_cap = None
    for r in comparison_results:
        if r['d_cm'] == selected_d_cm:
            w18_cap = r['w18']
            break
    e_eq_mpa = e_equivalent_psi / 145.038 if e_equivalent_psi > 0 else 0
    doc.add_paragraph(f"ความหนาที่เลือก: {selected_d_cm:.0f} ซม. ({selected_d_inch:.0f} นิ้ว)")
    doc.add_paragraph(f"ESAL ที่ต้องการ: {inputs['w18_design']:,.0f} ESALs")
    if w18_cap:
        doc.add_paragraph(f"ESAL ที่รองรับได้: {w18_cap:,.0f} ESALs")
    doc.add_paragraph(f"อัตราส่วน: {ratio:.2f}")
    doc.add_paragraph(f"ผลการตรวจสอบ: {'ผ่านเกณฑ์ ✓' if passed else 'ไม่ผ่านเกณฑ์ ✗'}")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_word_report_nomograph(params, img1_bytes, img2_bytes=None):
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        return None, "ไม่พบ library python-docx"
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'TH SarabunPSK'
    style.font.size = Pt(14)
    
    title = doc.add_heading('รายการคำนวณ Corrected Modulus of Subgrade Reaction', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'วันที่: {datetime.now().strftime("%d/%m/%Y %H:%M")}').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_heading('ส่วนที่ 1: การหาค่า Composite Modulus (k∞)', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = 'พารามิเตอร์', 'ค่า', 'หน่วย'
    for h in hdr:
        h.paragraphs[0].runs[0].bold = True
    data1 = [
        ('Roadbed Soil Resilient Modulus (MR)', f"{params.get('MR', 0):,.0f}", 'psi'),
        ('Subbase Elastic Modulus (ESB)', f"{params.get('ESB', 0):,.0f}", 'psi'),
        ('Subbase Thickness (DSB)', f"{params.get('DSB', 0):.1f}", 'inches'),
        ('Composite Modulus (k∞)', f"{params.get('k_inf', 0):,.0f}", 'pci'),
    ]
    for p, v, u in data1:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = p, v, u
    if img1_bytes:
        doc.add_paragraph()
        doc.add_picture(io.BytesIO(img1_bytes), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    doc.add_heading('ส่วนที่ 2: การปรับแก้ค่า Loss of Support (LS)', level=1)
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    hdr2[0].text, hdr2[1].text, hdr2[2].text = 'พารามิเตอร์', 'ค่า', 'หน่วย'
    for h in hdr2:
        h.paragraphs[0].runs[0].bold = True
    data2 = [
        ('Effective Modulus (k) - จากส่วนที่ 1', f"{params.get('k_inf', 0):,.0f}", 'pci'),
        ('Loss of Support Factor (LS)', f"{params.get('LS_factor', 0):.1f}", '-'),
        ('Corrected Modulus (k)', f"{params.get('k_corrected', 0):,.0f}", 'pci'),
    ]
    for p, v, u in data2:
        row = table2.add_row().cells
        row[0].text, row[1].text, row[2].text = p, v, u
    if img2_bytes:
        doc.add_paragraph()
        doc.add_picture(io.BytesIO(img2_bytes), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph("Reference: AASHTO Guide for Design of Pavement Structures 1993").style = 'List Bullet'
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, None

# ============================================================
# Main Application
# ============================================================

def main():
    st.set_page_config(page_title="AASHTO 1993 Rigid Pavement Design", page_icon="🛣️", layout="wide")
    st.title("🛣️ Rigid Pavement Design Calculator (AASHTO 1993)")
    st.markdown("**โปรแกรมออกแบบความหนาถนนคอนกรีต และหาค่า k-value พร้อมปรับแก้ Loss of Support**")
    
    # Initialize Session State
    for key, val in [('k_inf_result', 500), ('img1_bytes', None), ('img2_bytes', None), ('last_uploaded_file', None)]:
        if key not in st.session_state:
            st.session_state[key] = val
    
    # Sidebar: JSON Save/Load
    with st.sidebar:
        st.header("📁 จัดการโปรเจกต์")
        st.subheader("📂 โหลดไฟล์โปรเจกต์")
        uploaded_json = st.file_uploader("อัปโหลดไฟล์ .json", type=['json'], key='json_uploader')
        
        if uploaded_json is not None:
            try:
                # ตรวจสอบว่าเป็นไฟล์ใหม่หรือไม่
                file_id = f"{uploaded_json.name}_{uploaded_json.size}"
                if st.session_state.get('last_uploaded_file') != file_id:
                    st.session_state['last_uploaded_file'] = file_id
                    
                    # โหลดข้อมูลจาก JSON
                    loaded = load_project_from_json(uploaded_json)
                    if loaded:
                        # อัพเดท session_state สำหรับทุก input field
                        
                        # Project Info
                        st.session_state['calc_project_name'] = loaded.get('project_info', {}).get('project_name', '')
                        st.session_state['calc_pave_type'] = loaded.get('project_info', {}).get('pavement_type', 'JPCP')
                        
                        # Layers
                        st.session_state['calc_num_layers'] = loaded.get('layers', {}).get('num_layers', 5)
                        layers_data = loaded.get('layers', {}).get('layers_data', [])
                        for i, layer in enumerate(layers_data):
                            st.session_state[f'calc_layer_name_{i}'] = layer.get('name', '')
                            st.session_state[f'calc_layer_thick_{i}'] = layer.get('thickness_cm', 0)
                            layer_name = layer.get('name', '')
                            st.session_state[f'calc_layer_E_{i}_{layer_name}'] = layer.get('E_MPa', 100)
                        
                        # Design Parameters
                        dp = loaded.get('design_parameters', {})
                        st.session_state['calc_w18'] = dp.get('w18_design', 500000)
                        st.session_state['calc_pt'] = dp.get('pt', 2.0)
                        st.session_state['calc_reliability'] = dp.get('reliability', 90)
                        st.session_state['calc_so'] = dp.get('so', 0.35)
                        st.session_state['calc_k_eff'] = dp.get('k_eff', 200)
                        st.session_state['calc_ls'] = dp.get('ls_value', 1.0)
                        st.session_state['calc_fc'] = dp.get('fc_cube', 350)
                        st.session_state['calc_sc'] = dp.get('sc', 600)
                        st.session_state['calc_j'] = dp.get('j_value', 2.8)
                        st.session_state['calc_cd'] = dp.get('cd', 1.0)
                        st.session_state['calc_d'] = dp.get('d_cm_selected', 30)
                        
                        # Subgrade
                        st.session_state['calc_cbr'] = loaded.get('subgrade', {}).get('cbr_value', 4.0)
                        
                        # Nomograph
                        nomo = loaded.get('nomograph', {})
                        st.session_state['nomo_mr'] = nomo.get('mr_val', 7000)
                        st.session_state['nomo_esb'] = nomo.get('esb_val', 50000)
                        st.session_state['nomo_dsb'] = nomo.get('dsb_val', 6.0)
                        st.session_state['nomo_k_inf'] = nomo.get('k_inf_val', 400)
                        st.session_state['k_inf_result'] = nomo.get('k_inf_val', 400)
                        st.session_state['ls_select_box'] = nomo.get('ls_select', 1.0)
                        st.session_state['k_corr_input'] = nomo.get('k_corrected', 300)
                        
                        st.success("✅ โหลดข้อมูลสำเร็จ!")
                        st.rerun()
            except Exception as e:
                st.error(f"❌ ไม่สามารถอ่านไฟล์ได้: {e}")
        
        # แสดงสถานะโปรเจกต์ที่โหลด
        if st.session_state.get('calc_project_name'):
            st.info(f"📌 โปรเจกต์: {st.session_state.get('calc_project_name', 'ไม่ระบุ')}")
            if st.button("🗑️ ล้างข้อมูลที่โหลด"):
                # ล้าง session_state ทั้งหมด
                keys_to_clear = [key for key in st.session_state.keys() if key.startswith(('calc_', 'nomo_', 'ls_select', 'k_corr', 'k_inf'))]
                for key in keys_to_clear:
                    del st.session_state[key]
                st.session_state['last_uploaded_file'] = None
                st.rerun()
        st.markdown("---")
    
    
    # Define Tabs
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "🔢 AASHTO Calculator", "📊 Nomograph: Composite k∞", "📉 Nomograph: Loss of Support",
        "💾 บันทึกโปรเจกต์", "📋 คู่มือการใช้งาน"
    ])
    
    # =========================================================
    # TAB 1: AASHTO Calculator
    # =========================================================
    with tab1:
        st.header("1️⃣ การออกแบบความหนาถนนคอนกรีต (AASHTO 1993)")
        col1, col2 = st.columns([1, 1])
        
        with col1:
            st.subheader("📥 ข้อมูลนำเข้า (Input)")
            project_name = st.text_input("🏗️ ชื่อโครงการ", value=st.session_state.get('calc_project_name', ''), key="calc_project_name")
            st.markdown("---")
            
            pave_options = list(J_VALUES.keys())
            current_pave_type = st.session_state.get('calc_pave_type', 'JPCP')
            default_pave_idx = pave_options.index(current_pave_type) if current_pave_type in pave_options else 1
            pavement_type = st.selectbox("ประเภทผิวทางคอนกรีต", pave_options, index=default_pave_idx, key="calc_pave_type")
            st.markdown("---")
            
            st.subheader("🔶 ชั้นโครงสร้างทาง")
            material_options = list(MATERIAL_MODULUS.keys())
            num_layers = st.slider("จำนวนชั้นวัสดุ", 1, 6, st.session_state.get('calc_num_layers', 5), key="calc_num_layers")
            
            default_layers = [
                {"name": "รองผิวทางคอนกรีตด้วย AC", "thickness_cm": 5},
                {"name": "พื้นทางซีเมนต์ CTB", "thickness_cm": 20},
                {"name": "หินคลุก CBR 80%", "thickness_cm": 15},
                {"name": "รองพื้นทางวัสดุมวลรวม CBR 25%", "thickness_cm": 25},
                {"name": "วัสดุคัดเลือก ก", "thickness_cm": 30},
                {"name": "ดินถมคันทาง / ดินเดิม", "thickness_cm": 0},
            ]
            
            layers_data = []
            for i in range(num_layers):
                st.markdown(f"**ชั้นที่ {i+1}**")
                col_a, col_b, col_c = st.columns([2, 1, 1])
                
                # ใช้ค่าจาก session_state หรือ default
                def_name = st.session_state.get(f'calc_layer_name_{i}', default_layers[i]["name"] if i < len(default_layers) else "กำหนดเอง...")
                def_thick = st.session_state.get(f'calc_layer_thick_{i}', default_layers[i]["thickness_cm"] if i < len(default_layers) else 20)
                def_idx = material_options.index(def_name) if def_name in material_options else len(material_options) - 1
                
                with col_a:
                    layer_name = st.selectbox("เลือกวัสดุ", material_options, index=def_idx, key=f"calc_layer_name_{i}")
                with col_b:
                    layer_thickness = st.number_input("ความหนา (ซม.)", 0, 100, def_thick, key=f"calc_layer_thick_{i}")
                rec_mod = MATERIAL_MODULUS.get(layer_name, 100)
                def_E = st.session_state.get(f'calc_layer_E_{i}_{layer_name}', rec_mod)
                with col_c:
                    layer_modulus = st.number_input("E (MPa)", 10, 10000, def_E, key=f"calc_layer_E_{i}_{layer_name}")
                layers_data.append({"name": layer_name, "thickness_cm": layer_thickness, "E_MPa": layer_modulus})
            
            total_layer_cm = sum(l['thickness_cm'] for l in layers_data)
            st.markdown(f"**รวมความหนา {total_layer_cm:.0f} ซม. ({round(total_layer_cm/2.54)} นิ้ว)**")
            
            # คำนวณ E_equivalent
            valid_layers = [l for l in layers_data if l['thickness_cm'] > 0 and l['E_MPa'] > 0]
            if valid_layers:
                sum_h_e_cbrt = sum(l['thickness_cm'] * (l['E_MPa'] ** (1/3)) for l in valid_layers)
                total_valid_cm = sum(l['thickness_cm'] for l in valid_layers)
                e_eq_mpa = (sum_h_e_cbrt / total_valid_cm) ** 3 if total_valid_cm > 0 else 0
                e_eq_psi = e_eq_mpa * 145.038
                st.info(f"โมดูลัสเทียบเท่า (E_equivalent) = **{e_eq_psi:,.0f} psi** ({e_eq_mpa:.1f} MPa)")
            st.markdown("---")
            
            st.subheader("1️⃣ ปริมาณจราจร 🚛 ")
            with st.expander("📊 ตัวช่วยประมาณ ESAL ตามประเภทถนน", expanded=False):
                st.markdown("""
                | ประเภทถนน | ESAL (ล้าน) |
                |-----------|-------------|
                | ทางหลวงพิเศษระหว่างเมือง | 50-200 |
                | ทางหลวงแผ่นดินสายหลัก | 20-80 |
                | ทางหลวงแผ่นดินสายรอง | 5-30 |
                | ถนนในเมือง | 1-10 |
                """)
            w18_design = st.number_input("ESAL ที่ต้องการรองรับ (W₁₈)", 10000, 500000000, st.session_state.get('calc_w18', 500000), 100000, key="calc_w18")
            esal_million = w18_design / 1_000_000
            st.info(f"**{esal_million:.2f} ล้าน ESALs**")
            st.markdown("---")
            
            st.subheader("2️⃣ Serviceability📉")
            pt = st.slider("Terminal Serviceability (Pt)", 1.5, 3.0, st.session_state.get('calc_pt', 2.0), 0.1, key="calc_pt")
            delta_psi = 4.5 - pt
            st.info(f"ΔPSI = 4.5 - {pt:.1f} = **{delta_psi:.1f}**")
            st.markdown("---")
            
            st.subheader("3️⃣ ความเชื่อมั่น📈")
            reliability = st.select_slider("Reliability (R)", [80, 85, 90, 95], st.session_state.get('calc_reliability', 90), key="calc_reliability")
            zr = get_zr_value(reliability)
            st.info(f"ZR = **{zr:.3f}**")
            so = st.number_input("Standard Deviation (So)", 0.30, 0.45, st.session_state.get('calc_so', 0.35), 0.01, "%.2f", key="calc_so")
            st.markdown("---")
            
            st.subheader("4️⃣ คุณสมบัติดินคันทาง")
            cbr_value = st.number_input("ค่า CBR (%)", 1.0, 100.0, st.session_state.get('calc_cbr', 4.0), 0.5, key="calc_cbr")
            mr_subgrade_psi = 1500 * cbr_value if cbr_value < 10 else 1000 + 555 * cbr_value
            mr_subgrade_mpa = mr_subgrade_psi / 145.038
            st.info(f"M_R = {mr_subgrade_psi:,.0f} psi ({mr_subgrade_mpa:.0f} MPa)")
            
            k_eff = st.number_input("Effective k (pci)", 50, 1000, st.session_state.get('calc_k_eff', 200), 25, key="calc_k_eff")
        
            with st.expander("📊 ตารางค่า Loss of Support แนะนำ (AASHTO 1993)"):
                st.markdown("""
                | ประเภทวัสดุ | Loss of Support (LS) |
                |------------|---------------------|
                | Cement Treated Granular Base | 0.0 - 1.0 |
                | Cement Aggregate Mixtures | 0.0 - 1.0 |
                | Asphalt Treated Base | 0.0 - 1.0 |
                | Bituminous Stabilized Mixtures | 0.0 - 1.0 |
                | Lime Stabilized | 1.0 - 3.0 |
                | Unbound Granular Materials | 1.0 - 3.0 |
                | Fine Grained or Natural Subgrade | 2.0 - 3.0 |
            
                **หมายเหตุ:** ค่า LS ใช้ปรับลดค่า k_eff เพื่อคำนึงถึงการสูญเสียการรองรับจากการกัดเซาะ
                """)
            ls_value = st.number_input("Loss of Support (LS)", 0.0, 3.0, st.session_state.get('calc_ls', 1.0), 0.5, "%.1f", key="calc_ls")
            st.markdown("---")
            
            st.subheader("5️⃣ คุณสมบัติคอนกรีต")
            fc_cube = st.number_input("กำลังอัด Cube (ksc)", 200, 600, st.session_state.get('calc_fc', 350), 10, key="calc_fc")
            fc_cylinder = convert_cube_to_cylinder(fc_cube)
            ec = calculate_concrete_modulus(fc_cylinder)
            st.info(f"f'c (Cyl) = **{fc_cylinder:.0f} ksc** | Ec = **{ec:,.0f} psi**")
            sc_auto = estimate_modulus_of_rupture(fc_cylinder)
            sc = st.number_input("Modulus of Rupture (Sc) psi", 400, 1000, st.session_state.get('calc_sc', int(sc_auto)), 10, key="calc_sc")
            st.markdown("---")
            
            st.subheader("6️⃣ Load Transfer🔗 และ Drainage💧")
            st.caption(f"ค่าแนะนำสำหรับ {pavement_type}: **J = {J_VALUES[pavement_type]}**")
            with st.expander("📊 ตารางค่า Load Transfer Coefficient (J)", expanded=False):
                st.markdown("""
                | ประเภทถนน | J (AC Shoulder_Yes) | J (AC Shoulder_No) | J (Tied P.C.C_Yes) | J (Tied P.C.C_No) |
                |-----------|---------------------|--------------------|--------------------|-------------------|
                | 1. JRCP/JPCP | 3.2 | 3.8-4.4 | 2.5-3.1 (Mid 2.8) | 3.6-4.2 |
                | 2. CRCP | 2.9-3.2 | N/A | 2.3-2.9 (Mid 2.5) | N/A |
                
                **หมายเหตุ:** ค่า J ต่ำ = การถ่ายแรงดี = รองรับ ESAL ได้มากขึ้น
                
                ค่า J สามารถปรับได้ตามเงื่อนไข:
                - มี Dowel Bar: ลดลง 0.2-0.3
                - มี Tied Shoulder: ลดลง 0.2
                - ไม่มี Dowel Bar: เพิ่มขึ้น 0.5-1.0
                """)
            j_auto = J_VALUES[pavement_type]
            j_value = st.number_input("Load Transfer (J)", 2.0, 4.5, st.session_state.get('calc_j', j_auto), 0.1, "%.1f", key="calc_j")
            cd = st.number_input("Drainage (Cd)", 0.7, 1.3, st.session_state.get('calc_cd', 1.0), 0.05, "%.2f", key="calc_cd")

            
        with col2:
            st.subheader("7️⃣ 👷 🚧 ความหนาที่ตรวจสอบ")
            st.caption("ความหนาผิวทางคอนกรีต D (ซม.)")
            d_cm_selected = st.slider("", 20, 40, st.session_state.get('calc_d', 30), 1, key="calc_d", label_visibility="collapsed")
            d_inch_selected = round(d_cm_selected / 2.54)
            st.success(f"**ความหนาผิวทางคอนกรีต D = {d_cm_selected} ซม. ≈ {d_inch_selected} นิ้ว**")
            st.markdown("---")
            st.subheader(f"🎯 ผลการตรวจสอบ D = {d_cm_selected} ซม.")
            log_w18_sel, w18_sel = calculate_aashto_rigid_w18(d_inch_selected, delta_psi, pt, zr, so, sc, cd, j_value, ec, k_eff)
            passed_sel, ratio_sel = check_design(w18_design, w18_sel)
            
            col_a, col_b = st.columns(2)
            with col_a:
                st.metric("log₁₀(W₁₈)", f"{log_w18_sel:.4f}")
                st.metric("W₁₈ รองรับได้", f"{w18_sel:,.0f}", f"{w18_sel - w18_design:+,.0f}")
            with col_b:
                st.metric("W₁₈ ที่ต้องการ", f"{w18_design:,.0f}")
                st.metric("อัตราส่วน", f"{ratio_sel:.2f}")
            
            if passed_sel:
                st.success(f"✅ **ผ่านเกณฑ์** อัตราส่วน = {ratio_sel:.2f}")
            else:
                st.error(f"❌ **ไม่ผ่านเกณฑ์** อัตราส่วน = {ratio_sel:.2f}")
            
            st.markdown("---")
            st.subheader("📊 ผลการคำนวณ")
            comparison_results = []
            thicknesses_cm = [20, 22, 25, 28, 30, 32, 35, 38, 40]

            for d_cm in thicknesses_cm:
                d_inch = round(d_cm / 2.54)
                log_w18, w18_capacity = calculate_aashto_rigid_w18(d_inch, delta_psi, pt, zr, so, sc, cd, j_value, ec, k_eff)
                passed, ratio = check_design(w18_design, w18_capacity)
                comparison_results.append({'d_cm': d_cm, 'd_inch': d_inch, 'log_w18': log_w18, 'w18': w18_capacity, 'passed': passed, 'ratio': ratio})
            
            import pandas as pd
            df = pd.DataFrame([{
                'D (ซม.)': r['d_cm'], 'D (นิ้ว)': r['d_inch'], 'log₁₀(W₁₈)': f"{r['log_w18']:.4f}",
                'W₁₈ รองรับได้': f"{r['w18']:,.0f}", 'อัตราส่วน': f"{r['ratio']:.2f}", 'ผล': "✅" if r['passed'] else "❌"
            } for r in comparison_results])
            st.dataframe(df, use_container_width=True, hide_index=True)
            
            st.markdown("---")
           
            fig_structure = create_pavement_structure_figure(layers_data, d_cm_selected)
            if fig_structure:
                st.pyplot(fig_structure)
                img_buf = save_figure_to_bytes(fig_structure)
                st.download_button("📥 ดาวน์โหลดรูปโครงสร้าง", img_buf, f"pavement_structure_{datetime.now().strftime('%Y%m%d_%H%M')}.png", "image/png")
                plt.close(fig_structure)
            
            st.markdown("---")
            if st.button("📥 สร้างรายงาน Word", type="primary"):
                with st.spinner("กำลังสร้างรายงาน..."):
                    inputs_dict = {'w18_design': w18_design, 'pt': pt, 'reliability': reliability, 'so': so,
                                   'k_eff': k_eff, 'ls': ls_value, 'fc_cube': fc_cube, 'sc': sc, 'j': j_value, 'cd': cd}
                    calc_dict = {'fc_cylinder': fc_cylinder, 'ec': ec, 'zr': zr, 'delta_psi': delta_psi}
                    subgrade_info = {'cbr': cbr_value, 'mr_psi': mr_subgrade_psi, 'mr_mpa': mr_subgrade_mpa}
                    fig_report = create_pavement_structure_figure(layers_data, d_cm_selected)
                    
                    total_cm = sum(l['thickness_cm'] for l in layers_data)
                    sum_h_e_cbrt = sum(l['thickness_cm'] * (l['E_MPa'] ** (1/3)) for l in layers_data if l['thickness_cm'] > 0 and l['E_MPa'] > 0)
                    e_eq_mpa = (sum_h_e_cbrt / total_cm) ** 3 if total_cm > 0 else 0
                    e_eq_psi = e_eq_mpa * 145.038
                    
                    buffer = create_word_report(pavement_type, inputs_dict, calc_dict, comparison_results, d_cm_selected,
                                                (passed_sel, ratio_sel), layers_data, project_name, fig_report, subgrade_info, e_eq_psi)
                    if fig_report:
                        plt.close(fig_report)
                    if buffer:
                        st.download_button("⬇️ ดาวน์โหลดรายงาน (.docx)", buffer, f"AASHTO_Design_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                                           "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    
    # =========================================================
    # TAB 2: Composite Modulus (Nomograph)
    # =========================================================
    with tab2:
        st.header("2️⃣ หาค่า Composite Modulus of Subgrade Reaction (k∞)")
        uploaded_file = st.file_uploader("📂 อัปโหลดภาพ Figure 3.3 (Composite k)", type=['png', 'jpg', 'jpeg'], key='uploader_1')
        
        if uploaded_file is not None:
            image = Image.open(uploaded_file).convert("RGB")
            width, height = image.size
            img_draw = image.copy()
            draw = ImageDraw.Draw(img_draw)
            
            col_ctrl, col_img = st.columns([1, 2])
            with col_ctrl:
                st.subheader("⚙️ ปรับเส้นอ่านค่า")
                with st.expander("1. เส้น Turning Line (เขียว)", expanded=True):
                    gx1 = st.slider("X เริ่ม", 0, width, 411, key="gx1")
                    gy1 = st.slider("Y เริ่ม", 0, height, 339, key="gy1")
                    gx2 = st.slider("X จบ", 0, width, 470, key="gx2")
                    gy2 = st.slider("Y จบ", 0, height, 397, key="gy2")
                    draw.line([(gx1, gy1), (gx2, gy2)], fill="green", width=5)
                    slope_green = (gy2 - gy1) / (gx2 - gx1) if (gx2 - gx1) != 0 else 0
                
                with st.expander("2. พารามิเตอร์ (ส้ม/แดง/น้ำเงิน)", expanded=True):
                    start_x = st.slider("ตำแหน่งแกน D_sb (ซ้าย)", 0, width, int(width*0.15), key="s1_sx")
                    stop_y_esb = st.slider("ระดับค่า ESB (บน)", 0, height, int(height*0.10), key="s1_sy_esb")
                    stop_y_mr = st.slider("ระดับค่า MR (ล่าง)", 0, height, int(height*0.55), key="s1_sy_mr")
                    constrained_x = int(gx1 + (stop_y_mr - gy1) / slope_green) if slope_green != 0 else gx1
                
                lw = 4
                draw_arrow_fixed(draw, (start_x, stop_y_esb), (constrained_x, stop_y_esb), "orange", lw)
                draw_arrow_fixed(draw, (start_x, stop_y_esb), (start_x, stop_y_mr), "red", lw)
                draw_arrow_fixed(draw, (start_x, stop_y_mr), (constrained_x, stop_y_mr), "darkblue", lw)
                draw_arrow_fixed(draw, (constrained_x, stop_y_mr), (constrained_x, stop_y_esb), "blue", lw)
                r = 8
                draw.ellipse([(constrained_x-r, stop_y_mr-r), (constrained_x+r, stop_y_mr+r)], fill="black", outline="white")
                
                st.markdown("---")
                st.subheader("📝 บันทึกค่าที่อ่านได้")
                mr_val = st.number_input("MR (psi)", value=st.session_state.get('nomo_mr', 7000), step=500, key="nomo_mr")
                esb_val = st.number_input("ESB (psi)", value=st.session_state.get('nomo_esb', 50000), step=1000, key="nomo_esb")
                dsb_val = st.number_input("DSB (inches)", value=st.session_state.get('nomo_dsb', 6.0), step=0.5, key="nomo_dsb")
                k_inf_val = st.number_input("ค่า k∞ ที่อ่านได้ (pci)", value=st.session_state.get('nomo_k_inf', 400), step=10, key="nomo_k_inf")
                st.session_state.k_inf_result = k_inf_val
                
                buf = io.BytesIO()
                img_draw.save(buf, format='PNG')
                st.session_state.img1_bytes = buf.getvalue()
            
            with col_img:
                st.image(img_draw, caption="Step 1: Nomograph Analysis", use_container_width=True)
    
    # =========================================================
    # TAB 3: Loss of Support (Nomograph)
    # =========================================================
    with tab3:
        st.header("3️⃣ ปรับแก้ Loss of Support (LS)")
        st.info("ใช้กราฟ Figure 3.4 เพื่อปรับค่า k∞ กรณีที่มีการสูญเสียการรองรับ (LS > 0)")
        uploaded_file_2 = st.file_uploader("📂 อัปโหลดภาพ Figure 3.4 (LS Correction)", type=['png', 'jpg', 'jpeg'], key='uploader_2')
        
        if uploaded_file_2 is not None:
            img2 = Image.open(uploaded_file_2).convert("RGB")
            w2, h2 = img2.size
            img2_draw = img2.copy()
            draw2 = ImageDraw.Draw(img2_draw)
            
            col_ctrl2, col_img2 = st.columns([1, 2])
            with col_ctrl2:
                st.subheader("⚙️ กำหนดเส้นกราฟ")
                st.write("#### 1. เลือกค่า LS (เส้นแดง)")
                ls_options = [0.0, 0.5, 1.0, 1.5, 2.0, 3.0]
                current_ls = st.session_state.get('ls_select_box', 1.0)
                default_ls_idx = ls_options.index(current_ls) if current_ls in ls_options else 2
                ls_select = st.selectbox("เลือกค่า LS", ls_options, index=default_ls_idx, key="ls_select_box")
                
                if 'last_ls_select' not in st.session_state or st.session_state.last_ls_select != ls_select:
                    st.session_state.last_ls_select = ls_select
                    coords = LS_PRESETS.get(ls_select, (150, 718, 903, 84))
                    st.session_state['_ls_x1'], st.session_state['_ls_y1'] = coords[0], coords[1]
                    st.session_state['_ls_x2'], st.session_state['_ls_y2'] = coords[2], coords[3]
                
                with st.expander("ปรับแต่งตำแหน่งเส้น LS ละเอียด", expanded=False):
                    ls_x1 = st.slider("จุดเริ่ม X", -100, w2+100, key="_ls_x1")
                    ls_y1 = st.slider("จุดเริ่ม Y", -100, h2+100, key="_ls_y1")
                    ls_x2 = st.slider("จุดจบ X", -100, w2+100, key="_ls_x2")
                    ls_y2 = st.slider("จุดจบ Y", -100, h2+100, key="_ls_y2")
                
                draw2.line([(ls_x1, ls_y1), (ls_x2, ls_y2)], fill="red", width=6)
                m_red = (ls_y2 - ls_y1) / (ls_x2 - ls_x1) if ls_x2 - ls_x1 != 0 else None
                c_red = ls_y1 - m_red * ls_x1 if m_red else 0
                
                st.markdown("---")
                st.write("#### 2. ค่า k และขอบเขตแกน (เส้นเขียว)")
                with st.expander("📍 ตั้งค่าตำแหน่งแกนกราฟ", expanded=True):
                    col_b1, col_b2 = st.columns(2)
                    with col_b1:
                        axis_left_x = st.number_input("ตำแหน่งแกน Y (ซ้ายสุด)", value=100, step=5, key="axis_left")
                    with col_b2:
                        axis_bottom_y = st.number_input("ตำแหน่งแกน X (ล่างสุด)", value=h2-50, step=5, key="axis_bottom")
                
                st.caption(f"ค่า k จาก Step 1 คือ: {st.session_state.k_inf_result} pci")
                k_input_x = st.slider("ตำแหน่ง k บนแกน X", 0, w2, int(w2*0.5), key="k_pos_x")
                intersect_y = int(m_red * k_input_x + c_red) if m_red else h2//2
                
                draw2.line([(k_input_x, axis_bottom_y), (k_input_x, intersect_y)], fill="springgreen", width=5)
                draw_arrow_fixed(draw2, (k_input_x, intersect_y), (axis_left_x, intersect_y), "springgreen", width=5)
                draw2.ellipse([(k_input_x-8, intersect_y-8), (k_input_x+8, intersect_y+8)], fill="black", outline="white", width=2)
                
                st.markdown("---")
                st.subheader("📝 บันทึกผลลัพธ์")
                k_corrected = st.number_input("Corrected k (pci)", value=st.session_state.get('k_corr_input', st.session_state.k_inf_result - 100), step=10, key="k_corr_input")
                
                buf2 = io.BytesIO()
                img2_draw.save(buf2, format='PNG')
                st.session_state.img2_bytes = buf2.getvalue()
                
                st.markdown("---")
                params = {
                    'MR': st.session_state.get('nomo_mr', 7000),
                    'ESB': st.session_state.get('nomo_esb', 50000),
                    'DSB': st.session_state.get('nomo_dsb', 6.0),
                    'k_inf': st.session_state.k_inf_result,
                    'LS_factor': ls_select,
                    'k_corrected': k_corrected
                }
                if st.button("📄 สร้างรายงาน Nomograph (Word)", key="btn_nomo_report"):
                    with st.spinner("กำลังสร้างรายงาน..."):
                        doc_file, err = generate_word_report_nomograph(params, st.session_state.get('img1_bytes'), st.session_state.get('img2_bytes'))
                        if err:
                            st.error(err)
                        else:
                            st.download_button("📥 ดาวน์โหลด Word Report", doc_file, f"AASHTO_Nomograph_{datetime.now().strftime('%Y%m%d')}.docx",
                                               "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
            with col_img2:
                st.image(img2_draw, caption=f"Step 2: LS Correction (LS={ls_select})", use_container_width=True)
        else:
            st.info("👆 กรุณาอัปโหลดภาพ Figure 3.4 เพื่อเริ่มใช้งาน")
    
    # =========================================================
    # TAB 4: Save Project
    # =========================================================
    with tab4:
        st.header("💾 บันทึกโปรเจกต์")
        st.info("บันทึกข้อมูลทั้งหมดเป็นไฟล์ JSON เพื่อโหลดกลับมาแก้ไขภายหลัง")
        
        if st.button("💾 สร้างไฟล์บันทึก", type="primary"):
            project_data = collect_design_data(
                project_name=st.session_state.get('calc_project_name', ''),
                pavement_type=st.session_state.get('calc_pave_type', 'JPCP'),
                num_layers=st.session_state.get('calc_num_layers', 5),
                layers_data=[{"name": st.session_state.get(f'calc_layer_name_{i}', ''),
                              "thickness_cm": st.session_state.get(f'calc_layer_thick_{i}', 0),
                              "E_MPa": st.session_state.get(f'calc_layer_E_{i}_{st.session_state.get(f"calc_layer_name_{i}", "")}', 100)}
                             for i in range(st.session_state.get('calc_num_layers', 5))],
                w18_design=st.session_state.get('calc_w18', 500000),
                pt=st.session_state.get('calc_pt', 2.0),
                reliability=st.session_state.get('calc_reliability', 90),
                so=st.session_state.get('calc_so', 0.35),
                k_eff=st.session_state.get('calc_k_eff', 200),
                ls_value=st.session_state.get('calc_ls', 1.0),
                fc_cube=st.session_state.get('calc_fc', 350),
                sc=st.session_state.get('calc_sc', 600),
                j_value=st.session_state.get('calc_j', 2.8),
                cd=st.session_state.get('calc_cd', 1.0),
                d_cm_selected=st.session_state.get('calc_d', 30),
                cbr_value=st.session_state.get('calc_cbr', 4.0),
                mr_val=st.session_state.get('nomo_mr', 7000),
                esb_val=st.session_state.get('nomo_esb', 50000),
                dsb_val=st.session_state.get('nomo_dsb', 6.0),
                k_inf_val=st.session_state.get('nomo_k_inf', 400),
                ls_select=st.session_state.get('ls_select_box', 1.0),
                k_corrected=st.session_state.get('k_corr_input', 300)
            )
            json_bytes = save_project_to_json(project_data)
            proj_name = project_data['project_info']['project_name'] or 'AASHTO_Project'
            st.download_button("📥 ดาวน์โหลดไฟล์ JSON", json_bytes, f"{proj_name}_{datetime.now().strftime('%Y%m%d_%H%M')}.json", "application/json")
            st.success("สร้างไฟล์บันทึกสำเร็จ!")
    
    # =========================================================
    # TAB 5: User Guide
    # =========================================================
    with tab5:
        st.header("📋 คู่มือการใช้งาน")
        st.markdown("""
        ### 🔢 Tab 1: AASHTO Calculator
        1. กรอกข้อมูลโครงการและชั้นโครงสร้างทาง
        2. ระบุ ESAL, Serviceability, Reliability
        3. ระบุคุณสมบัติดินและคอนกรีต
        4. เลือกความหนาที่ต้องการตรวจสอบ
        5. ดูผลการคำนวณและสร้างรายงาน
        
        ### 📊 Tab 2: Nomograph - Composite k∞
        1. อัปโหลดรูป **Figure 3.3**
        2. ปรับ **Turning Line (เส้นเขียว)** ให้ตรงกับเส้นบนกราฟ
        3. ปรับตำแหน่งลูกศรสีแดง/ส้ม ให้ตรงกับค่า **MR** และ **ESB**
        4. บันทึกค่า k∞ ที่อ่านได้
        
        ### 📉 Tab 3: Nomograph - Loss of Support
        1. อัปโหลดรูป **Figure 3.4**
        2. เลือกค่า **LS** จากตัวเลือก
        3. ตั้งค่าตำแหน่งแกนกราฟ
        4. เลื่อน Slider ตำแหน่ง k บนแกน X
        5. อ่านค่า Corrected k และบันทึก
        
        ### 💾 Tab 4: บันทึกโปรเจกต์
        - กดปุ่ม **สร้างไฟล์บันทึก** เพื่อบันทึกข้อมูลทั้งหมดเป็น JSON
        - ไฟล์ JSON สามารถอัปโหลดกลับมาได้ที่ **Sidebar**
        
        ---
        **Reference:** AASHTO Guide for Design of Pavement Structures 1993
        """)
    
    st.markdown("---")
    st.caption("พัฒนาโดย: รศ.ดร.อิทธิพล มีผล // ภาควิชาครุศาสตร์โยธา // มจพ.")

if __name__ == "__main__":
    main()
