"""
ESAL Calculator - AASHTO 1993 (Version 2.0)
โปรแกรมคำนวณปริมาณเพลาเดี่ยวมาตรฐานเทียบเท่า (Equivalent Single Axle Load)
สำหรับผิวทาง Rigid Pavement และ Flexible Pavement
ตามมาตรฐาน AASHTO Guide for Design of Pavement Structures (1993)

Features:
- รองรับ Export Excel และ Word ในรูปแบบมาตรฐาน
- Save/Load Project สำหรับแก้ไขภายหลัง
- คำนวณ ACC. ESAL (สะสม)

พัฒนาโดย: รศ.ดร.อิทธิพล มีผล ภาควิชาครุศาสตร์โยธา มจพ.
"""

import streamlit as st
import pandas as pd
import numpy as np
from io import BytesIO
import json
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils.dataframe import dataframe_to_rows

# ============================================================
# ข้อมูลรถบรรทุก 6 ชนิดตามกรมทางหลวงประเทศไทย
# ============================================================
TRUCKS = {
    'MB': {'desc': 'Medium Bus (รถโดยสารขนาดกลาง)', 'axles': [{'name': 'เพลาหน้า', 'load_ton': 4.0, 'type': 'Single'}, {'name': 'เพลาหลัง', 'load_ton': 11.0, 'type': 'Tandem'}]},
    'HB': {'desc': 'Heavy Bus (รถโดยสารขนาดใหญ่)', 'axles': [{'name': 'เพลาหน้า', 'load_ton': 5.0, 'type': 'Single'}, {'name': 'เพลาหลัง', 'load_ton': 20.0, 'type': 'Tandem'}]},
    'MT': {'desc': 'Medium Truck (รถบรรทุกขนาดกลาง)', 'axles': [{'name': 'เพลาหน้า', 'load_ton': 4.0, 'type': 'Single'}, {'name': 'เพลาหลัง', 'load_ton': 11.0, 'type': 'Single'}]},
    'HT': {'desc': 'Heavy Truck (รถบรรทุกขนาดใหญ่)', 'axles': [{'name': 'เพลาหน้า', 'load_ton': 5.0, 'type': 'Single'}, {'name': 'เพลาหลัง', 'load_ton': 20.0, 'type': 'Tandem'}]},
    'TR': {'desc': 'Full Trailer (รถพ่วง)', 'axles': [{'name': 'เพลาหน้า', 'load_ton': 5.0, 'type': 'Single'}, {'name': 'เพลาหลัง', 'load_ton': 20, 'type': 'Tandem'}, {'name': 'เพลาพ่วงหน้า', 'load_ton': 11, 'type': 'Single'}, {'name': 'เพลาพ่วงหลัง', 'load_ton': 11, 'type': 'Single'}]},
    'STR': {'desc': 'Semi-Trailer (รถกึ่งพ่วง)', 'axles': [{'name': 'เพลาหน้า', 'load_ton': 5.0, 'type': 'Single'}, {'name': 'เพลาหลัง', 'load_ton': 20.0, 'type': 'Tandem'}, {'name': 'เพลาพ่วงหลัง', 'load_ton': 20.0, 'type': 'Tandem'}]}
}

# ============================================================
# ตาราง Truck Factor คำนวณตาม AASHTO 1993
# ============================================================

# Rigid Pavement - pt = 2.0
TRUCK_FACTORS_RIGID_PT20 = {
    'MB':  {10: 3.718199, 11: 3.742581, 12: 3.754803, 13: 3.760977, 14: 3.764184, 15: 3.765855, 16: 3.766727},
    'HB':  {10: 6.125043, 11: 6.204343, 12: 6.247170, 13: 6.269632, 14: 6.281529, 15: 6.287867, 16: 6.291257},
    'MT':  {10: 3.718199, 11: 3.742581, 12: 3.754803, 13: 3.760977, 14: 3.764184, 15: 3.765855, 16: 3.766727},
    'HT':  {10: 6.125043, 11: 6.204343, 12: 6.247170, 13: 6.269632, 14: 6.281529, 15: 6.287867, 16: 6.291257},
    'TR':  {10: 13.466316, 11: 13.594592, 12: 13.661961, 13: 13.696817, 14: 13.715152, 15: 13.724934, 16: 13.730167},
    'STR': {10: 12.128867, 11: 12.287718, 12: 12.373488, 13: 12.418469, 14: 12.442292, 15: 12.454956, 16: 12.461738}
}

# Rigid Pavement - pt = 2.5
TRUCK_FACTORS_RIGID_PT25 = {
    'MB':  {10: 3.657799, 11: 3.711341, 12: 3.738346, 13: 3.752027, 14: 3.759145, 15: 3.762869, 16: 3.764817},
    'HB':  {10: 5.921064, 11: 6.092776, 12: 6.186668, 13: 6.236237, 14: 6.262582, 15: 6.276617, 16: 6.284134},
    'MT':  {10: 3.657799, 11: 3.711341, 12: 3.738346, 13: 3.752027, 14: 3.759145, 15: 3.762869, 16: 3.764817},
    'HT':  {10: 5.921064, 11: 6.092776, 12: 6.186668, 13: 6.236237, 14: 6.262582, 15: 6.276617, 16: 6.284134},
    'TR':  {10: 13.141034, 11: 13.420301, 12: 13.568419, 13: 13.645455, 14: 13.686091, 15: 13.707787, 16: 13.719438},
    'STR': {10: 11.720309, 11: 12.064293, 12: 12.252335, 13: 12.351598, 14: 12.404353, 15: 12.432524, 16: 12.447620}
}

# Rigid Pavement - pt = 3.0
TRUCK_FACTORS_RIGID_PT30 = {
    'MB':  {10: 3.581408, 11: 3.671458, 12: 3.717236, 13: 3.740520, 14: 3.752660, 15: 3.759033, 16: 3.762385},
    'HB':  {10: 5.668347, 11: 5.951971, 12: 6.109552, 13: 6.193451, 14: 6.238241, 15: 6.262146, 16: 6.274979},
    'MT':  {10: 3.581408, 11: 3.671458, 12: 3.717236, 13: 3.740520, 14: 3.752660, 15: 3.759033, 16: 3.762385},
    'HT':  {10: 5.668347, 11: 5.951971, 12: 6.109552, 13: 6.193451, 14: 6.238241, 15: 6.262146, 16: 6.274979},
    'TR':  {10: 12.734883, 11: 13.199416, 12: 13.448924, 13: 13.579571, 14: 13.648731, 15: 13.685766, 16: 13.705646},
    'STR': {10: 11.214096, 11: 11.782308, 12: 12.097912, 13: 12.265925, 14: 12.355613, 15: 12.403556, 16: 12.429280}
}

# Flexible Pavement - pt = 2.0
TRUCK_FACTORS_FLEX_PT20 = {
    'MB':  {4: 3.529011, 5: 3.598168, 6: 3.719257, 7: 3.810681, 8: 3.874256, 9: 3.916863},
    'HB':  {4: 3.332846, 5: 3.384895, 6: 3.458092, 7: 3.508785, 8: 3.541983, 9: 3.562854},
    'MT':  {4: 3.529011, 5: 3.598168, 6: 3.719257, 7: 3.810681, 8: 3.874256, 9: 3.916863},
    'HT':  {4: 3.332846, 5: 3.384895, 6: 3.458092, 7: 3.508785, 8: 3.541983, 9: 3.562854},
    'TR':  {4: 10.291092, 5: 10.488813, 6: 10.808050, 7: 11.043444, 8: 11.203523, 9: 11.310117},
    'STR': {4: 6.537851, 5: 6.649420, 6: 6.800056, 7: 6.903531, 8: 6.971366, 9: 7.014261}
}

# Flexible Pavement - pt = 2.5
TRUCK_FACTORS_FLEX_PT25 = {
    'MB':  {4: 3.069453, 5: 3.203842, 6: 3.451114, 7: 3.645241, 8: 3.779066, 9: 3.869188},
    'HB':  {4: 3.053625, 5: 3.157524, 6: 3.311765, 7: 3.421800, 8: 3.494667, 9: 3.541837},
    'MT':  {4: 3.069453, 5: 3.203842, 6: 3.451114, 7: 3.645241, 8: 3.779066, 9: 3.869188},
    'HT':  {4: 3.053625, 5: 3.157524, 6: 3.311765, 7: 3.421800, 8: 3.494667, 9: 3.541837},
    'TR':  {4: 9.069826, 5: 9.462000, 6: 10.120276, 7: 10.622935, 8: 10.967259, 9: 11.196528},
    'STR': {4: 5.955718, 5: 6.182789, 6: 6.501567, 7: 6.726542, 8: 6.874756, 9: 6.970223}
}

# Flexible Pavement - pt = 3.0
TRUCK_FACTORS_FLEX_PT30 = {
    'MB':  {4: 2.552540, 5: 2.742623, 6: 3.120508, 7: 3.433469, 8: 3.657896, 9: 3.812063},
    'HB':  {4: 2.728486, 5: 2.879854, 6: 3.125499, 7: 3.308196, 8: 3.432738, 9: 3.513580},
    'MT':  {4: 2.552540, 5: 2.742623, 6: 3.120508, 7: 3.433469, 8: 3.657896, 9: 3.812063},
    'HT':  {4: 2.728486, 5: 2.879854, 6: 3.125499, 7: 3.308196, 8: 3.432738, 9: 3.513580},
    'TR':  {4: 7.671306, 5: 8.245291, 6: 9.265343, 7: 10.082089, 8: 10.658949, 9: 11.046207},
    'STR': {4: 5.266321, 5: 5.609502, 6: 6.120685, 7: 6.495126, 8: 6.750547, 9: 6.915832}
}


def get_default_truck_factor(truck_code, pavement_type, pt, param):
    """ดึงค่า Truck Factor เริ่มต้นจากตาราง"""
    if pavement_type == 'rigid':
        if pt == 2.0:
            return TRUCK_FACTORS_RIGID_PT20[truck_code][param]
        elif pt == 2.5:
            return TRUCK_FACTORS_RIGID_PT25[truck_code][param]
        else:
            return TRUCK_FACTORS_RIGID_PT30[truck_code][param]
    else:
        if pt == 2.0:
            return TRUCK_FACTORS_FLEX_PT20[truck_code][param]
        elif pt == 2.5:
            return TRUCK_FACTORS_FLEX_PT25[truck_code][param]
        else:
            return TRUCK_FACTORS_FLEX_PT30[truck_code][param]


def calculate_esal_with_acc(traffic_df, truck_factors, lane_factor=0.5, direction_factor=1.0):
    """คำนวณ ESAL และ Accumulated ESAL จากข้อมูลปริมาณจราจร"""
    results = []
    acc_esal = 0
    
    for idx, row in traffic_df.iterrows():
        year = row.get('Year', idx + 1)
        year_data = {'Year': int(year)}
        
        # คำนวณ ADT รายประเภทรถ
        total_aadt = 0
        for code in TRUCKS.keys():
            if code in traffic_df.columns:
                aadt = int(row[code])
                year_data[code] = aadt
                total_aadt += aadt
        
        year_data['AADT'] = total_aadt
        
        # คำนวณ ESAL
        year_esal = 0
        for code in TRUCKS.keys():
            if code in traffic_df.columns:
                aadt = row[code]
                tf = truck_factors[code]
                esal = aadt * tf * lane_factor * direction_factor * 365
                year_esal += esal
        
        year_data['ESAL'] = int(round(year_esal))
        acc_esal += year_esal
        year_data['ACC_ESAL'] = int(round(acc_esal))
        
        results.append(year_data)
    
    return pd.DataFrame(results), int(round(acc_esal))


def create_template():
    """สร้าง Template Excel สำหรับอัพโหลดข้อมูล"""
    base = {'MB': 120, 'HB': 60, 'MT': 250, 'HT': 180, 'TR': 100, 'STR': 120}
    growth_rate = 1.045
    
    data = {'Year': list(range(1, 21))}
    for code in base.keys():
        data[code] = [int(base[code] * (growth_rate ** i)) for i in range(20)]
    
    return pd.DataFrame(data)


def create_excel_report(results_df, pavement_type, pt, param, lane_factor, direction_factor, 
                       total_esal, truck_factors, num_years):
    """สร้างรายงาน Excel ในรูปแบบมาตรฐาน"""
    wb = Workbook()
    ws = wb.active
    ws.title = "ESAL Report"
    
    # Styles
    header_font = Font(bold=True, size=14)
    title_font = Font(bold=True, size=16)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    header_fill = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')
    center_align = Alignment(horizontal='center', vertical='center')
    right_align = Alignment(horizontal='right', vertical='center')
    
    # Title
    pavement_text = "Rigid Pavement" if pavement_type == 'rigid' else "Flexible Pavement"
    ws.merge_cells('A1:I1')
    ws['A1'] = f"ปริมาณเพลามาตรฐาน (ESALs) ระยะเวลาออกแบบ {num_years} ปี"
    ws['A1'].font = title_font
    ws['A1'].alignment = center_align
    
    ws.merge_cells('A2:I2')
    ws['A2'] = f"ผิวทางแบบ{'แข็ง' if pavement_type == 'rigid' else 'ยืดหยุ่น'} ({pavement_text})"
    ws['A2'].font = header_font
    ws['A2'].alignment = center_align
    
    # Parameter Table (Row 4-10)
    param_label = f"D = {param}" if pavement_type == 'rigid' else f"SN = {param}"
    params_data = [
        ('รายการ', 'ค่า'),
        ('ประเภทผิวทาง', pavement_text),
        ('pt', str(pt)),
        ('พารามิเตอร์', param_label),
        ('Lane Factor', str(lane_factor)),
        ('Direction Factor', str(direction_factor)),
        ('ESAL รวม', f"{total_esal:,}"),
        ('จำนวนปี', str(num_years))
    ]
    
    for i, (label, value) in enumerate(params_data):
        row = 4 + i
        ws[f'A{row}'] = label
        ws[f'B{row}'] = value
        ws[f'A{row}'].border = border
        ws[f'B{row}'].border = border
        if i == 0:
            ws[f'A{row}'].fill = header_fill
            ws[f'B{row}'].fill = header_fill
            ws[f'A{row}'].font = Font(bold=True)
            ws[f'B{row}'].font = Font(bold=True)
    
    # Truck Factor Table (Row 4-10, Column D-F)
    ws['D4'] = 'รหัส'
    ws['E4'] = 'ประเภท'
    ws['F4'] = 'Truck Factor'
    for col in ['D', 'E', 'F']:
        ws[f'{col}4'].fill = header_fill
        ws[f'{col}4'].font = Font(bold=True)
        ws[f'{col}4'].border = border
        ws[f'{col}4'].alignment = center_align
    
    for i, code in enumerate(TRUCKS.keys()):
        row = 5 + i
        ws[f'D{row}'] = code
        ws[f'E{row}'] = TRUCKS[code]['desc']
        ws[f'F{row}'] = f"{truck_factors[code]:.4f}"
        ws[f'D{row}'].border = border
        ws[f'E{row}'].border = border
        ws[f'F{row}'].border = border
        ws[f'D{row}'].alignment = center_align
        ws[f'F{row}'].alignment = right_align
    
    # ESAL Table Header (Row 13+)
    start_row = 14
    ws.merge_cells(f'I{start_row-1}:I{start_row-1}')
    ws[f'I{start_row-1}'] = 'แสดงปริมาณสะสม'
    ws[f'I{start_row-1}'].font = Font(italic=True, size=9)
    
    headers = ['Year', 'MB', 'HB', 'MT', 'HT', 'TR', 'STR', 'AADT', 'ESAL', 'ACC. ESAL']
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=start_row, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = Font(bold=True)
        cell.border = border
        cell.alignment = center_align
    
    # ESAL Data
    for row_idx, row_data in results_df.iterrows():
        excel_row = start_row + 1 + row_idx
        for col_idx, header in enumerate(headers, 1):
            if header == 'ACC. ESAL':
                value = row_data.get('ACC_ESAL', 0)
            else:
                value = row_data.get(header, 0)
            
            cell = ws.cell(row=excel_row, column=col_idx, value=value)
            cell.border = border
            
            if header in ['ESAL', 'ACC. ESAL', 'AADT']:
                cell.number_format = '#,##0'
                cell.alignment = right_align
            elif header == 'Year':
                cell.alignment = center_align
            else:
                cell.alignment = right_align
    
    # Column widths
    ws.column_dimensions['A'].width = 18
    ws.column_dimensions['B'].width = 15
    ws.column_dimensions['C'].width = 3
    ws.column_dimensions['D'].width = 8
    ws.column_dimensions['E'].width = 35
    ws.column_dimensions['F'].width = 14
    for col in ['G', 'H', 'I', 'J']:
        ws.column_dimensions[col].width = 14
    
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return output


def create_word_report(results_df, pavement_type, pt, param, lane_factor, direction_factor, 
                      total_esal, truck_factors, num_years):
    """สร้างรายงาน Word ในรูปแบบมาตรฐาน"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
    except ImportError:
        return None
    
    doc = Document()
    
    # Set page to landscape
    section = doc.sections[0]
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    
    pavement_text = "Rigid Pavement" if pavement_type == 'rigid' else "Flexible Pavement"
    pavement_thai = "แข็ง" if pavement_type == 'rigid' else "ยืดหยุ่น"
    param_label = f"D = {param}" if pavement_type == 'rigid' else f"SN = {param}"
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"ปริมาณเพลามาตรฐาน (ESALs) ระยะเวลาออกแบบ {num_years} ปี")
    run.bold = True
    run.font.size = Pt(16)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"ผิวทางแบบ{pavement_thai} ({pavement_text})")
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    # Parameter Table
    param_table = doc.add_table(rows=8, cols=2)
    param_table.style = 'Table Grid'
    
    param_data = [
        ('รายการ', 'ค่า'),
        ('ประเภทผิวทาง', pavement_text),
        ('pt', str(pt)),
        ('พารามิเตอร์', param_label),
        ('Lane Factor', str(lane_factor)),
        ('Direction Factor', str(direction_factor)),
        ('ESAL รวม', f"{total_esal:,}"),
        ('จำนวนปี', str(num_years))
    ]
    
    for i, (label, value) in enumerate(param_data):
        row = param_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        if i == 0:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    doc.add_paragraph()
    
    # Truck Factor Table
    tf_table = doc.add_table(rows=7, cols=3)
    tf_table.style = 'Table Grid'
    
    # Header
    hdr = tf_table.rows[0]
    hdr.cells[0].text = 'รหัส'
    hdr.cells[1].text = 'ประเภท'
    hdr.cells[2].text = 'Truck Factor'
    for cell in hdr.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    for i, code in enumerate(TRUCKS.keys()):
        row = tf_table.rows[i + 1]
        row.cells[0].text = code
        row.cells[1].text = TRUCKS[code]['desc']
        row.cells[2].text = f"{truck_factors[code]:.4f}"
    
    doc.add_paragraph()
    
    # ESAL Table
    headers = ['Year', 'MB', 'HB', 'MT', 'HT', 'TR', 'STR', 'AADT', 'ESAL', 'ACC. ESAL']
    esal_table = doc.add_table(rows=len(results_df) + 1, cols=len(headers))
    esal_table.style = 'Table Grid'
    
    # Header row
    hdr = esal_table.rows[0]
    for j, header in enumerate(headers):
        hdr.cells[j].text = header
        for paragraph in hdr.cells[j].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
    
    # Data rows
    for i, row_data in results_df.iterrows():
        row = esal_table.rows[i + 1]
        for j, header in enumerate(headers):
            if header == 'ACC. ESAL':
                value = row_data.get('ACC_ESAL', 0)
            else:
                value = row_data.get(header, 0)
            
            if header in ['ESAL', 'ACC. ESAL', 'AADT']:
                row.cells[j].text = f"{int(value):,}"
            else:
                row.cells[j].text = str(int(value))
            
            for paragraph in row.cells[j].paragraphs:
                if header == 'Year':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def save_project(pavement_type, pt, param, lane_factor, direction_factor, truck_factors, traffic_df):
    """บันทึก Project เป็น JSON"""
    project = {
        'version': '2.0',
        'created': datetime.now().isoformat(),
        'pavement_type': pavement_type,
        'pt': pt,
        'param': param,
        'lane_factor': lane_factor,
        'direction_factor': direction_factor,
        'truck_factors': truck_factors,
        'traffic_data': traffic_df.to_dict('records')
    }
    return json.dumps(project, ensure_ascii=False, indent=2)


def load_project(json_data):
    """โหลด Project จาก JSON"""
    try:
        project = json.loads(json_data)
        return project
    except:
        return None


def get_all_truck_factors_table(pavement_type, pt):
    """สร้างตาราง Truck Factor ทั้งหมด"""
    data = []
    
    if pavement_type == 'rigid':
        params = [10, 11, 12, 13, 14, 15, 16]
        param_label = 'D'
        if pt == 2.0:
            tf_table = TRUCK_FACTORS_RIGID_PT20
        elif pt == 2.5:
            tf_table = TRUCK_FACTORS_RIGID_PT25
        else:
            tf_table = TRUCK_FACTORS_RIGID_PT30
    else:
        params = [4, 5, 6, 7, 8, 9]
        param_label = 'SN'
        if pt == 2.0:
            tf_table = TRUCK_FACTORS_FLEX_PT20
        elif pt == 2.5:
            tf_table = TRUCK_FACTORS_FLEX_PT25
        else:
            tf_table = TRUCK_FACTORS_FLEX_PT30
    
    for code in TRUCKS.keys():
        row = {'ประเภท': code, 'รายละเอียด': TRUCKS[code]['desc']}
        for p in params:
            col_name = f'{param_label}={p}"' if pavement_type == 'rigid' else f'{param_label}={p}'
            row[col_name] = f"{tf_table[code][p]:.4f}"
        data.append(row)
    
    return pd.DataFrame(data)


# ============================================================
# Streamlit App
# ============================================================
def main():
    st.set_page_config(
        page_title="ESAL Calculator - AASHTO 1993",
        page_icon="🛣️",
        layout="wide"
    )
    
    st.markdown("""
    <style>
    .main-header { font-size: 2.5rem; font-weight: bold; color: #1E3A5F; text-align: center; margin-bottom: 0.5rem; }
    .sub-header { font-size: 1.2rem; color: #4A6FA5; text-align: center; margin-bottom: 2rem; }
    .metric-box { background: linear-gradient(135deg, #1E3A5F 0%, #4A6FA5 100%); padding: 1.5rem; border-radius: 10px; color: white; text-align: center; margin: 0.5rem 0; }
    .metric-value { font-size: 2rem; font-weight: bold; }
    .metric-label { font-size: 0.9rem; opacity: 0.9; }
    </style>
    """, unsafe_allow_html=True)
    
    st.markdown('<p class="main-header">🛣️ ESAL Calculator v2.0</p>', unsafe_allow_html=True)
    st.markdown('<p class="sub-header">คำนวณปริมาณเพลาเดี่ยวมาตรฐานเทียบเท่า ตามมาตรฐาน AASHTO 1993</p>', unsafe_allow_html=True)
    
    # Initialize session state
    if 'traffic_df' not in st.session_state:
        st.session_state['traffic_df'] = None
    
    # Sidebar
    with st.sidebar:
        st.header("⚙️ พารามิเตอร์การคำนวณ")
        
        # Project Load/Save
        st.subheader("📁 Project")
        
        uploaded_project = st.file_uploader("📥 โหลด Project", type=['json'], key='load_project')
        if uploaded_project is not None:
            try:
                # ตรวจสอบว่าเป็นไฟล์ใหม่
                file_id = f"{uploaded_project.name}_{uploaded_project.size}"
                if st.session_state.get('last_uploaded_file') != file_id:
                    st.session_state['last_uploaded_file'] = file_id
                    
                    project = load_project(uploaded_project.read().decode('utf-8'))
                    if project:
                        # อัพเดท session_state สำหรับทุก field
                        st.session_state['input_pavement_type'] = project.get('pavement_type', 'rigid')
                        st.session_state['input_pt'] = project.get('pt', 2.5)
                        st.session_state['input_param'] = project.get('param', 12)
                        st.session_state['input_lane_factor'] = project.get('lane_factor', 0.5)
                        st.session_state['input_direction_factor'] = project.get('direction_factor', 0.9)
                        st.session_state['loaded_tf'] = project.get('truck_factors', {})
                        
                        loaded_traffic = project.get('traffic_data', None)
                        if loaded_traffic:
                            st.session_state['traffic_df'] = pd.DataFrame(loaded_traffic)
                        
                        st.success("✅ โหลด Project สำเร็จ!")
                        st.rerun()
                    else:
                        st.error("❌ ไม่สามารถอ่านไฟล์ได้")
            except Exception as e:
                st.error(f"❌ เกิดข้อผิดพลาด: {e}")
        
        # กำหนดค่า default จาก session_state
        default_pavement = st.session_state.get('input_pavement_type', 'rigid')
        default_pt = st.session_state.get('input_pt', 2.5)
        default_param = st.session_state.get('input_param', 12 if default_pavement == 'rigid' else 7)
        default_lane = st.session_state.get('input_lane_factor', 0.5)
        default_dir = st.session_state.get('input_direction_factor', 0.9)
        loaded_tf = st.session_state.get('loaded_tf', {})
        
        st.divider()
        
        # Pavement Type
        pavement_options = ['rigid', 'flexible']
        pavement_idx = pavement_options.index(default_pavement) if default_pavement in pavement_options else 0
        pavement_type = st.selectbox(
            "ประเภทผิวทาง",
            options=pavement_options,
            index=pavement_idx,
            format_func=lambda x: '🧱 Rigid Pavement (คอนกรีต)' if x == 'rigid' else '🛤️ Flexible Pavement (ลาดยาง)',
            key="input_pavement_type"
        )
        
        pt_options = [2.0, 2.5, 3.0]
        pt_idx = pt_options.index(default_pt) if default_pt in pt_options else 1
        pt = st.selectbox(
            "Terminal Serviceability (pt)",
            options=pt_options,
            index=pt_idx,
            format_func=lambda x: f"pt = {x}",
            key="input_pt"
        )
        
        if pavement_type == 'rigid':
            param_options = [10, 11, 12, 13, 14, 15, 16]
            # ถ้า param จาก session_state ไม่อยู่ใน options ให้ใช้ default
            if default_param not in param_options:
                default_param = 12
            default_idx = param_options.index(default_param) if default_param in param_options else 2
            param = st.selectbox(
                "ความหนาพื้นคอนกรีต (D)",
                options=param_options,
                index=default_idx,
                format_func=lambda x: f"D = {x} นิ้ว",
                key="input_param_rigid"
            )
            param_label = f"D = {param} นิ้ว"
        else:
            param_options = [4, 5, 6, 7, 8, 9]
            # ถ้า param จาก session_state ไม่อยู่ใน options ให้ใช้ default
            if default_param not in param_options:
                default_param = 7
            default_idx = param_options.index(default_param) if default_param in param_options else 3
            param = st.selectbox(
                "Structural Number (SN)",
                options=param_options,
                index=default_idx,
                format_func=lambda x: f"SN = {x}",
                key="input_param_flex"
            )
            param_label = f"SN = {param}"
        
        st.divider()
        
        st.subheader("🚗 ค่าสัดส่วน")
        lane_factor = st.slider(
            "Lane Distribution Factor", 
            0.1, 1.0, 
            value=st.session_state.get('input_lane_factor', default_lane), 
            step=0.05,
            key="input_lane_factor"
        )
        direction_factor = st.slider(
            "Directional Factor", 
            0.5, 1.0, 
            value=st.session_state.get('input_direction_factor', default_dir), 
            step=0.1,
            key="input_direction_factor"
        )
        
        st.divider()
        
        # Truck Factor
        st.subheader("🚛 ค่า Truck Factor")
        
        tf_key = f"tf_{pavement_type}_{pt}_{param}"
        
        # ถ้ามี loaded_tf จาก JSON ให้ใช้ค่านั้น
        if loaded_tf and tf_key not in st.session_state:
            st.session_state[tf_key] = {}
            for code in TRUCKS.keys():
                if code in loaded_tf:
                    st.session_state[tf_key][code] = loaded_tf[code]
                else:
                    st.session_state[tf_key][code] = get_default_truck_factor(code, pavement_type, pt, param)
        elif tf_key not in st.session_state:
            st.session_state[tf_key] = {}
            for code in TRUCKS.keys():
                st.session_state[tf_key][code] = get_default_truck_factor(code, pavement_type, pt, param)
        
        if st.button("🔄 Reset เป็นค่า Default", use_container_width=True):
            for code in TRUCKS.keys():
                st.session_state[tf_key][code] = get_default_truck_factor(code, pavement_type, pt, param)
            st.rerun()
        
        st.caption("กรอกค่า Truck Factor (แก้ไขได้)")
        
        truck_factors = {}
        for code in TRUCKS.keys():
            default_val = get_default_truck_factor(code, pavement_type, pt, param)
            current_val = st.session_state[tf_key].get(code, default_val)
            
            new_val = st.number_input(
                f"{code}",
                min_value=0.0,
                max_value=50.0,
                value=float(current_val),
                step=0.0001,
                format="%.4f",
                key=f"input_{tf_key}_{code}",
                help=f"{TRUCKS[code]['desc']} | Default: {default_val:.4f}"
            )
            
            st.session_state[tf_key][code] = new_val
            truck_factors[code] = new_val
        
        st.divider()
        
        st.subheader("📥 ดาวน์โหลด Template")
        template_df = create_template()
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            template_df.to_excel(writer, index=False, sheet_name='Traffic Data')
        st.download_button(
            label="📄 ดาวน์โหลด Template Excel",
            data=output.getvalue(),
            file_name="traffic_template.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )
    
    # Main Content
    tab1, tab2, tab3 = st.tabs(["📊 คำนวณ ESAL", "🚛 ข้อมูล Truck Factor", "📘 คู่มือ"])
    
    with tab1:
        col1, col2 = st.columns([1, 2])
        
        with col1:
            st.subheader("📤 อัพโหลดข้อมูลปริมาณจราจร")
            
            uploaded_file = st.file_uploader(
                "เลือกไฟล์ Excel",
                type=['xlsx', 'xls'],
                help="อัพโหลดไฟล์ Excel (หน่วย: คัน/วัน)"
            )
            
            if 'use_sample' not in st.session_state:
                st.session_state['use_sample'] = False
            
            if uploaded_file is not None:
                try:
                    traffic_df = pd.read_excel(uploaded_file)
                    st.session_state['traffic_df'] = traffic_df
                    st.success("✅ อัพโหลดสำเร็จ!")
                    st.session_state['use_sample'] = False
                except Exception as e:
                    st.error(f"❌ เกิดข้อผิดพลาด: {e}")
                    traffic_df = st.session_state.get('traffic_df', None)
            elif st.session_state.get('traffic_df') is not None:
                traffic_df = st.session_state['traffic_df']
            else:
                st.info("📌 อัพโหลดไฟล์ Excel หรือใช้ข้อมูลตัวอย่าง")
                
                if st.button("🔄 ใช้ข้อมูลตัวอย่าง", use_container_width=True):
                    st.session_state['use_sample'] = True
                    st.session_state['traffic_df'] = create_template()
                
                traffic_df = st.session_state.get('traffic_df', None)
            
            if traffic_df is not None:
                st.write("**ข้อมูลปริมาณจราจร (คัน/วัน):**")
                st.dataframe(traffic_df, use_container_width=True, height=350)
        
        with col2:
            st.subheader("📈 ผลการคำนวณ ESAL")
            
            if traffic_df is not None:
                results_df, total_esal = calculate_esal_with_acc(
                    traffic_df, truck_factors, lane_factor, direction_factor
                )
                
                # Metrics
                col_m1, col_m2, col_m3, col_m4 = st.columns(4)
                
                with col_m1:
                    st.markdown(f"""
                    <div class="metric-box">
                        <div class="metric-value">{total_esal:,}</div>
                        <div class="metric-label">ESAL รวมทั้งหมด</div>
                    </div>
                    """, unsafe_allow_html=True)
                
                with col_m2:
                    st.markdown(f"""
                    <div class="metric-box">
                        <div class="metric-value">{len(traffic_df)}</div>
                        <div class="metric-label">จำนวนปี</div>
                    </div>
                    """, unsafe_allow_html=True)
                
                with col_m3:
                    pavement_label = "Rigid" if pavement_type == 'rigid' else "Flexible"
                    st.markdown(f"""
                    <div class="metric-box">
                        <div class="metric-value">{pavement_label}</div>
                        <div class="metric-label">ประเภทผิวทาง</div>
                    </div>
                    """, unsafe_allow_html=True)
                
                with col_m4:
                    st.markdown(f"""
                    <div class="metric-box">
                        <div class="metric-value">{param_label}</div>
                        <div class="metric-label">พารามิเตอร์</div>
                    </div>
                    """, unsafe_allow_html=True)
                
                st.divider()
                
                # Truck Factor Table
                st.write("**🚛 ค่า Truck Factor ที่ใช้:**")
                tf_display = []
                for code, tf in truck_factors.items():
                    default_tf = get_default_truck_factor(code, pavement_type, pt, param)
                    status = "✅" if abs(tf - default_tf) < 0.0001 else "✏️ แก้ไข"
                    tf_display.append({
                        'รหัส': code, 
                        'ประเภท': TRUCKS[code]['desc'], 
                        'Truck Factor': f"{tf:.4f}",
                        'Default': f"{default_tf:.4f}",
                        'สถานะ': status
                    })
                st.dataframe(pd.DataFrame(tf_display), use_container_width=True, hide_index=True)
                
                st.divider()
                
                # ESAL Results Table
                st.write("**📊 ESAL รายปี:**")
                display_df = results_df.copy()
                display_df.columns = ['ปีที่', 'MB', 'HB', 'MT', 'HT', 'TR', 'STR', 'AADT', 'ESAL', 'ACC. ESAL']
                st.dataframe(display_df, use_container_width=True, height=400)
                
                st.divider()
                
                # Download buttons
                st.write("**📥 ดาวน์โหลดรายงาน:**")
                col_dl1, col_dl2, col_dl3 = st.columns(3)
                
                with col_dl1:
                    excel_report = create_excel_report(
                        results_df, pavement_type, pt, param, lane_factor, direction_factor,
                        total_esal, truck_factors, len(traffic_df)
                    )
                    st.download_button(
                        label="📊 ดาวน์โหลด Excel",
                        data=excel_report.getvalue(),
                        file_name=f"ESAL_Report_{pavement_type}_{param}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True
                    )
                
                with col_dl2:
                    word_report = create_word_report(
                        results_df, pavement_type, pt, param, lane_factor, direction_factor,
                        total_esal, truck_factors, len(traffic_df)
                    )
                    if word_report:
                        st.download_button(
                            label="📝 ดาวน์โหลด Word",
                            data=word_report.getvalue(),
                            file_name=f"ESAL_Report_{pavement_type}_{param}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    else:
                        st.warning("⚠️ ไม่สามารถสร้างรายงาน Word ได้ (กรุณาติดตั้ง python-docx)")
                
                with col_dl3:
                    project_json = save_project(
                        pavement_type, pt, param, lane_factor, direction_factor,
                        truck_factors, traffic_df
                    )
                    st.download_button(
                        label="💾 บันทึก Project",
                        data=project_json,
                        file_name=f"ESAL_Project_{pavement_type}_{param}.json",
                        mime="application/json",
                        use_container_width=True
                    )
            else:
                st.warning("⚠️ กรุณาอัพโหลดข้อมูลหรือใช้ข้อมูลตัวอย่าง")
    
    with tab2:
        st.subheader("🚛 ข้อมูลรถบรรทุก 6 ประเภทตามกรมทางหลวง")
        
        truck_details = []
        for code, truck in TRUCKS.items():
            axle_info = []
            for axle in truck['axles']:
                axle_info.append(f"{axle['name']}: {axle['load_ton']} ตัน ({axle['type']})")
            truck_details.append({'รหัส': code, 'ประเภท': truck['desc'], 'ข้อมูลเพลา': ' | '.join(axle_info)})
        
        st.dataframe(pd.DataFrame(truck_details), use_container_width=True, hide_index=True)
        
        st.divider()
        st.subheader("📊 ตาราง Truck Factor (ค่า Default ตาม AASHTO 1993)")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.write("**🧱 Rigid Pavement (pt = 2.0)**")
            st.dataframe(get_all_truck_factors_table('rigid', 2.0), use_container_width=True, hide_index=True)
            
            st.write("**🧱 Rigid Pavement (pt = 2.5)**")
            st.dataframe(get_all_truck_factors_table('rigid', 2.5), use_container_width=True, hide_index=True)
            
            st.write("**🧱 Rigid Pavement (pt = 3.0)**")
            st.dataframe(get_all_truck_factors_table('rigid', 3.0), use_container_width=True, hide_index=True)
        
        with col2:
            st.write("**🛤️ Flexible Pavement (pt = 2.0)**")
            st.dataframe(get_all_truck_factors_table('flexible', 2.0), use_container_width=True, hide_index=True)
            
            st.write("**🛤️ Flexible Pavement (pt = 2.5)**")
            st.dataframe(get_all_truck_factors_table('flexible', 2.5), use_container_width=True, hide_index=True)
            
            st.write("**🛤️ Flexible Pavement (pt = 3.0)**")
            st.dataframe(get_all_truck_factors_table('flexible', 3.0), use_container_width=True, hide_index=True)
    
    with tab3:
        st.subheader("📘 คู่มือการใช้งาน")
        
        st.markdown("""
        ### 1️⃣ เตรียมไฟล์ Excel
        
        | คอลัมน์ | คำอธิบาย |
        |---------|----------|
        | `Year` | ปีที่ (1, 2, 3, ... n) |
        | `MB` | Medium Bus (คัน/วัน) |
        | `HB` | Heavy Bus (คัน/วัน) |
        | `MT` | Medium Truck (คัน/วัน) |
        | `HT` | Heavy Truck (คัน/วัน) |
        | `STR` | Semi-Trailer (คัน/วัน) |
        | `TR` | Full Trailer (คัน/วัน) |
        
        ### 2️⃣ ตั้งค่าพารามิเตอร์
        
        - **Rigid:** D = 10-16 นิ้ว
        - **Flexible:** SN = 4-9
        - **pt:** 2.0, 2.5 หรือ 3.0
        
        ### 3️⃣ ฟีเจอร์ใหม่ในเวอร์ชัน 2.0
        
        - **ACC. ESAL:** แสดงค่า ESAL สะสม
        - **Export Excel:** รายงานในรูปแบบมาตรฐาน
        - **Export Word:** รายงานสำหรับเอกสาร
        - **Save/Load Project:** บันทึกค่าที่ตั้งไว้และโหลดกลับมาแก้ไขได้
        
        ### 4️⃣ สูตรคำนวณ ESAL
        """)
        
        st.latex(r'ESAL = \sum_{i=1}^{n} \sum_{j=1}^{6} (ADT_{ij} \times TF_j \times LF \times DF \times 365)')
        
        st.markdown("""
        ### 📚 อ้างอิง
        - AASHTO Guide for Design of Pavement Structures (1993)
        - กรมทางหลวง กระทรวงคมนาคม
        """)
    
    st.divider()
    st.markdown("""
    <div style="text-align: center; color: #888;">
        พัฒนาเพื่อการเรียนการสอนโดย รศ.ดร.อิทธิพล มีผล ภาควิชาครุศาสตร์โยธา มจพ. | ESAL Calculator v2.0
    </div>
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
