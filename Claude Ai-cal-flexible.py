"""
================================================================================
AASHTO 1993 Flexible Pavement Design - Streamlit Application (Version 3)
================================================================================
แอปพลิเคชันสำหรับออกแบบ Flexible Pavement ตามวิธี AASHTO 1993
ปรับปรุงตามมาตรฐานกรมทางหลวง (DOH Thailand)

Features:
- Material database ตามมาตรฐาน ทล.
- Step-by-step thickness calculation (หาความหนาแต่ละชั้น)
- Drainage coefficient default = 1.0
- ปรับปรุงรูปภาพตัดขวางให้มีรูปแบบเดียวกับ Rigid Pavement

Author: Civil Engineering Department
Version: 3.0
================================================================================
"""

import streamlit as st
import numpy as np
from scipy.optimize import brentq
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.font_manager as fm
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ================================================================================
# PAGE CONFIGURATION
# ================================================================================

st.set_page_config(
    page_title="AASHTO 1993 Flexible Pavement Design (DOH)",
    page_icon="🛣️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ================================================================================
# MATERIAL DATABASE - ตามมาตรฐานกรมทางหลวง (DOH Thailand)
# ================================================================================

MATERIALS = {
    # ============ ชั้นผิวทาง (Surface) ============
    "ผิวทางลาดยาง AC": {
        "layer_coeff": 0.40,
        "drainage_coeff": 1.0,
        "mr_psi": 362500,
        "mr_mpa": 2500,
        "layer_type": "surface",
        "color": "#1C1C1C",  # สีดำ (Black)
        "short_name": "AC",
        "english_name": "Asphalt Concrete"
    },
    "ผิวทางลาดยาง PMA": {
        "layer_coeff": 0.40,
        "drainage_coeff": 1.0,
        "mr_psi": 536500,
        "mr_mpa": 3700,
        "layer_type": "surface",
        "color": "#2C2C2C",  # สีดำเข้ม (Dark Black)
        "short_name": "PMA",
        "english_name": "Polymer Modified Asphalt"
    },
    
    # ============ ชั้นพื้นทาง (Base) ============
    "พื้นทางซีเมนต์ CTB": {
        "layer_coeff": 0.18,
        "drainage_coeff": 1.0,
        "mr_psi": 174000,
        "mr_mpa": 1200,
        "layer_type": "base",
        "color": "#78909C",  # สีเทา (Gray)
        "short_name": "CTB",
        "english_name": "Cement Treated Base"
    },
    "พื้นทางหินคลุกผสมซีเมนต์ UCS 24.5 ksc.": {
        "layer_coeff": 0.15,
        "drainage_coeff": 1.0,
        "mr_psi": 123250,
        "mr_mpa": 850,
        "layer_type": "base",
        "color": "#607D8B",  # สีเทาเข้ม
        "short_name": "SCAB",
        "english_name": "Soil Cement Aggregate Base"
    },
    "พื้นทางหินคลุก CBR 80%": {
        "layer_coeff": 0.13,
        "drainage_coeff": 1.0,
        "mr_psi": 50750,
        "mr_mpa": 350,
        "layer_type": "base",
        "color": "#795548",  # สีน้ำตาล
        "short_name": "CAB",
        "english_name": "Crushed Rock Base"
    },
    "พื้นทางดินซีเมนต์ UCS 17.5 ksc.": {
        "layer_coeff": 0.13,
        "drainage_coeff": 1.0,
        "mr_psi": 50750,
        "mr_mpa": 350,
        "layer_type": "base",
        "color": "#8D6E63",  # สีน้ำตาลอ่อน
        "short_name": "SCB",
        "english_name": "Soil Cement Base"
    },
    "พื้นทางวัสดุหมุนเวียน (Recycling)": {
        "layer_coeff": 0.15,
        "drainage_coeff": 1.0,
        "mr_psi": 123250,
        "mr_mpa": 850,
        "layer_type": "base",
        "color": "#5D4037",  # สีน้ำตาลเข้ม
        "short_name": "RAP",
        "english_name": "Recycled Asphalt Pavement"
    },
    
    # ============ ชั้นรองพื้นทาง (Subbase) - วัสดุมวลรวม ============
    "รองพื้นทางวัสดุมวลรวม CBR 25%": {
        "layer_coeff": 0.10,
        "drainage_coeff": 1.0,
        "mr_psi": 21750,
        "mr_mpa": 150,
        "layer_type": "subbase",
        "color": "#FFB74D",  # สีส้มอ่อน (Light Orange) - วัสดุมวลรวม
        "short_name": "GSB",
        "english_name": "Aggregate Subbase"
    },
    
    # ============ วัสดุคัดเลือก (Selected Material) - ทราย ============
    "วัสดุคัดเลือก ก": {
        "layer_coeff": 0.08,
        "drainage_coeff": 1.0,
        "mr_psi": 14504,
        "mr_mpa": 100,
        "layer_type": "selected",
        "color": "#FFF176",  # สีเหลือง (Yellow) - ทราย
        "short_name": "SM-A",
        "english_name": "Selected Material"
    },
    
    # ============ ไม่ใช้วัสดุ (Skip layer) ============
    "ไม่ใช้วัสดุคัดเลือก (ใช้ดินทางทรพ)": {
        "layer_coeff": 0.00,
        "drainage_coeff": 1.0,
        "mr_psi": 0,
        "mr_mpa": 0,
        "layer_type": "none",
        "color": "#D7CCC8",
        "short_name": "NONE",
        "english_name": "None"
    }
}

# ================================================================================
# RELIABILITY TABLE: Zr VALUES
# ================================================================================

RELIABILITY_ZR = {
    50: -0.000,
    60: -0.253,
    70: -0.524,
    75: -0.674,
    80: -0.841,
    85: -1.037,
    90: -1.282,
    91: -1.340,
    92: -1.405,
    93: -1.476,
    94: -1.555,
    95: -1.645,
    96: -1.751,
    97: -1.881,
    98: -2.054,
    99: -2.327,
    99.9: -3.090
}

# ================================================================================
# CORE CALCULATION FUNCTIONS
# ================================================================================

def aashto_1993_equation(SN: float, W18: float, Zr: float, So: float, 
                          delta_psi: float, Mr: float) -> float:
    """
    AASHTO 1993 Main Design Equation for Flexible Pavement
    
    สมการออกแบบหลักของ AASHTO 1993
    
    log₁₀(W₁₈) = Zr×So + 9.36×log₁₀(SN+1) - 0.20 
                 + log₁₀(ΔPSI/(4.2-1.5)) / (0.4 + 1094/(SN+1)^5.19)
                 + 2.32×log₁₀(Mr) - 8.07
    """
    log_W18 = np.log10(W18)
    
    term1 = Zr * So
    term2 = 9.36 * np.log10(SN + 1) - 0.20
    
    numerator = np.log10(delta_psi / (4.2 - 1.5))
    denominator = 0.4 + (1094 / ((SN + 1) ** 5.19))
    term3 = numerator / denominator
    
    term4 = 2.32 * np.log10(Mr) - 8.07
    
    right_side = term1 + term2 + term3 + term4
    
    return right_side - log_W18


def calculate_sn_for_layer(W18: float, Zr: float, So: float, 
                            delta_psi: float, Mr: float) -> float:
    """
    Calculate required SN for a given subgrade/layer modulus
    
    คำนวณค่า SN ที่ต้องการสำหรับค่า Mr ที่กำหนด
    """
    def f(SN):
        return aashto_1993_equation(SN, W18, Zr, So, delta_psi, Mr)
    
    try:
        SN_required = brentq(f, 0.01, 25.0, xtol=1e-6, maxiter=100)
        return round(SN_required, 2)
    except ValueError:
        return None


def calculate_w18_supported(SN: float, Zr: float, So: float, 
                            delta_psi: float, Mr: float) -> float:
    """
    Calculate W18 that can be supported by a given SN
    
    คำนวณค่า W₁₈ ที่โครงสร้างรองรับได้จากค่า SN ที่ออกแบบ
    """
    term1 = Zr * So
    term2 = 9.36 * np.log10(SN + 1) - 0.20
    
    numerator = np.log10(delta_psi / (4.2 - 1.5))
    denominator = 0.4 + (1094 / ((SN + 1) ** 5.19))
    term3 = numerator / denominator
    
    term4 = 2.32 * np.log10(Mr) - 8.07
    
    log_W18 = term1 + term2 + term3 + term4
    
    W18_supported = 10 ** log_W18
    
    return W18_supported


def calculate_layer_thicknesses(W18: float, Zr: float, So: float, delta_psi: float,
                                 subgrade_mr: float, layers: list) -> dict:
    """
    Calculate minimum thickness for each layer using AASHTO 1993 method
    
    คำนวณความหนาขั้นต่ำของแต่ละชั้น ตามวิธี AASHTO 1993
    """
    results = {
        'layers': [],
        'sn_values': [],
        'subgrade_mr': subgrade_mr,
        'total_sn_required': None,
        'total_sn_provided': 0
    }
    
    # Get active layers (exclude "ไม่ใช้วัสดุ")
    active_layers = [l for l in layers if l['material'] != "ไม่ใช้วัสดุคัดเลือก (ใช้ดินทางทรพ)"]
    
    if not active_layers:
        return results
    
    num_layers = len(active_layers)
    sn_values = []
    
    for i in range(num_layers):
        if i == num_layers - 1:
            mr_below = subgrade_mr
        else:
            mat_below = MATERIALS[active_layers[i + 1]['material']]
            mr_below = mat_below['mr_psi']
        
        sn_i = calculate_sn_for_layer(W18, Zr, So, delta_psi, mr_below)
        sn_values.append({
            'layer_index': i + 1,
            'mr_below': mr_below,
            'sn_required': sn_i
        })
    
    results['sn_values'] = sn_values
    results['total_sn_required'] = calculate_sn_for_layer(W18, Zr, So, delta_psi, subgrade_mr)
    
    cumulative_sn = 0
    
    for i, layer in enumerate(active_layers):
        mat = MATERIALS[layer['material']]
        a_i = mat['layer_coeff']
        m_i = layer.get('drainage_coeff', 1.0)
        
        sn_required_at_layer = sn_values[i]['sn_required'] if sn_values[i]['sn_required'] else 0
        
        if a_i > 0 and m_i > 0:
            remaining_sn = max(0, sn_required_at_layer - cumulative_sn)
            min_thickness_inch = remaining_sn / (a_i * m_i)
            min_thickness_cm = min_thickness_inch * 2.54
        else:
            min_thickness_inch = 0
            min_thickness_cm = 0
        
        design_thickness_cm = layer['thickness_cm']
        design_thickness_inch = design_thickness_cm / 2.54
        
        sn_contribution = a_i * design_thickness_inch * m_i
        cumulative_sn += sn_contribution
        
        is_ok = design_thickness_cm >= min_thickness_cm
        
        results['layers'].append({
            'layer_no': i + 1,
            'material': layer['material'],
            'short_name': mat['short_name'],
            'english_name': mat.get('english_name', mat['short_name']),
            'mr_psi': mat['mr_psi'],
            'mr_mpa': mat['mr_mpa'],
            'a_i': a_i,
            'm_i': m_i,
            'sn_required_at_layer': sn_required_at_layer,
            'min_thickness_inch': round(min_thickness_inch, 2),
            'min_thickness_cm': round(min_thickness_cm, 1),
            'design_thickness_cm': design_thickness_cm,
            'design_thickness_inch': round(design_thickness_inch, 2),
            'sn_contribution': round(sn_contribution, 4),
            'cumulative_sn': round(cumulative_sn, 2),
            'is_ok': is_ok,
            'color': mat['color']
        })
    
    results['total_sn_provided'] = round(cumulative_sn, 2)
    
    return results


def check_design(sn_required: float, sn_provided: float) -> dict:
    """Check if design is adequate"""
    if sn_required is None:
        return {
            'status': 'ERROR',
            'passed': False,
            'message': 'Cannot calculate SN_required',
            'safety_margin': None
        }
    
    safety_margin = sn_provided - sn_required
    passed = sn_provided >= sn_required
    
    return {
        'status': 'OK' if passed else 'NG',
        'passed': passed,
        'safety_margin': round(safety_margin, 2),
        'message': f"SN_provided ({sn_provided:.2f}) {'≥' if passed else '<'} SN_required ({sn_required:.2f})"
    }



# ================================================================================
# VISUALIZATION FUNCTIONS - ปรับปรุงใหม่ตามรูปแบบ Rigid Pavement
# ================================================================================

def plot_pavement_section(layers_result: list, subgrade_mr: float = None, 
                          subgrade_cbr: float = None) -> plt.Figure:
    """
    Draw vertical pavement section diagram - รูปแบบเดียวกับ Rigid Pavement
    
    Layout:
    - ซ้าย: ชื่อวัสดุ (Material name)
    - กลาง: ความหนา (Thickness)
    - ขวา: E = xxx MPa
    - ขวาสุด: เส้นแสดงความหนารวม (Total thickness arrow)
    - ล่าง: กรอบ Total Pavement Thickness
    """
    
    plt.rcParams['font.family'] = 'DejaVu Sans'
    
    if not layers_result:
        fig, ax = plt.subplots(figsize=(12, 8))
        ax.text(0.5, 0.5, 'No layers defined', ha='center', va='center', fontsize=14)
        ax.axis('off')
        return fig
    
    # คำนวณความหนารวม
    total_thickness = sum([l['design_thickness_cm'] for l in layers_result])
    num_layers = len(layers_result)
    
    # สร้าง figure
    fig_height = max(8, total_thickness * 0.12 + 3)
    fig, ax = plt.subplots(figsize=(12, fig_height))
    
    # Scale factor - ปรับให้ความสูงสัมพันธ์กับความหนาจริง
    # กำหนด minimum height เพื่อให้อ่านข้อความได้
    min_layer_height = 0.8
    scale = 0.08
    
    # คำนวณความสูงแต่ละชั้น (สัมพันธ์กับความหนาจริง)
    layer_heights = []
    for layer in layers_result:
        h = max(layer['design_thickness_cm'] * scale, min_layer_height)
        layer_heights.append(h)
    total_height = sum(layer_heights)
    
    # ตำแหน่ง x
    layer_x_start = 3.5
    layer_width = 2.5
    layer_x_end = layer_x_start + layer_width
    
    # วาดแต่ละชั้น
    current_y = total_height
    
    for i, layer in enumerate(layers_result):
        thickness_cm = layer['design_thickness_cm']
        layer_height = layer_heights[i]
        color = layer.get('color', '#888888')
        english_name = layer.get('english_name', layer['short_name'])
        mr_mpa = layer.get('mr_mpa', 0)
        
        # วาดสี่เหลี่ยม
        rect = mpatches.Rectangle(
            (layer_x_start, current_y - layer_height),
            layer_width, layer_height,
            facecolor=color,
            edgecolor='black',
            linewidth=1.5
        )
        ax.add_patch(rect)
        
        # กำหนดสีข้อความ
        dark_colors = ['#1C1C1C', '#2C2C2C', '#78909C', '#607D8B', '#795548', 
                       '#8D6E63', '#5D4037', '#6D4C41', '#455A64']
        text_color = 'white' if color in dark_colors else 'black'
        
        # ตำแหน่งกลางชั้น
        layer_center_y = current_y - layer_height / 2
        
        # ข้อความกลางชั้น: ความหนา
        ax.text(
            layer_x_start + layer_width / 2,
            layer_center_y,
            f'{thickness_cm:.0f} cm',
            ha='center', va='center',
            fontsize=12, fontweight='bold',
            color=text_color
        )
        
        # ซ้าย: ชื่อวัสดุ
        ax.text(
            layer_x_start - 0.2,
            layer_center_y,
            english_name,
            ha='right', va='center',
            fontsize=10, fontweight='bold',
            color='#1565C0'
        )
        
        # ขวา: E = xxx MPa
        ax.text(
            layer_x_end + 0.2,
            layer_center_y,
            f'E = {mr_mpa:,} MPa',
            ha='left', va='center',
            fontsize=10,
            color='#546E7A'
        )
        
        current_y -= layer_height
    
    # เส้นแสดงความหนารวม (ขวาสุด)
    arrow_x = layer_x_end + 1.8
    top_y = total_height
    bottom_y = 0
    
    ax.annotate(
        '', 
        xy=(arrow_x, bottom_y), 
        xytext=(arrow_x, top_y),
        arrowprops=dict(
            arrowstyle='<->',
            color='#E65100',
            lw=2,
            shrinkA=0,
            shrinkB=0
        )
    )
    
    # ข้อความ Total
    ax.text(
        arrow_x + 0.15,
        (top_y + bottom_y) / 2,
        f'Total\n{total_thickness:.0f} cm',
        ha='left', va='center',
        fontsize=11, fontweight='bold',
        color='#E65100'
    )
    
    # กรอบ Total Pavement Thickness (ล่าง)
    box_text = f'Total Pavement Thickness: {total_thickness:.0f} cm'
    box_y = -0.6
    
    ax.text(
        layer_x_start + layer_width / 2,
        box_y,
        box_text,
        ha='center', va='center',
        fontsize=11, fontweight='bold',
        color='#1565C0',
        bbox=dict(
            boxstyle='round,pad=0.4',
            facecolor='#FFF9C4',
            edgecolor='#FFC107',
            linewidth=2
        )
    )
    
    # Title
    ax.text(
        layer_x_start + layer_width / 2,
        total_height + 0.5,
        'Pavement Structure',
        ha='center', va='center',
        fontsize=14, fontweight='bold',
        color='#37474F'
    )
    
    # ตั้งค่าแกน
    ax.set_xlim(0, arrow_x + 1.5)
    ax.set_ylim(box_y - 0.5, total_height + 0.8)
    ax.axis('off')
    
    plt.tight_layout()
    return fig


def plot_pavement_section_thai(layers_result: list, subgrade_mr: float = None,
                                subgrade_cbr: float = None) -> plt.Figure:
    """
    Draw vertical pavement section diagram - รูปแบบภาษาไทย
    
    Layout:
    - ซ้าย: ชื่อวัสดุ (ภาษาไทย)
    - กลาง: ความหนา (Thickness)
    - ขวา: E = xxx MPa
    - ขวาสุด: เส้นแสดงความหนารวม
    - ล่าง: กรอบ ความหนารวมโครงสร้างชั้นทาง
    """
    
    # ตั้งค่า Thai font
    thai_font_path = '/usr/share/fonts/truetype/tlwg/Garuda.ttf'
    try:
        thai_font = fm.FontProperties(fname=thai_font_path)
        thai_font_bold = fm.FontProperties(fname=thai_font_path, weight='bold')
    except:
        thai_font = fm.FontProperties()
        thai_font_bold = fm.FontProperties(weight='bold')
    
    plt.rcParams['font.family'] = 'DejaVu Sans'
    
    if not layers_result:
        fig, ax = plt.subplots(figsize=(12, 8))
        ax.text(0.5, 0.5, 'ไม่มีข้อมูลชั้นทาง', ha='center', va='center', 
                fontsize=14, fontproperties=thai_font)
        ax.axis('off')
        return fig
    
    # คำนวณความหนารวม
    total_thickness = sum([l['design_thickness_cm'] for l in layers_result])
    num_layers = len(layers_result)
    
    # สร้าง figure
    fig_height = max(8, total_thickness * 0.12 + 3)
    fig, ax = plt.subplots(figsize=(14, fig_height))
    
    # Scale factor - ปรับให้ความสูงสัมพันธ์กับความหนาจริง
    min_layer_height = 0.8
    scale = 0.08
    
    # คำนวณความสูงแต่ละชั้น
    layer_heights = []
    for layer in layers_result:
        h = max(layer['design_thickness_cm'] * scale, min_layer_height)
        layer_heights.append(h)
    total_height = sum(layer_heights)
    
    # ตำแหน่ง x - เลื่อนขวาเพื่อเว้นที่สำหรับชื่อภาษาไทย
    layer_x_start = 5.0
    layer_width = 2.5
    layer_x_end = layer_x_start + layer_width
    
    # วาดแต่ละชั้น
    current_y = total_height
    
    for i, layer in enumerate(layers_result):
        thickness_cm = layer['design_thickness_cm']
        layer_height = layer_heights[i]
        color = layer.get('color', '#888888')
        thai_name = layer.get('material', layer['short_name'])
        mr_mpa = layer.get('mr_mpa', 0)
        
        # วาดสี่เหลี่ยม
        rect = mpatches.Rectangle(
            (layer_x_start, current_y - layer_height),
            layer_width, layer_height,
            facecolor=color,
            edgecolor='black',
            linewidth=1.5
        )
        ax.add_patch(rect)
        
        # กำหนดสีข้อความ
        dark_colors = ['#1C1C1C', '#2C2C2C', '#78909C', '#607D8B', '#795548', 
                       '#8D6E63', '#5D4037', '#6D4C41', '#455A64']
        text_color = 'white' if color in dark_colors else 'black'
        
        # ตำแหน่งกลางชั้น
        layer_center_y = current_y - layer_height / 2
        
        # ข้อความกลางชั้น: ความหนา
        ax.text(
            layer_x_start + layer_width / 2,
            layer_center_y,
            f'{thickness_cm:.0f} cm',
            ha='center', va='center',
            fontsize=20, fontweight='bold',
            color=text_color
        )
        
        # ซ้าย: ชื่อวัสดุภาษาไทย
        ax.text(
            layer_x_start - 0.2,
            layer_center_y,
            thai_name,
            ha='right', va='center',
            fontsize=20,
            fontproperties=thai_font_bold,
            color='#1565C0'
        )
        
        # ขวา: E = xxx MPa
        ax.text(
            layer_x_end + 0.2,
            layer_center_y,
            f'E = {mr_mpa:,} MPa',
            ha='left', va='center',
            fontsize=10,
            color='#546E7A'
        )
        
        current_y -= layer_height
    
    # เส้นแสดงความหนารวม (ขวาสุด)
    arrow_x = layer_x_end + 1.8
    top_y = total_height
    bottom_y = 0
    
    ax.annotate(
        '', 
        xy=(arrow_x, bottom_y), 
        xytext=(arrow_x, top_y),
        arrowprops=dict(
            arrowstyle='<->',
            color='#E65100',
            lw=2,
            shrinkA=0,
            shrinkB=0
        )
    )
    
    # ข้อความ Total
    ax.text(
        arrow_x + 0.15,
        (top_y + bottom_y) / 2,
        f'รวม\n{total_thickness:.0f} cm',
        ha='left', va='center',
        fontsize=11, fontweight='bold',
        fontproperties=thai_font_bold,
        color='#E65100'
    )
    
    # กรอบ Total Pavement Thickness (ล่าง)
    box_text = f'ความหนารวมโครงสร้างชั้นทาง: {total_thickness:.0f} cm'
    box_y = -0.6
    
    ax.text(
        layer_x_start + layer_width / 2,
        box_y,
        box_text,
        ha='center', va='center',
        fontsize=11, fontweight='bold',
        fontproperties=thai_font_bold,
        color='#1565C0',
        bbox=dict(
            boxstyle='round,pad=0.4',
            facecolor='#FFF9C4',
            edgecolor='#FFC107',
            linewidth=2
        )
    )
    
    # Title
    ax.text(
        layer_x_start + layer_width / 2,
        total_height + 0.5,
        'รูปตัดโครงสร้างชั้นทาง',
        ha='center', va='center',
        fontsize=14, fontweight='bold',
        fontproperties=thai_font_bold,
        color='#37474F'
    )
    
    # ตั้งค่าแกน
    ax.set_xlim(0, arrow_x + 1.5)
    ax.set_ylim(box_y - 0.5, total_height + 0.8)
    ax.axis('off')
    
    plt.tight_layout()
    return fig

def get_figure_as_bytes(fig: plt.Figure) -> BytesIO:
    """Convert matplotlib figure to bytes"""
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    return buf


# ================================================================================
# WORD EXPORT FUNCTION
# ================================================================================

def create_word_report(project_title: str, inputs: dict, calc_results: dict,
                       design_check: dict, fig: plt.Figure) -> BytesIO:
    """Create Word document report with step-by-step calculations"""
    
    doc = Document()
    
    # ========================================
    # TITLE
    # ========================================
    title = doc.add_heading('รายงานการออกแบบ Flexible Pavement', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading(f'โครงการ: {project_title}', level=1)
    doc.add_paragraph(f'วันที่ออกแบบ: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    
    # ========================================
    # SECTION 1: Design Method
    # ========================================
    doc.add_heading('1. วิธีการออกแบบ', level=2)
    doc.add_paragraph(
        'การออกแบบโครงสร้างถนนใช้วิธี AASHTO 1993 Guide for Design of Pavement Structures '
        'ตามมาตรฐานกรมทางหลวง โดยใช้สมการหลักดังนี้:'
    )
    
    # Main equation
    eq_para = doc.add_paragraph()
    eq_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq_run = eq_para.add_run(
        'log₁₀(W₁₈) = Zᵣ·Sₒ + 9.36·log₁₀(SN+1) - 0.20 + '
        'log₁₀(ΔPSI/2.7) / [0.4 + 1094/(SN+1)⁵·¹⁹] + 2.32·log₁₀(Mᵣ) - 8.07'
    )
    eq_run.italic = True
    eq_run.font.size = Pt(11)
    
    # ========================================
    # SECTION 2: Input Parameters
    # ========================================
    doc.add_heading('2. ข้อมูลนำเข้า (Design Inputs)', level=2)
    
    input_table = doc.add_table(rows=1, cols=3)
    input_table.style = 'Table Grid'
    
    headers = ['พารามิเตอร์', 'ค่า', 'หน่วย']
    for i, header in enumerate(headers):
        cell = input_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    input_data = [
        ('Design ESALs (W₁₈)', f'{inputs["W18"]:,.0f}', '18-kip ESAL'),
        ('Reliability (R)', f'{inputs["reliability"]}', '%'),
        ('Standard Normal Deviate (Zᵣ)', f'{inputs["Zr"]:.3f}', '-'),
        ('Overall Standard Deviation (Sₒ)', f'{inputs["So"]:.2f}', '-'),
        ('Initial Serviceability (P₀)', f'{inputs["P0"]:.1f}', '-'),
        ('Terminal Serviceability (Pₜ)', f'{inputs["Pt"]:.1f}', '-'),
        ('ΔPSI = P₀ - Pₜ', f'{inputs["delta_psi"]:.1f}', '-'),
        ('Subgrade CBR', f'{inputs.get("CBR", "-")}', '%'),
        ('Subgrade Mᵣ = 1500 × CBR', f'{inputs["Mr"]:,.0f}', 'psi'),
    ]
    
    for param, value, unit in input_data:
        row = input_table.add_row()
        row.cells[0].text = param
        row.cells[1].text = value
        row.cells[2].text = unit
    
    # ========================================
    # SECTION 3: Material Properties
    # ========================================
    doc.add_heading('3. คุณสมบัติวัสดุชั้นทาง', level=2)
    
    mat_table = doc.add_table(rows=1, cols=6)
    mat_table.style = 'Table Grid'
    
    mat_headers = ['ชั้น', 'วัสดุ', 'aᵢ', 'mᵢ', 'Mᵣ (psi)', 'E (MPa)']
    for i, header in enumerate(mat_headers):
        cell = mat_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    for layer in calc_results['layers']:
        row = mat_table.add_row()
        row.cells[0].text = str(layer['layer_no'])
        row.cells[1].text = layer['material']
        row.cells[2].text = f'{layer["a_i"]:.2f}'
        row.cells[3].text = f'{layer["m_i"]:.2f}'
        row.cells[4].text = f'{layer["mr_psi"]:,}'
        row.cells[5].text = f'{layer["mr_mpa"]:,}'
    
    # ========================================
    # SECTION 4: Step-by-Step Calculation
    # ========================================
    doc.add_heading('4. ขั้นตอนการคำนวณความหนาชั้นทาง', level=2)
    
    doc.add_paragraph(
        'การคำนวณความหนาขั้นต่ำของแต่ละชั้น ใช้หลักการว่า Structural Number (SN) '
        'ที่จุดใดๆ ต้องมากกว่าหรือเท่ากับ SN ที่ต้องการ โดยคำนวณจากค่า Mᵣ ของชั้นถัดไป'
    )
    
    for layer in calc_results['layers']:
        # Layer header
        doc.add_heading(f'ชั้นที่ {layer["layer_no"]}: {layer["material"]}', level=3)
        
        # Material properties
        doc.add_paragraph(f'ข้อมูลวัสดุ:')
        props_para = doc.add_paragraph()
        props_para.add_run(f'    • Mᵣ = {layer["mr_psi"]:,} psi = {layer["mr_mpa"]:,} MPa\n')
        props_para.add_run(f'    • Layer Coefficient (a{layer["layer_no"]}) = {layer["a_i"]:.2f}\n')
        props_para.add_run(f'    • Drainage Coefficient (m{layer["layer_no"]}) = {layer["m_i"]:.2f}')
        
        # SN calculation
        doc.add_paragraph(f'การคำนวณ SN:')
        sn_para = doc.add_paragraph()
        sn_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sn_run = sn_para.add_run(f'จากสมการ AASHTO 1993:  SN{layer["layer_no"]} = {layer["sn_required_at_layer"]:.2f}')
        sn_run.bold = True
        
        # Thickness calculation
        doc.add_paragraph(f'การคำนวณความหนาขั้นต่ำ:')
        
        if layer['layer_no'] == 1:
            formula_para = doc.add_paragraph()
            formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            formula_text = f'D₁ ≥ SN₁ / (a₁ × m₁) = {layer["sn_required_at_layer"]:.2f} / ({layer["a_i"]:.2f} × {layer["m_i"]:.2f})'
            formula_para.add_run(formula_text).italic = True
        else:
            prev_sn = calc_results['layers'][layer['layer_no']-2]['cumulative_sn']
            
            formula_para = doc.add_paragraph()
            formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            formula_text = f'D{layer["layer_no"]} ≥ (SN{layer["layer_no"]} - SNₚᵣₑᵥ) / (a{layer["layer_no"]} × m{layer["layer_no"]}) = ({layer["sn_required_at_layer"]:.2f} - {prev_sn:.2f}) / ({layer["a_i"]:.2f} × {layer["m_i"]:.2f})'
            formula_para.add_run(formula_text).italic = True
        
        # Results
        result_para = doc.add_paragraph()
        result_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        result_para.add_run(f'D{layer["layer_no"]}(min) = {layer["min_thickness_inch"]:.2f} นิ้ว = {layer["min_thickness_cm"]:.1f} ซม.').bold = True
        
        # Design thickness selection
        doc.add_paragraph(f'เลือกใช้ความหนา:')
        design_para = doc.add_paragraph()
        design_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        design_para.add_run(f'D{layer["layer_no"]}(design) = {layer["design_thickness_cm"]:.0f} ซม. ({layer["design_thickness_inch"]:.2f} นิ้ว)').bold = True
        
        # SN contribution
        doc.add_paragraph(f'SN contribution:')
        contrib_para = doc.add_paragraph()
        contrib_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contrib_text = f'ΔSN{layer["layer_no"]} = a{layer["layer_no"]} × D{layer["layer_no"]} × m{layer["layer_no"]} = {layer["a_i"]:.2f} × {layer["design_thickness_inch"]:.2f} × {layer["m_i"]:.2f} = {layer["sn_contribution"]:.3f}'
        contrib_para.add_run(contrib_text)
        
        # Cumulative SN
        cum_para = doc.add_paragraph()
        cum_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cum_para.add_run(f'ΣSN = {layer["cumulative_sn"]:.2f}').bold = True
        
        # Check status
        status_text = '✓ OK' if layer['is_ok'] else '✗ NG - ต้องเพิ่มความหนา'
        status_para = doc.add_paragraph()
        status_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        status_run = status_para.add_run(f'สถานะ: {status_text}')
        status_run.bold = True
        
        doc.add_paragraph()  # Spacing
    
    # ========================================
    # SECTION 5: SN Summary Table
    # ========================================
    doc.add_heading('5. ตารางสรุปการคำนวณ Structural Number', level=2)
    
    sn_table = doc.add_table(rows=1, cols=8)
    sn_table.style = 'Table Grid'
    
    sn_headers = ['ชั้น', 'วัสดุ', 'aᵢ', 'mᵢ', 'Dᵢ (นิ้ว)', 'Dᵢ (ซม.)', 'ΔSNᵢ', 'ΣSN']
    for i, header in enumerate(sn_headers):
        cell = sn_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    for layer in calc_results['layers']:
        row = sn_table.add_row()
        row.cells[0].text = str(layer['layer_no'])
        row.cells[1].text = layer['material']
        row.cells[2].text = f'{layer["a_i"]:.2f}'
        row.cells[3].text = f'{layer["m_i"]:.2f}'
        row.cells[4].text = f'{layer["design_thickness_inch"]:.2f}'
        row.cells[5].text = f'{layer["design_thickness_cm"]:.0f}'
        row.cells[6].text = f'{layer["sn_contribution"]:.3f}'
        row.cells[7].text = f'{layer["cumulative_sn"]:.2f}'
    
    # Formula
    doc.add_paragraph()
    formula_p = doc.add_paragraph()
    formula_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_p.add_run('สูตร: SN = Σ(aᵢ × Dᵢ × mᵢ)').italic = True
    
    # ========================================
    # SECTION 6: Design Verification
    # ========================================
    doc.add_heading('6. ผลการตรวจสอบการออกแบบ', level=2)
    
    result_table = doc.add_table(rows=4, cols=2)
    result_table.style = 'Table Grid'
    
    result_data = [
        ('SN Required (จากสมการ AASHTO)', f'{calc_results["total_sn_required"]:.2f}'),
        ('SN Provided (จากชั้นทาง)', f'{calc_results["total_sn_provided"]:.2f}'),
        ('Safety Margin', f'{design_check["safety_margin"]:.2f}'),
        ('ผลการตรวจสอบ', 'ผ่าน (OK)' if design_check['passed'] else 'ไม่ผ่าน (NG)'),
    ]
    
    for i, (param, value) in enumerate(result_data):
        result_table.rows[i].cells[0].text = param
        result_table.rows[i].cells[1].text = value
    
    # Conclusion
    doc.add_paragraph()
    if design_check['passed']:
        conclusion = doc.add_paragraph()
        conclusion.add_run(
            f'สรุป: การออกแบบผ่านเกณฑ์ เนื่องจาก SN_provided ({calc_results["total_sn_provided"]:.2f}) ≥ '
            f'SN_required ({calc_results["total_sn_required"]:.2f})'
        ).bold = True
    else:
        conclusion = doc.add_paragraph()
        conclusion.add_run(
            f'สรุป: การออกแบบไม่ผ่านเกณฑ์ กรุณาปรับเพิ่มความหนาชั้นทาง'
        ).bold = True
    
    # ========================================
    # SECTION 7: Figure
    # ========================================
    doc.add_heading('7. ภาพตัดขวางโครงสร้างถนน', level=2)
    fig_bytes = get_figure_as_bytes(fig)
    doc.add_picture(fig_bytes, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ========================================
    # Save document
    # ========================================
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes


# ================================================================================
# STREAMLIT USER INTERFACE
# ================================================================================

def main():
    """Main Streamlit application"""
    
    # Header
    st.title("🛣️ AASHTO 1993 Flexible Pavement Design")
    st.markdown("""
    **แอปพลิเคชันออกแบบโครงสร้างทางแบบยืดหยุ่น ตามมาตรฐานกรมทางหลวง (DOH Thailand)**
    
    ✅ คำนวณความหนาขั้นต่ำแต่ละชั้น | ✅ แสดงขั้นตอนการคำนวณ | ✅ Export รายงาน Word
    """)
    
    st.markdown("---")
    
    # ========================================
    # SIDEBAR: Project Info & Material Database
    # ========================================
    with st.sidebar:
        st.header("📋 ข้อมูลโครงการ")
        project_title = st.text_input("ชื่อโครงการ", value="โครงการออกแบบถนน")
        
        st.markdown("---")
        
        # ตัวเลือกภาษาสำหรับรูปภาพ
        st.header("🖼️ ตั้งค่ารูปภาพ")
        figure_language = st.radio(
            "ภาษาในรูปภาพ",
            options=["English", "ภาษาไทย"],
            index=0,
            help="เลือกภาษาสำหรับแสดงในรูปตัดขวาง"
        )
        
        st.markdown("---")
        st.header("📚 ฐานข้อมูลวัสดุ (ทล.)")
        
        with st.expander("ดูค่า สปส. วัสดุทั้งหมด"):
            st.markdown("**ค่า สปส. สำหรับออกแบบ**")
            for mat_name, props in MATERIALS.items():
                if props['layer_coeff'] > 0:
                    st.markdown(f"**{mat_name}**")
                    st.markdown(f"- a = {props['layer_coeff']}, m = {props['drainage_coeff']}")
                    st.markdown(f"- MR = {props['mr_psi']:,} psi ({props['mr_mpa']:,} MPa)")
                    st.markdown("---")
    
    # ========================================
    # MAIN CONTENT: Two columns
    # ========================================
    col1, col2 = st.columns([1, 1])
    
    # ========================================
    # COLUMN 1: Design Inputs
    # ========================================
    with col1:
        st.header("📝 Design Inputs")
        
        # Traffic
        st.subheader("1️⃣ Traffic & Reliability")
        
        W18 = st.number_input(
            "Design ESALs (W₁₈)",
            min_value=100000,
            max_value=250000000,
            value=5000000,
            step=100000,
            format="%d",
            help="จำนวน 18-kip ESAL ตลอดอายุการใช้งาน (สูงสุด 250 ล้าน)"
        )
        
        # แสดงค่า ESAL เป็นล้าน (ภาษาไทย)
        esal_million = W18 / 1000000
        st.caption(f"💡 W₁₈ = **{esal_million:,.2f} ล้าน** ESALs")
        
        reliability = st.selectbox(
            "Reliability Level (R)",
            options=list(RELIABILITY_ZR.keys()),
            index=list(RELIABILITY_ZR.keys()).index(90),
        )
        Zr = RELIABILITY_ZR[reliability]
        st.info(f"Zᵣ = {Zr:.3f}")
        
        So = st.number_input(
            "Overall Standard Deviation (Sₒ)",
            min_value=0.30,
            max_value=0.60,
            value=0.45,
            step=0.01,
            format="%.2f"
        )
        
        # Serviceability
        st.subheader("2️⃣ Serviceability")
        
        col1a, col1b = st.columns(2)
        with col1a:
            P0 = st.number_input("P₀ (Initial)", min_value=3.0, max_value=5.0, value=4.2, step=0.1)
        with col1b:
            Pt = st.number_input("Pₜ (Terminal)", min_value=1.5, max_value=3.5, value=2.5, step=0.1)
        
        delta_psi = P0 - Pt
        st.success(f"**ΔPSI = {delta_psi:.1f}**")
        
        # Subgrade
        st.subheader("3️⃣ Subgrade (ดินเดิม/ดินถม)")
        
        CBR = st.number_input(
            "CBR (%)",
            min_value=1.0,
            max_value=30.0,
            value=5.0,
            step=0.5,
            help="ค่า CBR ของดินเดิมหรือดินถมคันทาง"
        )
        
        # Mr = 1500 × CBR (ตามมาตรฐาน ทล.)
        Mr = int(1500 * CBR)
        st.info(f"**Mᵣ = 1,500 × CBR = 1,500 × {CBR:.1f} = {Mr:,} psi**")
    
    # ========================================
    # COLUMN 2: Layer Configuration
    # ========================================
    with col2:
        st.header("🏗️ Layer Configuration")
        
        # จำนวนชั้นทาง
        num_layers = st.slider(
            "จำนวนชั้นทาง",
            min_value=2,
            max_value=6,
            value=4,
            help="เลือกจำนวนชั้นทาง (2-6 ชั้น)"
        )
        
        # สร้าง list วัสดุทั้งหมด (ยกเว้น "ไม่ใช้")
        all_materials = [m for m, p in MATERIALS.items() if p['layer_type'] != 'none']
        
        # สร้าง list วัสดุสำหรับชั้นที่ 2-6 (รวม "ไม่ใช้ชั้นนี้")
        optional_materials = all_materials + ["ไม่ใช้ชั้นนี้"]
        
        # เก็บข้อมูลชั้นทาง
        layer_data = []
        
        # เก็บ placeholders สำหรับแสดงสถานะ
        status_placeholders = {}
        
        # ========== ชั้นที่ 1: ผิวทาง (บังคับ) ==========
        st.subheader("4️⃣ ชั้นที่ 1: ผิวทาง (Surface)")
        
        surface_materials = [m for m, p in MATERIALS.items() if p['layer_type'] == 'surface']
        
        layer1_mat = st.selectbox(
            "เลือกวัสดุ",
            options=surface_materials,
            index=0,
            key="layer1_mat"
        )
        
        col_a, col_b = st.columns(2)
        with col_a:
            layer1_thick = st.number_input(
                "ความหนา (cm)", min_value=1.0, max_value=30.0, value=5.0, step=1.0,
                key="layer1_thick"
            )
        with col_b:
            layer1_m = st.number_input(
                "m₁", min_value=0.5, max_value=1.5, value=1.0, step=0.05,
                key="layer1_m"
            )
        
        mat_props = MATERIALS[layer1_mat]
        st.caption(f"a₁ = {mat_props['layer_coeff']}, E = {mat_props['mr_mpa']:,} MPa")
        
        # Placeholder สำหรับแสดงสถานะชั้นที่ 1
        status_placeholders[1] = st.empty()
        
        layer_data.append({
            'material': layer1_mat,
            'thickness_cm': layer1_thick,
            'drainage_coeff': layer1_m
        })
        
        # ========== ชั้นที่ 2-6: เลือกวัสดุได้ทุกชนิด ==========
        default_materials = [
            "พื้นทางซีเมนต์ CTB",
            "รองพื้นทางวัสดุมวลรวม CBR 25%",
            "วัสดุคัดเลือก ก",
            "วัสดุคัดเลือก ก",
            "วัสดุคัดเลือก ก"
        ]
        default_thickness = [15.0, 15.0, 30.0, 30.0, 30.0]
        
        for i in range(2, num_layers + 1):
            st.markdown("---")
            layer_icons = ['5️⃣', '6️⃣', '7️⃣', '8️⃣', '9️⃣']
            st.subheader(f"{layer_icons[i-2]} ชั้นที่ {i}")
            
            # Default index
            default_idx = all_materials.index(default_materials[i-2]) if default_materials[i-2] in all_materials else 0
            
            layer_mat = st.selectbox(
                f"เลือกวัสดุชั้นที่ {i}",
                options=all_materials,
                index=min(default_idx, len(all_materials)-1),
                key=f"layer{i}_mat"
            )
            
            col_c, col_d = st.columns(2)
            with col_c:
                layer_thick = st.number_input(
                    "ความหนา (cm)",
                    min_value=1.0, max_value=150.0, value=default_thickness[i-2], step=5.0,
                    key=f"layer{i}_thick"
                )
            with col_d:
                layer_m = st.number_input(
                    f"m{i}",
                    min_value=0.5, max_value=1.5, value=1.0, step=0.05,
                    key=f"layer{i}_m"
                )
            
            mat_props = MATERIALS[layer_mat]
            st.caption(f"a{i} = {mat_props['layer_coeff']}, E = {mat_props['mr_mpa']:,} MPa")
            
            # Placeholder สำหรับแสดงสถานะชั้นที่ i
            status_placeholders[i] = st.empty()
            
            layer_data.append({
                'material': layer_mat,
                'thickness_cm': layer_thick,
                'drainage_coeff': layer_m
            })
    
    # ========================================
    # BUILD LAYERS LIST
    # ========================================
    layers = layer_data
    
    # Store inputs
    inputs = {
        'W18': W18, 'reliability': reliability, 'Zr': Zr, 'So': So,
        'P0': P0, 'Pt': Pt, 'delta_psi': delta_psi, 'CBR': CBR, 'Mr': Mr
    }
    
    # ========================================
    # CALCULATION & RESULTS
    # ========================================
    st.markdown("---")
    st.header("📊 ผลการคำนวณ (Calculation Results)")
    
    # Calculate layer thicknesses
    calc_results = calculate_layer_thicknesses(W18, Zr, So, delta_psi, Mr, layers)
    
    # Design check
    design_check = check_design(
        calc_results['total_sn_required'],
        calc_results['total_sn_provided']
    )
    
    # ========================================
    # FILL STATUS PLACEHOLDERS
    # ========================================
    for layer in calc_results['layers']:
        layer_no = layer['layer_no']
        if layer_no in status_placeholders:
            with status_placeholders[layer_no]:
                if layer['is_ok']:
                    st.success(f"✅ ผ่าน (ต้องการ ≥ {layer['min_thickness_cm']:.1f} cm)")
                else:
                    shortage = layer['min_thickness_cm'] - layer['design_thickness_cm']
                    st.error(f"❌ ไม่ผ่าน (ต้องเพิ่มอีก {shortage:.1f} cm)")
    
    # ========================================
    # STEP-BY-STEP CALCULATION DISPLAY
    # ========================================
    st.subheader("🔢 ขั้นตอนการคำนวณความหนาแต่ละชั้น")
    
    for layer in calc_results['layers']:
        with st.container():
            # Header with colored background
            layer_status = "✅" if layer['is_ok'] else "❌"
            st.markdown(f"### {layer_status} ชั้นที่ {layer['layer_no']}: {layer['material']}")
            
            col_a, col_b = st.columns([1, 1])
            
            with col_a:
                st.markdown("**ข้อมูลวัสดุ:**")
                st.markdown(f"- E (MPa) = **{layer['mr_mpa']:,}**")
                st.markdown(f"- Mᵣ (psi) = **{layer['mr_psi']:,}**")
                st.markdown(f"- Layer Coefficient (a{layer['layer_no']}) = **{layer['a_i']:.2f}**")
                st.markdown(f"- Drain Coefficient (m{layer['layer_no']}) = **{layer['m_i']:.2f}**")
            
            with col_b:
                st.markdown("**จากสมการ AASHTO:**")
                
                # Show SN calculation
                sn_at_layer = layer['sn_required_at_layer']
                
                if layer['layer_no'] == 1:
                    st.latex(f"SN_{{{layer['layer_no']}}} = {sn_at_layer:.2f}")
                else:
                    st.latex(f"SN_{{{layer['layer_no']}}} = {sn_at_layer:.2f}")
            
            # Thickness calculation formula
            st.markdown("**คำนวณความหนาผิวทาง:**")
            
            if layer['layer_no'] == 1:
                # First layer formula
                st.latex(f"D_{{1}} \\geq \\frac{{SN_{{1}}}}{{a_{{1}} \\times m_{{1}}}} = \\frac{{{sn_at_layer:.2f}}}{{{layer['a_i']:.2f} \\times {layer['m_i']:.2f}}} = {layer['min_thickness_inch']:.2f} \\text{{ นิ้ว}}")
            else:
                # Subsequent layers
                prev_sn = calc_results['layers'][layer['layer_no']-2]['cumulative_sn']
                st.latex(f"D_{{{layer['layer_no']}}} \\geq \\frac{{SN_{{{layer['layer_no']}}} - SN_{{prev}}}}{{a_{{{layer['layer_no']}}} \\times m_{{{layer['layer_no']}}}}} = \\frac{{{sn_at_layer:.2f} - {prev_sn:.2f}}}{{{layer['a_i']:.2f} \\times {layer['m_i']:.2f}}} = {layer['min_thickness_inch']:.2f} \\text{{ นิ้ว}}")
            
            # Results table
            result_cols = st.columns(4)
            
            with result_cols[0]:
                st.metric("ความหนาขั้นต่ำ", f"{layer['min_thickness_cm']:.1f} cm")
            
            with result_cols[1]:
                st.metric("ความหนาที่เลือก", f"{layer['design_thickness_cm']:.0f} cm", 
                         delta=f"{layer['design_thickness_cm'] - layer['min_thickness_cm']:.1f} cm")
            
            with result_cols[2]:
                st.metric("SN contribution", f"{layer['sn_contribution']:.3f}")
            
            with result_cols[3]:
                st.metric("Cumulative SN", f"{layer['cumulative_sn']:.2f}")
            
            # Status
            if layer['is_ok']:
                st.success(f"✅ **OK** - ความหนาเพียงพอ ({layer['design_thickness_cm']:.0f} ≥ {layer['min_thickness_cm']:.1f} cm)")
            else:
                st.error(f"❌ **NG** - ต้องเพิ่มความหนาอีก {layer['min_thickness_cm'] - layer['design_thickness_cm']:.1f} cm")
            
            st.markdown("---")
    
    # ========================================
    # SUMMARY RESULTS
    # ========================================
    st.subheader("📈 สรุปผลการออกแบบ")
    
    res_col1, res_col2, res_col3, res_col4 = st.columns(4)
    
    with res_col1:
        st.metric("SN Required", f"{calc_results['total_sn_required']:.2f}")
    
    with res_col2:
        st.metric("SN Provided", f"{calc_results['total_sn_provided']:.2f}")
    
    with res_col3:
        if design_check['passed']:
            st.metric("Safety Margin", f"{design_check['safety_margin']:.2f}", delta="OK")
        else:
            st.metric("Safety Margin", f"{design_check['safety_margin']:.2f}", delta="NG", delta_color="inverse")
    
    with res_col4:
        if design_check['passed']:
            st.success("**PASS** ✅")
        else:
            st.error("**FAIL** ❌")
    
    # W18 Supported calculation
    w18_supported = calculate_w18_supported(
        calc_results['total_sn_provided'], Zr, So, delta_psi, Mr
    )
    w18_supported_million = w18_supported / 1_000_000
    w18_diff_percent = ((w18_supported - W18) / W18) * 100
    
    st.markdown("---")
    
    w18_col1, w18_col2 = st.columns(2)
    
    with w18_col1:
        st.metric(
            "W₁₈ ออกแบบ",
            f"{W18/1_000_000:,.2f} ล้าน"
        )
    
    with w18_col2:
        delta_str = f"{w18_diff_percent:+.1f}%"
        if w18_diff_percent >= 0:
            st.metric(
                "W₁₈ รองรับได้",
                f"{w18_supported_million:,.2f} ล้าน",
                delta=delta_str
            )
        else:
            st.metric(
                "W₁₈ รองรับได้",
                f"{w18_supported_million:,.2f} ล้าน",
                delta=delta_str,
                delta_color="inverse"
            )
    
    # Status message
    if design_check['passed']:
        st.success(f"✅ การออกแบบผ่านเกณฑ์: {design_check['message']}")
    else:
        st.error(f"❌ การออกแบบไม่ผ่าน: {design_check['message']}")
    
    # ========================================
    # PAVEMENT SECTION FIGURE
    # ========================================
    st.subheader("📐 ภาพตัดขวางโครงสร้างถนน")
    
    # เลือกฟังก์ชันวาดรูปตามภาษาที่เลือก
    if figure_language == "English":
        fig = plot_pavement_section(calc_results['layers'], Mr, CBR)
    else:
        fig = plot_pavement_section_thai(calc_results['layers'], Mr, CBR)
    
    st.pyplot(fig)
    
    # ========================================
    # SN CALCULATION TABLE
    # ========================================
    with st.expander("📋 ตารางสรุปการคำนวณ SN"):
        st.markdown("### SN Contribution Table")
        
        table_data = []
        for layer in calc_results['layers']:
            table_data.append({
                'ชั้น': layer['layer_no'],
                'วัสดุ': layer['short_name'],
                'aᵢ': layer['a_i'],
                'Dᵢ (cm)': layer['design_thickness_cm'],
                'Dᵢ (in)': layer['design_thickness_inch'],
                'mᵢ': layer['m_i'],
                'E (MPa)': layer['mr_mpa'],
                'SN contrib.': layer['sn_contribution'],
                'SN cumul.': layer['cumulative_sn']
            })
        
        st.table(table_data)
        
        st.markdown(f"""
        **สูตรการคำนวณ:**
        
        $$SN = \\sum_{{i=1}}^{{n}} a_i \\times D_i \\times m_i$$
        
        **ผลลัพธ์:**
        - SN_provided = {calc_results['total_sn_provided']:.2f}
        - SN_required = {calc_results['total_sn_required']:.2f}
        """)
    
    # ========================================
    # EXPORT
    # ========================================
    st.subheader("📄 Export Report")
    
    col_exp1, col_exp2 = st.columns(2)
    
    with col_exp1:
        if st.button("📝 Generate Word Report", type="primary"):
            with st.spinner("กำลังสร้างรายงาน..."):
                # ใช้รูปภาษาไทยในรายงาน Word
                fig_thai = plot_pavement_section_thai(calc_results['layers'], Mr, CBR)
                doc_bytes = create_word_report(
                    project_title, inputs, calc_results, design_check, fig_thai
                )
                
                st.download_button(
                    label="⬇️ Download Word Report",
                    data=doc_bytes,
                    file_name=f"AASHTO_Flexible_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    
    with col_exp2:
        fig_bytes = get_figure_as_bytes(fig)
        st.download_button(
            label="📸 Download Section Diagram (PNG)",
            data=fig_bytes,
            file_name=f"Pavement_Section_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
            mime="image/png"
        )
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; color: gray;'>
    <p>AASHTO 1993 Flexible Pavement Design Application v3.0</p>
    <p>พัฒนาตามมาตรฐานกรมทางหลวง (DOH Thailand)</p>
    </div>
    """, unsafe_allow_html=True)


# ================================================================================
# ENTRY POINT
# ================================================================================

if __name__ == "__main__":
    main()
