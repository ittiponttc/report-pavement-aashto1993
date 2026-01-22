import streamlit as st
from PIL import Image, ImageDraw, ImageFont
import io
import os
import math
from datetime import datetime

# ============================================
# Helper Functions
# ============================================

def interpolate_log_scale(pixel, p1, v1, p2, v2):
    """Interpolate value on logarithmic scale."""
    if v1 <= 0 or v2 <= 0: return 0
    log_v1 = math.log10(v1)
    log_v2 = math.log10(v2)
    t = (pixel - p1) / (p2 - p1) if p2 != p1 else 0
    log_v = log_v1 + t * (log_v2 - log_v1)
    return 10 ** log_v

def draw_arrow_fixed(draw, start, end, color, width=4, arrow_size=15):
    """Draw a line with an arrow head at the end."""
    draw.line([start, end], fill=color, width=width)
    
    # Calculate direction vector
    dx = end[0] - start[0]
    dy = end[1] - start[1]
    length = math.sqrt(dx*dx + dy*dy)
    if length > 0:
        dx /= length
        dy /= length
        
        # Perpendicular vector
        px = -dy
        py = dx
        
        # Arrow points
        x3 = end[0] - width*dx + arrow_size*dx  # Tip (adjust to not overshoot too much if needed)
        # Actually end point is tip
        x3 = end[0] 
        y3 = end[1]
        
        # Back base of triangle logic needs to be relative to tip
        # Move back from tip
        base_x = end[0] - arrow_size * dx
        base_y = end[1] - arrow_size * dy
        
        x4 = base_x + arrow_size * 0.5 * px
        y4 = base_y + arrow_size * 0.5 * py
        x5 = base_x - arrow_size * 0.5 * px
        y5 = base_y - arrow_size * 0.5 * py
        
        draw.polygon([(x3, y3), (x4, y4), (x5, y5)], fill=color)

# ============================================
# Word Report Generation
# ============================================

def generate_word_report(params, img1_bytes, img2_bytes=None):
    """Generate Word report with calculation steps for both graphs."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        return None, "ไม่พบ library python-docx กรุณาติดตั้งด้วย: pip install python-docx"
    
    doc = Document()
    
    # Set Thai font
    style = doc.styles['Normal']
    style.font.name = 'TH SarabunPSK'
    style.font.size = Pt(14)
    
    # Header
    title = doc.add_heading('รายการคำนวณ Corrected Modulus of Subgrade Reaction', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'วันที่: {datetime.now().strftime("%d/%m/%Y %H:%M")}', style='Normal').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # ---------------------------------------------------------
    # PART 1: Composite Modulus (k_infinity)
    # ---------------------------------------------------------
    doc.add_heading('ส่วนที่ 1: การหาค่า Composite Modulus (k∞)', level=1)
    
    # Table 1
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['พารามิเตอร์', 'ค่า', 'หน่วย']
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        
    data_1 = [
        ('Roadbed Soil Resilient Modulus (MR)', f"{params.get('MR', 0):,.0f}", 'psi'),
        ('Subbase Elastic Modulus (ESB)', f"{params.get('ESB', 0):,.0f}", 'psi'),
        ('Subbase Thickness (DSB)', f"{params.get('DSB', 0):.1f}", 'inches'),
        ('Composite Modulus (k∞)', f"{params.get('k_inf', 0):,.0f}", 'pci'),
    ]
    for p, v, u in data_1:
        row = table.add_row().cells
        row[0].text = p
        row[1].text = v
        row[2].text = u
        
    doc.add_paragraph()
    if img1_bytes:
        doc.add_picture(io.BytesIO(img1_bytes), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("รูปที่ 1: Nomograph for Composite Modulus (Figure 3.3)", style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------------------------------------------------------
    # PART 2: Loss of Support Correction
    # ---------------------------------------------------------
    doc.add_page_break()
    doc.add_heading('ส่วนที่ 2: การปรับแก้ค่า Loss of Support (LS)', level=1)
    
    # Table 2
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells2 = table2.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells2[i].text = h
        hdr_cells2[i].paragraphs[0].runs[0].bold = True
        
    data_2 = [
        ('Effective Modulus (k) - จากส่วนที่ 1', f"{params.get('k_inf', 0):,.0f}", 'pci'),
        ('Loss of Support Factor (LS)', f"{params.get('LS_factor', 0):.1f}", '-'),
        ('Corrected Modulus (k)', f"{params.get('k_corrected', 0):,.0f}", 'pci'),
    ]
    for p, v, u in data_2:
        row = table2.add_row().cells
        row[0].text = p
        row[1].text = v
        row[2].text = u

    doc.add_paragraph()
    if img2_bytes:
        doc.add_picture(io.BytesIO(img2_bytes), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("รูปที่ 2: Correction for Loss of Support (Figure 3.4)", style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Footer Note
    doc.add_paragraph()
    note = doc.add_paragraph("Reference: AASHTO Guide for Design of Pavement Structures 1993")
    note.style = 'List Bullet'

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, None

# ============================================
# Main Application
# ============================================

def main():
    st.set_page_config(page_title="AASHTO Rigid Pavement Calculator", page_icon="🛣️", layout="wide")
    
    st.title("🛣️ AASHTO 1993 Rigid Pavement Design Calculator")
    st.markdown("**เครื่องมือหาค่า k-value และปรับแก้ Loss of Support (LS)**")
    
    # Define Tabs
    tab1, tab2, tab3 = st.tabs(["📊 ขั้นตอนที่ 1: Composite k∞", "📉 ขั้นตอนที่ 2: Loss of Support", "📋 คู่มือการใช้งาน"])
    
    # Initialize Session State
    if 'k_inf_result' not in st.session_state:
        st.session_state.k_inf_result = 500
    if 'img1_bytes' not in st.session_state:
        st.session_state.img1_bytes = None
    if 'img2_bytes' not in st.session_state:
        st.session_state.img2_bytes = None
        
    # --- Config: LS Presets from User Images ---
    # Format: (x1, y1, x2, y2)
    LS_PRESETS = {
        0.0: (138, 715, 753, 84),
        0.5: (129, 728, 908, 0),
        1.0: (150, 718, 903, 84),
        1.5: (153, 721, 928, 138),
        2.0: (164, 718, 929, 220),
        3.0: (212, 719, 929, 328)
    }

    # =========================================================
    # TAB 1: Composite Modulus
    # =========================================================
    with tab1:
        st.header("1️⃣ หาค่า Composite Modulus of Subgrade Reaction (k∞)")
        
        uploaded_file = st.file_uploader("📂 อัปโหลดภาพ Figure 3.3 (Composite k)", type=['png', 'jpg', 'jpeg'], key='uploader_1')
        
        if uploaded_file is not None:
            image = Image.open(uploaded_file).convert("RGB")
            width, height = image.size
            img_draw = image.copy()
            draw = ImageDraw.Draw(img_draw)
            
            col_ctrl, col_img = st.columns([1, 2])
            
            with col_ctrl:
                st.subheader("⚙️ ปรับเส้นอ่านค่า")
                
                with st.expander("1. เส้น Turning Line (เขียว)", expanded=True):
                    # --- UPDATED DEFAULTS FOR GREEN LINE ---
                    # Using values from user image: 411, 339, 470, 397
                    gx1 = st.slider("X เริ่ม", 0, width, 411, key="gx1")
                    gy1 = st.slider("Y เริ่ม", 0, height, 339, key="gy1")
                    gx2 = st.slider("X จบ", 0, width, 470, key="gx2")
                    gy2 = st.slider("Y จบ", 0, height, 397, key="gy2")
                    
                    draw.line([(gx1, gy1), (gx2, gy2)], fill="green", width=5)
                    slope_green = (gy2 - gy1) / (gx2 - gx1) if (gx2 - gx1) != 0 else 0

                with st.expander("2. พารามิเตอร์ (ส้ม/แดง/น้ำเงิน)", expanded=True):
                    start_x = st.slider("ตำแหน่งแกน D_sb (ซ้าย)", 0, width, int(width*0.15), key="s1_sx")
                    stop_y_esb = st.slider("ระดับค่า ESB (บน)", 0, height, int(height*0.10), key="s1_sy_esb")
                    stop_y_mr = st.slider("ระดับค่า MR (ล่าง)", 0, height, int(height*0.55), key="s1_sy_mr")

                    target_y = stop_y_mr
                    if slope_green != 0:
                        constrained_x = gx1 + (target_y - gy1) / slope_green
                    else:
                        constrained_x = gx1
                    constrained_x = int(constrained_x)
                
                # Draw Box Logic
                lw = 4
                draw_arrow_fixed(draw, (start_x, stop_y_esb), (constrained_x, stop_y_esb), "orange", lw)
                draw_arrow_fixed(draw, (start_x, stop_y_esb), (start_x, stop_y_mr), "red", lw)
                draw_arrow_fixed(draw, (start_x, stop_y_mr), (constrained_x, stop_y_mr), "darkblue", lw)
                draw_arrow_fixed(draw, (constrained_x, stop_y_mr), (constrained_x, stop_y_esb), "blue", lw)
                
                r = 8
                draw.ellipse([(constrained_x-r, stop_y_mr-r), (constrained_x+r, stop_y_mr+r)], fill="black", outline="white")

                st.markdown("---")
                st.subheader("📝 บันทึกค่าที่อ่านได้")
                mr_val = st.number_input("MR (psi)", value=7000, step=500)
                esb_val = st.number_input("ESB (psi)", value=50000, step=1000)
                dsb_val = st.number_input("DSB (inches)", value=6.0, step=0.5)
                k_inf_val = st.number_input("ค่า k∞ ที่อ่านได้ (pci)", value=400, step=10)
                
                st.session_state.k_inf_result = k_inf_val
                
                buf = io.BytesIO()
                img_draw.save(buf, format='PNG')
                st.session_state.img1_bytes = buf.getvalue()

            with col_img:
                st.image(img_draw, caption="Step 1: Nomograph Analysis", use_container_width=True)

    # =========================================================
    # TAB 2: Loss of Support (IMPROVED)
    # =========================================================
    with tab2:
        st.header("2️⃣ ปรับแก้ Loss of Support (LS)")
        st.info("ใช้กราฟ Figure 3.4 เพื่อปรับค่า k∞ กรณีที่มีการสูญเสียการรองรับ (LS > 0)")
        
        uploaded_file_2 = st.file_uploader("📂 อัปโหลดภาพ Figure 3.4 (LS Correction)", type=['png', 'jpg', 'jpeg'], key='uploader_2')
        
        if uploaded_file_2 is not None:
            img2 = Image.open(uploaded_file_2).convert("RGB")
            w2, h2 = img2.size
            img2_draw = img2.copy()
            draw2 = ImageDraw.Draw(img2_draw)
            
            col_ctrl2, col_img2 = st.columns([1, 2])
            
            with col_ctrl2:
                st.subheader("⚙️ กำหนดเส้นกราฟ")
                
                # --- 1. LS Selection & Red Line ---
                st.write("#### 1. เลือกค่า LS (เส้นแดง)")
                ls_select = st.selectbox("เลือกค่า LS", [0.0, 0.5, 1.0, 1.5, 2.0, 3.0], index=2)
                
                # Logic to update sliders when dropdown changes
                if 'last_ls_select' not in st.session_state or st.session_state.last_ls_select != ls_select:
                    st.session_state.last_ls_select = ls_select
                    coords = LS_PRESETS.get(ls_select, (150, 718, 903, 84))
                    # Set defaults in session state so sliders pick them up
                    st.session_state['_ls_x1'] = coords[0]
                    st.session_state['_ls_y1'] = coords[1]
                    st.session_state['_ls_x2'] = coords[2]
                    st.session_state['_ls_y2'] = coords[3]

                with st.expander("ปรับแต่งตำแหน่งเส้น LS ละเอียด", expanded=False):
                    ls_x1 = st.slider("จุดเริ่ม X", -100, w2+100, key="_ls_x1")
                    ls_y1 = st.slider("จุดเริ่ม Y", -100, h2+100, key="_ls_y1")
                    ls_x2 = st.slider("จุดจบ X", -100, w2+100, key="_ls_x2")
                    ls_y2 = st.slider("จุดจบ Y", -100, h2+100, key="_ls_y2")

                # Draw LS Line (Red)
                draw2.line([(ls_x1, ls_y1), (ls_x2, ls_y2)], fill="red", width=6)
                
                # Calculate Slope of Red Line
                if ls_x2 - ls_x1 != 0:
                    m_red = (ls_y2 - ls_y1) / (ls_x2 - ls_x1)
                    c_red = ls_y1 - m_red * ls_x1
                else:
                    m_red = None # Vertical line
                    c_red = 0

                # --- 2. Green Line Logic (UPDATED) ---
                st.markdown("---")
                st.write("#### 2. ค่า k และขอบเขตแกน (เส้นเขียว)")
                
                # Boundary Settings (To make lines stop exactly at axes)
                with st.expander("📍 ตั้งค่าตำแหน่งแกนกราฟ (เพื่อให้เส้นไม่ล้น)", expanded=True):
                    st.caption("ปรับให้ตรงกับเส้นแกนสีดำของรูปภาพ")
                    col_b1, col_b2 = st.columns(2)
                    with col_b1:
                        # Axis Left X: Where the vertical Y-axis is located. 
                        # Default ~100 based on user images
                        axis_left_x = st.number_input("ตำแหน่งแกน Y (ซ้ายสุด)", value=100, step=5)
                    with col_b2:
                        # Axis Bottom Y: Where the horizontal X-axis is located. 
                        # Default ~730 based on user images
                        axis_bottom_y = st.number_input("ตำแหน่งแกน X (ล่างสุด)", value=h2-50, step=5)
                
                st.caption(f"ค่า k จาก Step 1 คือ: {st.session_state.k_inf_result} pci")
                k_input_x = st.slider("ตำแหน่ง k บนแกน X", 0, w2, int(w2*0.5), key="k_pos_x")
                
                # Calculate Intersection
                if m_red is not None:
                    intersect_y = m_red * k_input_x + c_red
                else:
                    intersect_y = h2/2
                intersect_y = int(intersect_y)
                
                # Draw Green Lines (Perpendicular)
                # 1. Vertical Up (From Bottom Axis to Red Line)
                # เริ่มจาก axis_bottom_y ขึ้นไปหา intersect_y
                draw2.line([(k_input_x, axis_bottom_y), (k_input_x, intersect_y)], fill="springgreen", width=5)
                
                # 2. Horizontal Left (From Red Line to Left Axis)
                # เริ่มจาก k_input_x ไปหา axis_left_x (ทางซ้าย)
                draw_arrow_fixed(draw2, (k_input_x, intersect_y), (axis_left_x, intersect_y), "springgreen", width=5)
                
                # Intersection Dot
                draw2.ellipse([(k_input_x-8, intersect_y-8), (k_input_x+8, intersect_y+8)], fill="black", outline="white", width=2)

                # --- 3. Results ---
                st.markdown("---")
                st.subheader("📝 บันทึกผลลัพธ์")
                k_corrected = st.number_input("Corrected k (pci) [อ่านจากแกน Y]", value=st.session_state.k_inf_result - 100, step=10)
                
                # Save image 2
                buf2 = io.BytesIO()
                img2_draw.save(buf2, format='PNG')
                st.session_state.img2_bytes = buf2.getvalue()
                
                # Report Generation
                st.markdown("---")
                params = {
                    'MR': mr_val if 'mr_val' in locals() else 0,
                    'ESB': esb_val if 'esb_val' in locals() else 0,
                    'DSB': dsb_val if 'dsb_val' in locals() else 0,
                    'k_inf': st.session_state.k_inf_result,
                    'LS_factor': ls_select,
                    'k_corrected': k_corrected
                }
                
                if st.button("📄 สร้างรายงาน (Word)", key="btn_report"):
                    with st.spinner("กำลังสร้างรายงาน..."):
                        img1_b = st.session_state.get('img1_bytes', None)
                        img2_b = st.session_state.get('img2_bytes', None)
                        
                        doc_file, err = generate_word_report(params, img1_b, img2_b)
                        if err:
                            st.error(err)
                        else:
                            st.download_button(
                                label="📥 ดาวน์โหลด Word Report",
                                data=doc_file,
                                file_name=f"AASHTO_Design_{datetime.now().strftime('%Y%m%d')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )

            with col_img2:
                st.image(img2_draw, caption=f"Step 2: LS Correction (LS={ls_select})", use_container_width=True)
                
        else:
            st.info("👆 กรุณาอัปโหลดภาพ Figure 3.4 เพื่อเริ่มใช้งานในส่วนนี้")

    # =========================================================
    # TAB 3: Guide
    # =========================================================
    with tab3:
        st.header("📋 คู่มือการใช้งาน")
        st.markdown("""
        ### ขั้นตอนที่ 1: หาค่า Composite k∞
        1. อัปโหลดรูป **Figure 3.3**
        2. ปรับ **Turning Line (เส้นเขียว)** ให้ตรงกับเส้นบนกราฟ (ตั้งค่าเริ่มต้นให้แล้ว)
        3. ปรับตำแหน่งลูกศรสีแดง/ส้ม ให้ตรงกับค่า **MR** และ **ESB**
        4. จุดตัดจะแสดงค่า **k∞** (แกนขวา)

        ### ขั้นตอนที่ 2: ปรับแก้ Loss of Support (LS)
        1. อัปโหลดรูป **Figure 3.4**
        2. **เลือกค่า LS** จากกล่องตัวเลือก (0, 0.5, 1.0, ...) เส้นสีแดงจะขยับไปตำแหน่งมาตรฐานตามรูปที่ส่งมา
        3. **ตั้งค่าตำแหน่งแกนกราฟ** (ครั้งแรก) โดยปรับ "ตำแหน่งแกน Y (ซ้ายสุด)" และ "ตำแหน่งแกน X (ล่างสุด)" ให้ตรงกับขอบแกนดำของรูปภาพ เพื่อให้ลูกศรชี้ได้แม่นยำและไม่ล้น
        4. เลื่อน **Slider ตำแหน่ง k บนแกน X** ให้ตรงกับค่า k ที่ได้จากขั้นตอนที่ 1
        5. โปรแกรมจะวาดเส้นสีเขียวตั้งฉากจากแกนล่าง ขึ้นไปชนเส้นแดง และเลี้ยวซ้ายไปชี้ที่แกน Y
        6. อ่านค่า Corrected k และกดสร้างรายงาน
        """)

if __name__ == "__main__":
    main()
