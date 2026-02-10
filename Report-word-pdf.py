# -*- coding: utf-8 -*-
"""
โปรแกรมรวมไฟล์ Word รายงานออกแบบโครงสร้างชั้นทาง
Pavement Design Report Merger
Version 2.2

โดย: ภาควิชาครุศาสตร์โยธา มจพ.

การปรับปรุง v2.2:
- เพิ่มหน้าสรุปรูปภาพโครงสร้างชั้นทางในรายงาน Word (ก่อนเนื้อหา)
- แสดงรูปภาพจากแต่ละไฟล์แบบ Grid ในตาราง 2 คอลัมน์
- รองรับ Caption ใต้รูปภาพแต่ละรูป
"""

import streamlit as st
import os
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import io
import re
import zipfile
from lxml import etree

# ตั้งค่าหน้าเว็บ
st.set_page_config(
    page_title="โปรแกรมรวมรายงานออกแบบโครงสร้างชั้นทาง",
    page_icon="🛣️",
    layout="wide"
)

# CSS สำหรับตกแต่งหน้าเว็บ
st.markdown("""
<style>
    .main-header {
        font-size: 28px;
        font-weight: bold;
        color: #1E3A5F;
        text-align: center;
        padding: 20px;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 10px;
        margin-bottom: 20px;
    }
    .sub-header {
        font-size: 18px;
        color: #4A5568;
        text-align: center;
        margin-bottom: 30px;
    }
    .file-section {
        background-color: #F7FAFC;
        padding: 15px;
        border-radius: 10px;
        margin-bottom: 10px;
        border-left: 4px solid #667eea;
    }
    .file-section-sub {
        background-color: #EDF2F7;
        padding: 10px 15px;
        border-radius: 8px;
        margin: 5px 0 5px 20px;
        border-left: 3px solid #A0AEC0;
    }
    .section-header {
        background-color: #C6F6D5;
        padding: 10px 15px;
        border-radius: 8px;
        margin: 15px 0 10px 0;
        font-weight: bold;
        color: #276749;
        border-left: 4px solid #38A169;
    }
    .success-box {
        background-color: #C6F6D5;
        padding: 15px;
        border-radius: 10px;
        border-left: 4px solid #38A169;
    }
    .warning-box {
        background-color: #FEFCBF;
        padding: 15px;
        border-radius: 10px;
        border-left: 4px solid #D69E2E;
    }
    .stButton>button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        font-weight: bold;
        padding: 10px 30px;
        border-radius: 25px;
        border: none;
        font-size: 16px;
    }
    .stButton>button:hover {
        background: linear-gradient(135deg, #764ba2 0%, #667eea 100%);
    }
</style>
""", unsafe_allow_html=True)


def set_thai_font(run, font_name="TH Sarabun New", font_size=15):
    """ตั้งค่าฟอนต์ไทยและขนาด"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)


def set_page_margins(section):
    """ตั้งค่าหน้ากระดาษ A4 แนวตั้ง กั้นหน้า-หลัง 2.5 cm"""
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


def set_cell_border(cell, top=True, bottom=True, left=True, right=True, color="999999", size=4):
    """ตั้งค่า border ของ cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    sides = []
    if top: sides.append('top')
    if bottom: sides.append('bottom')
    if left: sides.append('left')
    if right: sides.append('right')
    
    for side in sides:
        border_el = OxmlElement(f'w:{side}')
        border_el.set(qn('w:val'), 'single')
        border_el.set(qn('w:sz'), str(size))
        border_el.set(qn('w:color'), color)
        tcBorders.append(border_el)
    
    tcPr.append(tcBorders)


def extract_images_from_docx(file_bytes):
    """
    ดึงรูปภาพทั้งหมดจากไฟล์ docx
    Returns: list of (image_bytes, image_ext) tuples
    """
    images = []
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes), 'r') as z:
            media_files = [f for f in z.namelist() if f.startswith('word/media/')]
            # เรียงลำดับให้แน่นอน
            media_files.sort()
            for media_file in media_files:
                ext = media_file.rsplit('.', 1)[-1].lower()
                if ext in ('png', 'jpg', 'jpeg', 'gif', 'bmp', 'wmf', 'emf'):
                    img_bytes = z.read(media_file)
                    images.append((img_bytes, ext))
    except Exception as e:
        pass
    return images


def get_first_image_from_docx(file_bytes):
    """ดึงรูปภาพแรกจากไฟล์ docx"""
    images = extract_images_from_docx(file_bytes)
    if images:
        return images[0]
    return None


def get_largest_image_from_docx(file_bytes):
    """ดึงรูปภาพที่ใหญ่ที่สุด (น่าจะเป็นรูปโครงสร้างชั้นทาง) จากไฟล์ docx"""
    images = extract_images_from_docx(file_bytes)
    if not images:
        return None
    # เลือกรูปที่มีขนาด bytes ใหญ่ที่สุด (มักเป็นรูปหลัก)
    largest = max(images, key=lambda x: len(x[0]))
    return largest


def add_image_to_cell(cell, img_bytes, img_ext, max_width_cm=7.5, max_height_cm=9.0):
    """
    เพิ่มรูปภาพลงใน cell โดยรักษาสัดส่วน
    """
    try:
        from PIL import Image as PILImage
        import struct

        # วิเคราะห์ขนาดรูปภาพ
        img_stream = io.BytesIO(img_bytes)
        
        # ลองใช้ PIL ก่อน
        try:
            pil_img = PILImage.open(img_stream)
            orig_w, orig_h = pil_img.size  # pixels
            dpi = pil_img.info.get('dpi', (96, 96))
            dpi_x = dpi[0] if isinstance(dpi, tuple) else 96
            # แปลงเป็น cm
            orig_w_cm = orig_w / dpi_x * 2.54
            orig_h_cm = orig_h / dpi_x * 2.54
        except:
            # ถ้า PIL ไม่ได้ ใช้ขนาด default
            orig_w_cm = 10.0
            orig_h_cm = 8.0

        # คำนวณขนาดที่พอดี
        scale_w = max_width_cm / orig_w_cm if orig_w_cm > max_width_cm else 1.0
        scale_h = max_height_cm / orig_h_cm if orig_h_cm > max_height_cm else 1.0
        scale = min(scale_w, scale_h)
        
        final_w_cm = orig_w_cm * scale
        final_h_cm = orig_h_cm * scale
        
        # เพิ่มรูปลงใน cell
        img_stream.seek(0)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(img_stream, width=Cm(final_w_cm))
        
        return True
    except Exception as e:
        # ลองวิธีที่ 2 ถ้า PIL ไม่ได้
        try:
            img_stream = io.BytesIO(img_bytes)
            img_stream.seek(0)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(img_stream, width=Cm(max_width_cm * 0.85))
            return True
        except Exception as e2:
            return False


def add_pavement_summary_page(merged_doc, uploaded_files, section_titles, summary_images):
    """
    สร้างหน้าสรุปโครงสร้างชั้นทาง พร้อมรูปภาพในตาราง 2 คอลัมน์
    เหมือนรูปแบบในรายงานมาตรฐาน
    
    Parameters:
        summary_images: dict {key: (img_bytes, ext, caption)} รูปภาพสำหรับแต่ละส่วน
    """
    # หัวข้อหน้าสรุป
    title_para = merged_doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("สรุปโครงสร้างชั้นทางที่ออกแบบด้วยวิธี AASHTO 1993")
    set_thai_font(title_run, font_size=18)
    title_run.font.bold = True

    # คำอธิบาย
    desc_para = merged_doc.add_paragraph()
    desc_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    desc_run = desc_para.add_run(
        "จากการคำนวณ/ออกแบบตามวิธี AASHTO 1993 สำหรับการออกแบบโครงสร้างชั้นทาง"
        " (Asphalt Concrete) และผิวทางคอนกรีต (Concrete Pavement)"
        " สามารถสรุปรูปแบบโครงสร้างชั้นทาง ดังแสดงในรูปที่ด้านล่าง"
    )
    set_thai_font(desc_run, font_size=15)

    merged_doc.add_paragraph()  # ระยะห่าง

    # กรองเฉพาะรายการที่มีรูปภาพ
    valid_items = [(key, img_data) for key, img_data in summary_images.items() if img_data is not None]

    if not valid_items:
        no_img_para = merged_doc.add_paragraph()
        no_img_run = no_img_para.add_run("ไม่พบรูปภาพโครงสร้างชั้นทางในไฟล์ที่อัปโหลด")
        set_thai_font(no_img_run, font_size=14)
        no_img_run.font.italic = True
        return

    # สร้างตาราง 2 คอลัมน์
    # A4 content width = 21 - 2.5 - 2.5 = 16 cm
    # แบ่งเป็น 2 คอลัมน์ละ ~7.8 cm + กั้นกลาง 0.4 cm
    
    # จัดกลุ่มเป็นคู่ (2 รูปต่อแถว)
    pairs = []
    for i in range(0, len(valid_items), 2):
        row_items = valid_items[i:i+2]
        pairs.append(row_items)

    fig_counter = 11  # เริ่มต้น figure number (ปรับได้)
    
    for pair in pairs:
        # สร้างตาราง 1 แถว x 2 คอลัมน์
        tbl = merged_doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # กำหนด width รวม 16 cm (แบ่งเท่าๆ กัน)
        col_width_dxa = int(16 * 567)  # 16 cm in DXA (1 cm = 567 DXA)
        col_half = col_width_dxa // 2
        
        tbl_element = tbl._tbl
        tblPr = tbl_element.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl_element.insert(0, tblPr)
        
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(col_width_dxa))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)
        
        # ตั้ง column widths
        tblGrid = OxmlElement('w:tblGrid')
        for _ in range(2):
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), str(col_half))
            tblGrid.append(gridCol)
        tbl_element.insert(1, tblGrid)
        
        row = tbl.rows[0]
        
        for col_idx, (key, img_data) in enumerate(pair):
            img_bytes, img_ext, caption = img_data
            cell = row.cells[col_idx]
            
            # ตั้งค่า cell width
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(col_half))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            
            # ตั้งค่า vertical alignment
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'top')
            tcPr.append(vAlign)
            
            # เพิ่ม border ให้ cell
            set_cell_border(cell, color="CCCCCC", size=4)
            
            # margin ใน cell
            tcMar = OxmlElement('w:tcMar')
            for side in ['top', 'bottom', 'left', 'right']:
                mar = OxmlElement(f'w:{side}')
                mar.set(qn('w:w'), '100')
                mar.set(qn('w:type'), 'dxa')
                tcMar.append(mar)
            tcPr.append(tcMar)
            
            # เพิ่มรูปภาพ (ลบ paragraph เดิมออก)
            cell.paragraphs[0].clear()
            img_para = cell.paragraphs[0]
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # ขนาดรูปภาพ: กว้าง max 7.2 cm, สูง max 9.5 cm
            img_added = False
            try:
                img_stream = io.BytesIO(img_bytes)
                run = img_para.add_run()
                run.add_picture(img_stream, width=Cm(7.2))
                img_added = True
            except Exception:
                try:
                    img_stream = io.BytesIO(img_bytes)
                    run = img_para.add_run()
                    run.add_picture(img_stream, width=Cm(6.5))
                    img_added = True
                except Exception:
                    pass
            
            if not img_added:
                err_run = img_para.add_run("[ไม่สามารถแสดงรูปภาพได้]")
                set_thai_font(err_run, font_size=12)
                err_run.font.italic = True
            
            # Caption ใต้รูป
            cap_para = cell.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_para.add_run(f"รูปที่ 2-{fig_counter} {caption}")
            set_thai_font(cap_run, font_size=13)
            cap_run.font.italic = True
            fig_counter += 1
        
        # ถ้าแถวนี้มีแค่ 1 รูป ให้เติม cell ว่างทางขวา
        if len(pair) == 1:
            empty_cell = row.cells[1]
            tc = empty_cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(col_half))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
        
        merged_doc.add_paragraph()  # ระยะห่างระหว่างแถว

    merged_doc.add_page_break()


def collect_summary_images(uploaded_files, section_titles, summary_captions):
    """
    ดึงรูปภาพจากแต่ละ uploaded file เพื่อใช้ในหน้าสรุป
    Returns: dict {key: (img_bytes, ext, caption) or None}
    """
    summary_images = {}
    
    # กำหนดว่า key ใดที่ต้องการดึงรูปสำหรับสรุป (ส่วนออกแบบ)
    design_keys = ['ac_design', 'jpcp_jrcp_design', 'crcp_design']
    
    for key in design_keys:
        file = uploaded_files.get(key)
        if file is None:
            summary_images[key] = None
            continue
        
        try:
            file_bytes = file.read()
            file.seek(0)
            
            # ดึงรูปภาพจากไฟล์
            img_data = get_largest_image_from_docx(file_bytes)
            
            if img_data:
                img_bytes, img_ext = img_data
                caption = summary_captions.get(key, section_titles.get(key, key))
                summary_images[key] = (img_bytes, img_ext, caption)
            else:
                summary_images[key] = None
        except Exception:
            summary_images[key] = None
    
    return summary_images


def append_document(master_doc, source_doc):
    """
    คัดลอกเนื้อหาจากเอกสารต้นทางไปยังเอกสารปลายทาง
    รองรับรูปภาพ ตาราง และการจัดรูปแบบ
    """
    # คัดลอก relationships สำหรับรูปภาพ
    if source_doc.part.rels:
        for rel_id, rel in source_doc.part.rels.items():
            if "image" in rel.reltype:
                try:
                    image_part = rel.target_part
                    new_rel = master_doc.part.relate_to(image_part, rel.reltype)
                except:
                    pass
    
    # คัดลอก elements จาก body
    for element in source_doc.element.body:
        if element.tag.endswith('sectPr'):
            continue
        new_element = deepcopy(element)
        master_doc.element.body.append(new_element)


def merge_documents(uploaded_files, section_titles, project_name, report_date,
                    include_summary_page=True, summary_captions=None):
    """รวมเอกสารทั้งหมดเป็นไฟล์เดียว (รองรับรูปภาพ ตาราง และหน้าสรุป)"""
    
    if summary_captions is None:
        summary_captions = {}
    
    # ─────────────────────────────────────────────────────────────────
    # 1. ดึงรูปภาพสำหรับหน้าสรุป (ก่อนเริ่มอ่านไฟล์จริง)
    # ─────────────────────────────────────────────────────────────────
    summary_images = {}
    if include_summary_page:
        summary_images = collect_summary_images(uploaded_files, section_titles, summary_captions)
    
    # ─────────────────────────────────────────────────────────────────
    # 2. สร้างเอกสารหลัก
    # ─────────────────────────────────────────────────────────────────
    merged_doc = Document()
    section = merged_doc.sections[0]
    set_page_margins(section)
    
    # ─────────────────────────────────────────────────────────────────
    # 3. หน้าปก
    # ─────────────────────────────────────────────────────────────────
    title_para = merged_doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("\n\n\n\n\n")
    
    main_title = merged_doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    main_run = main_title.add_run("รายงานการออกแบบโครงสร้างชั้นทาง")
    set_thai_font(main_run, font_size=24)
    main_run.font.bold = True
    
    if project_name:
        project_para = merged_doc.add_paragraph()
        project_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        project_run = project_para.add_run(f"\n{project_name}")
        set_thai_font(project_run, font_size=20)
        project_run.font.bold = True
    
    date_para = merged_doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"\n\n\n\n{report_date}")
    set_thai_font(date_run, font_size=16)
    
    merged_doc.add_page_break()
    
    # ─────────────────────────────────────────────────────────────────
    # 4. สารบัญ
    # ─────────────────────────────────────────────────────────────────
    toc_title = merged_doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_run = toc_title.add_run("สารบัญ")
    set_thai_font(toc_run, font_size=18)
    toc_run.font.bold = True
    
    merged_doc.add_paragraph()
    
    toc_items = []
    section_num = 1
    for key, file in uploaded_files.items():
        if file is not None:
            toc_items.append((section_num, section_titles[key]))
            section_num += 1
    
    for num, title in toc_items:
        toc_para = merged_doc.add_paragraph()
        toc_run = toc_para.add_run(f"{num}. {title}")
        set_thai_font(toc_run, font_size=15)
    
    merged_doc.add_page_break()
    
    # ─────────────────────────────────────────────────────────────────
    # 5. หน้าสรุปโครงสร้างชั้นทาง (ใหม่ v2.2)
    # ─────────────────────────────────────────────────────────────────
    if include_summary_page:
        has_any_image = any(v is not None for v in summary_images.values())
        if has_any_image:
            add_pavement_summary_page(merged_doc, uploaded_files, section_titles, summary_images)
    
    # ─────────────────────────────────────────────────────────────────
    # 6. รวมเนื้อหาจากแต่ละไฟล์
    # ─────────────────────────────────────────────────────────────────
    section_num = 1
    for key, file in uploaded_files.items():
        if file is not None:
            file_bytes = file.read()
            file.seek(0)
            
            section_title_para = merged_doc.add_paragraph()
            section_title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            section_run = section_title_para.add_run(f"{section_num}. {section_titles[key]}")
            set_thai_font(section_run, font_size=18)
            section_run.font.bold = True
            
            merged_doc.add_paragraph()
            
            source_doc = Document(io.BytesIO(file_bytes))
            append_document(merged_doc, source_doc)
            
            merged_doc.add_page_break()
            section_num += 1
    
    return merged_doc


def main():
    # หัวข้อหลัก
    st.markdown('<div class="main-header">🛣️ โปรแกรมรวมรายงานออกแบบโครงสร้างชั้นทาง</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-header">Pavement Structure Design Report Merger v2.2</div>', unsafe_allow_html=True)
    
    # ข้อมูลโครงการ
    st.markdown("### 📋 ข้อมูลโครงการ")
    col1, col2 = st.columns(2)
    with col1:
        project_name = st.text_input("ชื่อโครงการ", placeholder="กรอกชื่อโครงการ")
    with col2:
        report_date = st.date_input("วันที่รายงาน", datetime.now())
        report_date_str = report_date.strftime("%d/%m/%Y")
    
    st.markdown("---")
    
    # ─────────────────────────────────────────────────────────────────
    # ตั้งค่าหน้าสรุปโครงสร้างชั้นทาง (v2.2)
    # ─────────────────────────────────────────────────────────────────
    st.markdown("### 📸 ตั้งค่าหน้าสรุปโครงสร้างชั้นทาง")
    col_opt1, col_opt2 = st.columns([1, 2])
    with col_opt1:
        include_summary = st.checkbox(
            "✅ เพิ่มหน้าสรุปรูปภาพโครงสร้างชั้นทาง",
            value=True,
            help="เมื่อเปิดใช้งาน จะดึงรูปภาพโครงสร้างชั้นทางจากไฟล์ออกแบบ AC, JPCP/JRCP, CRCP มาแสดงรวมกันก่อนเนื้อหา"
        )
    with col_opt2:
        if include_summary:
            st.info("💡 โปรแกรมจะดึงรูปภาพที่ใหญ่ที่สุดจากไฟล์ออกแบบ AC, JPCP/JRCP, และ CRCP มาแสดงในตาราง 2 คอลัมน์ เหมือนรูปแบบรายงานมาตรฐาน")
    
    # กำหนด Caption สำหรับแต่ละรูปในหน้าสรุป
    summary_captions = {}
    if include_summary:
        with st.expander("⚙️ ปรับแต่ง Caption ของรูปในหน้าสรุป (ไม่บังคับ)"):
            summary_captions['ac_design'] = st.text_input(
                "Caption รูป AC Design",
                value="โครงสร้างชั้นทางรูปแบบที่ 1 ผิวทางลาดยาง แบบ AC",
                help="คำบรรยายใต้รูปโครงสร้างชั้นทาง AC"
            )
            summary_captions['jpcp_jrcp_design'] = st.text_input(
                "Caption รูป JPCP/JRCP Design",
                value="โครงสร้างชั้นทางรูปแบบที่ 2 ผิวทางคอนกรีต แบบ JPCP",
                help="คำบรรยายใต้รูปโครงสร้างชั้นทาง JPCP/JRCP"
            )
            summary_captions['crcp_design'] = st.text_input(
                "Caption รูป CRCP Design",
                value="โครงสร้างชั้นทางรูปแบบที่ 3 ผิวทางคอนกรีต แบบ CRCP",
                help="คำบรรยายใต้รูปโครงสร้างชั้นทาง CRCP"
            )
    else:
        # ค่า default ถ้าไม่ได้เปิด expander
        summary_captions = {
            'ac_design': "โครงสร้างชั้นทางรูปแบบที่ 1 ผิวทางลาดยาง แบบ AC",
            'jpcp_jrcp_design': "โครงสร้างชั้นทางรูปแบบที่ 2 ผิวทางคอนกรีต แบบ JPCP",
            'crcp_design': "โครงสร้างชั้นทางรูปแบบที่ 3 ผิวทางคอนกรีต แบบ CRCP",
        }
    
    st.markdown("---")
    
    # คำอธิบายส่วนต่างๆ
    section_titles = {
        'truck_factor': 'การคำนวณ Truck Factor',
        'esals_ac': 'การคำนวณ ESALs สำหรับผิวทางลาดยาง (Flexible Pavement)',
        'esals_concrete': 'การคำนวณ ESALs สำหรับผิวทางคอนกรีต (Rigid Pavement)',
        'cbr_analysis': 'การวิเคราะห์ค่า CBR ที่เปอร์เซ็นต์ไทล์',
        'ac_design': 'การออกแบบผิวทางลาดยาง (Flexible Pavement)',
        'jpcp_jrcp_design': 'การออกแบบผิวทางคอนกรีต JPCP/JRCP',
        'crcp_design': 'การออกแบบผิวทางคอนกรีต CRCP',
        'k_value_jpcp_jrcp': 'การคำนวณ Corrected Modulus of Subgrade Reaction (k-value) สำหรับ JPCP/JRCP',
        'k_value_crcp': 'การคำนวณ Corrected Modulus of Subgrade Reaction (k-value) สำหรับ CRCP',
        'cost_estimate': 'การประมาณราคาค่าก่อสร้าง'
    }
    
    st.markdown("### 📁 อัปโหลดไฟล์รายงาน")
    st.info("💡 อัปโหลดไฟล์ Word (.docx) สำหรับแต่ละส่วนของรายงาน ไฟล์ที่มีเครื่องหมาย (ถ้ามี) สามารถเว้นว่างได้")
    
    uploaded_files = {}
    
    # ═══════════════════════════════════════════════════════════════
    # ส่วนที่ 1: Truck Factor
    # ═══════════════════════════════════════════════════════════════
    st.markdown('<div class="section-header">📊 1. การคำนวณ Truck Factor</div>', unsafe_allow_html=True)
    st.markdown('<div class="file-section">', unsafe_allow_html=True)
    st.markdown("**การคำนวณ Truck Factor** (ถ้ามี)")
    uploaded_files['truck_factor'] = st.file_uploader(
        "เลือกไฟล์ Truck Factor",
        type=['docx'],
        key='truck_factor',
        help="ไฟล์รายงานการคำนวณ Truck Factor"
    )
    st.markdown('</div>', unsafe_allow_html=True)
    
    # ═══════════════════════════════════════════════════════════════
    # ส่วนที่ 2: ESALs
    # ═══════════════════════════════════════════════════════════════
    st.markdown('<div class="section-header">📈 2. การคำนวณ ESALs (Equivalent Single Axle Loads)</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown('<div class="file-section">', unsafe_allow_html=True)
        st.markdown("**2.1 ESALs สำหรับผิวทางลาดยาง** (Flexible Pavement)")
        uploaded_files['esals_ac'] = st.file_uploader(
            "เลือกไฟล์ ESALs ผิวทางลาดยาง",
            type=['docx'],
            key='esals_ac',
            help="ไฟล์รายงานการคำนวณ ESALs สำหรับผิวทางลาดยาง (AC)"
        )
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col2:
        st.markdown('<div class="file-section">', unsafe_allow_html=True)
        st.markdown("**2.2 ESALs สำหรับผิวทางคอนกรีต** (Rigid Pavement)")
        uploaded_files['esals_concrete'] = st.file_uploader(
            "เลือกไฟล์ ESALs ผิวทางคอนกรีต",
            type=['docx'],
            key='esals_concrete',
            help="ไฟล์รายงานการคำนวณ ESALs สำหรับผิวทางคอนกรีต"
        )
        st.markdown('</div>', unsafe_allow_html=True)
    
    # ═══════════════════════════════════════════════════════════════
    # ส่วนที่ 3: CBR
    # ═══════════════════════════════════════════════════════════════
    st.markdown('<div class="section-header">🔬 3. การวิเคราะห์ค่า CBR ที่เปอร์เซ็นต์ไทล์</div>', unsafe_allow_html=True)
    st.markdown('<div class="file-section">', unsafe_allow_html=True)
    st.markdown("**การวิเคราะห์ค่า CBR ที่เปอร์เซ็นต์ไทล์**")
    uploaded_files['cbr_analysis'] = st.file_uploader(
        "เลือกไฟล์วิเคราะห์ CBR",
        type=['docx'],
        key='cbr_analysis',
        help="ไฟล์รายงานการวิเคราะห์ค่า CBR ที่เปอร์เซ็นต์ไทล์ (Percentile Analysis)"
    )
    st.markdown('</div>', unsafe_allow_html=True)
    
    # ═══════════════════════════════════════════════════════════════
    # ส่วนที่ 4: AC Design
    # ═══════════════════════════════════════════════════════════════
    st.markdown('<div class="section-header">🛤️ 4. การออกแบบผิวทางลาดยาง (Flexible Pavement)</div>', unsafe_allow_html=True)
    st.markdown('<div class="file-section">', unsafe_allow_html=True)
    st.markdown("**การออกแบบผิวทางลาดยาง (AC)**")
    if include_summary:
        st.caption("🖼️ รูปภาพจากไฟล์นี้จะถูกดึงไปใช้ในหน้าสรุปโครงสร้างชั้นทาง")
    uploaded_files['ac_design'] = st.file_uploader(
        "เลือกไฟล์ออกแบบ AC",
        type=['docx'],
        key='ac_design',
        help="ไฟล์รายงานการออกแบบผิวทางแอสฟัลต์ตามวิธี AASHTO 1993"
    )
    st.markdown('</div>', unsafe_allow_html=True)
    
    # ═══════════════════════════════════════════════════════════════
    # ส่วนที่ 5: Rigid Pavement
    # ═══════════════════════════════════════════════════════════════
    st.markdown('<div class="section-header">🏗️ 5. การออกแบบผิวทางคอนกรีต (Rigid Pavement)</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown('<div class="file-section">', unsafe_allow_html=True)
        st.markdown("**5.1 การออกแบบ JPCP/JRCP**")
        st.caption("Jointed Plain/Reinforced Concrete Pavement")
        if include_summary:
            st.caption("🖼️ รูปภาพจากไฟล์นี้จะถูกดึงไปใช้ในหน้าสรุปโครงสร้างชั้นทาง")
        uploaded_files['jpcp_jrcp_design'] = st.file_uploader(
            "เลือกไฟล์ออกแบบ JPCP/JRCP",
            type=['docx'],
            key='jpcp_jrcp_design',
            help="ไฟล์รายงานการออกแบบผิวทาง JPCP หรือ JRCP"
        )
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col2:
        st.markdown('<div class="file-section">', unsafe_allow_html=True)
        st.markdown("**5.2 การออกแบบ CRCP**")
        st.caption("Continuously Reinforced Concrete Pavement")
        if include_summary:
            st.caption("🖼️ รูปภาพจากไฟล์นี้จะถูกดึงไปใช้ในหน้าสรุปโครงสร้างชั้นทาง")
        uploaded_files['crcp_design'] = st.file_uploader(
            "เลือกไฟล์ออกแบบ CRCP",
            type=['docx'],
            key='crcp_design',
            help="ไฟล์รายงานการออกแบบผิวทาง CRCP"
        )
        st.markdown('</div>', unsafe_allow_html=True)
    
    # ═══════════════════════════════════════════════════════════════
    # ส่วนที่ 6: k-value
    # ═══════════════════════════════════════════════════════════════
    st.markdown('<div class="section-header">📐 6. การคำนวณ Corrected Modulus of Subgrade Reaction (k-value)</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown('<div class="file-section">', unsafe_allow_html=True)
        st.markdown("**6.1 k-value สำหรับ JPCP/JRCP**")
        uploaded_files['k_value_jpcp_jrcp'] = st.file_uploader(
            "เลือกไฟล์ k-value JPCP/JRCP",
            type=['docx'],
            key='k_value_jpcp_jrcp',
            help="ไฟล์รายการคำนวณ Corrected k-value สำหรับ JPCP/JRCP"
        )
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col2:
        st.markdown('<div class="file-section">', unsafe_allow_html=True)
        st.markdown("**6.2 k-value สำหรับ CRCP**")
        uploaded_files['k_value_crcp'] = st.file_uploader(
            "เลือกไฟล์ k-value CRCP",
            type=['docx'],
            key='k_value_crcp',
            help="ไฟล์รายการคำนวณ Corrected k-value สำหรับ CRCP"
        )
        st.markdown('</div>', unsafe_allow_html=True)
    
    # ═══════════════════════════════════════════════════════════════
    # ส่วนที่ 7: ประมาณราคา
    # ═══════════════════════════════════════════════════════════════
    st.markdown('<div class="section-header">💰 7. การประมาณราคาค่าก่อสร้าง</div>', unsafe_allow_html=True)
    st.markdown('<div class="file-section">', unsafe_allow_html=True)
    st.markdown("**การประมาณราคาค่าก่อสร้าง** (ถ้ามี)")
    uploaded_files['cost_estimate'] = st.file_uploader(
        "เลือกไฟล์ประมาณราคา",
        type=['docx'],
        key='cost_estimate',
        help="ไฟล์รายงานการประมาณราคาค่าก่อสร้าง"
    )
    st.markdown('</div>', unsafe_allow_html=True)
    
    st.markdown("---")
    
    # ═══════════════════════════════════════════════════════════════
    # แสดงสถานะไฟล์
    # ═══════════════════════════════════════════════════════════════
    st.markdown("### 📊 สถานะไฟล์ที่อัปโหลด")
    
    file_count = sum(1 for f in uploaded_files.values() if f is not None)
    
    status_data = {
        'หมวด': [
            '1. Truck Factor',
            '2.1 ESALs (Flexible)',
            '2.2 ESALs (Rigid)',
            '3. CBR Analysis',
            '4. AC Design ⭐',
            '5.1 JPCP/JRCP ⭐',
            '5.2 CRCP ⭐',
            '6.1 k-value (JPCP/JRCP)',
            '6.2 k-value (CRCP)',
            '7. Cost Estimate'
        ],
        'สถานะ': []
    }
    
    file_keys = ['truck_factor', 'esals_ac', 'esals_concrete', 'cbr_analysis', 'ac_design', 
                 'jpcp_jrcp_design', 'crcp_design', 'k_value_jpcp_jrcp', 
                 'k_value_crcp', 'cost_estimate']
    
    for key in file_keys:
        if uploaded_files[key] is not None:
            status_data['สถานะ'].append('✅ อัปโหลดแล้ว')
        else:
            status_data['สถานะ'].append('⬜ ยังไม่อัปโหลด')
    
    cols = st.columns(3)
    for i, (name, status) in enumerate(zip(status_data['หมวด'], status_data['สถานะ'])):
        with cols[i % 3]:
            if '✅' in status:
                st.success(f"{name}: {status}")
            else:
                st.warning(f"{name}: {status}")
    
    if include_summary:
        st.caption("⭐ = ไฟล์ที่จะดึงรูปภาพไปใช้ในหน้าสรุปโครงสร้างชั้นทาง")
    
    st.markdown(f"### 📈 อัปโหลดแล้ว: **{file_count}** จาก **10** ไฟล์")
    
    st.markdown("---")
    
    # ═══════════════════════════════════════════════════════════════
    # ปุ่มรวมไฟล์
    # ═══════════════════════════════════════════════════════════════
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        merge_button = st.button("🔄 รวมไฟล์และสร้างรายงาน", use_container_width=True)
    
    if merge_button:
        if file_count == 0:
            st.error("❌ กรุณาอัปโหลดไฟล์อย่างน้อย 1 ไฟล์")
        else:
            with st.spinner("กำลังรวมไฟล์และสร้างรายงาน..."):
                try:
                    merged_doc = merge_documents(
                        uploaded_files,
                        section_titles,
                        project_name,
                        report_date_str,
                        include_summary_page=include_summary,
                        summary_captions=summary_captions
                    )
                    
                    output = io.BytesIO()
                    merged_doc.save(output)
                    output.seek(0)
                    
                    base_filename = "รายงานออกแบบโครงสร้างชั้นทาง"
                    if project_name:
                        base_filename = f"รายงานออกแบบ_{project_name.replace(' ', '_')}"
                    
                    st.markdown('<div class="success-box">', unsafe_allow_html=True)
                    st.success(f"✅ รวมไฟล์เรียบร้อยแล้ว! ({file_count} ไฟล์)")
                    if include_summary:
                        design_keys = ['ac_design', 'jpcp_jrcp_design', 'crcp_design']
                        img_count = sum(1 for k in design_keys if uploaded_files.get(k) is not None)
                        if img_count > 0:
                            st.info(f"🖼️ เพิ่มหน้าสรุปโครงสร้างชั้นทางพร้อมรูปภาพ {img_count} รูปเรียบร้อยแล้ว")
                    st.markdown('</div>', unsafe_allow_html=True)
                    
                    st.markdown("### 📥 ดาวน์โหลดรายงาน")
                    
                    st.download_button(
                        label="📄 ดาวน์โหลดไฟล์ Word (.docx)",
                        data=output.getvalue(),
                        file_name=f"{base_filename}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                
                except Exception as e:
                    st.error(f"❌ เกิดข้อผิดพลาด: {str(e)}")
                    st.exception(e)
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #718096; font-size: 14px;">
        <p>พัฒนาโดย ภาควิชาครุศาสตร์โยธา คณะครุศาสตร์อุตสาหกรรม</p>
        <p>มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ</p>
        <p>© 2025 - Pavement Design Report Merger v2.2</p>
    </div>
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
