"""
แอปพลิเคชันวิเคราะห์ความคุ้มค่าโครงสร้างชั้นทาง (AASHTO 1993)
Version 5.1 - Bug Fixed
พัฒนาโดย: Claude AI สำหรับ อ.อิทธิพล - KMUTNB

CHANGELOG v5.1:
  BUG FIX 1: generate_word_report_table() — qty คูณ road_length ซ้ำ (layer['quantity'] ถูก calc ด้วย road_length แล้ว)
  BUG FIX 2: generate_word_report_table() — pre-allocate num_rows ไม่ถูกต้อง → ใช้ dynamic row append แทน
  BUG FIX 3: render_joint_editor() CRCP — ใช้ area_per_km param แทน session_state ที่อาจยังไม่ถูก set
  BUG FIX 4: _parse_json_details_to_layers() — qty_unit='cu.m' ไม่สอดคล้องกับ render_layer_editor() → แก้เป็น 'sq.m' เสมอ
  CLEANUP: generate_word_report_table() ถูก dead code → ลบออก (ไม่ถูกเรียกใช้ที่ไหน)
"""

import streamlit as st
import pandas as pd
import numpy as np
import json
from datetime import datetime
import io

# Import with error handling
try:
    import plotly.graph_objects as go
    from plotly.subplots import make_subplots
    PLOTLY_AVAILABLE = True
except ImportError:
    PLOTLY_AVAILABLE = False
    st.warning("⚠️ Plotly ไม่สามารถใช้งานได้ กราฟบางส่วนอาจไม่แสดง")

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.warning("⚠️ python-docx ไม่สามารถใช้งานได้ การสร้างรายงาน Word อาจไม่ทำงาน")

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    st.warning("⚠️ openpyxl ไม่สามารถใช้งานได้ การ Upload/Download Excel อาจไม่ทำงาน")

# ตั้งค่าหน้าเว็บ
st.set_page_config(
    page_title="วิเคราะห์ค่าก่อสร้างโครงสร้างชั้นทาง",
    page_icon="🛣️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS
st.markdown("""
<style>
    .main-header {
        font-size: 2rem;
        font-weight: bold;
        color: #1E3A5F;
        text-align: center;
        padding: 1rem;
        background: linear-gradient(90deg, #E8F4FD, #D1E9FA);
        border-radius: 10px;
        margin-bottom: 2rem;
    }
    .cost-box {
        background: #f0f8ff;
        padding: 10px;
        border-radius: 8px;
        border-left: 4px solid #2E86AB;
        margin: 10px 0;
    }

    /* === Selectbox: เขียวอ่อน === */
    div[data-baseweb="select"] > div {
        background-color: #e8f5e9 !important;
        border-color: #66bb6a !important;
        border-radius: 8px !important;
    }
    div[data-baseweb="select"] > div:hover {
        background-color: #c8e6c9 !important;
        border-color: #43a047 !important;
    }
    div[data-baseweb="select"] svg {
        fill: #2e7d32 !important;
    }
    div[data-baseweb="menu"] {
        background-color: #f1f8e9 !important;
    }
    div[data-baseweb="menu"] li {
        background-color: #f1f8e9 !important;
        color: #1b5e20 !important;
    }
    div[data-baseweb="menu"] li:hover {
        background-color: #c8e6c9 !important;
    }
    div[data-baseweb="select"] span {
        color: #1b5e20 !important;
        font-weight: 500 !important;
    }
</style>
""", unsafe_allow_html=True)


# ===== Library ราคาวัสดุ (Price Library) =====
AC_PRICE_TABLE = {
    'PMA Wearing Course': {
        2.5: 170, 3: 203, 4: 268, 5: 333, 6: 406, 7: 471, 8: 536, 9: 601, 10: 667
    },
    'AC Wearing Course': {
        2.5: 128, 3: 152, 4: 202, 5: 250, 6: 306, 7: 355, 8: 403, 9: 452, 10: 502
    },
    'AC Binder Course': {
        2.5: 129, 3: 154, 4: 202, 5: 251, 6: 308, 7: 356, 8: 405, 9: 454, 10: 503
    },
    'AC Base Course': {
        2.5: 129, 3: 154, 4: 202, 5: 251, 6: 308, 7: 356, 8: 405, 9: 454, 10: 503
    },
}

CONCRETE_PRICE_TABLE = {
    'JRCP': {25: 924, 28: 1002, 30: 0, 32: 1106, 35: 1184},
    'JPCP': {25: 928, 28: 1000, 30: 0, 32: 1095, 35: 1167},
    'CRCP': {25: 1245, 28: 1358, 30: 0, 32: 1509, 35: 1622},
}

CONCRETE_EXCL_JOINT = {
    'JRCP': 830,
    'JPCP': 764,
    'CRCP': 1204,
}

BASE_MATERIAL_PRICES = {
    'Crushed Rock Base Course': 583,
    'Cement Modified Crushed Rock Base (UCS 24.5 ksc)': 864,
    'Cement Treated Base (UCS 40 ksc)': 1096,
    'Soil Aggregate Subbase': 375,
    'Soil Cement Subbase (UCS 7 ksc)': 854,
    'Selected Material A': 375,
    'Embankment': 0,
    'Prime Coat': 37.47,  # บาท/ตร.ม. (ราดก่อนปู AC Interlayer)
}

MATERIAL_LIBRARY = {
    'ผิวทาง': {
        'ผิวทางลาดยาง AC': {'unit_cost': 480, 'cost_unit': 'บาท/ตร.ม.'},
        'ผิวทางลาดยาง PMA': {'unit_cost': 550, 'cost_unit': 'บาท/ตร.ม.'},
        'คอนกรีต 350 Ksc.': {'unit_cost': 800, 'cost_unit': 'บาท/ตร.ม.'},
    },
    'พื้นทาง': {
        'Crushed Rock Base Course': {'unit_cost': 583, 'cost_unit': 'บาท/ลบ.ม.'},
        'Cement Modified Crushed Rock Base (UCS 24.5 ksc)': {'unit_cost': 864, 'cost_unit': 'บาท/ลบ.ม.'},
        'Cement Treated Base (UCS 40 ksc)': {'unit_cost': 1096, 'cost_unit': 'บาท/ลบ.ม.'},
        'Soil Cement Subbase (UCS 7 ksc)': {'unit_cost': 854, 'cost_unit': 'บาท/ลบ.ม.'},
    },
    'รองพื้นทาง': {
        'Soil Aggregate Subbase': {'unit_cost': 375, 'cost_unit': 'บาท/ลบ.ม.'},
        'Selected Material A': {'unit_cost': 375, 'cost_unit': 'บาท/ลบ.ม.'},
    },
    'วัสดุอื่นๆ': {
        'Tack Coat': {'unit_cost': 20, 'cost_unit': 'บาท/ตร.ม.'},
        'Prime Coat': {'unit_cost': 30, 'cost_unit': 'บาท/ตร.ม.'},
        'Non Woven Geotextile': {'unit_cost': 78, 'cost_unit': 'บาท/ตร.ม.'},
    },
}


# ===== BUG FIX 4: _parse_json_details_to_layers() =====
# เดิม: กำหนด qty_unit='cu.m' สำหรับ base materials → ไม่สอดคล้องกับ render_layer_editor() ที่ใช้ 'sq.m' เสมอ
# แก้:  ใช้ qty_unit='sq.m' ทุก layer เสมอ (ตาม convention ใน render_layer_editor)
def _parse_json_details_to_layers(details):
    """แปลง JSON details → (layers, joints) format ที่ app ใช้ภายใน
    
    FIX v5.1: qty_unit='sq.m' เสมอ (เหมือน render_layer_editor) — ไม่ใช้ 'cu.m' อีกต่อไป
    เพราะ calculate_layer_cost() ทำ conversion เองจาก unit_cost และ cost_cum ที่เก็บไว้
    """
    layers, joints = [], []
    for item in details:
        name = item.get('รายการ', '')
        qty = item.get('ปริมาณ', 22000)
        unit_cost = item.get('ราคา/หน่วย', 0)

        if 'Joint' in name or item.get('หน่วย', '') == 'm':
            joints.append({
                'name': name,
                'quantity': qty,
                'qty_unit': 'm',
                'unit_cost': unit_cost,
            })
            continue

        thick_str = str(item.get('ความหนา', '1'))
        try:
            parts = thick_str.split()
            thick_val = float(parts[0])
            unit_val = parts[1] if len(parts) > 1 else 'cm'
        except Exception:
            thick_val = 1.0
            unit_val = 'cm'

        # FIX: qty_unit='sq.m' เสมอ (ไม่แยก cu.m อีกต่อไป)
        layers.append({
            'name': name,
            'thickness': thick_val,
            'unit': unit_val,
            'quantity': qty,
            'qty_unit': 'sq.m',   # ← FIX: เดิมเป็น 'cu.m' สำหรับ base materials
            'unit_cost': unit_cost,
        })
    return layers, joints


def get_default_ac1_layers():
    """AC1: แอสฟัลต์บนหินคลุก"""
    _d = st.session_state.get('loaded_project', {}).get('construction', {}).get('AC1', {})
    if _d.get('layers'):
        return _d['layers']
    if _d.get('details'):
        layers, _ = _parse_json_details_to_layers(_d['details'])
        if layers: return layers
    return [
        {'name': 'Wearing Course', 'thickness': 7, 'unit': 'cm', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 480},
        {'name': 'Binder Course', 'thickness': 7, 'unit': 'cm', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 480},
        {'name': 'Asphalt Base Course', 'thickness': 10, 'unit': 'cm', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 600},
        {'name': 'Tack Coat', 'thickness': 2, 'unit': 'Layer', 'quantity': 44000, 'qty_unit': 'sq.m', 'unit_cost': 20},
        {'name': 'Prime Coat', 'thickness': 1, 'unit': 'Layer', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 30},
        {'name': 'Crushed Rock Base', 'thickness': 20, 'unit': 'cm', 'quantity': 4400, 'qty_unit': 'sq.m', 'unit_cost': 714},
        {'name': 'Soil Aggregate Subbase', 'thickness': 30, 'unit': 'cm', 'quantity': 6600, 'qty_unit': 'sq.m', 'unit_cost': 714},
        {'name': 'Sand Embankment', 'thickness': 40, 'unit': 'cm', 'quantity': 8800, 'qty_unit': 'sq.m', 'unit_cost': 361},
    ]

def get_default_ac2_layers():
    """AC2: แอสฟัลต์บนหินคลุกผสมซีเมนต์"""
    _d = st.session_state.get('loaded_project', {}).get('construction', {}).get('AC2', {})
    if _d.get('layers'):
        return _d['layers']
    if _d.get('details'):
        layers, _ = _parse_json_details_to_layers(_d['details'])
        if layers: return layers
    return [
        {'name': 'Wearing Course', 'thickness': 5, 'unit': 'cm', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 400},
        {'name': 'Binder Course', 'thickness': 5, 'unit': 'cm', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 400},
        {'name': 'Tack Coat', 'thickness': 1, 'unit': 'Layer', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 20},
        {'name': 'Prime Coat', 'thickness': 1, 'unit': 'Layer', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 30},
        {'name': 'Cement Modified Crushed Rock', 'thickness': 20, 'unit': 'cm', 'quantity': 4400, 'qty_unit': 'sq.m', 'unit_cost': 914},
        {'name': 'Soil Aggregate Subbase', 'thickness': 20, 'unit': 'cm', 'quantity': 4400, 'qty_unit': 'sq.m', 'unit_cost': 714},
        {'name': 'Sand Embankment', 'thickness': 30, 'unit': 'cm', 'quantity': 6600, 'qty_unit': 'sq.m', 'unit_cost': 361},
    ]

def get_default_jrcp1_layers():
    """JPCP ชุดที่ 1 (key v6: JPCP1, key เก่า: JRCP1)"""
    _c = st.session_state.get('loaded_project', {}).get('construction', {})
    _d = _c.get('JPCP1') or _c.get('JRCP1') or {}  # รองรับทั้ง key v6 และ key เก่า
    if _d.get('layers'):
        return _d['layers']
    if _d.get('details'):
        layers, _ = _parse_json_details_to_layers(_d['details'])
        if layers: return layers
    return [
        {'name': '350 Ksc. Cubic Type Concrete', 'thickness': 28, 'unit': 'cm', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 800},
        {'name': 'Wire Mesh', 'thickness': 1, 'unit': 'ชั้น', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 100},
        {'name': 'Non Woven Geotextile', 'thickness': 1, 'unit': 'ชั้น', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 78},
        {'name': 'Soil Cement Base', 'thickness': 20, 'unit': 'cm', 'quantity': 4400, 'qty_unit': 'sq.m', 'unit_cost': 621},
        {'name': 'Sand Embankment', 'thickness': 60, 'unit': 'cm', 'quantity': 13200, 'qty_unit': 'sq.m', 'unit_cost': 361},
    ]

def get_default_jrcp1_joints():
    _c = st.session_state.get('loaded_project', {}).get('construction', {})
    _d = _c.get('JPCP1') or _c.get('JRCP1') or {}
    if _d.get('joints'):
        return _d['joints']
    if _d.get('details'):
        _, joints = _parse_json_details_to_layers(_d['details'])
        if joints: return joints
    return [
        {'name': 'Transverse Joint @10m', 'quantity': 2200, 'qty_unit': 'm', 'unit_cost': 430},
        {'name': 'Longitudinal Joint', 'quantity': 4000, 'qty_unit': 'm', 'unit_cost': 120},
    ]

def get_default_jrcp2_layers():
    """JRCP ชุดที่ 1 (key v6: JRCP1, key เก่า: JRCP2)"""
    _c = st.session_state.get('loaded_project', {}).get('construction', {})
    _d = _c.get('JRCP1') or _c.get('JRCP2') or {}
    if _d.get('layers'):
        return _d['layers']
    if _d.get('details'):
        layers, _ = _parse_json_details_to_layers(_d['details'])
        if layers: return layers
    return [
        {'name': '350 Ksc. Cubic Type Concrete', 'thickness': 28, 'unit': 'cm', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 800},
        {'name': 'Wire Mesh', 'thickness': 1, 'unit': 'ชั้น', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 100},
        {'name': 'Non Woven Geotextile', 'thickness': 1, 'unit': 'ชั้น', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 78},
        {'name': 'Cement Modified Crushed Rock', 'thickness': 20, 'unit': 'cm', 'quantity': 4400, 'qty_unit': 'sq.m', 'unit_cost': 914},
        {'name': 'Sand Embankment', 'thickness': 50, 'unit': 'cm', 'quantity': 11000, 'qty_unit': 'sq.m', 'unit_cost': 361},
    ]

def get_default_jrcp2_joints():
    _c = st.session_state.get('loaded_project', {}).get('construction', {})
    _d = _c.get('JRCP1') or _c.get('JRCP2') or {}
    if _d.get('joints'):
        return _d['joints']
    if _d.get('details'):
        _, joints = _parse_json_details_to_layers(_d['details'])
        if joints: return joints
    return [
        {'name': 'Transverse Joint @10m', 'quantity': 2200, 'qty_unit': 'm', 'unit_cost': 430},
        {'name': 'Longitudinal Joint', 'quantity': 4000, 'qty_unit': 'm', 'unit_cost': 120},
    ]

def get_default_crcp1_layers():
    """CRCP ชุดที่ 1 (key v6: CRCP1)"""
    _c = st.session_state.get('loaded_project', {}).get('construction', {})
    _d = _c.get('CRCP1') or {}
    if _d.get('layers'):
        return _d['layers']
    if _d.get('details'):
        layers, _ = _parse_json_details_to_layers(_d['details'])
        if layers: return layers
    return [
        {'name': '350 Ksc. Cubic Type Concrete', 'thickness': 25, 'unit': 'cm', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 850},
        {'name': 'Wire Mesh', 'thickness': 1, 'unit': 'ชั้น', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 100},
        {'name': 'Non Woven Geotextile', 'thickness': 1, 'unit': 'ชั้น', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 78},
        {'name': 'Soil Cement Base', 'thickness': 15, 'unit': 'cm', 'quantity': 3300, 'qty_unit': 'sq.m', 'unit_cost': 621},
        {'name': 'Sand Embankment', 'thickness': 50, 'unit': 'cm', 'quantity': 11000, 'qty_unit': 'sq.m', 'unit_cost': 361},
    ]

def get_default_crcp2_layers():
    """CRCP ชุดที่ 2 (key v6: CRCP2)"""
    _c = st.session_state.get('loaded_project', {}).get('construction', {})
    _d = _c.get('CRCP2') or {}
    if _d.get('layers'):
        return _d['layers']
    if _d.get('details'):
        layers, _ = _parse_json_details_to_layers(_d['details'])
        if layers: return layers
    return [
        {'name': '350 Ksc. Cubic Type Concrete', 'thickness': 25, 'unit': 'cm', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 850},
        {'name': 'Wire Mesh', 'thickness': 1, 'unit': 'ชั้น', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 100},
        {'name': 'Non Woven Geotextile', 'thickness': 1, 'unit': 'ชั้น', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 78},
        {'name': 'Cement Modified Crushed Rock', 'thickness': 15, 'unit': 'cm', 'quantity': 3300, 'qty_unit': 'sq.m', 'unit_cost': 914},
        {'name': 'Sand Embankment', 'thickness': 40, 'unit': 'cm', 'quantity': 8800, 'qty_unit': 'sq.m', 'unit_cost': 361},
    ]

def get_default_crcp1_joints():
    _c = st.session_state.get('loaded_project', {}).get('construction', {})
    _d = _c.get('CRCP1') or {}
    if _d.get('joints'):
        return _d['joints']
    if _d.get('details'):
        _, joints = _parse_json_details_to_layers(_d['details'])
        if joints: return joints
    return [
        {'name': 'Longitudinal Steel (CRCP)', 'quantity': 4000, 'qty_unit': 'm', 'unit_cost': 200},
        {'name': 'Transverse Joint (End)',     'quantity': 0,    'qty_unit': 'm', 'unit_cost': 500},
    ]

def get_default_crcp2_joints():
    _c = st.session_state.get('loaded_project', {}).get('construction', {})
    _d = _c.get('CRCP2') or {}
    if _d.get('joints'):
        return _d['joints']
    if _d.get('details'):
        _, joints = _parse_json_details_to_layers(_d['details'])
        if joints: return joints
    return [
        {'name': 'Longitudinal Steel (CRCP)', 'quantity': 4000, 'qty_unit': 'm', 'unit_cost': 200},
        {'name': 'Transverse Joint (End)',     'quantity': 0,    'qty_unit': 'm', 'unit_cost': 500},
    ]


def calculate_quantity(thickness_cm, width_m, length_km, qty_unit):
    """คำนวณปริมาณจากความหนา ความกว้าง และความยาว"""
    area = width_m * length_km * 1000
    if qty_unit == 'sq.m':
        return area
    elif qty_unit == 'cu.m':
        return area * thickness_cm / 100
    return area


def calculate_layer_cost(layers, road_length_km=1.0):
    """คำนวณค่าก่อสร้างจากชั้นโครงสร้าง
    หมายเหตุ: layer['quantity'] = ปริมาณรวมทั้งโครงการ (ตร.ม.) คำนวณโดย render_layer_editor()
    ดังนั้น cost = quantity * unit_cost โดยตรง ไม่ต้องคูณ road_length อีก
    """
    total = 0
    details = []

    BASE_KEYWORDS = ['crushed rock', 'soil aggregate', 'soil cement', 'cement modified',
                     'cement treated', 'selected material', 'sand embankment']

    for layer in layers:
        qty_raw = layer['quantity']      # ตร.ม. (รวมทั้งโครงการ)
        unit_cost = layer['unit_cost']   # บาท/ตร.ม.
        cost = qty_raw * unit_cost
        total += cost

        name_lower = layer['name'].lower()
        is_base = any(kw in name_lower for kw in BASE_KEYWORDS)

        if is_base:
            thick_cm = float(layer.get('thickness', 1))
            u = layer.get('unit', 'cm').lower()
            if 'cost_cum' in layer and layer['cost_cum']:
                price_cum = layer['cost_cum']
            elif thick_cm > 0 and u in ('cm', 'ซม.', 'ซ.ม.'):
                price_cum = unit_cost / (thick_cm / 100)
            else:
                price_cum = unit_cost

            if thick_cm > 0 and u in ('cm', 'ซม.', 'ซ.ม.'):
                qty_display = qty_raw * thick_cm / 100
            else:
                qty_display = qty_raw

            display_unit = 'ลบ.ม.'
            display_price_str = f"{price_cum:,.0f}"
            display_price_label = 'บาท/ลบ.ม.'
            qty_show = qty_display
        else:
            qty_show = qty_raw
            display_unit = 'ตร.ม.'
            display_price_str = f"{unit_cost:,.0f}"
            display_price_label = 'บาท/ตร.ม.'

        details.append({
            'รายการ': layer['name'],
            'ความหนา': f"{layer['thickness']} {layer['unit']}",
            'ปริมาณ': qty_show,
            'หน่วย': display_unit,
            'ราคา/หน่วย': unit_cost,
            'ราคา/หน่วย (แสดง)': display_price_str,
            'หน่วยราคา': display_price_label,
            'มูลค่า (บาท)': cost,
        })

    return total, details


def calculate_joint_cost(joints, road_length_km=1.0, include_joints=True):
    """คำนวณค่ารอยต่อ
    หมายเหตุ: joint['quantity'] = ปริมาณรวมทั้งโครงการ (ม.) แล้ว
    road_length_km ยังคงส่งมาเพื่อ compatibility แต่ไม่ใช้คูณซ้ำ
    """
    total = 0
    details = []

    for joint in joints:
        qty = joint['quantity']   # ม. รวมทั้งโครงการ
        cost = qty * joint['unit_cost']if include_joints else 0
        total += cost

        unit_th = 'ม.' if joint.get('qty_unit', 'm') == 'm' else joint.get('qty_unit', 'm')
        details.append({
            'รายการ': joint['name'],
            'ความหนา': '-',
            'ปริมาณ': qty,
            'หน่วย': unit_th,
            'ราคา/หน่วย': joint['unit_cost'],
            'ราคา/หน่วย (แสดง)': f"{joint['unit_cost']:,.0f}",
            'หน่วยราคา': 'บาท/ม.',
            'มูลค่า (บาท)': cost,
        })

    return total, details


def get_price_from_library(layer_name, thickness):
    """ดึงราคาจาก Library ตามชื่อและความหนา"""
    if 'price_library' not in st.session_state:
        return None

    lib = st.session_state['price_library']
    name_lower = layer_name.lower()

    if 'pma' in name_lower and 'wearing' in name_lower:
        return lib['ac_prices'].get('PMA Wearing Course', {}).get(thickness)
    elif 'wearing' in name_lower:
        return lib['ac_prices'].get('AC Wearing Course', {}).get(thickness)
    elif 'binder' in name_lower:
        return lib['ac_prices'].get('AC Binder Course', {}).get(thickness)
    elif 'asphalt' in name_lower and 'base' in name_lower:
        return lib['ac_prices'].get('AC Base Course', {}).get(thickness)
    elif 'jrcp' in name_lower or ('concrete' in name_lower and 'jrcp' in str(thickness)):
        return lib['concrete_prices'].get('JRCP', {}).get(int(thickness))
    elif 'jpcp' in name_lower:
        return lib['concrete_prices'].get('JPCP', {}).get(int(thickness))
    elif 'crcp' in name_lower:
        return lib['concrete_prices'].get('CRCP', {}).get(int(thickness))
    elif 'crushed rock' in name_lower and 'cement' not in name_lower:
        return lib['base_prices'].get('Crushed Rock Base Course')
    elif 'cement modified' in name_lower or 'cmcr' in name_lower:
        return lib['base_prices'].get('Cement Modified Crushed Rock Base (UCS 24.5 ksc)')
    elif 'cement treated' in name_lower or 'ctb' in name_lower:
        return lib['base_prices'].get('Cement Treated Base (UCS 40 ksc)')
    elif 'soil aggregate' in name_lower:
        return lib['base_prices'].get('Soil Aggregate Subbase')
    elif 'soil cement' in name_lower:
        return lib['base_prices'].get('Soil Cement Subbase (UCS 7 ksc)')
    elif 'selected' in name_lower:
        return lib['base_prices'].get('Selected Material A')

    return None


def render_layer_editor(layers, key_prefix, total_width, road_length, v=0, ptype='AC'):
    """แสดง UI สำหรับแก้ไขโครงสร้างชั้นทาง พร้อมคำนวณปริมาณอัตโนมัติ

    ptype : 'AC' | 'JPCP' | 'JRCP' | 'CRCP'
            ใช้กำหนดชื่อ concrete layer แบบตายตัว (ไม่มี dropdown อีกต่อไป)
    การเปลี่ยนแปลง v6.1:
      - ลบ dropdown concrete_options ออก → แสดงชื่อตายตัวตาม ptype
      - Wearing course ยังเลือก AC/PMA ได้ (เป็นทางเลือกจริง)
      - Binder / AC Base แสดงชื่อตายตัว (ไม่มีทางเลือกอื่น)
    """
    updated_layers = []
    area_per_km = total_width * 1000

    # concrete_label ตายตัวตาม ptype
    concrete_label = f"350 Ksc. Cubic Type Concrete ({ptype})"

    surface_layers = []
    base_layers = []

    for layer in layers:
        name_lower = layer['name'].lower()
        if any(x in name_lower for x in [
            'wearing', 'binder', 'asphalt', 'concrete', 'tack', 'prime',
            'geotextile', 'steel', 'wire', 'ac base', 'ac wearing', 'ac binder',
        ]):
            surface_layers.append(layer)
        else:
            base_layers.append(layer)

    st.markdown("**ผิวทาง** (หน่วย: ตร.ม.)")
    cols = st.columns([3, 1, 1.5])
    cols[0].markdown("รายการ")
    cols[1].markdown("หนา (cm)")
    cols[2].markdown("ราคา (บาท/ตร.ม.)")

    wearing_options = ['AC Wearing Course', 'PMA Wearing Course']

    for i, layer in enumerate(surface_layers):
        name_lower = layer['name'].lower()
        is_wearing  = 'wearing' in name_lower
        is_binder   = 'binder' in name_lower
        is_ac_base  = ('asphalt' in name_lower and 'base' in name_lower) or \
                      ('ac base' in name_lower) or ('interlayer' in name_lower)
        is_concrete = 'concrete' in name_lower or 'ksc' in name_lower
        is_geotextile = 'geotextile' in name_lower
        is_wire_mesh  = 'wire' in name_lower

        # JPCP ไม่มี Wire Mesh
        if is_wire_mesh and ptype == 'JPCP':
            continue

        # Non Woven Geotextile — แสดงเป็น checkbox แยกออกมา
        if is_geotextile:
            geo_cols = st.columns([2.5, 1, 1.5])
            with geo_cols[0]:
                use_geo = st.checkbox(
                    "Non Woven Geotextile",
                    value=True,
                    key=f"{key_prefix}_use_geo_{i}_v{v}",
                )
            if not use_geo:
                continue   # ไม่ใส่ → ข้ามไปเลย
            with geo_cols[1]:
                st.markdown("**1 ชั้น**")   # ปริมาณ 1 ตร.ม. เสมอ
            with geo_cols[2]:
                geo_cost = st.number_input(
                    "ราคา (บาท/ตร.ม.)",
                    value=float(layer['unit_cost']),
                    min_value=0.0, step=5.0,
                    key=f"{key_prefix}_geo_price_{i}_v{v}",
                    label_visibility="collapsed",
                )
            auto_qty = area_per_km * road_length
            updated_layers.append({
                'name': 'Non Woven Geotextile',
                'thickness': 1, 'unit': 'ชั้น',
                'quantity': auto_qty, 'qty_unit': 'sq.m',
                'unit_cost': geo_cost, 'cost_per_sqm': geo_cost,
            })
            continue   # จัดการครบแล้ว ข้ามไป loop ถัดไป

        cols = st.columns([3, 1, 1.5])

        with cols[0]:
            if is_wearing:
                # Wearing: ยังเลือก AC/PMA ได้ (ทางเลือกจริง)
                default_idx = 1 if 'pma' in name_lower else 0
                selected_material = st.selectbox(
                    "วัสดุ", wearing_options, index=default_idx,
                    key=f"{key_prefix}_mat_{i}_v{v}", label_visibility="collapsed"
                )
            elif is_concrete:
                # Concrete: ชื่อตายตัวตาม ptype — ไม่มี dropdown
                st.markdown(f"**{concrete_label}**")
                selected_material = concrete_label
            elif is_binder:
                st.markdown("**AC Binder Course**")
                selected_material = 'AC Binder Course'
            elif is_ac_base:
                st.markdown("**AC Base Course**")
                selected_material = 'AC Base Course'
            else:
                st.text(layer['name'])
                selected_material = layer['name']

        with cols[1]:
            thick = st.number_input(
                "หนา", value=float(layer['thickness']),
                key=f"{key_prefix}_st_{i}_v{v}", label_visibility="collapsed",
                min_value=0.0, step=1.0
            )

        _unit_low = layer.get('unit', 'cm').lower()
        auto_qty = area_per_km * road_length * thick if _unit_low == 'layer' and thick > 1 \
                   else area_per_km * road_length

        # ดึงราคาจาก Library
        lib_price = None
        if 'price_library' in st.session_state:
            lib = st.session_state['price_library']
            if is_wearing:
                prices = lib['ac_prices'].get(selected_material, {})
                lib_price = prices.get(thick) or (
                    lib['ac_prices'][selected_material].get(
                        min(prices.keys(), key=lambda x: abs(x - thick))
                    ) if prices else None
                )
            elif is_binder:
                prices = lib['ac_prices'].get('AC Binder Course', {})
                lib_price = prices.get(thick) or (
                    prices.get(min(prices.keys(), key=lambda x: abs(x - thick))) if prices else None
                )
            elif is_ac_base:
                prices = lib['ac_prices'].get('AC Base Course', {})
                lib_price = prices.get(thick) or (
                    prices.get(min(prices.keys(), key=lambda x: abs(x - thick))) if prices else None
                )
            elif is_concrete:
                prices = lib['concrete_prices'].get(ptype, {})
                lib_price = prices.get(int(thick)) or (
                    prices.get(min(prices.keys(), key=lambda x: abs(x - thick))) if prices else None
                )

        default_cost = lib_price if lib_price else layer['unit_cost']

        with cols[2]:
            if is_concrete or is_wearing or is_binder or is_ac_base:
                # ราคาดึงจาก Library อัตโนมัติ → อ่านอย่างเดียว
                st.markdown(f"**{default_cost:,.2f}**")
            else:
                # geotextile / wire mesh / tack / prime → แก้ไขได้
                default_cost = st.number_input(
                    "ราคา", value=float(default_cost), min_value=0.0, step=5.0,
                    key=f"{key_prefix}_price_{i}_v{v}", label_visibility="collapsed"
                )

        if thick == 0:
            continue
        updated_layers.append({
            'name': selected_material, 'thickness': thick, 'unit': layer['unit'],
            'quantity': auto_qty, 'qty_unit': 'sq.m', 'unit_cost': default_cost,
            'cost_per_sqm': default_cost,
        })

    # ===== ส่วนพื้นทาง/รองพื้นทาง =====
    st.markdown("---")
    st.markdown("**พื้นทาง/รองพื้นทาง** (ราคาแสดงเป็น บาท/ตร.ม.)")

    # FIX v6.1: ใช้ ptype แทน key_prefix เพื่อตรวจว่าเป็นคอนกรีตหรือไม่
    is_concrete_pavement = ptype in ('JPCP', 'JRCP', 'CRCP')

    # ===== AC Interlayer — checkbox แยกต่างหาก =====
    if is_concrete_pavement:
        col_acil1, col_acil2 = st.columns([2, 2])
        with col_acil1:
            use_ac_interlayer = st.checkbox(
                "มี AC Interlayer รองใต้คอนกรีต",
                value=True,
                key=f"{key_prefix}_use_acil_v{v}",
                help="ชั้น AC ที่รองใต้แผ่นคอนกรีต ทั่วไปใช้ 5 cm"
            )
        if use_ac_interlayer:
            with col_acil2:
                # ดึงราคาจาก Library
                if 'price_library' in st.session_state:
                    _ac_prices = st.session_state['price_library']['ac_prices'].get('AC Base Course', {})
                    _acil_price_default = _ac_prices.get(5.0, 251)
                else:
                    _acil_price_default = 251

                acil_thick = st.number_input(
                    "ความหนา AC Interlayer (cm)",
                    value=5.0, min_value=1.0, max_value=10.0, step=1.0,
                    key=f"{key_prefix}_acil_thick_v{v}",
                )

            # คำนวณราคา AC Interlayer ตามความหนาที่เลือก
            if 'price_library' in st.session_state:
                _ac_prices = st.session_state['price_library']['ac_prices'].get('AC Binder Course', {})
                acil_cost_sqm = _ac_prices.get(acil_thick, 0)
                if acil_cost_sqm == 0 and _ac_prices:
                    acil_cost_sqm = _ac_prices.get(
                        min(_ac_prices.keys(), key=lambda x: abs(x - acil_thick)), 251
                    )
            else:
                acil_cost_sqm = _acil_price_default

            acil_qty = area_per_km * road_length
            st.markdown(
                f'<div class="cost-box" style="border-left-color:#378ADD;padding:6px 10px;margin:4px 0;">'
                f'AC Interlayer {acil_thick:.0f} cm &nbsp;|&nbsp; '
                f'ราคา <b>{acil_cost_sqm:,.2f}</b> บาท/ตร.ม.</div>',
                unsafe_allow_html=True
            )
            # เพิ่มเข้า updated_layers ทันที
            updated_layers.append({
                'name': f'AC Interlayer ({acil_thick:.0f} cm)',
                'thickness': acil_thick, 'unit': 'cm',
                'quantity': acil_qty, 'qty_unit': 'sq.m',
                'unit_cost': acil_cost_sqm, 'cost_per_sqm': acil_cost_sqm,
            })

            # ===== Prime Coat — แสดงเมื่อมี AC Interlayer เท่านั้น =====
            col_pc1, col_pc2 = st.columns([2, 2])
            with col_pc1:
                use_prime_coat = st.checkbox(
                    "มี Prime Coat",
                    value=True,
                    key=f"{key_prefix}_use_pc_v{v}",
                    help="ราดบนชั้น Base ก่อนปู AC Interlayer (บาท/ตร.ม.)"
                )
            if use_prime_coat:
                # ดึงราคาจาก Library
                if 'price_library' in st.session_state:
                    _pc_default = st.session_state['price_library']['base_prices'].get('Prime Coat', 37.47)
                else:
                    _pc_default = 37.47

                with col_pc2:
                    pc_cost = st.number_input(
                        "ราคา Prime Coat (บาท/ตร.ม.)",
                        value=float(_pc_default),
                        min_value=0.0, step=1.0,
                        key=f"{key_prefix}_pc_price_v{v}",
                    )
                st.markdown(
                    f'<div class="cost-box" style="border-left-color:#FFA500;padding:6px 10px;margin:4px 0;">'
                    f'Prime Coat &nbsp;|&nbsp; '
                    f'ราคา <b>{pc_cost:,.2f}</b> บาท/ตร.ม.</div>',
                    unsafe_allow_html=True
                )
                updated_layers.append({
                    'name': 'Prime Coat',
                    'thickness': 1, 'unit': 'Layer',
                    'quantity': area_per_km * road_length, 'qty_unit': 'sq.m',
                    'unit_cost': pc_cost, 'cost_per_sqm': pc_cost,
                })

    # ── วัสดุพื้นทาง/รองพื้นทาง (ไม่รวม AC Interlayer อีกต่อไป) ──
    _HARDCODED_BASE = {
        'Crushed Rock Base Course',
        'Cement Modified Crushed Rock Base (UCS 24.5 ksc)',
        'Cement Treated Base (UCS 40 ksc)',
        'Soil Cement Subbase (UCS 7 ksc)',
        'Soil Aggregate Subbase',
        'Selected Material A',
    }
    if 'price_library' in st.session_state:
        base_lib = st.session_state['price_library']['base_prices']
        base_materials = {
            'Crushed Rock Base Course':                          {'unit_cost_cum': base_lib.get('Crushed Rock Base Course', 583),                          'is_ac': False},
            'Cement Modified Crushed Rock Base (UCS 24.5 ksc)': {'unit_cost_cum': base_lib.get('Cement Modified Crushed Rock Base (UCS 24.5 ksc)', 864),  'is_ac': False},
            'Cement Treated Base (UCS 40 ksc)':                 {'unit_cost_cum': base_lib.get('Cement Treated Base (UCS 40 ksc)', 1096),                 'is_ac': False},
            'Soil Cement Subbase (UCS 7 ksc)':                  {'unit_cost_cum': base_lib.get('Soil Cement Subbase (UCS 7 ksc)', 854),                   'is_ac': False},
            'Soil Aggregate Subbase':                            {'unit_cost_cum': base_lib.get('Soil Aggregate Subbase', 375),                            'is_ac': False},
            'Selected Material A':                               {'unit_cost_cum': base_lib.get('Selected Material A', 375),                               'is_ac': False},
        }
        # เพิ่มวัสดุที่ไม่ได้อยู่ใน hardcode list (Embankment, Custom Materials)
        for _mat, _price in base_lib.items():
            if _mat not in _HARDCODED_BASE:
                try:
                    base_materials[_mat] = {'unit_cost_cum': float(_price), 'is_ac': False}
                except (ValueError, TypeError):
                    pass
    else:
        base_materials = {
            'Crushed Rock Base Course':                          {'unit_cost_cum': 583,  'is_ac': False},
            'Cement Modified Crushed Rock Base (UCS 24.5 ksc)': {'unit_cost_cum': 864,  'is_ac': False},
            'Cement Treated Base (UCS 40 ksc)':                 {'unit_cost_cum': 1096, 'is_ac': False},
            'Soil Cement Subbase (UCS 7 ksc)':                  {'unit_cost_cum': 854,  'is_ac': False},
            'Soil Aggregate Subbase':                            {'unit_cost_cum': 375,  'is_ac': False},
            'Selected Material A':                               {'unit_cost_cum': 375,  'is_ac': False},
            'Embankment':                                        {'unit_cost_cum': 0,    'is_ac': False},
        }

    material_names = list(base_materials.keys())

    # กรอง base_layers ที่ไม่ใช่ AC Interlayer (เพราะ Interlayer ถูกจัดการแยกด้านบนแล้ว)
    base_layers_filtered = [
        bl for bl in base_layers
        if 'interlayer' not in bl['name'].lower()
    ]
    num_base_default = len(base_layers_filtered) if base_layers_filtered else 0
    num_base = st.number_input("จำนวนชั้นพื้นทาง/รองพื้นทาง", value=num_base_default,
                                min_value=0, max_value=5, key=f"{key_prefix}_num_base_v{v}")

    cols = st.columns([3, 1, 1.2, 1.2, 1.2])
    cols[0].markdown("วัสดุ")
    cols[1].markdown("หนา (cm)")
    cols[2].markdown("ปริมาณ (ตร.ม.)")
    cols[3].markdown("ราคา (บาท/ลบ.ม.)")
    cols[4].markdown("ราคา (บาท/ตร.ม.)")

    for i in range(int(num_base)):
        cols = st.columns([3, 1, 1.2, 1.2, 1.2])

        if i < len(base_layers_filtered):
            default_name = base_layers_filtered[i]['name']
            default_thick = base_layers_filtered[i]['thickness']
        else:
            default_name = material_names[0]
            default_thick = 20.0

        try:
            default_idx = material_names.index(default_name)
        except ValueError:
            default_idx = 0
            dn_lower = default_name.lower()
            for mi, mn in enumerate(material_names):
                if mn.lower() in dn_lower or dn_lower in mn.lower():
                    default_idx = mi
                    break

        with cols[0]:
            selected = st.selectbox("วัสดุ", material_names, index=default_idx,
                key=f"{key_prefix}_bm_{i}_v{v}", label_visibility="collapsed")
        with cols[1]:
            if base_materials[selected].get('is_ac', False):
                default_thick_val = base_materials[selected].get('default_thick', 5)
                thick = st.number_input("หนา", value=float(default_thick_val),
                    key=f"{key_prefix}_bt_{i}_v{v}", label_visibility="collapsed",
                    min_value=0.0, step=1.0)
            else:
                thick = st.number_input("หนา", value=float(default_thick),
                    key=f"{key_prefix}_bt_{i}_v{v}", label_visibility="collapsed",
                    min_value=0.0, step=5.0)

        auto_qty = area_per_km * road_length  # ตร.ม. รวมทั้งโครงการ

        if base_materials[selected].get('is_ac', False):
            if 'price_library' in st.session_state:
                ac_prices = st.session_state['price_library']['ac_prices'].get('AC Base Course', {})
                cost_per_sqm = ac_prices.get(thick, 0)
                if cost_per_sqm == 0 and ac_prices:
                    closest = min(ac_prices.keys(), key=lambda x: abs(x - thick))
                    cost_per_sqm = ac_prices.get(closest, 251)
            else:
                cost_per_sqm = 251
            lib_cost_cum = cost_per_sqm
        else:
            lib_cost_cum = base_materials[selected]['unit_cost_cum']
            cost_per_sqm = lib_cost_cum * thick / 100

        with cols[2]:
            st.text(f"{auto_qty:,.0f}")
        with cols[3]:
            if base_materials[selected].get('is_ac', False):
                st.markdown("**-**")
            else:
                st.markdown(f"**{lib_cost_cum:,.2f}**")
        with cols[4]:
            st.markdown(f"**{cost_per_sqm:,.2f}**")

        if thick == 0:
            continue
        updated_layers.append({
            'name': selected, 'thickness': thick, 'unit': 'cm',
            'quantity': auto_qty, 'qty_unit': 'sq.m', 'unit_cost': cost_per_sqm,
            'cost_per_sqm': cost_per_sqm,
            'cost_cum': lib_cost_cum,
        })

    return updated_layers


# ===== BUG FIX 3: render_joint_editor() =====
# เดิม (CRCP branch): ใช้ `_pi = st.session_state.get('project_info', {})` เพื่อหา lane_width
#   → session_state['project_info'] ถูก set หลัง render_layer_editor() ใน Tab 2
#   → ตอน first render อาจยังไม่มี → _lane_w = 3.5 (ยัง OK เพราะมี default)
#   แต่ปัญหาจริงคือ width_m ควรใช้ area_per_km ที่รับมาเป็น param โดยตรง
# แก้: คำนวณ width_m = area_per_km / 1000 (ซึ่งถูกส่งมาแล้ว) แทนการอ่าน session_state
def render_joint_editor(joints, key_prefix, area_per_km, road_length, v=0, ptype='JRCP'):
    """แสดง UI สำหรับแก้ไขรอยต่อ

    v6.1: รับ ptype param โดยตรง — ไม่ต้องอ่าน ctype จาก session_state อีกต่อไป
    ptype : 'JPCP' | 'JRCP' | 'CRCP'
    """
    st.markdown("---")

    concrete_type = ptype  # ← ใช้ param โดยตรง แทนการ scan session_state

    if concrete_type == 'JPCP':
        joint_spacing = 4
        joint_label = '@4m'
    else:
        joint_spacing = 10
        joint_label = '@10m'

    adjusted_joints = []
    for j in joints:
        jname = j['name']

        if concrete_type == 'CRCP':
            width_m = area_per_km / 1000
            lane_w = st.session_state.get('project_info', {}).get('lane_width', 3.5)
            if not lane_w or lane_w <= 0:
                lane_w = 3.5

            if 'steel' in jname.lower() or 'longitudinal' in jname.lower():
                # Longitudinal Steel = คำนวณเหมือน Longitudinal Joint
                # จำนวนแถวเหล็ก = จำนวนรอยต่อตามยาว = round(กว้าง/ช่อง) - 1
                num_lj = max(1, round(width_m / lane_w) - 1)
                adj_qty = num_lj * road_length * 1000   # ม.
                jname = 'Longitudinal Steel (CRCP)'
            else:
                # Transverse Joint ของ CRCP = ให้ผู้ใช้กรอกเอง (default 0)
                adj_qty = j['quantity']  # ใช้ค่าที่บันทึกไว้ (default=0)
                jname = 'Transverse Joint (End)'

        elif 'Transverse Joint' in jname:
            # JPCP/JRCP: คำนวณจาก spacing
            jname = f"Transverse Joint {joint_label}"
            width_m = area_per_km / 1000
            adj_qty = (road_length * 1000 / joint_spacing) * width_m

        elif 'Longitudinal' in jname:
            # JPCP/JRCP: Longitudinal Joint คำนวณจากจำนวนรอยต่อตามยาว
            width_m = area_per_km / 1000
            lane_w = st.session_state.get('project_info', {}).get('lane_width', 3.5)
            if not lane_w or lane_w <= 0:
                lane_w = 3.5
            num_lj = max(1, round(width_m / lane_w) - 1)
            adj_qty = num_lj * road_length * 1000
        else:
            adj_qty = j['quantity']

        adjusted_joints.append({**j, 'name': jname, 'quantity': adj_qty})

    col_header = st.columns([3, 1])
    with col_header[0]:
        if concrete_type == 'CRCP':
            st.markdown("**Longitudinal Steel & Transverse Joint (CRCP)**")
        else:
            st.markdown(f"**รอยต่อ (Joints) — {concrete_type} ระยะ {joint_spacing} ม.**")
    with col_header[1]:
        _cb_label = "รวมราคา Longitudinal Steel & Joints" if concrete_type == 'CRCP' else "รวมราคา Joints"
        include_joints = st.checkbox(_cb_label, value=True, key=f"{key_prefix}_include_joints_v{v}")

    cols = st.columns([3, 1.5, 1.5, 1.5])
    cols[0].markdown("รายการ")
    cols[1].markdown("ปริมาณ (m)")
    cols[2].markdown("ราคา/ม. (บาท)")
    cols[3].markdown("ราคา (บาท/ตร.ม. รวม)")

    updated_joints = []
    total_area = area_per_km * road_length

    for i, joint in enumerate(adjusted_joints):
        cols = st.columns([3, 1.5, 1.5, 1.5])

        with cols[0]:
            st.text(joint['name'])

        with cols[1]:
            qty = st.number_input(
                "ปริมาณ (m)", value=float(joint['quantity']),
                key=f"{key_prefix}_jq_{i}_s{joint_spacing}_v{v}", label_visibility="collapsed",
                min_value=0.0, step=100.0
            )

        with cols[2]:
            cost = st.number_input(
                "ราคา/ม.", value=float(joint['unit_cost']),
                key=f"{key_prefix}_jc_{i}_s{joint_spacing}_v{v}", label_visibility="collapsed",
                min_value=0.0, step=10.0
            )

        joint_total = qty * cost
        cost_per_sqm = joint_total / total_area if total_area > 0 else 0

        with cols[3]:
            st.markdown(f"**{cost_per_sqm:.2f}**")

        updated_joints.append({
            'name': joint['name'],
            'quantity': qty,
            'qty_unit': joint['qty_unit'],
            'unit_cost': cost,
            'cost_per_sqm': cost_per_sqm,
        })

    return updated_joints, include_joints


# ===== BUG FIX 1 & 2: generate_word_report_table() ถูก REMOVED (dead code) =====
# ฟังก์ชันนี้ถูก define ใน v5.0 แต่ไม่มีที่ไหนเรียกใช้เลย (Tab 3 เรียก generate_word_report_materials_only
# และ generate_word_report_consultant แทน)
# นอกจากนี้ยังมี Bug:
#   - qty_raw = layer['quantity'] * road_length  → คูณ road_length ซ้ำ (layer['quantity'] รวม road_length แล้ว)
#   - num_rows = 2 + len(surface_layers) + 1 → อาจไม่ตรงกับจำนวนแถวที่ใส่จริง → index error
# การลบออกเป็น clean fix ที่สุดเพราะไม่ได้ใช้งาน
# (ถ้าต้องการใช้ในอนาคต ให้สร้างใหม่โดยใช้ dynamic row append แทน pre-allocate)


def generate_word_report_materials_only(project_info, all_details):
    """สร้างรายงาน Word - เฉพาะวัสดุและราคา (ไม่มี NPV) พร้อมตารางสรุปแยกชนิด"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx ไม่สามารถใช้งานได้")

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'TH SarabunPSK'
    style.font.size = Pt(16)

    doc.add_heading('รายงานวัสดุและราคาโครงสร้างชั้นทาง', 0)

    doc.add_heading('1. ข้อมูลโครงการ', level=1)
    doc.add_paragraph(f"ชื่อโครงการ: {project_info.get('name', '-')}")
    doc.add_paragraph(f"ความยาวถนน: {project_info.get('length', 1):.2f} กม.")
    doc.add_paragraph(f"ความกว้างรวม: {project_info.get('total_width', 0):.2f} ม.")
    doc.add_paragraph(f"จำนวนช่องจราจร: {project_info.get('num_lanes', 2)} ช่อง")

    doc.add_heading('2. รายละเอียดวัสดุและราคา', level=1)

    summary_data = []
    length = project_info.get('length', 1)

    for ptype, data in all_details.items():
        structure_name = data.get('name', ptype)          # ชื่อสั้น: AC / JPCP / JRCP / CRCP
        name_detail    = data.get('name_detail', '')      # ชื่อยาว: รายละเอียด
        details = data.get('details', [])

        doc.add_heading(structure_name, level=2)
        if name_detail and name_detail != structure_name:
            p_sub = doc.add_paragraph(name_detail)
            p_sub.runs[0].italic = True
            p_sub.runs[0].font.size = Pt(13)
        if details:
            table = doc.add_table(rows=len(details) + 1, cols=5)
            table.style = 'Table Grid'

            headers = ['รายการ', 'ปริมาณ', 'หน่วย', 'ราคา/หน่วย (บาท)', 'มูลค่า (บาท)']
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = h
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            subtotal = 0
            for i, d in enumerate(details):
                unit_display = d.get('หน่วย', 'ตร.ม.')
                price_display = d.get('ราคา/หน่วย (แสดง)', f"{d['ราคา/หน่วย']:,.0f}")
                _fallback_price_label = {'ม.': 'บาท/ม.', 'ลบ.ม.': 'บาท/ลบ.ม.'}.get(unit_display, 'บาท/ตร.ม.')
                price_label = d.get('หน่วยราคา', _fallback_price_label)

                table.rows[i + 1].cells[0].text = str(d['รายการ'])
                table.rows[i + 1].cells[1].text = f"{d['ปริมาณ']:,.0f}"
                table.rows[i + 1].cells[2].text = unit_display
                table.rows[i + 1].cells[3].text = f"{price_display} ({price_label})"
                table.rows[i + 1].cells[4].text = f"{d['มูลค่า (บาท)']:,.0f}"
                subtotal += d['มูลค่า (บาท)']

            doc.add_paragraph(f"รวม {structure_name}: {subtotal:,.0f} บาท", style='Intense Quote')
            doc.add_paragraph()

            cost_per_km_million = data.get('cost_per_km', 0)
            cost_per_sqm = data.get('cost_sqm', 0)

            summary_data.append({
                'name': structure_name,
                'total_value': subtotal,
                'cost_per_km_million': cost_per_km_million,
                'cost_per_sqm': cost_per_sqm,
            })

    doc.add_heading('3. สรุปค่าใช้จ่าย', level=1)

    if summary_data:
        table = doc.add_table(rows=len(summary_data) + 1, cols=4)
        table.style = 'Table Grid'

        headers = ['ชนิดโครงสร้าง', 'มูลค่ารวม/กม. (บาท)', 'ราคา/กม. (ล้านบาท)', 'ราคา/ตร.ม. (บาท)']
        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = h
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        for i, item in enumerate(summary_data):
            total_per_km = item['total_value'] / length if length > 0 else 0
            table.rows[i + 1].cells[0].text = item['name']
            table.rows[i + 1].cells[1].text = f"{total_per_km:,.0f}"
            table.rows[i + 1].cells[2].text = f"{item['cost_per_km_million']:.2f}"
            table.rows[i + 1].cells[3].text = f"{item['cost_per_sqm']:,.2f}"

    doc.add_paragraph()
    doc.add_paragraph(f"รายงานสร้างเมื่อ: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    return doc


def generate_word_report_consultant(project_info, all_details, chapter_num="4", section_start="4.7", intro_text=""):
    """สร้างรายงาน Word แบบที่ปรึกษา - มีหมายเลขหัวข้อและบทเกริ่นนำ"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx ไม่สามารถใช้งานได้")

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'TH SarabunPSK'
    style.font.size = Pt(16)

    def _set_run_font(run, size=16, bold=False):
        run.font.name = 'TH SarabunPSK'
        run.font.size = Pt(size)
        run.font.bold = bold
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')
        rFonts.set(qn('w:ascii'), 'TH SarabunPSK')
        rFonts.set(qn('w:hAnsi'), 'TH SarabunPSK')

    def _add_heading_para(text, size=16, bold=True, underline=False, space_before=6, space_after=3):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after = Pt(space_after)
        run = para.add_run(text)
        _set_run_font(run, size=size, bold=bold)
        run.underline = underline
        return para

    sec_main    = section_start
    sec_info    = f"{section_start}.1"
    sec_detail  = f"{section_start}.2"
    sec_summary = f"{section_start}.3"

    _add_heading_para(
        f"{sec_main} รายงานวัสดุและราคาโครงสร้างชั้นทาง",
        size=18, bold=True, underline=True, space_before=12, space_after=6
    )

    _add_heading_para(
        f"{sec_info} ข้อมูลของถนน",
        size=16, bold=True, underline=True, space_before=8
    )

    if intro_text:
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Cm(1.0)
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(intro_text)
        _set_run_font(run, size=16)

    fields = [
        ("ความยาวถนน",     f"{project_info.get('length', 1):.2f} กม."),
        ("ความกว้างรวม",   f"{project_info.get('total_width', 0):.2f} ม."),
        ("จำนวนช่องจราจร", f"{project_info.get('num_lanes', 2)} ช่อง"),
    ]
    for label, value in fields:
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Cm(1.0)
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        run_label = para.add_run(f"{label}: ")
        _set_run_font(run_label, size=16, bold=True)
        run_value = para.add_run(value)
        _set_run_font(run_value, size=16, bold=False)

    _add_heading_para(
        f"{sec_detail} รายละเอียดวัสดุและราคา",
        size=16, bold=True, underline=True, space_before=10
    )

    summary_data = []
    length = project_info.get('length', 1)

    for ptype, data in all_details.items():
        structure_name = data.get('name', ptype)          # ชื่อสั้น: AC / JPCP / JRCP / CRCP
        name_detail    = data.get('name_detail', '')      # ชื่อยาว: รายละเอียด
        details = data.get('details', [])

        # header: ชื่อสั้น เช่น "JPCP (ชุดที่ 1)"
        _add_heading_para(f"ผิวทางประเภท {structure_name}", size=16, bold=True, space_before=6, space_after=2)
        # subtitle: ชื่อยาว (ถ้ามี)
        if name_detail and name_detail != structure_name:
            _sub = doc.add_paragraph()
            _sub.paragraph_format.space_before = Pt(0)
            _sub.paragraph_format.space_after  = Pt(4)
            _run = _sub.add_run(name_detail)
            _set_run_font(_run, size=14, bold=False)
            _run.italic = True

        if details:
            table = doc.add_table(rows=len(details) + 2, cols=5)
            table.style = 'Table Grid'
            col_widths_t = [Cm(6.5), Cm(2.5), Cm(1.8), Cm(3.5), Cm(3.5)]
            for row in table.rows:
                for idx, cell in enumerate(row.cells):
                    cell.width = col_widths_t[idx]

            headers_t = ['รายการ', 'ปริมาณ', 'หน่วย', 'ราคา/หน่วย (บาท)', 'มูลค่า (บาท)']
            for j, h in enumerate(headers_t):
                cell = table.rows[0].cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(h)
                _set_run_font(run, size=15, bold=True)

            subtotal = 0
            for i, d in enumerate(details):
                rc = table.rows[i + 1].cells
                rc[0].text = ''
                run = rc[0].paragraphs[0].add_run(str(d['รายการ']))
                _set_run_font(run, size=15)

                rc[1].text = ''
                rc[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = rc[1].paragraphs[0].add_run(f"{d['ปริมาณ']:,.0f}")
                _set_run_font(run, size=15)

                unit_display = d.get('หน่วย', 'ตร.ม.')
                rc[2].text = ''
                rc[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = rc[2].paragraphs[0].add_run(unit_display)
                _set_run_font(run, size=15)

                price_display = d.get('ราคา/หน่วย (แสดง)', f"{d['ราคา/หน่วย']:,.0f}")
                price_label = d.get('หน่วยราคา', 'บาท/ตร.ม.')
                rc[3].text = ''
                rc[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = rc[3].paragraphs[0].add_run(f"{price_display}")
                _set_run_font(run, size=15)
                run2 = rc[3].paragraphs[0].add_run(f" ({price_label})")
                _set_run_font(run2, size=12)

                rc[4].text = ''
                rc[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = rc[4].paragraphs[0].add_run(f"{d['มูลค่า (บาท)']:,.0f}")
                _set_run_font(run, size=15)
                subtotal += d['มูลค่า (บาท)']

            last_row = table.rows[len(details) + 1]
            last_row.cells[0].merge(last_row.cells[3])
            last_row.cells[0].text = ''
            p_sum = last_row.cells[0].paragraphs[0]
            p_sum.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p_sum.add_run(f"รวม {structure_name}")
            _set_run_font(run, size=15, bold=True)
            last_row.cells[4].text = ''
            last_row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = last_row.cells[4].paragraphs[0].add_run(f"{subtotal:,.0f}")
            _set_run_font(run, size=15, bold=True)

            doc.add_paragraph()
            summary_data.append({
                'name': structure_name,
                'total_value': subtotal,
                'cost_per_km_million': data.get('cost_per_km', 0),
                'cost_per_sqm': data.get('cost_sqm', 0),
            })

    _add_heading_para(f"{sec_summary} สรุปค่าใช้จ่าย", size=16, bold=True, underline=True, space_before=10)

    if summary_data:
        sum_table = doc.add_table(rows=len(summary_data) + 1, cols=4)
        sum_table.style = 'Table Grid'
        col_widths_s = [Cm(7), Cm(3.5), Cm(3), Cm(3.5)]
        for row in sum_table.rows:
            for idx, cell in enumerate(row.cells):
                cell.width = col_widths_s[idx]

        hdrs = ['ชนิดโครงสร้าง', 'มูลค่ารวม/กม. (บาท)', 'ราคา/กม. (ล้านบาท)', 'ราคา/ตร.ม. (บาท)']
        for j, h in enumerate(hdrs):
            cell = sum_table.rows[0].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            _set_run_font(run, size=15, bold=True)

        for i, item in enumerate(summary_data):
            total_per_km = item['total_value'] / length if length > 0 else 0
            r = sum_table.rows[i + 1]
            r.cells[0].text = ''
            run = r.cells[0].paragraphs[0].add_run(item['name'])
            _set_run_font(run, size=15)
            for cidx, val in enumerate([
                f"{total_per_km:,.0f}",
                f"{item['cost_per_km_million']:.2f}",
                f"{item['cost_per_sqm']:,.2f}",
            ]):
                r.cells[cidx + 1].text = ''
                r.cells[cidx + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = r.cells[cidx + 1].paragraphs[0].add_run(val)
                _set_run_font(run, size=15)

    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_para.add_run(f"รายงานสร้างเมื่อ: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    _set_run_font(run, size=14)

    return doc


# ===== Main Application =====

@st.cache_data
def generate_excel_template():
    """สร้าง Excel Template และ cache ไว้เพื่อประสิทธิภาพ"""
    template_data = {
        'AC_Prices': pd.DataFrame({
            'Material': ['PMA Wearing Course', 'AC Wearing Course', 'AC Binder Course', 'AC Base Course'],
            '2.5cm': [170, 128, 129, 129],
            '3cm':   [203, 152, 154, 154],
            '4cm':   [268, 202, 202, 202],
            '5cm':   [333, 250, 251, 251],
            '6cm':   [406, 306, 308, 308],
            '7cm':   [471, 355, 356, 356],
            '8cm':   [536, 403, 405, 405],
            '9cm':   [601, 452, 454, 454],
            '10cm':  [667, 502, 503, 503],
        }),
        'Concrete_Prices': pd.DataFrame({
            'Type':  ['JRCP', 'JPCP', 'CRCP'],
            '25cm':  [924, 928, 1245],
            '28cm':  [1002, 1000, 1358],
            '30cm':  [0, 0, 0],
            '32cm':  [1106, 1095, 1509],
            '35cm':  [1184, 1167, 1622],
        }),
        'Base_Materials': pd.DataFrame({
            'Material':             list(BASE_MATERIAL_PRICES.keys()),
            'Price (Baht/cu.m)':   list(BASE_MATERIAL_PRICES.values()),
        }),
    }

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        template_data['AC_Prices'].to_excel(writer, sheet_name='AC_Prices', index=False)
        template_data['Concrete_Prices'].to_excel(writer, sheet_name='Concrete_Prices', index=False)
        template_data['Base_Materials'].to_excel(writer, sheet_name='Base_Materials', index=False)
    output.seek(0)
    return output.getvalue()


def main():
    st.markdown('<div class="main-header">🛣️ ระบบวิเคราะห์ค่าก่อสร้างโครงสร้างชั้นทาง</div>', unsafe_allow_html=True)
    st.markdown("##### ตามแนวทาง AASHTO 1993 - รองรับ AC, JPCP/JRCP, CRCP")
    st.markdown("""
    <div style='text-align: center; color: #666; font-size: 0.9rem; margin-top: -10px; margin-bottom: 20px;'>
        พัฒนาโดย <b>รศ.ดร.อิทธิพล มีผล</b><br>
        ภาควิชาครุศาสตร์โยธา คณะครุศาสตร์อุตสาหกรรม<br>
        มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ (มจพ.)
    </div>
    """, unsafe_allow_html=True)

    # Sidebar
    with st.sidebar:
        st.header("📋 ข้อมูลโครงการ")

        st.subheader("📂 โหลดโครงการ")
        uploaded_json = st.file_uploader(
            "เลือกไฟล์ JSON", type=['json'],
            help="อัพโหลดไฟล์โครงการที่บันทึกไว้", key="upload_json"
        )

        if uploaded_json is not None:
            try:
                import hashlib
                file_bytes = uploaded_json.read()
                file_hash = hashlib.md5(file_bytes).hexdigest()
                loaded_data = json.loads(file_bytes.decode('utf-8'))
                st.success("✅ โหลดไฟล์สำเร็จ!")

                if 'project_info' in loaded_data:
                    st.info(f"📌 โครงการ: {loaded_data['project_info'].get('name', '-')}")
                    st.info(f"📅 บันทึกเมื่อ: {loaded_data.get('saved_at', '-')}")

                if st.button("📥 นำเข้าข้อมูล", key="import_json"):
                    if 'project_info' in loaded_data:
                        if st.session_state.get('loaded_json_hash') != file_hash:
                            st.session_state['loaded_project'] = loaded_data
                            st.session_state['loaded_json_hash'] = file_hash
                            new_v = st.session_state.get('json_version', 0) + 1
                            st.session_state['json_version'] = new_v
                            # ล้าง widget keys ทั้งหมดที่ขึ้นกับ version
                            # ครอบคลุม: ความหนาผิวทาง (_st_), ความหนาพื้นทาง (_bt_),
                            # วัสดุพื้นทาง (_bm_), จำนวนชั้น (_num_base_),
                            # ราคา (_price_), ชื่อโครงสร้าง (_name_), AC interlayer (_acil_)
                            keys_to_clear = [k for k in st.session_state.keys()
                                             if any(p in k for p in [
                                                 '_ctype_', '_fullname_', '_name_suffix_',
                                                 'jrcp1_fullname', 'jrcp2_fullname',
                                                 '_bt_', '_bm_', '_st_',
                                                 '_num_base_', '_price_',
                                                 '_name_v', '_acil_',
                                                 '_mat_', '_show',
                                             ])]
                            for k in keys_to_clear:
                                del st.session_state[k]
                        st.rerun()
            except Exception as e:
                st.error(f"❌ ไม่สามารถอ่านไฟล์ได้: {e}")

        st.divider()

        st.subheader("💰 Price Library (Excel)")
        st.caption("อัพโหลด Excel เพื่อแทนที่ราคา Default")

        uploaded_price_excel = st.file_uploader(
            "เลือกไฟล์ Excel Price Library", type=['xlsx', 'xls'],
            help="อัพโหลดไฟล์ราคาที่ Download จาก Tab 1", key="sidebar_upload_price"
        )

        if uploaded_price_excel is not None:
            try:
                ac_df = pd.read_excel(uploaded_price_excel, sheet_name='AC_Prices')
                concrete_df = pd.read_excel(uploaded_price_excel, sheet_name='Concrete_Prices')
                base_df = pd.read_excel(uploaded_price_excel, sheet_name='Base_Materials')

                uploaded_ac_prices = {}
                for _, row in ac_df.iterrows():
                    material = row['Material']
                    prices = {}
                    for col in ac_df.columns[1:]:
                        try:
                            thickness = float(col.replace('cm', '').strip())
                            val = row[col]
                            if pd.notna(val):
                                prices[thickness] = float(val)
                        except (ValueError, TypeError):
                            pass
                    if prices:
                        uploaded_ac_prices[material] = prices
                # เติม key ที่ขาดจาก default
                for mat, default_prices in AC_PRICE_TABLE.items():
                    if mat not in uploaded_ac_prices:
                        uploaded_ac_prices[mat] = dict(default_prices)
                    else:
                        for thk, dp in default_prices.items():
                            uploaded_ac_prices[mat].setdefault(thk, dp)

                uploaded_concrete_prices = {}
                for _, row in concrete_df.iterrows():
                    conc_type = row['Type']
                    prices = {}
                    for col in concrete_df.columns[1:]:
                        try:
                            thickness = int(float(col.replace('cm', '').strip()))
                            val = row[col]
                            if pd.notna(val):
                                prices[thickness] = float(val)
                        except (ValueError, TypeError):
                            pass
                    if prices:
                        uploaded_concrete_prices[conc_type] = prices
                # เติม key ที่ขาดจาก default
                for ct, default_prices in CONCRETE_PRICE_TABLE.items():
                    if ct not in uploaded_concrete_prices:
                        uploaded_concrete_prices[ct] = dict(default_prices)
                    else:
                        for thk, dp in default_prices.items():
                            uploaded_concrete_prices[ct].setdefault(thk, dp)

                uploaded_base_prices = {}
                for _, row in base_df.iterrows():
                    try:
                        mat = row['Material']
                        val = row['Price (Baht/cu.m)']
                        if pd.notna(mat) and pd.notna(val):
                            uploaded_base_prices[str(mat)] = float(val)
                    except (ValueError, TypeError):
                        pass
                # เติม key ที่ขาดจาก default (รวม Embankment)
                for mat, dp in BASE_MATERIAL_PRICES.items():
                    uploaded_base_prices.setdefault(mat, dp)

                st.session_state['uploaded_price_library'] = {
                    'ac_prices': uploaded_ac_prices,
                    'concrete_prices': uploaded_concrete_prices,
                    'base_prices': uploaded_base_prices,
                }

                import hashlib
                file_hash_p = hashlib.md5(uploaded_price_excel.getvalue()).hexdigest()[:8]
                st.session_state['price_upload_version'] = file_hash_p

                st.success("✅ โหลด Price Library สำเร็จ!")
                st.caption(f"📊 {len(uploaded_ac_prices)} AC types, {len(uploaded_concrete_prices)} Concrete types")

                with st.expander("🔍 ตัวอย่างราคาที่อ่านได้"):
                    st.write("**AC Wearing Course (7cm):**", uploaded_ac_prices.get('AC Wearing Course', {}).get(7.0, 'N/A'))
                    st.write("**JPCP (25cm):**", uploaded_concrete_prices.get('JPCP', {}).get(25, 'N/A'))
                    st.write("**Crushed Rock:**", uploaded_base_prices.get('Crushed Rock Base Course', 'N/A'))

            except Exception as e:
                st.error(f"❌ อ่านไฟล์ไม่สำเร็จ: {str(e)}")

        st.divider()

        loaded_project = st.session_state.get('loaded_project', {})
        loaded_info = loaded_project.get('project_info', {})
        v_sb = st.session_state.get('json_version', 0)

        project_name = st.text_input("ชื่อโครงการ",
            value=loaded_info.get('name', "โครงการก่อสร้างทางหลวง"),
            key=f"sidebar_project_name_v{v_sb}")
        road_length = st.number_input("ความยาวถนน (กม.)",
            value=float(loaded_info.get('length', 1.0)),
            min_value=0.1, step=0.1, key=f"sidebar_road_length_v{v_sb}")

        st.divider()
        st.header("📐 ขนาดถนน")
        lane_width = st.number_input("ความกว้างช่องจราจร (ม.)",
            value=float(loaded_info.get('lane_width', 3.5)),
            min_value=2.5, max_value=4.0, step=0.25, key=f"sidebar_lane_width_v{v_sb}")

        default_lanes_total = loaded_info.get('num_lanes', 4)
        default_lanes_per_dir = default_lanes_total // 2
        lanes_per_dir_options = [2, 3, 4]
        lanes_per_dir_idx = lanes_per_dir_options.index(default_lanes_per_dir) if default_lanes_per_dir in lanes_per_dir_options else 0
        lanes_per_direction = st.selectbox("จำนวนช่องต่อทิศทาง (เลน/ทิศทาง)",
            options=lanes_per_dir_options, index=lanes_per_dir_idx,
            key=f"sidebar_lanes_per_dir_v{v_sb}")

        num_lanes = lanes_per_direction * 2

        shoulder_left = st.number_input("ไหล่ทางซ้าย (ม.)",
            value=float(loaded_info.get('shoulder_left', 2.5)),
            min_value=0.0, max_value=3.5, step=0.25, key=f"sidebar_shoulder_left_v{v_sb}")
        shoulder_right = st.number_input("ไหล่ทางขวา (ม.)",
            value=float(loaded_info.get('shoulder_right', 1.5)),
            min_value=0.0, max_value=3.5, step=0.25, key=f"sidebar_shoulder_right_v{v_sb}")

        road_surface_width = lane_width * num_lanes
        total_shoulders = (shoulder_left + shoulder_right) * 2
        total_width = road_surface_width + total_shoulders
        st.info(f"📏 จำนวนช่องรวม (2 ทิศทาง): {num_lanes} ช่อง\n"
                f"📏 ความกว้างผิวจราจร: {road_surface_width:.2f} ม.\n"
                f"📏 ความกว้างไหล่ทาง (2 ทิศทาง): {total_shoulders:.2f} ม.\n"
                f"📏 ความกว้างรวม: {total_width:.2f} ม.")

    project_info = {
        'name': project_name,
        'length': road_length,
        'lane_width': lane_width,
        'shoulder_left': shoulder_left,
        'shoulder_right': shoulder_right,
        'num_lanes': num_lanes,
        'total_width': total_width,
    }

    area_per_km = total_width * 1000

    tab1, tab2, tab3, tab4 = st.tabs([
        "📊 Library ราคา",
        "🏗️ โครงสร้างชั้นทาง",
        "📄 รายงาน",
        "📷 วิเคราะห์จากรูปภาพ",
    ])

    # ===== Tab 1: Library ราคา =====
    with tab1:
        st.header("📊 ตารางราคาเปรียบเทียบโครงสร้างชั้นทาง")
        st.info("💡 สามารถปรับเปลี่ยนราคาได้ตามต้องการ หรือ Upload ไฟล์ Excel ใน **Sidebar** เพื่ออัพเดทราคาทั้งหมด")

        st.subheader("📥 ดาวน์โหลด Template Excel")
        col1, col2, col3 = st.columns([2, 1, 2])
        with col2:
            template_bytes = generate_excel_template()
            st.download_button(
                label="⬇️ ดาวน์โหลด Template",
                data=template_bytes,
                file_name=f"Price_Library_Template_{datetime.now().strftime('%Y%m%d')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
        st.caption("📌 ดาวน์โหลด Template → แก้ไขราคา → Upload ใน Sidebar ด้านซ้าย")
        st.divider()

        if 'uploaded_price_library' in st.session_state:
            st.session_state['price_library'] = st.session_state['uploaded_price_library'].copy()
        elif 'price_library' not in st.session_state:
            st.session_state['price_library'] = {
                'ac_prices': {
                    'PMA Wearing Course': dict(AC_PRICE_TABLE['PMA Wearing Course']),
                    'AC Wearing Course': dict(AC_PRICE_TABLE['AC Wearing Course']),
                    'AC Binder Course': dict(AC_PRICE_TABLE['AC Binder Course']),
                    'AC Base Course': dict(AC_PRICE_TABLE['AC Base Course']),
                },
                'concrete_prices': {
                    'JRCP': dict(CONCRETE_PRICE_TABLE['JRCP']),
                    'JPCP': dict(CONCRETE_PRICE_TABLE['JPCP']),
                    'CRCP': dict(CONCRETE_PRICE_TABLE['CRCP']),
                },
                'base_prices': dict(BASE_MATERIAL_PRICES),
            }

        if 'uploaded_price_library' in st.session_state:
            st.info("📋 **กำลังใช้ราคาจากไฟล์ Excel ที่ Upload ใน Sidebar**")
            ac_7 = st.session_state['price_library']['ac_prices'].get('AC Wearing Course', {}).get(7.0, 'N/A')
            st.caption(f"ตัวอย่าง: AC Wearing Course 7cm = {ac_7} บาท")
        else:
            st.caption("💡 กำลังใช้ราคา Default (Upload Excel ใน Sidebar เพื่อเปลี่ยนราคา)")

        upload_version = st.session_state.get('price_upload_version', 'default')

        st.subheader("🔵 ผิวทาง Asphalt Concrete (บาท/ตร.ม.)")
        ac_cols = st.columns(4)
        ac_types = ['PMA Wearing Course', 'AC Wearing Course', 'AC Binder Course', 'AC Base Course']
        thicknesses = [2.5, 3, 4, 5, 6, 7, 8, 9, 10]

        for col_idx, ac_type in enumerate(ac_types):
            with ac_cols[col_idx]:
                st.markdown(f"**{ac_type}**")
                for thk in thicknesses:
                    current_price = st.session_state['price_library']['ac_prices'][ac_type].get(thk, 0)
                    price = st.number_input(
                        f"{thk} cm", value=float(current_price),
                        key=f"ac_{ac_type}_{thk}_{upload_version}", step=10.0
                    )
                    st.session_state['price_library']['ac_prices'][ac_type][thk] = price

        st.divider()

        st.subheader("🟠 ผิวทางคอนกรีต (บาท/ตร.ม.)")
        conc_cols = st.columns(3)
        conc_types = ['JRCP', 'JPCP', 'CRCP']
        conc_thicknesses = [25, 28, 30, 32, 35]

        for col_idx, conc_type in enumerate(conc_types):
            with conc_cols[col_idx]:
                st.markdown(f"**{conc_type}**")
                for thk in conc_thicknesses:
                    current_price = st.session_state['price_library']['concrete_prices'][conc_type].get(thk, 0)
                    price = st.number_input(
                        f"{thk} cm", value=float(current_price),
                        key=f"conc_{conc_type}_{thk}_{upload_version}", step=10.0
                    )
                    st.session_state['price_library']['concrete_prices'][conc_type][thk] = price
                st.markdown("---")
                st.number_input(
                    f"{conc_type} (excl. Joint)",
                    value=float(CONCRETE_EXCL_JOINT[conc_type]),
                    key=f"conc_excl_{conc_type}_{upload_version}", step=10.0
                )

        st.divider()

        st.subheader("🟤 วัสดุพื้นทาง/รองพื้นทาง (บาท/ลบ.ม.)")
        base_cols = st.columns(3)
        base_materials_list = list(BASE_MATERIAL_PRICES.keys())

        for i, mat in enumerate(base_materials_list):
            with base_cols[i % 3]:
                current_price = st.session_state['price_library']['base_prices'].get(mat, 0)
                price = st.number_input(
                    mat, value=float(current_price),
                    key=f"base_{mat}_{upload_version}", step=10.0
                )
                st.session_state['price_library']['base_prices'][mat] = price

        st.markdown("---")
        st.markdown("**✨ วัสดุกำหนดเอง** (เพิ่มวัสดุของคุณเอง)")
        custom_cols = st.columns(3)

        for i in range(1, 4):
            with custom_cols[i - 1]:
                custom_key = f"custom_material_{i}"
                if 'custom_materials' not in st.session_state:
                    st.session_state['custom_materials'] = {}

                existing_data = st.session_state['custom_materials'].get(custom_key, {'name': '', 'price': 0.0})
                material_name = st.text_input(
                    f"ชื่อวัสดุ {i}", value=existing_data['name'],
                    key=f"custom_name_{i}_{upload_version}", placeholder=f"ระบุชื่อวัสดุ {i}..."
                )

                if material_name:
                    material_price = st.number_input(
                        "ราคา (บาท/ลบ.ม.)", value=float(existing_data['price']),
                        key=f"custom_price_{i}_{upload_version}", step=10.0, min_value=0.0
                    )
                    st.session_state['custom_materials'][custom_key] = {
                        'name': material_name, 'price': material_price
                    }
                    st.session_state['price_library']['base_prices'][material_name] = material_price
                else:
                    if custom_key in st.session_state['custom_materials']:
                        old_name = st.session_state['custom_materials'][custom_key]['name']
                        if old_name in st.session_state['price_library']['base_prices']:
                            del st.session_state['price_library']['base_prices'][old_name]
                        del st.session_state['custom_materials'][custom_key]

        st.divider()
        st.subheader("📥 ดาวน์โหลดตารางราคา")
        col_dl1, col_dl2 = st.columns(2)

        with col_dl1:
            if st.button("📊 สร้างไฟล์ Excel", key="btn_excel_price", use_container_width=True):
                output = io.BytesIO()
                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    ac_data = []
                    for ac_type in ac_types:
                        for thk in thicknesses:
                            ac_data.append({'ประเภท': ac_type, 'ความหนา (cm)': thk,
                                            'ราคา (บาท/ตร.ม.)': st.session_state['price_library']['ac_prices'][ac_type][thk]})
                    pd.DataFrame(ac_data).to_excel(writer, sheet_name='AC Prices', index=False)

                    conc_data = []
                    for conc_type in conc_types:
                        for thk in conc_thicknesses:
                            conc_data.append({'ประเภท': conc_type, 'ความหนา (cm)': thk,
                                              'ราคา (บาท/ตร.ม.)': st.session_state['price_library']['concrete_prices'][conc_type][thk]})
                    pd.DataFrame(conc_data).to_excel(writer, sheet_name='Concrete Prices', index=False)

                    base_data = [{'วัสดุ': k, 'ราคา (บาท/ลบ.ม.)': v}
                                 for k, v in st.session_state['price_library']['base_prices'].items()]
                    pd.DataFrame(base_data).to_excel(writer, sheet_name='Base Materials', index=False)

                output.seek(0)
                st.download_button(
                    label="⬇️ Download Excel", data=output,
                    file_name="ราคาเปรียบเทียบโครงสร้างชั้นทาง.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )

        with col_dl2:
            if st.button("📄 สร้างไฟล์ Word", key="btn_word_price", use_container_width=True):
                doc = Document()
                doc.add_heading('ตารางราคาเปรียบเทียบโครงสร้างชั้นทาง', 0)
                doc.add_heading('1. ผิวทาง Asphalt Concrete (บาท/ตร.ม.)', level=1)
                table = doc.add_table(rows=len(thicknesses) + 1, cols=5)
                table.style = 'Table Grid'
                for j, h in enumerate(['ความหนา (cm)'] + ac_types):
                    table.rows[0].cells[j].text = h
                for i, thk in enumerate(thicknesses):
                    table.rows[i + 1].cells[0].text = str(thk)
                    for j, ac_type in enumerate(ac_types):
                        table.rows[i + 1].cells[j + 1].text = f"{st.session_state['price_library']['ac_prices'][ac_type][thk]:,.0f}"

                doc.add_heading('2. ผิวทางคอนกรีต (บาท/ตร.ม.)', level=1)
                table = doc.add_table(rows=len(conc_thicknesses) + 1, cols=4)
                table.style = 'Table Grid'
                for j, h in enumerate(['ความหนา (cm)'] + conc_types):
                    table.rows[0].cells[j].text = h
                for i, thk in enumerate(conc_thicknesses):
                    table.rows[i + 1].cells[0].text = str(thk)
                    for j, conc_type in enumerate(conc_types):
                        table.rows[i + 1].cells[j + 1].text = f"{st.session_state['price_library']['concrete_prices'][conc_type][thk]:,.0f}"

                doc.add_heading('3. วัสดุพื้นทาง/รองพื้นทาง (บาท/ลบ.ม.)', level=1)
                table = doc.add_table(rows=len(base_materials_list) + 1, cols=2)
                table.style = 'Table Grid'
                table.rows[0].cells[0].text = 'วัสดุ'
                table.rows[0].cells[1].text = 'ราคา (บาท/ลบ.ม.)'
                for i, mat in enumerate(base_materials_list):
                    table.rows[i + 1].cells[0].text = mat
                    table.rows[i + 1].cells[1].text = f"{st.session_state['price_library']['base_prices'][mat]:,.0f}"

                doc_output = io.BytesIO()
                doc.save(doc_output)
                doc_output.seek(0)
                st.download_button(
                    label="⬇️ Download Word", data=doc_output,
                    file_name="ราคาเปรียบเทียบโครงสร้างชั้นทาง.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    # ===== Tab 2: โครงสร้างชั้นทาง (v6 — 2 ชุด) =====
    with tab2:
        st.header("กำหนดโครงสร้างชั้นทาง")

        v = st.session_state.get('json_version', 0)

        if st.session_state.get('loaded_project'):
            loaded_name = st.session_state['loaded_project'].get('project_info', {}).get('name', '-')
            st.success(f"✅ กำลังใช้ข้อมูลจาก: **{loaded_name}**")

        area_per_km = total_width * 1000

        # ===== Toggle ชุดที่ 2 =====
        col_hdr1, col_hdr2 = st.columns([3, 1])
        with col_hdr1:
            st.markdown(
                "**ชุดที่ 1** = โครงสร้างหลักสำหรับโครงการ &nbsp;|&nbsp; "
                "**ชุดที่ 2** = เปรียบเทียบ (เปิด/ปิดได้)"
            )
        with col_hdr2:
            show_set2 = st.toggle("แสดงชุดที่ 2", value=False, key="toggle_set2")

        st.divider()

        # ── helper: render 1 โครงสร้าง (AC/JPCP/JRCP/CRCP) ──────────────────
        def _render_structure(ptype, suffix, default_layers_fn, default_joints_fn,
                              label, color, set_label, v=0):
            """
            ptype       : 'AC' | 'JPCP' | 'JRCP' | 'CRCP'
            suffix      : 'a' (ชุด1) | 'b' (ชุด2)
            label       : ชื่อย่อแสดงใน expander
            color       : สีกล่อง CSS hex
            set_label   : 'ชุดที่ 1' | 'ชุดที่ 2'
            คืนค่า: (name, cost_per_km, cost_per_sqm, details, layers, joints, show)
            """
            key_prefix = f"{ptype.lower()}_{suffix}"
            is_concrete = ptype in ('JPCP', 'JRCP', 'CRCP')

            # ── badge ──
            badge_color      = '#E6F1FB' if suffix == 'a' else '#EAF3DE'
            badge_text_color = '#0C447C' if suffix == 'a' else '#3B6D11'
            st.markdown(
                f'<span style="background:{badge_color};color:{badge_text_color};'
                f'font-size:11px;font-weight:600;padding:2px 8px;'
                f'border-radius:99px;margin-right:6px;">{set_label}</span>',
                unsafe_allow_html=True
            )

            # default name — เรียบง่าย ไม่ซ้ำซ้อน
            _default_name = label

            struct_name = st.text_input(
                "ชื่อโครงสร้าง",
                value=_default_name,
                key=f"{key_prefix}_name_v{v}",
                label_visibility="collapsed"
            )
            show_flag = st.checkbox(
                "รวมในรายงาน", value=True,
                key=f"{key_prefix}_show"
            )

            with st.expander(f"▶ {struct_name}", expanded=True):
                # ส่ง ptype ไป render_layer_editor เพื่อกำหนดชื่อ concrete ตายตัว
                layers = render_layer_editor(
                    default_layers_fn(), key_prefix, total_width, road_length, v=v,
                    ptype=ptype
                )
                layer_cost, layer_details = calculate_layer_cost(layers, road_length)

                if is_concrete and default_joints_fn is not None:
                    # ส่ง ptype ไป render_joint_editor เพื่อกำหนด spacing ตายตัว
                    joints, include_joints = render_joint_editor(
                        default_joints_fn(), key_prefix, area_per_km, road_length, v=v,
                        ptype=ptype
                    )
                    joint_cost, joint_details = calculate_joint_cost(joints, road_length, include_joints=include_joints)
                    joints_sqm = sum(j.get('cost_per_sqm', 0) for j in joints)
                    total_cost = layer_cost + (joint_cost if include_joints else 0)
                    if include_joints:
                        cost_sqm = layer_cost / (area_per_km * road_length) + joints_sqm
                        note = "(รวม Joints)" if ptype != 'CRCP' else "(รวม Long. Steel & Joints)"
                    else:
                        cost_sqm = layer_cost / (area_per_km * road_length)
                        note = "(ไม่รวม Joints)" if ptype != 'CRCP' else "(ไม่รวม Long. Steel & Joints)"
                    all_details = layer_details + (joint_details if include_joints else [])
                else:
                    joints = None
                    total_cost = layer_cost
                    cost_sqm = layer_cost / (area_per_km * road_length)
                    note = ""
                    all_details = layer_details

                cost_per_km = total_cost / road_length / 1_000_000

                st.markdown(
                    f'<div class="cost-box" style="border-left-color:{color};">'
                    f'💰 <b>{cost_per_km:.2f}</b> ล้านบาท/กม. &nbsp;|&nbsp; '
                    f'<b>{cost_sqm:.2f}</b> บาท/ตร.ม. {note}</div>',
                    unsafe_allow_html=True
                )

            return struct_name, cost_per_km, cost_sqm, all_details, layers, joints, show_flag

        # ── ตารางสำหรับเก็บผลลัพธ์ทั้งหมด ──────────────────────────────────
        construction = {}

        # ── ตัวกำหนด default layers/joints ต่อประเภท ──
        STRUCT_CONFIG = {
            'AC':   (get_default_ac1_layers,   None,                     'แอสฟัลต์คอนกรีต AC ',           '#378ADD', 20),
            'JPCP': (get_default_jrcp1_layers, get_default_jrcp1_joints, 'คอนกรีต JPCP',                 '#E29A30', 20),
            'JRCP': (get_default_jrcp2_layers, get_default_jrcp2_joints, 'คอนกรีต JRCP',          '#E29A30', 20),
            'CRCP': (get_default_crcp1_layers, get_default_crcp1_joints, 'คอนกรีตเสริมเหล็กต่อเนื่อง CRCP', '#C94040', 20),
        }
        STRUCT_CONFIG_B = {
            'AC':   (get_default_ac1_layers,   None,                     'แอสฟัลต์คอนกรีต AC (เปรียบเทียบ)',       '#378ADD', 20),
            'JPCP': (get_default_jrcp1_layers, get_default_jrcp1_joints, 'คอนกรีต JPCP (เปรียบเทียบ)',         '#E29A30', 20),
            'JRCP': (get_default_jrcp2_layers, get_default_jrcp2_joints, 'คอนกรีต JRCP (เปรียบเทียบ)', '#E29A30', 20),
            'CRCP': (get_default_crcp1_layers, get_default_crcp1_joints, 'คอนกรีตเสริมเหล็กต่อเนื่อง CRCP (เปรียบเทียบ)',         '#C94040', 20),
        }

        ptypes = ['AC', 'JPCP', 'JRCP', 'CRCP']
        type_icons = {'AC': '🔵', 'JPCP': '🟠', 'JRCP': '🟠', 'CRCP': '🔴'}

        for ptype in ptypes:
            layers_fn_a, joints_fn_a, label_a, color_a, life_default = STRUCT_CONFIG[ptype]
            layers_fn_b, joints_fn_b, label_b, color_b, _            = STRUCT_CONFIG_B[ptype]

            # ── header: ชื่อประเภท + อายุออกแบบ (ชุดที่ 1 = ชุดที่ 2 เสมอ) ──
            hcol1, hcol2 = st.columns([3, 1])
            with hcol1:
                st.subheader(f"{type_icons[ptype]} {ptype}")
            with hcol2:
                life_a = st.number_input(
                    "อายุออกแบบ (ปี)",
                    value=int(life_default),
                    min_value=1, max_value=50, step=1,
                    key=f"life_{ptype}_v{v}",
                    help="ชุดที่ 1 และ 2 ใช้อายุออกแบบเดียวกัน"
                )

            if show_set2:
                col_a, col_b = st.columns(2)
            else:
                col_a = st.container()

            with col_a:
                (name_a, cpk_a, csqm_a, det_a, lay_a, jnt_a, show_a) = _render_structure(
                    ptype, 'a', layers_fn_a, joints_fn_a, label_a, color_a, 'ชุดที่ 1', v=v
                )
            construction[f'{ptype}_A'] = {
                'name': name_a, 'cost': cpk_a, 'cost_sqm': csqm_a,
                'details': det_a, 'layers': lay_a, 'joints': jnt_a,
                'show': show_a, 'life': life_a, 'set': 1,
            }

            if show_set2:
                with col_b:
                    (name_b, cpk_b, csqm_b, det_b, lay_b, jnt_b, show_b) = _render_structure(
                        ptype, 'b', layers_fn_b, joints_fn_b, label_b, color_b, 'ชุดที่ 2', v=v
                    )
                construction[f'{ptype}_B'] = {
                    'name': name_b, 'cost': cpk_b, 'cost_sqm': csqm_b,
                    'details': det_b, 'layers': lay_b, 'joints': jnt_b,
                    'show': show_b, 'life': life_a, 'set': 2,
                }
            else:
                # ชุดที่ 2 ไม่ได้เปิด → เก็บ placeholder ไว้ (show=False)
                construction[f'{ptype}_B'] = {
                    'name': f"{ptype} (ชุดที่ 2)", 'cost': 0, 'cost_sqm': 0,
                    'details': [], 'layers': [], 'joints': None,
                    'show': False, 'life': life_a, 'set': 2,
                }

            st.divider()

        # ── บันทึก session_state ──
        # รักษา key เดิม (AC1/AC2/JRCP1/JRCP2/CRCP1/CRCP2) สำหรับ Tab 3 / JSON
        # รักษา key เดิม (AC1/AC2/JRCP1/JRCP2/CRCP1/CRCP2) ให้ Tab 3 และ JSON ใช้ได้
        # mapping: ชุดที่ 1 → key ลงท้าย 1, ชุดที่ 2 → key ลงท้าย 2
        st.session_state['construction'] = {
            'AC1':   construction['AC_A'],
            'AC2':   construction['AC_B'],
            'JPCP1': construction['JPCP_A'],
            'JPCP2': construction['JPCP_B'],
            'JRCP1': construction['JRCP_A'],
            'JRCP2': construction['JRCP_B'],
            'CRCP1': construction['CRCP_A'],
            'CRCP2': construction['CRCP_B'],
        }
        st.session_state['project_info'] = project_info
        st.session_state['area_per_km']  = area_per_km

        # ===== Summary Tables =====
        st.subheader("📊 สรุปค่าก่อสร้างทุกโครงสร้าง")

        summary_rows = []
        for key, data in construction.items():
            if data['layers']:   # แสดงเฉพาะที่มีข้อมูล
                summary_rows.append({
                    'ชุด': f"ชุดที่ {data['set']}",
                    'รหัส': key,
                    'ชื่อโครงสร้าง': data['name'],
                    'ค่าก่อสร้าง (ล้านบาท/กม.)': data['cost'],
                    'ค่าก่อสร้าง (บาท/ตร.ม.)': data['cost_sqm'],
                    'อายุออกแบบ (ปี)': data['life'],
                    'รายงาน': '✅' if data['show'] else '❌',
                })

        if summary_rows:
            summary_df = pd.DataFrame(summary_rows)
            st.dataframe(
                summary_df.style.format({
                    'ค่าก่อสร้าง (ล้านบาท/กม.)': '{:.2f}',
                    'ค่าก่อสร้าง (บาท/ตร.ม.)': '{:.2f}',
                }),
                use_container_width=True, hide_index=True
            )

        st.divider()
        st.subheader("📋 รายละเอียดราคาแต่ละโครงสร้าง")

        # สร้างรายการ options เฉพาะที่มีข้อมูล
        detail_options = [k for k, d in construction.items() if d['layers']]
        if detail_options:
            selected_structure = st.selectbox(
                "เลือกดูรายละเอียด",
                options=detail_options,
                format_func=lambda x: f"[ชุดที่ {construction[x]['set']}] {construction[x]['name']}"
            )

            struct = construction[selected_structure]
            layers = struct['layers']
            joints = struct.get('joints')

            detail_data = []
            total_cost = 0
            st.markdown(f"**{struct['name']}**")

            for i, d in enumerate(struct['details']):
                total_cost += d['มูลค่า (บาท)']
                detail_data.append({
                    'ลำดับ': i + 1,
                    'รายการ': d['รายการ'],
                    'ความหนา': d['ความหนา'],
                    'ปริมาณ (ตร.ม.)': f"{d['ปริมาณ']:,.0f} {d['หน่วย']}",
                    'ราคา (บาท/ตร.ม.)': f"{d['ราคา/หน่วย (แสดง)']} ({d['หน่วยราคา']})",
                    'มูลค่า (บาท)': f"{d['มูลค่า (บาท)']:,.0f}",
                })

            st.dataframe(pd.DataFrame(detail_data), use_container_width=True, hide_index=True)

            area_km = area_per_km * road_length
            cost_per_sqm_det = total_cost / area_km if area_km > 0 else 0
            col_s1, col_s2, col_s3, col_s4 = st.columns(4)
            with col_s1:
                st.metric("💰 ราคารวม", f"{total_cost:,.0f} บาท")
            with col_s2:
                st.metric("📏 ราคาต่อ กม.", f"{total_cost / road_length:,.0f} บาท/กม.")
            with col_s3:
                st.metric("📊 ล้านบาท/กม.", f"{total_cost / road_length / 1_000_000:.2f}")
            with col_s4:
                st.metric("📐 บาท/ตร.ม.", f"{cost_per_sqm_det:.2f}")

    # ===== Tab 3: รายงาน =====
    with tab3:
        st.header("📄 สร้างรายงาน")
        st.info("💡 รายงานจะแสดงเฉพาะข้อมูลวัสดุและราคา (ไม่รวม NPV)")

        if 'construction' in st.session_state and st.session_state['construction']:
            constr = st.session_state.get('construction', {})

            # สร้าง all_details ก่อน แล้วค่อยตรวจว่ามีกี่ชุด
            _raw_details = {}
            for k, v_data in constr.items():
                if v_data.get('show', True) and v_data.get('details'):
                    _raw_details[k] = v_data

            # ตรวจว่ามีชุดที่ 2 อยู่ในรายงานด้วยหรือไม่
            _has_set2 = any(k.endswith('2') for k in _raw_details)

            # กำหนดชื่อสั้น: ถ้ามีแค่ชุดที่ 1 → ไม่ต้องใส่ "ชุดที่ 1"
            def _make_short_name(k):
                _base = {'AC1': 'AC', 'AC2': 'AC',
                         'JPCP1': 'JPCP', 'JPCP2': 'JPCP',
                         'JRCP1': 'JRCP', 'JRCP2': 'JRCP',
                         'CRCP1': 'CRCP', 'CRCP2': 'CRCP'}.get(k, k)
                if _has_set2:
                    _set_num = '1' if k.endswith('1') else '2'
                    return f"{_base} (ชุดที่ {_set_num})"
                return _base  # ชุดเดียว → ไม่ต้องบอก "ชุดที่ 1"

            all_details = {}
            for k, v_data in _raw_details.items():
                all_details[k] = {
                    'name':        _make_short_name(k),
                    'name_detail': v_data.get('name', k),
                    'details':     v_data.get('details', []),
                    'cost_per_km': v_data.get('cost', 0),
                    'cost_sqm':    v_data.get('cost_sqm', 0),
                }

            if not all_details:
                st.warning("⚠️ กรุณาเลือกอย่างน้อย 1 โครงสร้างที่ต้องการแสดงในรายงาน (tick ✅ แสดงในรายงาน ใน Tab 2)")
            else:
                st.subheader("📊 ข้อมูลที่จะรวมในรายงาน")
                for k, data in all_details.items():
                    if data['details']:
                        with st.expander(f"🔍 {data['name']}"):
                            # แสดงชื่อยาวเป็น caption
                            if data.get('name_detail') and data['name_detail'] != data['name']:
                                st.caption(f"รายละเอียด: {data['name_detail']}")
                            df_preview = pd.DataFrame(data['details'])
                            st.dataframe(df_preview, use_container_width=True, hide_index=True)

                st.divider()

                st.subheader("📋 สร้างรายงานแบบที่ปรึกษา")
                with st.expander("⚙️ ตั้งค่ารายงาน", expanded=True):
                    col_cfg1, col_cfg2 = st.columns(2)
                    with col_cfg1:
                        chapter_num = st.text_input("หมายเลขบทหลัก (เช่น 4, 5)", value="4", key="rpt_chapter_num")
                    with col_cfg2:
                        section_start = st.text_input("หมายเลขหัวข้อเริ่มต้น (เช่น 4.7)", value="4.7", key="rpt_section_start")

                    _pi = project_info
                    _default_intro = (
                        f"รายงานวิเคราะห์ต้นทุนค่าก่อสร้างโครงสร้างชั้นทางฉบับนี้ "
                        f"จัดทำขึ้นเพื่อเปรียบเทียบทางเลือกโครงสร้างชั้นทางประเภทต่าง ๆ "
                        f"สำหรับถนน {_pi.get('num_lanes', 4)} ช่องจราจร "
                        f"ความกว้างรวม {_pi.get('total_width', 0):.2f} เมตร "
                        f"ระยะทาง {_pi.get('length', 1):.2f} กิโลเมตร "
                        f"โดยครอบคลุมทั้งผิวทางแอสฟัลต์คอนกรีต (AC) และผิวทางคอนกรีตซีเมนต์ (JPCP, JRCP, CRCP) "
                        f"การวิเคราะห์อ้างอิงราคาวัสดุและค่าก่อสร้างตามมาตรฐานกรมกรมบัญชีกลาง "
                        f"เพื่อใช้เป็นข้อมูลประกอบการตัดสินใจเลือกโครงสร้างชั้นทางที่เหมาะสมกับสภาพโครงการ"
                    )
                    intro_text = st.text_area(
                        "บทเกริ่นนำ (แสดงใต้หัวข้อข้อมูลโครงการ)",
                        value=_default_intro, height=120, key=f"rpt_intro_text_{total_width:.0f}_{num_lanes}"
                    )

                if not DOCX_AVAILABLE:
                    st.warning("⚠️ ไม่สามารถสร้างรายงาน Word ได้ เนื่องจาก python-docx ไม่สามารถใช้งานได้")
                elif st.button("📋 สร้างรายงาน Word แบบที่ปรึกษา", type="primary", use_container_width=True, key="btn_consultant_report"):
                    try:
                        doc = generate_word_report_consultant(
                            st.session_state['project_info'], all_details,
                            chapter_num=chapter_num, section_start=section_start, intro_text=intro_text
                        )
                        buf = io.BytesIO()
                        doc.save(buf)
                        buf.seek(0)
                        _proj_name = st.session_state.get('project_info', {}).get('name', 'Project')
                        st.download_button(
                            "⬇️ ดาวน์โหลด Word แบบที่ปรึกษา", data=buf,
                            file_name=f"Cost Est {_proj_name}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True, key="dl_consultant_report"
                        )
                        st.success("✅ สร้างรายงานแบบที่ปรึกษาสำเร็จ!")
                    except Exception as e:
                        st.error(f"❌ เกิดข้อผิดพลาด: {str(e)}")

                st.divider()

                c1, c2 = st.columns(2)

                with c1:
                    if not DOCX_AVAILABLE:
                        st.warning("⚠️ ไม่สามารถสร้างรายงาน Word ได้")
                    elif st.button("📄 สร้างรายงาน Word แบบย่อ", use_container_width=True, key="btn_short_report"):
                        try:
                            doc = generate_word_report_materials_only(st.session_state['project_info'], all_details)
                            buf = io.BytesIO()
                            doc.save(buf)
                            buf.seek(0)
                            st.download_button(
                                "⬇️ ดาวน์โหลด Word แบบย่อ", data=buf,
                                file_name=f"Materials_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True, key="dl_short_report"
                            )
                            st.success("✅ สร้างรายงานสำเร็จ!")
                        except Exception as e:
                            st.error(f"❌ เกิดข้อผิดพลาดในการสร้างรายงาน: {str(e)}")

                with c2:
                    if st.button("💾 บันทึกโครงการ (JSON)", use_container_width=True):
                        # FIX: บันทึก 'show' และ 'cost_sqm' ด้วย เพื่อให้ load กลับมาครบถ้วน
                        data = {
                            'project_info': st.session_state['project_info'],
                            'construction': {
                                k: {
                                    'cost':     v_s.get('cost', 0),
                                    'cost_sqm': v_s.get('cost_sqm', 0),   # ← FIX: เพิ่ม cost_sqm
                                    'show':     v_s.get('show', True),     # ← FIX: เพิ่ม show flag
                                    'details':  v_s.get('details', []),
                                    'layers':   v_s.get('layers', []),
                                    'joints':   v_s.get('joints') or [],
                                } for k, v_s in st.session_state.get('construction', {}).items()
                            },
                            'saved_at': datetime.now().isoformat(),
                            'version': '5.1',   # ← เพิ่ม version marker
                        }
                        st.download_button(
                            "⬇️ ดาวน์โหลด JSON", data=json.dumps(data, ensure_ascii=False, indent=2),
                            file_name=f"Cost Est {st.session_state['project_info'].get('name', 'Project')}_{datetime.now().strftime('%Y%m%d_%H%M')}.json",
                            mime="application/json", use_container_width=True
                        )
                        st.success("✅ บันทึกโครงการสำเร็จ!")
        else:
            st.warning("⚠️ กรุณาเพิ่มข้อมูลโครงสร้างชั้นทางใน Tab 2 ก่อน")

    # ===== Tab 4: วิเคราะห์จากรูปภาพ =====
    with tab4:
        st.info("💡 Upload รูปภาพโครงสร้างชั้นทาง แล้วระบบจะวิเคราะห์และคำนวณราคาให้อัตโนมัติ")

        uploaded_image = st.file_uploader(
            "เลือกรูปภาพโครงสร้างชั้นทาง", type=['png', 'jpg', 'jpeg'],
            help="รองรับไฟล์ PNG, JPG, JPEG"
        )

        if uploaded_image is not None:
            col_img, col_result = st.columns([1, 1])

            with col_img:
                st.subheader("🖼️ รูปภาพที่ Upload")
                st.image(uploaded_image, use_container_width=True)

            with col_result:
                st.subheader("📋 กรอกข้อมูลโครงสร้างชั้นทาง")
                st.markdown("กรุณาตรวจสอบและแก้ไขข้อมูลที่อ่านจากรูปภาพ")

                structure_type = st.selectbox(
                    "ประเภทโครงสร้าง",
                    options=['AC Pavement', 'JPCP', 'JRCP', 'CRCP'],
                    key="img_structure_type"
                )

                num_layers = st.number_input(
                    "จำนวนชั้นโครงสร้าง", min_value=1, max_value=10, value=6,
                    key="img_num_layers"
                )

                st.divider()

                surface_materials = {
                    'AC Pavement': ['AC Wearing Course', 'PMA Wearing Course', 'AC Binder Course', 'AC Base Course', 'Tack Coat', 'Prime Coat'],
                    'JPCP': ['Concrete Slab (JPCP)', 'AC Interlayer', 'Non Woven Geotextile'],
                    'JRCP': ['Concrete Slab (JRCP)', 'AC Interlayer', 'Non Woven Geotextile'],
                    'CRCP': ['Concrete Slab (CRCP)', 'AC Interlayer', 'Steel Reinforcement', 'Non Woven Geotextile'],
                }

                base_materials_img = [
                    'Cement Treated Base (UCS 40 ksc)',
                    'Cement Modified Crushed Rock Base (UCS 24.5 ksc)',
                    'Crushed Rock Base Course',
                    'Soil Cement Subbase (UCS 7 ksc)',
                    'Soil Aggregate Subbase',
                    'Selected Material A',
                ]

                all_materials = surface_materials.get(structure_type, []) + base_materials_img

                if 'img_layers' not in st.session_state:
                    st.session_state['img_layers'] = []

                img_layers = []

                st.markdown("**รายละเอียดแต่ละชั้น:**")
                cols_h = st.columns([3, 1.5, 2])
                cols_h[0].markdown("**วัสดุ**")
                cols_h[1].markdown("**ความหนา (cm)**")
                cols_h[2].markdown("**ราคา (บาท/ตร.ม.)**")

                for i in range(int(num_layers)):
                    cols = st.columns([3, 1.5, 2])

                    default_materials_map = {
                        'AC Pavement': ['AC Wearing Course', 'AC Binder Course', 'AC Base Course', 'Cement Treated Base (UCS 40 ksc)', 'Soil Aggregate Subbase', 'Selected Material A'],
                        'JPCP': ['Concrete Slab (JPCP)', 'AC Interlayer', 'Cement Treated Base (UCS 40 ksc)', 'Crushed Rock Base Course', 'Soil Aggregate Subbase', 'Selected Material A'],
                        'JRCP': ['Concrete Slab (JRCP)', 'AC Interlayer', 'Cement Treated Base (UCS 40 ksc)', 'Crushed Rock Base Course', 'Soil Aggregate Subbase', 'Selected Material A'],
                        'CRCP': ['Concrete Slab (CRCP)', 'AC Interlayer', 'Cement Treated Base (UCS 40 ksc)', 'Crushed Rock Base Course', 'Soil Aggregate Subbase', 'Selected Material A'],
                    }
                    default_list = default_materials_map.get(structure_type, all_materials)
                    default_mat = default_list[i] if i < len(default_list) else all_materials[0]
                    try:
                        mat_idx = all_materials.index(default_mat)
                    except Exception:
                        mat_idx = 0

                    with cols[0]:
                        material = st.selectbox(
                            f"วัสดุชั้น {i+1}", options=all_materials, index=mat_idx,
                            key=f"img_mat_{i}", label_visibility="collapsed"
                        )

                    default_thicknesses_map = {
                        'AC Pavement': [5, 7, 8, 20, 25, 30],
                        'JPCP': [30, 5, 20, 15, 25, 30],
                        'JRCP': [30, 5, 20, 15, 25, 30],
                        'CRCP': [30, 5, 20, 15, 25, 30],
                    }
                    default_thick_list = default_thicknesses_map.get(structure_type, [20] * 10)
                    default_thick = default_thick_list[i] if i < len(default_thick_list) else 20

                    with cols[1]:
                        thickness = st.number_input(
                            f"หนา {i+1}", min_value=0.0, max_value=100.0,
                            value=float(default_thick), step=1.0,
                            key=f"img_thick_{i}", label_visibility="collapsed"
                        )

                    price_sqm = 0
                    mat_lower = material.lower()
                    if 'price_library' in st.session_state:
                        lib = st.session_state['price_library']
                        if 'ac wearing' in mat_lower:
                            prices = lib['ac_prices'].get('AC Wearing Course', {})
                            price_sqm = prices.get(thickness, 0)
                            if price_sqm == 0 and prices:
                                price_sqm = prices.get(min(prices.keys(), key=lambda x: abs(x - thickness)), 0)
                        elif 'pma' in mat_lower:
                            prices = lib['ac_prices'].get('PMA Wearing Course', {})
                            price_sqm = prices.get(thickness, 0)
                            if price_sqm == 0 and prices:
                                price_sqm = prices.get(min(prices.keys(), key=lambda x: abs(x - thickness)), 0)
                        elif 'binder' in mat_lower:
                            prices = lib['ac_prices'].get('AC Binder Course', {})
                            price_sqm = prices.get(thickness, 0)
                            if price_sqm == 0 and prices:
                                price_sqm = prices.get(min(prices.keys(), key=lambda x: abs(x - thickness)), 0)
                        elif 'ac base' in mat_lower or 'ac interlayer' in mat_lower:
                            prices = lib['ac_prices'].get('AC Base Course', {})
                            price_sqm = prices.get(thickness, 0)
                            if price_sqm == 0 and prices:
                                price_sqm = prices.get(min(prices.keys(), key=lambda x: abs(x - thickness)), 0)
                        elif 'tack' in mat_lower:
                            price_sqm = 20
                        elif 'prime' in mat_lower:
                            price_sqm = 30
                        elif 'geotextile' in mat_lower:
                            price_sqm = 78
                        elif 'steel' in mat_lower:
                            price_sqm = 200
                        elif 'concrete' in mat_lower or 'slab' in mat_lower:
                            if 'jpcp' in mat_lower:
                                prices = lib['concrete_prices'].get('JPCP', {})
                            elif 'jrcp' in mat_lower:
                                prices = lib['concrete_prices'].get('JRCP', {})
                            elif 'crcp' in mat_lower:
                                prices = lib['concrete_prices'].get('CRCP', {})
                            else:
                                prices = lib['concrete_prices'].get('JPCP', {})
                            price_sqm = prices.get(int(thickness), 0)
                            if price_sqm == 0 and prices:
                                price_sqm = prices.get(min(prices.keys(), key=lambda x: abs(x - thickness)), 0)
                        elif 'cement treated' in mat_lower:
                            price_sqm = lib['base_prices'].get('Cement Treated Base (UCS 40 ksc)', 1096) * thickness / 100
                        elif 'cement modified' in mat_lower:
                            price_sqm = lib['base_prices'].get('Cement Modified Crushed Rock Base (UCS 24.5 ksc)', 864) * thickness / 100
                        elif 'crushed rock' in mat_lower:
                            price_sqm = lib['base_prices'].get('Crushed Rock Base Course', 583) * thickness / 100
                        elif 'soil cement' in mat_lower:
                            price_sqm = lib['base_prices'].get('Soil Cement Subbase (UCS 7 ksc)', 854) * thickness / 100
                        elif 'soil aggregate' in mat_lower or 'aggregate subbase' in mat_lower:
                            price_sqm = lib['base_prices'].get('Soil Aggregate Subbase', 375) * thickness / 100
                        elif 'selected' in mat_lower:
                            price_sqm = lib['base_prices'].get('Selected Material A', 375) * thickness / 100

                    with cols[2]:
                        st.markdown(f"**{price_sqm:,.2f}**")

                    img_layers.append({'material': material, 'thickness': thickness, 'price_sqm': price_sqm})

                st.session_state['img_layers'] = img_layers

        if uploaded_image is not None and 'img_layers' in st.session_state and st.session_state['img_layers']:
            st.divider()
            st.subheader("📊 สรุปผลการวิเคราะห์")

            img_layers = st.session_state['img_layers']
            total_cost_sqm = sum(layer['price_sqm'] for layer in img_layers)

            summary_data_img = []
            for i, layer in enumerate(img_layers):
                summary_data_img.append({
                    'ลำดับ': i + 1,
                    'วัสดุ': layer['material'],
                    'ความหนา (cm)': layer['thickness'],
                    'ราคา (บาท/ตร.ม.)': f"{layer['price_sqm']:,.2f}",
                })

            st.dataframe(pd.DataFrame(summary_data_img), use_container_width=True, hide_index=True)

            col_m1, col_m2, col_m3 = st.columns(3)
            with col_m1:
                st.metric("💰 ราคารวม", f"{total_cost_sqm:,.2f} บาท/ตร.ม.")
            with col_m2:
                area_km = st.session_state.get('area_per_km', 22000)
                cost_per_km_img = total_cost_sqm * area_km / 1_000_000
                st.metric("📏 ราคาต่อ กม.", f"{cost_per_km_img:,.2f} ล้านบาท/กม.")
            with col_m3:
                structure_type_img = st.session_state.get('img_structure_type', 'JPCP')
                if 'AC' in structure_type_img:
                    design_life = 20
                elif 'CRCP' in structure_type_img:
                    design_life = 30
                else:
                    design_life = 25
                st.metric("⏱️ อายุออกแบบ", f"{design_life} ปี")

    st.divider()
    st.markdown("""
    <div style='text-align: center; color: #888; font-size: 0.85rem; padding: 20px;'>
        <b>พัฒนาโดย</b><br>
        รศ.ดร.อิทธิพล มีผล<br>
        ภาควิชาครุศาสตร์โยธา คณะครุศาสตร์อุตสาหกรรม<br>
        มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ (มจพ.)<br>
        <small style='color: #aaa;'>Pavement Structure Cost Analysis System v5.1 — Bug Fixed</small>
    </div>
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
